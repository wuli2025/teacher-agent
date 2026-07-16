# -*- coding: utf-8 -*-
"""第09讲：线性方程组、矩阵与线性变换 —— 60分钟课件包构建脚本

字号规范（deckkit 硬性下限）：
  · PPT 文字 >= 20pt（deckkit 构件内建）
  · 图内文字 >= 20pt（matplotlib 已被 monkeypatch 抬到 20）
  · 「不缩小」铁律：贴图宽度 >= PNG 画布宽度，故所有 new_fig/figsize 的宽度
    必须 <= 该图在幻灯片上的展示宽度；公式贴图（eqpic）一律 1:1 原尺寸贴。
"""
import sys, os
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon, FancyArrowPatch, FancyBboxPatch, Rectangle

OUT = '/mnt/c/Users/mi/Desktop/数学名师课件包/09_线性方程组与矩阵'
# 说明：本机无 LaTeX，matplotlib 只能用 mathtext（不支持 pmatrix/cases/array），
# 故矩阵型公式统一用下面的 eqrow() 精确绘制（括号+网格对齐），普通公式仍走 k.formula()。
FIG = os.path.join(OUT, 'figures')
os.makedirs(FIG, exist_ok=True)
F = lambda n: os.path.join(FIG, n)

ACC, ACC2, INK, GRN, RED = k.M_ACC, k.M_ACC2, k.M_INK, k.M_GRN, k.M_RED
RULE, PAPER, SLATE = k.M_RULE, k.M_PAPER, k.M_SLATE
FS = k.FIG_MIN_PT          # 20pt，图内统一字号

# ======================= 配图 =======================
figs = []

def _reg(fig, name):
    p = k.save_fig(fig, F(name)); figs.append(p); return p

def _reg_raw(fig, name):
    """按 1 数据单位 = 1 英寸 画的示意图，不能走 tight_layout。"""
    p = F(name)
    fig.savefig(p, dpi=200)
    plt.close(fig)
    figs.append(p)
    return p

def inch_fig(w, h):
    """返回 (fig, ax)，ax 铺满画布且 1 数据单位 = 1 英寸。"""
    fig = plt.figure(figsize=(w, h))
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis('off')
    ax.set_xlim(0, w); ax.set_ylim(0, h)
    ax.set_facecolor(PAPER)
    return fig, ax


# ---------- 公式排版引擎（矩阵/方程组，mathtext 无法胜任的部分） ----------
# 关键：eqrow 以 1 数据单位 = 1 英寸 作图，字号固定 20pt；eqpic 原尺寸贴图，
# 故幻灯片上的公式字号 = 20pt，绝不缩水。

def _mw(ax, rend, s, fs):
    """量出字符串在图上的宽度（英寸）。"""
    t = ax.text(0, -50, s, fontsize=fs, ha='left', va='center')
    bb = t.get_window_extent(rend).transformed(ax.transData.inverted())
    t.remove()
    return bb.width

def _draw_matrix(ax, rend, x, ycen, rows, fs, aug=None, br='()', col=INK):
    """画矩阵（或行列式），返回右边界 x。"""
    nr, nc = len(rows), len(rows[0])
    pad = 0.11
    ws = [max(_mw(ax, rend, str(rows[i][j]), fs) for i in range(nr)) + 2 * pad
          for j in range(nc)]
    rh = fs / 72.0 * 1.75
    top = ycen + nr * rh / 2
    bot = ycen - nr * rh / 2
    x0 = x + 0.16
    cx = x0
    for j in range(nc):
        for i in range(nr):
            yy = top - (i + 0.5) * rh
            c = ACC if (aug is not None and j >= aug) else col
            ax.text(cx + ws[j] / 2, yy, str(rows[i][j]), ha='center', va='center',
                    fontsize=fs, color=c)
        cx += ws[j]
        if aug is not None and j == aug - 1:
            ax.plot([cx + 0.07, cx + 0.07], [bot, top], color=RULE, lw=1.6, ls='--')
            cx += 0.14
    x1 = cx
    ext = 0.09
    for xb, s in ((x0 - 0.12, 1), (x1 + 0.12, -1)):
        if br == '()':
            ax.plot([xb + s * ext, xb, xb, xb + s * ext],
                    [top, top - 0.07, bot + 0.07, bot],
                    color=col, lw=1.8, solid_capstyle='round')
        else:                      # 行列式竖线
            ax.plot([xb, xb], [bot, top], color=col, lw=1.8)
    return x1 + 0.26

def _draw_sys(ax, rend, x, ycen, lines, fs, col=INK):
    """画带大括号的方程组。"""
    n = len(lines)
    rh = fs / 72.0 * 1.85
    top = ycen + n * rh / 2
    bot = ycen - n * rh / 2
    x0 = x + 0.26
    ax.plot([x0 - 0.10, x0 - 0.21, x0 - 0.21, x0 - 0.10],
            [top, top - 0.08, bot + 0.08, bot], color=col, lw=2.0, solid_capstyle='round')
    w = 0
    for i, ln in enumerate(lines):
        ax.text(x0, top - (i + 0.5) * rh, ln, ha='left', va='center', fontsize=fs, color=col)
        w = max(w, _mw(ax, rend, ln, fs))
    return x0 + w + 0.20

def eqrow(items, name, h=1.9, fs=FS):
    """items: ('t', mathtext[, color]) | ('m', rows[, aug]) | ('d', rows) 行列式
       | ('s', lines) 方程组 | ('a', 上标, 下标) 长箭头。
       1 数据单位 = 1 英寸，字号 20pt，输出透明 PNG（原尺寸贴片）。"""
    Wt = 40.0
    fig = plt.figure(figsize=(Wt, h))
    ax = fig.add_axes([0, 0, 1, 1]); ax.axis('off')
    ax.set_xlim(0, Wt); ax.set_ylim(0, h)
    fig.canvas.draw()
    rend = fig.canvas.get_renderer()
    x = 0.20; yc = h / 2
    for it in items:
        kind = it[0]
        if kind == 't':
            col = it[2] if len(it) > 2 else INK
            ax.text(x, yc, it[1], ha='left', va='center', fontsize=fs, color=col)
            x += _mw(ax, rend, it[1], fs) + 0.20
        elif kind == 'm':
            x = _draw_matrix(ax, rend, x, yc, it[1], fs,
                             aug=(it[2] if len(it) > 2 else None))
        elif kind == 'd':
            x = _draw_matrix(ax, rend, x, yc, it[1], fs, br='||')
        elif kind == 's':
            x = _draw_sys(ax, rend, x, yc, it[1], fs)
        elif kind == 'a':
            up, dn = it[1], it[2]
            lw_ = max(_mw(ax, rend, up, fs) if up else 0,
                      _mw(ax, rend, dn, fs) if dn else 0)
            wid = max(1.0, lw_ + 0.30)
            ax.add_patch(FancyArrowPatch((x + 0.10, yc), (x + wid, yc), arrowstyle='-|>',
                                         mutation_scale=18, color=ACC, lw=1.8))
            xm = x + (wid + 0.10) / 2
            if up:
                ax.text(xm, yc + 0.10, up, ha='center', va='bottom', fontsize=fs, color=ACC)
            if dn:
                ax.text(xm, yc - 0.10, dn, ha='center', va='top', fontsize=fs, color=ACC)
            x += wid + 0.24
    total = x + 0.20
    ax.set_xlim(0, total)
    fig.set_size_inches(total, h)
    p = F(name)
    fig.savefig(p, dpi=200, transparent=True)
    plt.close(fig)
    print("  eq %-24s  %.2f x %.2f in" % (name, total, h))
    return p


# --- fig01 二元方程组解的三种情形（展示宽度 10.2in） ---
fig, axes = plt.subplots(1, 3, figsize=(10.2, 3.3))
xg = np.linspace(-1, 5, 200)
data = [
    ("相交 → 唯一解", [(1, 1, 4, ACC), (1, -1, 0, ACC2)], (2, 2)),
    ("平行 → 无解", [(1, 1, 4, ACC), (1, 1, 1, ACC2)], None),
    ("重合 → 无穷多解", [(1, 1, 4, ACC), (2, 2, 8, ACC2)], None),
]
for ax, (t, lines, pt) in zip(axes, data):
    for i, (a, b, c, col) in enumerate(lines):
        ax.plot(xg, (c - a * xg) / b, color=col, lw=3.4 if i == 0 else 2.6,
                ls='-' if i == 0 else (0, (7, 4)), zorder=3)
    if pt:
        ax.plot(*pt, 'o', color=RED, ms=12, zorder=5)
    k.style_axes(ax)
    ax.set_xlim(-1, 5); ax.set_ylim(-1, 5)
    ax.set_xticks([2, 4]); ax.set_yticks([2, 4])
    ax.set_title(t, fontsize=FS, color=INK, pad=8)
_reg(fig, 'fig01_three_cases.png')


# --- fig02 行视角 vs 列视角（展示宽度 11.3in） ---
fig, axes = plt.subplots(1, 2, figsize=(11.3, 4.2))
ax = axes[0]
xs = np.linspace(-1, 4, 100)
ax.plot(xs, (4 - xs) / 2, color=ACC, lw=3.2)
ax.plot(xs, 3 * xs - 5, color=ACC2, lw=3.2)
ax.text(-0.9, 2.9, r"$x+2y=4$", fontsize=FS, color=ACC)
ax.text(2.55, -2.3, r"$3x-y=5$", fontsize=FS, color=ACC2)
ax.plot(2, 1, 'o', color=RED, ms=13, zorder=5)
ax.annotate(r"$(2,1)$", (2, 1), xytext=(2.5, 2.6), fontsize=FS, color=RED,
            arrowprops=dict(arrowstyle='->', color=RED, lw=1.8))
k.style_axes(ax); ax.set_xlim(-1, 4); ax.set_ylim(-3, 4)
ax.set_xticks([2]); ax.set_yticks([2])
ax.set_title("行视角：解 = 交点", fontsize=FS, color=ACC, pad=10)

ax = axes[1]
c1 = np.array([1, 3]); c2 = np.array([2, -1]); bvec = np.array([4, 5])
def arr(ax, s, e, col, lw=2.8, ls='-'):
    ax.add_patch(FancyArrowPatch(s, e, arrowstyle='-|>', mutation_scale=20,
                                 color=col, lw=lw, linestyle=ls, zorder=4))
arr(ax, (0, 0), 2 * c1, ACC, 3.4)
arr(ax, 2 * c1, 2 * c1 + c2, ACC2, 3.4)
arr(ax, (0, 0), bvec, RED, 3.6)
ax.text(0.15, 4.4, r"$2\mathbf{a}_1$", fontsize=FS, color=ACC)
ax.text(2.75, 6.0, r"$1\mathbf{a}_2$", fontsize=FS, color=ACC2)
ax.text(2.5, 2.0, r"$\mathbf{b}$", fontsize=FS, color=RED)
k.style_axes(ax); ax.set_xlim(-1, 6); ax.set_ylim(-2, 7.6)
ax.set_xticks([2, 4]); ax.set_yticks([2, 4, 6])
ax.set_title("列视角：解 = 配出 b 的系数", fontsize=FS, color=ACC2, pad=10)
_reg(fig, 'fig02_row_vs_col.png')


# --- fig03 列空间（展示宽度 6.0in） ---
fig, axes = plt.subplots(1, 2, figsize=(6.0, 3.0))
ax = axes[0]
c1 = np.array([1, 3]); c2 = np.array([2, -1])
for s in np.linspace(-1.2, 1.6, 13):
    for t in np.linspace(-1.2, 1.6, 13):
        ax.plot(*(s * c1 + t * c2), '.', color=RULE, ms=5, zorder=1)
arr(ax, (0, 0), c1, ACC); arr(ax, (0, 0), c2, ACC2)
arr(ax, (0, 0), np.array([4, 5]), RED, 3.0)
ax.text(1.2, 5.6, "b 有解", fontsize=FS, color=RED)
ax.set_xlim(-4, 7); ax.set_ylim(-5, 7.6)
ax.set_xticks([]); ax.set_yticks([])
for sp in ('top', 'right'):
    ax.spines[sp].set_visible(False)
for sp in ('left', 'bottom'):
    ax.spines[sp].set_position('zero'); ax.spines[sp].set_color(SLATE)
ax.set_title("不共线：铺满平面", fontsize=FS, color=GRN, pad=8)

ax = axes[1]
tt = np.linspace(-3, 3, 50)
arr(ax, (0, 0), np.array([1, 2]), ACC); arr(ax, (0, 0), np.array([2, 4]), ACC2)
ax.plot(tt, tt * 2, color=RULE, lw=8, alpha=0.9, zorder=1)
arr(ax, (0, 0), np.array([4, 5]), RED, 3.0)
ax.text(1.2, 5.6, "b 无解", fontsize=FS, color=RED)
ax.set_xlim(-4, 7); ax.set_ylim(-5, 7.6)
ax.set_xticks([]); ax.set_yticks([])
for sp in ('top', 'right'):
    ax.spines[sp].set_visible(False)
for sp in ('left', 'bottom'):
    ax.spines[sp].set_position('zero'); ax.spines[sp].set_color(SLATE)
ax.set_title("共线：压成直线", fontsize=FS, color=RED, pad=8)
_reg(fig, 'fig03_colspace.png')


# --- fig04 三种初等行变换（展示宽度 5.9in） ---
fig, ax = inch_fig(5.9, 2.6)
items = [("① 换行", r"$r_i \leftrightarrow r_j$", ACC),
         ("② 乘 k", r"$k\cdot r_i\ (k\neq 0)$", ACC2),
         ("③ 倍加", r"$r_i + k\, r_j$", GRN)]
for i, (t, tex_, col) in enumerate(items):
    x0 = 0.12 + i * 1.92
    ax.add_patch(FancyBboxPatch((x0, 0.30), 1.72, 1.55, boxstyle="round,pad=0.05",
                                fc='white', ec=col, lw=2))
    ax.text(x0 + 0.86, 1.52, t, ha='center', va='center', fontsize=FS, color=col, weight='bold')
    ax.text(x0 + 0.86, 0.85, tex_, ha='center', va='center', fontsize=FS, color=INK)
ax.text(2.95, 2.22, "三招：换 · 乘 · 倍加（解集不变）", ha='center', va='center',
        fontsize=FS, color=INK, weight='bold')
_reg_raw(fig, 'fig04_row_ops.png')


# --- fig05 高斯消元流程（展示宽度 12.0in，1 单位 = 1 英寸） ---
fig, ax = inch_fig(12.0, 4.4)
M = [
    ([[1, 2, 1, 8], [2, -1, 1, 3], [1, 1, -1, 0]], "增广矩阵", ACC2, ""),
    ([[1, 2, 1, 8], [0, -5, -1, -13], [0, -1, -2, -8]], "第1列消元", ACC, r"$r_2-2r_1,\ r_3-r_1$"),
    ([[1, 2, 1, 8], [0, -5, -1, -13], [0, 0, -9, -27]], "第2列消元", ACC, r"$r_3-\frac{1}{5}r_2$"),
    ([[1, 0, 0, 1], [0, 1, 0, 2], [0, 0, 1, 3]], "最简形", GRN, r"$x=1,\ y=2,\ z=3$"),
]
BW, BH = 2.50, 1.60
CX = [1.42, 4.47, 7.52, 10.57]
def draw_mat(ax, cx, cy, m, col, title, sub):
    ax.add_patch(FancyBboxPatch((cx - BW / 2, cy - BH / 2), BW, BH,
                                boxstyle="round,pad=0.06", fc='white', ec=col, lw=2))
    cw = (BW - 0.30 - 0.10) / 4.0
    rh = 0.44
    for i, row in enumerate(m):
        for j, v in enumerate(row):
            xx = cx - BW / 2 + 0.15 + (j + 0.5) * cw + (0.10 if j == 3 else 0)
            yy = cy + BH / 2 - 0.25 - i * rh
            isz = (v == 0 and j < 3 and i > j)
            ax.text(xx, yy, f"{v}", ha='center', va='center', fontsize=FS,
                    color=RED if isz else INK)
    xs = cx - BW / 2 + 0.15 + 3 * cw + 0.03
    ax.plot([xs, xs], [cy - BH / 2 + 0.10, cy + BH / 2 - 0.10], color=RULE, lw=1.6, ls='--')
    ax.text(cx, cy + BH / 2 + 0.22, title, ha='center', va='center', fontsize=FS,
            color=col, weight='bold')
    if sub:
        ax.text(cx, cy - BH / 2 - 0.28, sub, ha='center', va='center', fontsize=FS, color=SLATE)
for i, (m, t, c, s) in enumerate(M):
    draw_mat(ax, CX[i], 2.35, m, c, t, s)
    if i < 3:
        ax.add_patch(FancyArrowPatch((CX[i] + BW / 2 + 0.06, 2.35),
                                     (CX[i + 1] - BW / 2 - 0.06, 2.35),
                                     arrowstyle='-|>', mutation_scale=18, color=ACC, lw=2.2))
ax.text(6.0, 4.10, "高斯消元：把增广矩阵压成阶梯形（红 0 = 已消掉）", ha='center',
        va='center', fontsize=FS, color=INK, weight='bold')
ax.text(6.0, 0.30, "全程只用三种初等行变换 → 与原方程组同解", ha='center', va='center',
        fontsize=FS, color=GRN)
_reg_raw(fig, 'fig05_gauss_flow.png')


# --- fig06 阶梯形示意（展示宽度 >= 5.5in） ---
fig, ax = inch_fig(5.5, 3.4)
rows = [[('1', 1), ('2', 0), ('1', 0), ('|', -1), ('8', 0)],
        [('0', 2), ('-5', 1), ('-1', 0), ('|', -1), ('-13', 0)],
        [('0', 2), ('0', 2), ('-9', 1), ('|', -1), ('-27', 0)]]
for i, row in enumerate(rows):
    for j, (v, kind) in enumerate(row):
        x0 = 0.38 + j * 0.50; y0 = 2.30 - i * 0.55
        if kind == -1:
            ax.text(x0, y0, '|', ha='center', va='center', fontsize=FS, color=RULE); continue
        if kind == 1:
            ax.add_patch(Rectangle((x0 - 0.24, y0 - 0.22), 0.48, 0.44,
                                   fc='white', ec=ACC, lw=2.2))
        c = GRN if kind == 2 else (ACC if kind == 1 else INK)
        ax.text(x0, y0, v, ha='center', va='center', fontsize=FS, color=c,
                weight='bold' if kind == 1 else 'normal')
ax.plot([0.08, 0.66, 0.66, 1.16, 1.16, 1.66, 1.66],
        [2.64, 2.64, 2.10, 2.10, 1.55, 1.55, 1.00], color=RED, lw=2.4, ls='--')
ax.text(3.05, 2.30, "橙框 = 主元", ha='left', va='center', fontsize=FS, color=ACC)
ax.text(3.05, 1.75, "红线 = 阶梯", ha='left', va='center', fontsize=FS, color=RED)
ax.text(3.05, 1.20, "绿 0 = 已消", ha='left', va='center', fontsize=FS, color=GRN)
ax.text(2.75, 0.45, r"$3$ 个主元 $\Rightarrow r(A)=3$", ha='center', va='center',
        fontsize=FS, color=INK)
ax.text(2.75, 3.05, "行阶梯形：主元与秩", ha='center', va='center', fontsize=FS,
        color=INK, weight='bold')
_reg_raw(fig, 'fig06_echelon.png')


# --- fig07 三种阶梯形 ⇔ 解的三种情况（展示宽度 11.6in） ---
fig, axes = plt.subplots(1, 3, figsize=(11.6, 4.0))
cases = [
    ("唯一解", [["1", "*", "*", "|", "*"], ["0", "1", "*", "|", "*"], ["0", "0", "1", "|", "*"]],
     r"$r(A)=r(\bar A)=n$", GRN),
    ("无穷多解", [["1", "*", "*", "|", "*"], ["0", "1", "*", "|", "*"], ["0", "0", "0", "|", "0"]],
     r"$r(A)=r(\bar A)<n$", ACC),
    ("无解", [["1", "*", "*", "|", "*"], ["0", "1", "*", "|", "*"], ["0", "0", "0", "|", "5"]],
     r"$r(A)<r(\bar A)$", RED),
]
for ax, (t, m, cond, col) in zip(axes, cases):
    ax.axis('off'); ax.set_xlim(0, 100); ax.set_ylim(0, 100)
    ax.add_patch(FancyBboxPatch((8, 30), 84, 46, boxstyle="round,pad=2",
                                fc='white', ec=col, lw=2.2))
    for i, row in enumerate(m):
        for j, v in enumerate(row):
            xx = 21 + j * 15; yy = 66 - i * 15
            ax.text(xx, yy, v, ha='center', va='center', fontsize=FS,
                    color=RULE if v == '|' else (col if i == 2 else INK))
    ax.text(50, 90, t, ha='center', va='center', fontsize=FS, color=col, weight='bold')
    ax.text(50, 14, cond, ha='center', va='center', fontsize=FS, color=INK)
_reg(fig, 'fig07_rank_cases.png')


# --- fig08 秩的直观（展示宽度 6.5in） ---
fig, axes = plt.subplots(1, 2, figsize=(6.5, 3.4))
g = np.linspace(-2, 2, 9)
def grid_pts():
    P = []
    for a in g:
        P.append(np.array([[a, t] for t in np.linspace(-2, 2, 40)]))
        P.append(np.array([[t, a] for t in np.linspace(-2, 2, 40)]))
    return P
sq = np.array([[0, 0], [1, 0], [1, 1], [0, 1]])
for ax, (A, t, col) in zip(axes, [
        (np.array([[1., 0.6], [0.3, 1.2]]), r"满秩 $r=2$", GRN),
        (np.array([[1., 2.], [2., 4.]]), r"降秩 $r=1$", RED)]):
    for P in grid_pts():
        Q = (P * 1.6) @ A.T
        ax.plot(Q[:, 0], Q[:, 1], color=RULE, lw=1.2, zorder=1)
    for v, c in [(np.array([1, 0]), ACC), (np.array([0, 1]), ACC2)]:
        w = A @ v
        ax.add_patch(FancyArrowPatch((0, 0), 2.2 * w, arrowstyle='-|>', mutation_scale=18,
                                     color=c, lw=3.4, zorder=5))
    ax.set_xlim(-5, 5); ax.set_ylim(-5, 5); ax.set_aspect('equal')
    ax.set_xticks([]); ax.set_yticks([])
    ax.axhline(0, color=SLATE, lw=1.0); ax.axvline(0, color=SLATE, lw=1.0)
    for s in ax.spines.values():
        s.set_color(RULE)
    ax.set_title(t, fontsize=FS, color=col, pad=8)
_reg(fig, 'fig08_rank_squash.png')


# --- fig09 矩阵的列 = 基向量的像（展示宽度 5.6in） ---
fig, axes = plt.subplots(1, 2, figsize=(5.6, 3.2))
A = np.array([[2., 1.], [1., 3.]])
for ax, (P, t, col) in zip(axes, [(sq, "变换前", ACC2), (sq @ A.T, "变换后", ACC)]):
    ax.add_patch(Polygon(P, closed=True, fc=col, alpha=0.18, ec=col, lw=2.2, zorder=2))
    e1 = P[1]; e2 = P[3]
    ax.add_patch(FancyArrowPatch((0, 0), e1, arrowstyle='-|>', mutation_scale=18,
                                 color=ACC, lw=3, zorder=5))
    ax.add_patch(FancyArrowPatch((0, 0), e2, arrowstyle='-|>', mutation_scale=18,
                                 color=ACC2, lw=3, zorder=5))
    if col == ACC2:
        ax.text(0.55, -0.75, r"$\mathbf{e}_1$", fontsize=FS, color=ACC, ha='center')
        ax.text(-0.6, 1.1, r"$\mathbf{e}_2$", fontsize=FS, color=ACC2, ha='center')
    else:
        ax.text(2.3, 0.5, r"$(2,1)$", fontsize=FS, color=ACC)
        ax.text(-0.9, 3.2, r"$(1,3)$", fontsize=FS, color=ACC2)
    ax.set_xlim(-1.6, 4.2); ax.set_ylim(-1.4, 4.4); ax.set_aspect('equal')
    ax.set_xticks([]); ax.set_yticks([])
    ax.axhline(0, color=SLATE, lw=1.0); ax.axvline(0, color=SLATE, lw=1.0)
    for s in ax.spines.values():
        s.set_color(RULE)
    ax.set_title(t, fontsize=FS, color=col, pad=8)
_reg(fig, 'fig09_columns_are_images.png')


# --- fig10 ★ 网格变形四联图（展示宽度 12.3in，本讲高光） ---
th = np.pi / 6
mats = [
    (np.array([[np.cos(th), -np.sin(th)], [np.sin(th), np.cos(th)]]), "旋转 30°",
     r"$\det = 1$", "面积不变", ACC2),
    (np.array([[1., 1.2], [0., 1.]]), "剪切",
     r"$\det = 1$", "形状歪斜", ACC),
    (np.array([[1.8, 0.], [0., 0.6]]), "伸缩",
     r"$\det = 1.08$", "面积略变", GRN),
    (np.array([[1., 0.], [0., 0.]]), "投影",
     r"$\det = 0$", "压扁·不可逆", RED),
]
fig, axes = plt.subplots(2, 4, figsize=(12.3, 5.6))
gg = np.linspace(-2, 2, 9)
base_lines = []
for a in gg:
    base_lines.append(np.array([[a, t] for t in np.linspace(-2, 2, 60)]))
    base_lines.append(np.array([[t, a] for t in np.linspace(-2, 2, 60)]))
for j, (A, name, dtx, note, col) in enumerate(mats):
    for r in (0, 1):
        ax = axes[r][j]
        for P in base_lines:
            Q = P @ A.T if r == 1 else P
            ax.plot(Q[:, 0], Q[:, 1], color=RULE, lw=0.9, zorder=1)
        S = sq @ A.T if r == 1 else sq
        ax.add_patch(Polygon(S, closed=True, fc=col, alpha=0.22, ec=col, lw=2.2, zorder=3))
        for v, c in [(np.array([1., 0.]), ACC), (np.array([0., 1.]), ACC2)]:
            w = A @ v if r == 1 else v
            if np.linalg.norm(w) > 1e-9:
                ax.add_patch(FancyArrowPatch((0, 0), w, arrowstyle='-|>', mutation_scale=16,
                                             color=c, lw=2.6, zorder=6))
        ax.set_xlim(-2.6, 2.6); ax.set_ylim(-2.6, 2.6); ax.set_aspect('equal')
        ax.set_xticks([]); ax.set_yticks([])
        for s in ax.spines.values():
            s.set_color(RULE)
        ax.axhline(0, color=SLATE, lw=1.0); ax.axvline(0, color=SLATE, lw=1.0)
        if r == 0:
            ax.set_title(name, fontsize=FS, color=col, pad=8)
        else:
            ax.text(0.5, -0.12, dtx, transform=ax.transAxes, ha='center', va='top',
                    fontsize=FS, color=INK)
            ax.text(0.5, -0.34, note, transform=ax.transAxes, ha='center', va='top',
                    fontsize=FS, color=col)
fig.suptitle("矩阵 = 线性变换：同一张方格网的四种“捏法”", fontsize=FS, color=INK, y=0.975)
fig.text(0.5, 0.495, "上：变换前　　下：变换后", ha='center', va='center',
         fontsize=FS, color=SLATE)
fig.subplots_adjust(hspace=0.80, wspace=0.18, top=0.845, bottom=0.20, left=0.02, right=0.98)
fig.savefig(F('fig10_grid_transforms.png'), dpi=200)
plt.close(fig)
figs.append(F('fig10_grid_transforms.png'))


# --- fig11 复合变换 = 矩阵乘法（展示宽度 11.6in） ---
fig, axes = plt.subplots(1, 3, figsize=(11.6, 4.0))
Sm = np.array([[1., 1.0], [0., 1.]])
Rm = np.array([[0., -1.], [1., 0.]])
stages = [(np.eye(2), r"原图 $\mathbf{x}$", ACC2), (Sm, r"先剪切 $S\mathbf{x}$", ACC),
          (Rm @ Sm, r"再旋转 $(RS)\mathbf{x}$", GRN)]
for ax, (A, t, col) in zip(axes, stages):
    for P in base_lines:
        Q = P @ A.T
        ax.plot(Q[:, 0], Q[:, 1], color=RULE, lw=0.9, zorder=1)
    ax.add_patch(Polygon(sq @ A.T, closed=True, fc=col, alpha=0.22, ec=col, lw=2.2, zorder=3))
    ax.set_xlim(-3, 3); ax.set_ylim(-3, 3); ax.set_aspect('equal')
    ax.set_xticks([]); ax.set_yticks([])
    ax.axhline(0, color=SLATE, lw=1); ax.axvline(0, color=SLATE, lw=1)
    for s in ax.spines.values():
        s.set_color(RULE)
    ax.set_title(t, fontsize=FS, color=col, pad=8)
fig.suptitle("复合变换 = 矩阵乘法（右边先作用，顺序不能颠倒）", fontsize=FS, color=INK)
_reg(fig, 'fig11_composition.png')


# --- fig12 行列式 = 面积伸缩因子（展示宽度 10.4in） ---
fig, axes = plt.subplots(1, 2, figsize=(10.4, 4.0))
A = np.array([[3., 1.], [1., 2.]])
ax = axes[0]
ax.add_patch(Polygon(sq, closed=True, fc=ACC2, alpha=0.25, ec=ACC2, lw=2.6))
ax.text(2.5, 1.6, "面积 1", ha='center', va='center', fontsize=FS, color=ACC2)
k.style_axes(ax); ax.set_xlim(-0.8, 4.6); ax.set_ylim(-0.8, 3.8); ax.set_aspect('equal')
ax.set_xticks([2, 4]); ax.set_yticks([2])
ax.set_title("单位正方形", fontsize=FS, color=INK, pad=8)
ax = axes[1]
P = sq @ A.T
ax.add_patch(Polygon(P, closed=True, fc=ACC, alpha=0.25, ec=ACC, lw=2.6))
ax.text(2.0, 1.5, "面积 5", ha='center', va='center', fontsize=FS, color=ACC)
for v, c, lb in [(np.array([1., 0.]), ACC, r"$(3,1)$"), (np.array([0., 1.]), ACC2, r"$(1,2)$")]:
    w = A @ v
    ax.add_patch(FancyArrowPatch((0, 0), w, arrowstyle='-|>', mutation_scale=18,
                                 color=c, lw=2.8, zorder=6))
    ax.text(w[0] + 0.12, w[1] + 0.12, lb, fontsize=FS, color=c)
k.style_axes(ax); ax.set_xlim(-0.8, 4.6); ax.set_ylim(-0.8, 3.8); ax.set_aspect('equal')
ax.set_xticks([2, 4]); ax.set_yticks([2])
ax.set_title(r"像：面积 $=|\det A|$", fontsize=FS, color=INK, pad=8)
fig.suptitle(r"$\det A = 3\times 2 - 1\times 1 = 5$：面积放大 5 倍", fontsize=FS, color=INK)
_reg(fig, 'fig12_det_area.png')


# --- fig12s 面积伸缩（小版，展示宽度 4.2in，用于变式2） ---
fig, ax = plt.subplots(figsize=(4.2, 2.8))
ax.add_patch(Polygon(sq, closed=True, fc=ACC2, alpha=0.3, ec=ACC2, lw=2.4, zorder=3))
ax.add_patch(Polygon(sq @ A.T, closed=True, fc=ACC, alpha=0.2, ec=ACC, lw=2.4, zorder=2))
ax.text(0.5, 0.45, "1", ha='center', va='center', fontsize=FS, color=ACC2)
ax.text(2.4, 1.6, "5", ha='center', va='center', fontsize=FS, color=ACC)
k.style_axes(ax); ax.set_xlim(-0.6, 4.6); ax.set_ylim(-0.6, 3.6); ax.set_aspect('equal')
ax.set_xticks([2, 4]); ax.set_yticks([2])
ax.set_title(r"面积 $\times|\det A|$", fontsize=FS, color=INK, pad=8)
_reg(fig, 'fig12s_det_area_small.png')


# --- fig13 det<0 翻转 与 det=0 压扁（展示宽度 11.6in） ---
Fshape = np.array([[0, 0], [0, 1.6], [1.0, 1.6], [1.0, 1.25], [0.35, 1.25],
                   [0.35, 0.95], [0.85, 0.95], [0.85, 0.6], [0.35, 0.6], [0.35, 0]])
cfg = [
    (np.eye(2), r"$\det=+1$", GRN, "保向"),
    (np.array([[1., 0.], [0., -1.]]), r"$\det=-1$", ACC, "翻面"),
    (np.array([[1., 2.], [0.5, 1.]]), r"$\det=0$", RED, "压扁成线"),
]
def draw_detsign(figsize, with_sup, short=False):
    fig, axes = plt.subplots(1, 3, figsize=figsize)
    for ax, (A, t, col, note) in zip(axes, cfg):
        if short:
            t, note = {"保向": ("+1", "保向"), "翻面": ("−1", "翻面"),
                       "压扁成线": ("0", "压扁")}[note]
        Q = Fshape @ A.T
        ax.add_patch(Polygon(Q, closed=True, fc=col, alpha=0.3, ec=col, lw=2.4, zorder=3))
        ax.add_patch(Polygon(Fshape, closed=True, fc='none', ec=RULE, lw=1.6, ls='--', zorder=2))
        ax.set_xlim(-2.4, 4.4); ax.set_ylim(-2.8, 2.8); ax.set_aspect('equal')
        ax.set_xticks([]); ax.set_yticks([])
        ax.axhline(0, color=SLATE, lw=1); ax.axvline(0, color=SLATE, lw=1)
        for s in ax.spines.values():
            s.set_color(RULE)
        ax.set_title(t + "  " + note, fontsize=FS, color=col, pad=8)
    if with_sup:
        fig.suptitle("正 = 保向，负 = 翻面，零 = 压扁（虚线为原图 F）", fontsize=FS, color=INK)
    return fig
_reg(draw_detsign((11.6, 4.0), True), 'fig13_det_sign.png')
_reg(draw_detsign((4.6, 2.4), False, short=True), 'fig13s_det_sign_small.png')


# --- fig14 四件事串一条线（展示宽度 12.0in，1 单位 = 1 英寸） ---
fig, ax = inch_fig(12.0, 4.0)
nodes = [(1.55, r"$\det A=0$", RED), (4.60, "空间被压扁", ACC),
         (7.65, r"$A$ 不可逆", ACC2), (10.70, "无唯一解", GRN)]
NW, NH = 2.30, 0.95
for cx, t, col in nodes:
    ax.add_patch(FancyBboxPatch((cx - NW / 2, 2.10 - NH / 2), NW, NH,
                                boxstyle="round,pad=0.06", fc='white', ec=col, lw=2.4))
    ax.text(cx, 2.10, t, ha='center', va='center', fontsize=FS, color=col, weight='bold')
for i in range(3):
    x0 = nodes[i][0] + NW / 2 + 0.06; x1 = nodes[i + 1][0] - NW / 2 - 0.06
    ax.add_patch(FancyArrowPatch((x0, 2.10), (x1, 2.10), arrowstyle='<|-|>',
                                 mutation_scale=12, color=INK, lw=1.6))
    ax.text((x0 + x1) / 2, 2.52, "⇔", ha='center', va='center', fontsize=FS, color=INK)
ax.text(6.0, 3.55, "一条线串起四件事（等价）", ha='center', va='center', fontsize=FS,
        color=INK, weight='bold')
ax.text(6.0, 0.85, r"反过来：$\det A\neq 0\ \Leftrightarrow$ 不压扁 $\Leftrightarrow$ 可逆"
                   r"$\ \Leftrightarrow\ \mathbf{x}=A^{-1}\mathbf{b}$ 唯一",
        ha='center', va='center', fontsize=FS, color=GRN)
_reg_raw(fig, 'fig14_four_equiv.png')


# --- fig15 逆矩阵 = 逆变换（展示宽度 6.4in） ---
fig, axes = plt.subplots(1, 3, figsize=(6.4, 2.8))
Ai_A = np.array([[2., 1.], [1., 1.5]])
for ax, (M_, t, col) in zip(axes, [(np.eye(2), "原图", ACC2), (Ai_A, r"作用 $A$", ACC),
                                   (np.eye(2), r"再作用 $A^{-1}$", GRN)]):
    for P in base_lines:
        Q = P @ M_.T
        ax.plot(Q[:, 0], Q[:, 1], color=RULE, lw=0.9)
    ax.add_patch(Polygon(sq @ M_.T, closed=True, fc=col, alpha=0.25, ec=col, lw=2.4, zorder=3))
    ax.set_xlim(-3, 3.5); ax.set_ylim(-3, 3.5); ax.set_aspect('equal')
    ax.set_xticks([]); ax.set_yticks([])
    ax.axhline(0, color=SLATE, lw=1); ax.axvline(0, color=SLATE, lw=1)
    for s in ax.spines.values():
        s.set_color(RULE)
    ax.set_title(t, fontsize=FS, color=col, pad=6)
_reg(fig, 'fig15_inverse.png')


# --- fig16 特征向量（展示宽度 6.5in） ---
fig, ax = plt.subplots(figsize=(6.5, 4.0))
A = np.array([[2., 1.], [1., 2.]])
for ang in np.linspace(0, 2 * np.pi, 13)[:-1]:
    v = np.array([np.cos(ang), np.sin(ang)]) * 1.6
    w = A @ v
    ax.add_patch(FancyArrowPatch((0, 0), v, arrowstyle='-|>', mutation_scale=10,
                                 color=RULE, lw=1.4, zorder=2))
    ax.add_patch(FancyArrowPatch(v, w, arrowstyle='-|>', mutation_scale=10,
                                 color=SLATE, lw=1.0, ls=(0, (3, 3)), zorder=2))
for v, lam, col, off in [(np.array([1.6, 1.6]), 3, RED, (0.2, 0.3)),
                         (np.array([1.6, -1.6]), 1, GRN, (0.4, -0.9))]:
    w = A @ v
    ax.plot([-w[0], w[0]], [-w[1], w[1]], color=col, lw=1.4, ls='--', alpha=0.6, zorder=1)
    ax.add_patch(FancyArrowPatch((0, 0), v, arrowstyle='-|>', mutation_scale=16,
                                 color=col, lw=2.4, zorder=6))
    ax.add_patch(FancyArrowPatch((0, 0), w, arrowstyle='-|>', mutation_scale=20,
                                 color=col, lw=3.4, zorder=5))
    ax.text(w[0] + off[0], w[1] + off[1], rf"$\lambda={lam}$", fontsize=FS, color=col)
k.style_axes(ax, xlabel="x", ylabel=""); ax.set_xlim(-6.5, 6.5); ax.set_ylim(-6.5, 6.5)
ax.set_aspect('equal')
ax.set_xticks([-4, 4]); ax.set_yticks([-4, 4])
ax.set_title(r"$A\mathbf{v}=\lambda\mathbf{v}$：方向不变的轴", fontsize=FS, color=INK, pad=14)
_reg(fig, 'fig16_eigen.png')


# --- fig17 知识地图（展示宽度 11.4in，1 单位 = 1 英寸） ---
fig, ax = inch_fig(11.4, 4.2)
ax.add_patch(FancyBboxPatch((4.55, 1.75), 2.30, 0.90, boxstyle="round,pad=0.08",
                            fc=ACC, ec=ACC, lw=2))
ax.text(5.70, 2.20, "矩阵 = 变换", ha='center', va='center', fontsize=FS,
        color='white', weight='bold')
leaves = [
    (1.90, 3.72, "列 = 基的像", ACC2, 1.15),
    (5.70, 3.72, "高斯消元 → 阶梯形", GRN, 1.75),
    (9.50, 3.72, "秩 = 压掉几维", ACC2, 1.30),
    (1.65, 0.48, r"$\det$ = 面积因子", RED, 1.40),
    (5.70, 0.48, r"$\det=0\ \Leftrightarrow$ 压扁 $\Leftrightarrow$ 无唯一解", RED, 2.45),
    (9.75, 0.48, "特征向量（下讲）", SLATE, 1.50),
]
for cx, cy, t, col, hw in leaves:
    ax.plot([5.70, cx], [2.20, cy], color=RULE, lw=1.6, zorder=0)
    ax.add_patch(FancyBboxPatch((cx - hw, cy - 0.32), 2 * hw, 0.64,
                                boxstyle="round,pad=0.06", fc='white', ec=col, lw=2))
    ax.text(cx, cy, t, ha='center', va='center', fontsize=FS, color=col)
_reg_raw(fig, 'fig17_map.png')


# --- fig17m 知识地图（小版，展示宽度 5.6in） ---
fig, ax = inch_fig(5.6, 3.4)
ax.add_patch(FancyBboxPatch((1.75, 1.40), 2.10, 0.72, boxstyle="round,pad=0.06",
                            fc=ACC, ec=ACC, lw=2))
ax.text(2.80, 1.76, "矩阵 = 变换", ha='center', va='center', fontsize=FS,
        color='white', weight='bold')
mini = [(1.05, 2.90, "列 = 像", ACC2), (4.20, 2.90, "秩", ACC2),
        (1.05, 0.55, r"$\det$ = 面积", RED), (4.20, 0.55, "消元", GRN)]
for cx, cy, t, col in mini:
    hw = 0.95
    ax.plot([2.80, cx], [1.76, cy], color=RULE, lw=1.6, zorder=0)
    ax.add_patch(FancyBboxPatch((cx - hw, cy - 0.30), 2 * hw, 0.60,
                                boxstyle="round,pad=0.05", fc='white', ec=col, lw=2))
    ax.text(cx, cy, t, ha='center', va='center', fontsize=FS, color=col)
_reg_raw(fig, 'fig17m_map_mini.png')


# --- fig18 应用：例4 的方程组（展示宽度 6.3in，1 单位 = 1 英寸） ---
fig, ax = inch_fig(6.3, 3.2)
ax.text(3.15, 2.95, "例4：三种票 → 3×3 方程组", ha='center', va='center', fontsize=FS,
        color=INK, weight='bold')
ax.add_patch(FancyBboxPatch((0.15, 0.55), 6.0, 2.05, boxstyle="round,pad=0.06",
                            fc='white', ec=ACC2, lw=2))
for j, hd in enumerate([r"$x$", r"$y$", r"$z$"]):
    ax.text(1.55 + j * 0.75, 2.30, hd, ha='center', va='center', fontsize=FS, color=ACC2)
rows = [("张数", "1", "1", "1", "= 100"),
        ("收入", "30", "50", "80", "= 5400"),
        ("约束", "1", "-2", "0", "= 0")]
for i, (nm, a, b, c, rhs) in enumerate(rows):
    y = 1.80 - i * 0.48
    ax.text(0.60, y, nm, ha='center', va='center', fontsize=FS, color=SLATE)
    for j, v in enumerate([a, b, c]):
        ax.text(1.55 + j * 0.75, y, v, ha='center', va='center', fontsize=FS, color=INK)
    ax.text(4.85, y, rhs, ha='center', va='center', fontsize=FS, color=ACC)
ax.text(3.15, 0.22, r"$\det A\neq 0\ \Rightarrow$ 唯一解", ha='center', va='center',
        fontsize=FS, color=GRN)
_reg_raw(fig, 'fig18_application.png')

print("figures:", len(figs))

# ======================= PPT =======================
prs = k.new_deck()
TMP = os.path.join(FIG, '_tex')
os.makedirs(TMP, exist_ok=True)
_c = [0]

from PIL import Image
def eqpic(slide, path, y, maxw=12.5, x=None):
    """公式贴图：一律原尺寸（1:1）贴入，绝不缩小，故图内 20pt = 幻灯片 20pt。"""
    iw, ih = Image.open(path).size
    w = iw / 200.0
    if w > maxw:
        raise ValueError(f"{os.path.basename(path)} 天然宽 {w:.2f}in > 可用 {maxw:.2f}in，请拆行")
    xx = (13.333 - w) / 2 if x is None else x
    return slide.shapes.add_picture(path, k.Inches(xx), k.Inches(y), width=k.Inches(w))

def tex(slide, t, **kw):
    _c[0] += 1
    return k.formula(slide, t, out=os.path.join(TMP, f'tex{_c[0]:02d}.png'), **kw)

# 1 封面
k.title_slide(prs, "第09讲　线性方程组、矩阵与线性变换",
              "矩阵不是数表——它是一次“对空间的动作”",
              "数学名师课件包", "60 分钟 · 高中衔接/大学线性代数入门")

# 2 学习目标
s = k.content_slide(prs, "学习目标", "导入")
k.bullets(s, [
    "会用高斯消元法规范求解 3×3 线性方程组，掌握三种初等行变换",
    "能由行阶梯形判别解的存在唯一性（唯一 / 无穷 / 无解）",
    "★ 建立“矩阵 = 线性变换”的核心视角，会读矩阵的列",
    ("列 j 就是基向量 e_j 被送到的位置", 1),
    "理解行列式的几何意义：面积（体积）伸缩因子",
    "★ 把 det=0 ⟺ 压扁 ⟺ 不可逆 ⟺ 无唯一解 串成一条线",
    "初识逆矩阵（逆变换）与特征向量（方向不变的轴）",
], y=1.55, w=6.4, size=17)
k.picture(s, F('fig17m_map_mini.png'), x=7.3, y=2.3, w=5.6)

# 3 幕1
k.section_slide(prs, "第 1 幕", "从两条直线到两个向量：行视角 vs 列视角", "0–8 min")

# 4 情境导入 + 三种情形
s = k.content_slide(prs, "情境与结论：解的三种情形", "8 min")
k.bullets(s, [
    "剧场售两种票，列出二元一次方程组——初中的消元、代入你都会",
    "但未知数变成 3 个、10 个、1000 个呢？必须把消元“机械化”",
    "先看几何结论：相交 → 唯一解；平行 → 无解；重合 → 无穷多解",
    "关键提问：解方程组到底在问什么？看你按“行”读还是按“列”读",
], y=1.5, w=12.2, size=17)
k.full_picture(s, F('fig01_three_cases.png'), y=3.75, w=10.2)

# 5 ★ 行视角 vs 列视角
s = k.content_slide(prs, "★ 同一个方程组，两副眼镜", "核心")
k.full_picture(s, F('fig02_row_vs_col.png'), y=1.45, w=11.3)
k.callout(s, "行视角：每行一条直线，解 = 交点（你熟悉的）　|　列视角：解 = 用两列向量“配”出 b 的系数（矩阵的灵魂）",
          x=0.9, y=5.8, w=11.5, h=0.85, kind="key")

# 6 列视角公式
s = k.content_slide(prs, "列视角的代数写法", "推演")
p = eqrow([('t', r'$x$'), ('m', [[1], [3]]), ('t', r'$+\ y$'), ('m', [[2], [-1]]),
           ('t', r'$=$'), ('m', [[4], [5]]), ('t', r'$\Leftrightarrow$'),
           ('m', [[1, 2], [3, -1]]), ('m', [['x'], ['y']]), ('t', r'$=$'), ('m', [[4], [5]]),
           ('t', r'$\Leftrightarrow\ A\mathbf{x}=\mathbf{b}$')], 'eq01_colview.png', h=1.7)
eqpic(s, p, y=1.5)
k.bullets(s, [
    "Ax = 把 A 的各列按 x 的分量加权求和",
    "于是 Ax=b 变成一个几何问题：",
    ("b 能否由 A 的列组合出来？", 1),
    ("能→有解；不能→无解；不唯一→无穷多解", 1),
], y=3.15, w=6.1, size=17)
k.picture(s, F('fig03_colspace.png'), x=7.0, y=3.2, w=6.0)
k.callout(s, "两列不共线 → 列空间铺满平面，任何 b 都有解；两列共线 → 压成一条直线，线外的 b 无解。这是“det=0”的前身。",
          x=0.85, y=5.0, w=6.0, h=1.2, kind="note")

# 7 幕2
k.section_slide(prs, "第 2 幕", "高斯消元：把方程组交给算法", "8–22 min")

# 8 矩阵与增广矩阵
s = k.content_slide(prs, "把方程组“脱水”成矩阵", "概念")
p = eqrow([('s', [r'$x+2y+z=8$', r'$2x-y+z=3$', r'$x+y-z=0$']),
           ('a', '脱水', ''),
           ('t', r'$\bar A=(A\,|\,\mathbf{b})=$'),
           ('m', [[1, 2, 1, 8], [2, -1, 1, 3], [1, 1, -1, 0]], 3)], 'eq02_augmented.png', h=1.95)
eqpic(s, p, y=1.45)
k.bullets(s, [
    "系数矩阵 A：只留系数，未知数的名字是多余的",
    "增广矩阵 Ā = (A | b)：把常数列也搬进来，竖线只是分隔",
    "解方程组 = 对增广矩阵做“合法的化简”",
    "唯一合法的三招：换、乘、倍加——都可逆，故解集不变",
], y=3.5, w=5.9, size=17)
k.picture(s, F('fig04_row_ops.png'), x=7.0, y=3.7, w=5.9)

# 9 高斯消元流程
s = k.content_slide(prs, "高斯消元的规范流程", "推演")
k.full_picture(s, F('fig05_gauss_flow.png'), y=1.45, w=12.0)
k.bullets(s, [
    "从左到右逐列：选主元 → 把主元下方全部消成 0 → 移到下一列下一行",
], y=6.0, w=12.2, size=16)

# 10 例题1
s = k.content_slide(prs, "例1　高斯消元解 3×3（板演）", "例题")
p = eqrow([('m', [[1, 2, 1, 8], [2, -1, 1, 3], [1, 1, -1, 0]], 3),
           ('a', r'$r_2-2r_1$', r'$r_3-r_1$'),
           ('m', [[1, 2, 1, 8], [0, -5, -1, -13], [0, -1, -2, -8]], 3),
           ('a', r'$r_3-\frac{1}{5}r_2$', r'$\times 5$'),
           ('m', [[1, 2, 1, 8], [0, -5, -1, -13], [0, 0, -9, -27]], 3)],
          'eq03_ex1_elim.png', h=2.0)
eqpic(s, p, y=1.45)
k.bullets(s, [
    "第 1 列：以 a₁₁=1 为主元，r₂−2r₁、r₃−r₁ 清零",
    "第 2 列：以 −5 为主元，r₃−(1/5)r₂ 清零，再 ×5 去分母",
    "三个主元 1、−5、−9 非零 ⟹ r(A)=3，唯一解",
], y=3.6, w=6.4, size=17)
k.picture(s, F('fig06_echelon.png'), x=7.4, y=3.5, w=5.5)

# 11 例1 回代与检验
s = k.content_slide(prs, "例1　回代与检验", "例题")
p = eqrow([('m', [[1, 2, 1, 8], [0, -5, -1, -13], [0, 0, -9, -27]], 3),
           ('t', r'$\Rightarrow\ z=3,\ \ y=2,\ \ x=1$', k.M_GRN)],
          'eq04_ex1_back.png', h=1.9)
eqpic(s, p, y=1.45, x=0.8)
tex(s, "检验：1+4+3=8　2−2+3=3　1+2−3=0", x=0.8, y=3.4, w=6.4, size=0.5, color=k.M_GRN)
k.bullets(s, [
    "回代顺序：从最后一行往上，逐个解出 z → y → x",
    "解为 x = 1，y = 2，z = 3（三式全部成立）",
    "消元结束后必须回代原方程组验证，不验证不得分",
    "常见错误：把行变换写成 r₂ ← 2r₁ − r₂（改变主元行）",
], y=4.3, w=6.4, size=16)
k.picture(s, F('fig06_echelon.png'), x=7.4, y=3.4, w=5.5)

# 12 解的三种情况判别
s = k.content_slide(prs, "解的存在唯一性：看阶梯形的“长相”", "核心")
k.full_picture(s, F('fig07_rank_cases.png'), y=1.5, w=11.6)
k.callout(s, "r(A) < r(Ā) → 无解；r(A) = r(Ā) = n → 唯一解；r(A) = r(Ā) < n → 无穷多解（自由未知量 n − r 个）",
          x=0.9, y=5.7, w=11.5, h=0.9, kind="key")

# 13 秩的直观
s = k.content_slide(prs, "秩的直观：变换有没有“压扁”空间", "核心")
k.picture(s, F('fig08_rank_squash.png'), x=6.5, y=1.9, w=6.5)
k.bullets(s, [
    "秩 r(A) = 主元个数 = 变换后“像”还剩几维",
    "满秩：平面还是平面 → 可逆 → 唯一解",
    "降秩：平面被压成直线/点 → 信息丢失",
    ("b 在压扁后的像里 → 无穷多解", 1),
    ("b 不在里面 → 无解", 1),
], y=1.9, w=5.5, size=17)
k.callout(s, "把“秩”翻译成人话：这次变换，把空间压掉了几维。",
          x=0.85, y=5.7, w=5.6, h=0.9, kind="note")

# 14 例2
s = k.content_slide(prs, "例2　含参数：判断解的情况", "例题")
p = eqrow([('s', [r'$x+y+z=1$', r'$x+2y+az=2$', r'$x+ay+2z=3$']),
           ('t', r'　　$\det A=$'),
           ('d', [[1, 1, 1], [1, 2, 'a'], [1, 'a', 2]]),
           ('t', r'$=-(a-2)(a+1)$')], 'eq05_ex2.png', h=2.0)
eqpic(s, p, y=1.45)
tex(s, r"$\det A\neq 0\ \Leftrightarrow\ a\neq 2,\ a\neq -1\ \Rightarrow\ $唯一解",
    x=2.5, y=3.55, w=8.3, size=0.5, color=k.M_GRN)
k.bullets(s, [
    "a = 2：消元后出现 0 = 0 行 → r(A)=r(Ā)=2 < 3 → 无穷多解",
    "a = −1：出现 0 = 3 的矛盾行 → r(A)=2 < r(Ā)=3 → 无解",
    "方法论：先用 det 抓“唯一解”的条件，再对例外值逐一消元检查",
], y=4.5, w=12.2, size=17)
k.callout(s, "行列式先行——它是判别唯一解的“总开关”，下一幕解释它为什么这么灵。",
          x=0.9, y=6.2, w=11.5, h=0.8, kind="key")

# 15 幕3
k.section_slide(prs, "第 3 幕", "★ 矩阵即线性变换：把方格网“捏”一下", "22–38 min")

# 16 线性变换定义
s = k.content_slide(prs, "什么叫线性变换", "概念")
tex(s, r"$T(\mathbf{u}+\mathbf{v})=T(\mathbf{u})+T(\mathbf{v}),\quad T(k\mathbf{u})=kT(\mathbf{u})$",
    x=0.85, y=1.5, w=6.4, size=0.42)
k.bullets(s, [
    "几何判据（更好记）：",
    ("① 网格线保持是直线，② 保持平行且等距，③ 原点不动", 1),
    "关键推论：只要知道 e₁、e₂ 被送到哪，整个变换就定死了",
], y=2.5, w=6.3, size=17)
tex(s, r"$T(\mathbf{x})=A\mathbf{x},\ \ A=[\,T(\mathbf{e}_1)\ \ T(\mathbf{e}_2)\,]$",
    x=0.85, y=4.75, w=6.4, size=0.42)
k.picture(s, F('fig09_columns_are_images.png'), x=7.3, y=2.2, w=5.6)
k.callout(s, "读矩阵的正确姿势：先把两列画成箭头，你就“看见”了这个变换。",
          x=0.85, y=5.75, w=6.3, h=0.85, kind="key")

# 17 ★ 高光：网格变形四联图
s = k.content_slide(prs, "★ 四个矩阵，四种“捏法”", "高光")
k.full_picture(s, F('fig10_grid_transforms.png'), y=1.4, w=12.3)

# 18 变式1
s = k.content_slide(prs, "变式1　看矩阵，说几何作用", "变式")
p = eqrow([('t', r'$A_1=$'), ('m', [[0, -1], [1, 0]]),
           ('t', r'$,\ A_2=$'), ('m', [[1, 0], [0, -1]]),
           ('t', r'$,\ A_3=$'), ('m', [[2, 0], [0, 2]]),
           ('t', r'$,\ A_4=$'), ('m', [[1, 1], [1, 1]])], 'eq06_var1.png', h=1.6)
eqpic(s, p, y=1.5)
k.bullets(s, [
    "A₁：e₁→(0,1)，e₂→(−1,0)：旋转 90°，det = 1",
    "A₂：e₂ 翻到 (0,−1)：x 轴镜像，det = −1（翻面）",
    "A₃：两列放大 2 倍：位似，det = 4（面积 ×4）",
    "A₄：两列相同：压到直线 y = x，det = 0（不可逆）",
], y=3.25, w=7.4, size=17)
k.picture(s, F('fig13s_det_sign_small.png'), x=8.4, y=3.6, w=4.6)
k.callout(s, "训练：先画列，再报 det 的符号——形与数同时到位。",
          x=0.85, y=6.25, w=7.4, h=0.8, kind="note")

# 19 复合变换
s = k.content_slide(prs, "复合变换 = 矩阵乘法（顺序要命）", "推演")
k.full_picture(s, F('fig11_composition.png'), y=1.5, w=11.6)
tex(s, r"$(RS)\mathbf{x}=R(S\mathbf{x}),\qquad RS\neq SR$", x=3.0, y=5.9, w=7.3, size=0.5)

# 20 幕4
k.section_slide(prs, "第 4 幕", "行列式：一个数，量出空间被拉伸了多少", "38–52 min")

# 21 det = 面积伸缩因子
s = k.content_slide(prs, "行列式的几何意义", "核心")
k.full_picture(s, F('fig12_det_area.png'), y=1.5, w=10.4)
p = eqrow([('t', r'$\det$'), ('d', [['a', 'b'], ['c', 'd']]), ('t', r'$=ad-bc$'),
           ('t', r'　　$|\det A|$ = 像的面积 ÷ 原面积')], 'eq07_det.png', h=1.3)
eqpic(s, p, y=5.75)

# 22 符号与零点
s = k.content_slide(prs, "det 的符号与零点", "核心")
k.full_picture(s, F('fig13_det_sign.png'), y=1.5, w=11.6)
k.callout(s, "det > 0：保持定向；det < 0：把平面“翻了个面”；det = 0：整个平面被压扁成线（或点）——面积归零。",
          x=0.9, y=5.7, w=11.5, h=0.85, kind="key")

# 23 ★ 四件事串一条线
s = k.content_slide(prs, "★ 一条线串起四件事", "高光")
k.full_picture(s, F('fig14_four_equiv.png'), y=1.5, w=12.0)
k.callout(s, "本讲最值钱的一句话——考试判断题的总开关。",
          x=0.9, y=5.8, w=11.5, h=0.8, kind="key")

# 24 逆矩阵
s = k.content_slide(prs, "逆矩阵 = 逆变换（撤销键）", "概念")
k.picture(s, F('fig15_inverse.png'), x=6.5, y=2.4, w=6.4)
p = eqrow([('t', r'$A^{-1}=\dfrac{1}{ad-bc}$'), ('m', [['d', '-b'], ['-c', 'a']]),
           ('t', r'$,\ AA^{-1}=I$')], 'eq08_inv.png', h=1.6)
eqpic(s, p, y=1.5, x=0.85, maxw=5.7)
k.bullets(s, [
    "det A = 0 时 A⁻¹ 不存在：压扁了就撤不回来",
    "Ax = b 且 A 可逆 ⟹ x = A⁻¹b（唯一解）",
    "求逆通法：对 (A | I) 做高斯消元 → (I | A⁻¹)",
], y=3.2, w=5.5, size=17)
k.callout(s, "口诀：能撤销 ⟺ 没压扁 ⟺ det≠0。",
          x=0.85, y=6.0, w=5.6, h=0.8, kind="key")

# 25 例3 求逆
s = k.content_slide(prs, "例3　求逆矩阵（增广消元法）", "例题")
p = eqrow([('m', [[2, 1, 1, 0], [1, 1, 0, 1]], 2),
           ('a', r'$r_1\leftrightarrow r_2$', ''),
           ('m', [[1, 1, 0, 1], [2, 1, 1, 0]], 2),
           ('a', r'$r_2-2r_1$', ''),
           ('m', [[1, 1, 0, 1], [0, -1, 1, -2]], 2)], 'eq09a_ex3.png', h=1.7)
eqpic(s, p, y=1.5)
p = eqrow([('a', r'$-r_2$', r'$r_1-r_2$'),
           ('m', [[1, 0, 1, -1], [0, 1, -1, 2]], 2),
           ('t', r'$\Rightarrow\ A^{-1}=$'), ('m', [[1, -1], [-1, 2]]),
           ('t', '　验：'), ('m', [[2, 1], [1, 1]]), ('m', [[1, -1], [-1, 2]]),
           ('t', r'$=I$', k.M_GRN)], 'eq09b_ex3.png', h=1.7)
eqpic(s, p, y=3.25)
k.bullets(s, [
    "左边化成 I 的同时，右边自动变成 A⁻¹（同一串行变换）",
    "2×2 也可套公式：det = 2·1 − 1·1 = 1",
], y=5.0, w=12.2, size=17)
k.callout(s, "若消元中左边出现全零行 → det = 0 → 该矩阵不可逆，立刻停手。",
          x=0.9, y=6.2, w=11.5, h=0.75, kind="warn")

# 26 例4 应用题
s = k.content_slide(prs, "例4　用矩阵解应用题", "例题")
k.picture(s, F('fig18_application.png'), x=6.6, y=2.1, w=6.3)
k.bullets(s, [
    "设 A、B、C 三种票各 x、y、z 张",
    "总张数 100；单价 30/50/80，总收入 5400 元",
    "另有约束：A 是 B 的 2 倍",
    "写成 Ax = b，det A = 130 ≠ 0 ⟹ 唯一解",
    "得 x = 40, y = 20, z = 40（代回三式成立）",
], y=2.0, w=5.6, size=17)
k.callout(s, "应用题三步走：设元 → 写成 Ax=b → 先算 det 判解，再求解。",
          x=0.85, y=5.7, w=5.6, h=0.9, kind="note")

# 27 变式2
s = k.content_slide(prs, "变式2　几何与代数的往返", "变式")
k.bullets(s, [
    "(1) 变换把 e₁ 送到 (1,2)、e₂ 送到 (3,4)，写出 A 并求 det。",
    ("A = [[1,3],[2,4]]，det = −2：面积 ×2 且翻面。", 1),
    "(2) 求 A 的逆变换矩阵；若 b = (5,6)，解 Ax = b。",
    ("A⁻¹ = [[−2,1.5],[1,−0.5]]，x = A⁻¹b = (−1, 2)。", 1),
    "(3) 若把 A 的第二列改成 (2,4)，det 变成多少？还有唯一解吗？",
    ("det = 0：两列共线，平面被压成直线 → 无唯一解。", 1),
], y=1.7, w=7.7, size=17)
k.picture(s, F('fig12s_det_area_small.png'), x=8.7, y=2.6, w=4.2)

# 28 幕5
k.section_slide(prs, "第 5 幕", "伏笔：变换中方向不变的轴 · 小结与作业", "52–60 min")

# 29 特征向量
s = k.content_slide(prs, "特征向量：下一讲的钥匙", "伏笔")
k.picture(s, F('fig16_eigen.png'), x=6.5, y=1.9, w=6.5)
tex(s, r"$A\mathbf{v}=\lambda\mathbf{v}\quad(\mathbf{v}\neq\mathbf{0})$",
    x=0.85, y=1.6, w=5.2, size=0.42)
k.bullets(s, [
    "多数向量被 A 一作用就“转向”了",
    "少数向量只被拉长/压短，方向不变——特征向量",
    "拉伸倍数 λ 就是特征值",
    "求法预告：det(A − λI) = 0",
    "意义：找到变换自己的“坐标轴”",
], y=2.6, w=5.5, size=17)
k.callout(s, "det 又出现了——它是判断“有没有非零向量被压成 0”的那把尺。",
          x=0.85, y=5.9, w=5.5, h=0.8, kind="note")

# 30 小结
s = k.content_slide(prs, "课堂小结：一张图收束", "小结")
k.full_picture(s, F('fig17_map.png'), y=1.5, w=11.4)
k.callout(s, "带走一句话：矩阵是动作，不是数表；det=0 就是这个动作把空间压扁了，压扁了就撤不回来，方程也就没了唯一解。",
          x=0.9, y=5.85, w=11.5, h=0.85, kind="key")

# 31 分层作业
s = k.content_slide(prs, "分层作业", "作业")
k.bullets(s, [
    "【基础 · 必做】",
    ("1. 用高斯消元解：x+y+z=6, 2x−y+z=3, x+2y−z=2（写出每一步）", 1),
    ("2. 求 [[3,2],[1,4]] 的行列式与逆矩阵，并验证 AA⁻¹=I", 1),
    "【提高 · 选做】",
    ("3. 画出 [[1,0],[0,−1]]、[[0,1],[1,0]] 对方格网的作用，解释 det 符号", 1),
    ("4. 讨论参数 k：kx+y=1, x+ky=1 何时无解 / 唯一解 / 无穷多解", 1),
    "【拓展 · 挑战】",
    ("5. 证明：若 A 的两列共线，则 det A = 0，并给出几何解释", 1),
    ("6. 猜想并验证 det(AB) = det A · det B", 1),
], y=1.5, w=12.2, size=16)

# 32 板书提纲
s = k.content_slide(prs, "板书提纲", "板书")
k.bullets(s, [
    "左板（概念）：Ax=b 的两副眼镜 —— 行=交点 / 列=组合",
    "中板（算法）：增广矩阵 → 三种行变换 → 阶梯形 → 回代",
    ("例1 全过程留板，例2 参数讨论列表", 1),
    "右板（几何）：方格网 + 四种捏法；det = 面积伸缩因子",
    ("红框：det=0 ⟺ 压扁 ⟺ 不可逆 ⟺ 无唯一解", 1),
    "角落（伏笔）：Av = λv",
], y=1.6, w=12.2, size=17)
k.callout(s, "板书节奏：例1 边讲边写，学生同步在草稿纸上跟做；四种捏法用彩色粉笔画在右板并全程保留。",
          x=0.9, y=5.5, w=11.5, h=1.0, kind="note")

path = k.save(prs, os.path.join(OUT, '09_线性方程组与矩阵.pptx'))
print("slides:", len(prs.slides.__iter__.__self__._sldIdLst), "->", path)

# ======================= 教案 DOCX =======================
from docx import Document
from docx.shared import Pt as DPt, RGBColor as DRGB, Inches as DIn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Noto Serif CJK SC'
st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')

def H(t, lv=1):
    p = doc.add_heading(t, level=lv)
    for r in p.runs:
        r.font.name = 'Noto Serif CJK SC'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
    return p

def P(t, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    return p

ttl = doc.add_heading('教案　第09讲　线性方程组、矩阵与线性变换', level=0)
for r in ttl.runs:
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')

info = doc.add_table(rows=2, cols=4); info.style = 'Table Grid'
for (i, j), v in zip([(0,0),(0,1),(0,2),(0,3),(1,0),(1,1),(1,2),(1,3)],
                     ['课题', '线性方程组、矩阵与线性变换', '课时', '1 课时（60 分钟）',
                      '课型', '新授课（概念 + 算法 + 几何直观）', '授课对象', '高中衔接 / 大一线性代数入门']):
    info.cell(i, j).text = v

H('一、教材分析', 1)
P('本讲处于“代数运算”向“空间变换”过渡的枢纽位置。传统教材把矩阵作为“数表”引入，学生虽会算，却不知其所以然；'
  '本设计以“矩阵 = 线性变换”为统摄性观念（big idea）重组内容：先以二元方程组的行/列双视角撬动认知冲突，'
  '再引入高斯消元的规范算法，最后用方格网变形、行列式的面积伸缩因子把“代数判别式”与“几何压扁”打通，'
  '并以特征向量为下一讲埋设伏笔。核心链条为：det=0 ⟺ 压扁 ⟺ 不可逆 ⟺ 无唯一解。')

H('二、学情分析', 1)
P('学生已掌握二元/三元一次方程组的消元与代入解法，具备向量的基本运算与坐标表示经验。主要困难有三：'
  '（1）只会“行视角”（两直线求交点），对“列视角”（向量线性组合）完全陌生；'
  '（2）把矩阵当作无意义的数字方阵，行列式仅记住 ad−bc 的口诀；'
  '（3）对“无解 / 无穷多解”的判别停留在死记条件，缺乏几何图像支撑。'
  '教学须以图形直观为脚手架，反复在“代数式—几何图”之间往返。')

H('三、教学目标（三维）', 1)
P('知识与技能：', True)
P('1. 掌握增广矩阵与三种初等行变换，能规范地用高斯消元法求解 3×3 线性方程组并回代检验；\n'
  '2. 能由行阶梯形的主元结构判别解的存在唯一性（r(A)、r(Ā) 与 n 的关系）；\n'
  '3. 理解矩阵作为线性变换的意义，会由矩阵的列还原变换的几何作用；\n'
  '4. 理解行列式是面积（体积）伸缩因子，会求 2×2 逆矩阵并解 Ax=b。')
P('过程与方法：', True)
P('经历“具体方程组 → 双视角对照 → 算法机械化 → 几何解释 → 统一判别”的完整建构过程，'
  '体会数形结合、化归与算法化思想；通过方格网变形的观察、比较、归纳，发展几何直观与推理能力。')
P('情感态度与价值观：', True)
P('感受“换一副眼镜看旧问题”带来的思维跃迁，体会数学抽象的威力与统一之美；'
  '在“四件事串成一条线”的收束中获得结构化认知的成就感。')

H('四、教学重难点', 1)
P('重点：高斯消元的规范算法；矩阵即线性变换的核心视角；行列式的几何意义。', True)
P('难点：列视角（Ax=b ⟺ b 是否落在 A 的列空间中）的建立；'
  '把 det=0、空间被压扁、矩阵不可逆、方程组无唯一解四件事理解为等价命题。', True)
P('突破策略：以“行视角 vs 列视角”对照图制造认知冲突；以方格网变形四联图（旋转/剪切/伸缩/投影）'
  '把抽象矩阵可视化；以“面积归零”这一可见事实统摄四个等价命题。')

H('五、教法与学法', 1)
P('教法：问题驱动 + 直观演示 + 变式训练。以“解方程组到底在问什么”为核心问题贯穿全课，'
  '关键处（列视角、方格网变形、四件事等价）采用板演 + 图形对照 + 追问的方式推进。')
P('学法：观察—猜想—验证—归纳。学生在草稿纸上同步跟做例 1 的消元全过程；'
  '在变式 1 中用“先画两列箭头”的方法自主判读矩阵的几何作用。')

H('六、教学准备', 1)
P('多媒体课件（32 页 PPT，含 18 张精确作图）；彩色粉笔（区分主元、零元、特征方向）；'
  '学生草稿纸与直尺；预留右侧板面全程保留方格网变形示意图。')

H('七、教学过程（分钟级时间轴）', 1)
rows = [
    ("0–3", "封面·目标（PPT 1–2）",
     "开课设问：“解方程组，究竟在问什么？”出示学习目标与知识地图。",
     "回忆初中消元法，带着问题进入。",
     "以核心问题统摄全课，明确本课的“新眼镜”。"),
    ("3–8", "第1幕 情境与三种情形（PPT 3–4）",
     "呈现售票情境，抽象出二元方程组；出示相交/平行/重合三子图。",
     "回答三种情形对应的解的个数。",
     "从熟悉经验切入，建立“解 = 交点”的行视角基线。"),
    ("8–13", "★ 行视角 vs 列视角（PPT 5–6）",
     "同一方程组画两幅图：左为两直线求交；右为两列向量加权凑出 b。追问：“为什么两幅图算出同一个 (2,1)？”"
     "再引列空间图：b 在不在张成里。",
     "对照观察，尝试用向量语言复述方程组；小组讨论 30 秒。",
     "制造认知冲突，打通本讲最关键的视角转换（难点突破 1）。"),
    ("13–17", "第2幕 增广矩阵与三种行变换（PPT 7–8）",
     "把方程组“脱水”为增广矩阵；讲清换、乘、加三招及其可逆性。",
     "口答：为什么这三招不改变解集？",
     "揭示算法的合法性根据，避免机械记忆。"),
    ("17–22", "高斯消元流程 + 例1 板演（PPT 9–11）",
     "出示矩阵变形流程图；完整板演例 1 的消元、回代与检验，红笔标出主元与被消成 0 的位置。",
     "在草稿纸同步跟做，逐步核对；指出老师故意留下的检验陷阱。",
     "算法规范化；强化“不检验不得分”的习惯。"),
    ("22–27", "解的判别 + 秩的直观 + 例2（PPT 12–14）",
     "由三种阶梯形长相读出解的三种情况；用“压扁”解释秩；讲解含参数 a 的例 2。",
     "先独立判断 a=2、a=−1 两种例外情形，再交流。",
     "把死记的秩条件转译为可视的几何事实（难点突破 2 前置）。"),
    ("27–33", "★ 第3幕 矩阵即线性变换（PPT 15–17）",
     "给出线性变换的几何判据；证明“列 = 基向量的像”；出示方格网变形四联图（旋转/剪切/伸缩/投影），"
     "逐幅追问：面积变了吗？方向翻了吗？",
     "观察四幅图，归纳每种变换的矩阵特征；口答投影矩阵的 det。",
     "本讲高光：把抽象矩阵变为可见的“动作”，为行列式与可逆性铺路。"),
    ("33–38", "变式1 + 复合变换（PPT 18–19）",
     "变式 1：给四个矩阵，说几何作用与 det 符号；讲解复合变换即矩阵乘法，强调顺序。",
     "先画列向量箭头再作答；举例说明 RS ≠ SR。",
     "变式训练，形数同时到位；顺带交代乘法不可交换的几何原因。"),
    ("38–44", "第4幕 行列式的几何意义（PPT 20–22）",
     "单位正方形→平行四边形，面积即 |det A|；对比 det>0、det<0（翻面）、det=0（压扁）。",
     "计算 det 并预测面积；观察字母 F 的翻转。",
     "建立“行列式 = 面积伸缩因子”的核心直观。"),
    ("44–48", "★ 四件事串成一条线（PPT 23）",
     "出示等价链：det=0 ⟺ 压扁 ⟺ 不可逆 ⟺ 无唯一解；逐环追问理由，回扣第 2 幕的秩与例 2。",
     "复述等价链并举反例；回答“为什么例 2 先算 det”。",
     "结构化收束，难点突破 2 完成，形成可迁移的判断总开关。"),
    ("48–54", "逆矩阵与例3、例4（PPT 24–26）",
     "逆矩阵即逆变换；示范 (A|I)→(I|A⁻¹) 的求逆消元；讲解应用题的“设元→Ax=b→先算 det”三步走。",
     "完成例 3 的验证 AA⁻¹=I；参与例 4 的列式与求解。",
     "算法与应用落地，呼应行列式判别。"),
    ("54–57", "变式2 + 特征向量伏笔（PPT 27–29）",
     "变式 2 做几何↔代数往返；演示特征向量图：多数向量被掰弯，少数只被拉伸。",
     "指认图中方向不变的两条轴，猜测 λ 的值。",
     "巩固迁移，并为下一讲设置悬念。"),
    ("57–60", "小结·作业·板书（PPT 30–32）",
     "用知识地图收束；布置三层作业；强调板书右侧的方格网图请自行誊抄。",
     "复述“矩阵是动作，不是数表”；记录作业。",
     "结构化回顾，分层落实，减负提效。"),
]
t = doc.add_table(rows=1, cols=5); t.style = 'Table Grid'
for c, h in zip(t.rows[0].cells, ['时间', '环节 / 对应 PPT', '教师活动', '学生活动', '设计意图']):
    c.text = ''
    r = c.paragraphs[0].add_run(h); r.bold = True
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for row in rows:
    cells = t.add_row().cells
    for c, v in zip(cells, row):
        c.text = ''
        r = c.paragraphs[0].add_run(v)
        r.font.size = DPt(9)
        r.font.name = 'Noto Serif CJK SC'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
P('　')
P('时间合计：3+5+5+4+5+5+6+5+6+4+6+3+3 = 60 分钟。', True)

H('八、板书设计', 1)
P('左板（概念）：Ax=b 的两副眼镜——行视角：每行一条直线，解=交点；列视角：解=凑出 b 的配方。')
P('中板（算法）：增广矩阵 (A|b) → 三种初等行变换（换、乘、加）→ 行阶梯形 → 回代 → 检验。'
  '例 1 全过程留板；例 2 的参数讨论用表格分栏（a≠2 且 a≠−1 / a=2 / a=−1）。')
P('右板（几何，全程保留）：单位方格网 + 旋转 / 剪切 / 伸缩 / 投影 四幅示意；下方红框写：')
P('　　det A = 0 ⟺ 空间被压扁 ⟺ A 不可逆 ⟺ Ax=b 无唯一解', True)
P('角落（伏笔）：Av = λv —— 方向不变的轴。')

H('九、分层作业', 1)
P('基础（必做）：1. 高斯消元解 x+y+z=6, 2x−y+z=3, x+2y−z=2，写出每一步行变换；'
  '2. 求 [[3,2],[1,4]] 的 det 与逆矩阵，并验证 AA⁻¹=I。')
P('提高（选做）：3. 画出 [[1,0],[0,−1]] 与 [[0,1],[1,0]] 对单位方格网的作用，解释 det 符号；'
  '4. 讨论参数 k：kx+y=1, x+ky=1 何时无解 / 唯一解 / 无穷多解。')
P('拓展（挑战）：5. 证明两列共线 ⟹ det A = 0，并用“压扁”给出几何解释；'
  '6. 猜想并验证 det(AB) = det A · det B（提示：面积伸缩因子相乘）。')

H('十、教学反思（课后填写）', 1)
for q in ['1. 列视角的引入，学生的认知冲突是否被真正激活？有多少学生能主动用“向量组合”复述方程组？',
          '2. 方格网变形四联图的观察时间是否充分？投影（det=0）一幅是否讲透？',
          '3. “四件事串一条线”的收束，学生能否独立复述并举例？',
          '4. 例 1 的检验陷阱是否达到预期的警示效果？',
          '5. 时间分配偏差与下次调整：']:
    P(q)
    P('　\n　\n________________________________________________________________')

dpath = os.path.join(OUT, '教案_09_线性方程组与矩阵.docx')
doc.save(dpath)
print("docx:", dpath)
print("OK figs=", len(figs))
