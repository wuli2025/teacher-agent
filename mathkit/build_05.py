# -*- coding: utf-8 -*-
"""第05讲：条件概率与贝叶斯定理 —— 60分钟课件包生成器"""
import sys, os
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, Rectangle, FancyArrowPatch, Ellipse

OUT = "/mnt/c/Users/mi/Desktop/数学名师课件包/05_条件概率与贝叶斯定理"
FIG = os.path.join(OUT, "figures")
os.makedirs(FIG, exist_ok=True)
F = lambda n: os.path.join(FIG, n)

# ============================ 配图 ============================
# 画布宽度 = 该图在幻灯片上的展示宽度（英寸），保证贴图不缩小，图内字 >= 20pt
def fig01_guess():
    """导入：学生直觉猜测 vs 真实答案"""
    fig, ax = k.new_fig(6.2, 3.6)
    labels = ["99%", "90%", "50%", "<20%", "真实"]
    vals = [46, 24, 18, 12, 9]
    cols = [k.M_SLATE, k.M_SLATE, k.M_SLATE, k.M_SLATE, k.M_RED]
    b = ax.bar(labels, vals, color=cols, width=0.62, zorder=3)
    for r, v in zip(b, vals):
        ax.text(r.get_x() + r.get_width()/2, v + 2, f"{v}",
                ha="center", color=k.M_INK, fontweight="bold")
    ax.set_ylim(0, 62)
    ax.set_yticks([0, 20, 40, 60])
    k.style_axes(ax, "", "占比 %", origin=False)
    ax.grid(True, axis="y", color=k.M_RULE, lw=0.8, zorder=0)
    ax.set_title("课堂投票 vs 真实答案", color=k.M_INK, pad=10)
    ax.annotate("差 10 倍", xy=(4, 11), xytext=(2.4, 40),
                color=k.M_RED, fontweight="bold",
                arrowprops=dict(arrowstyle="->", color=k.M_RED, lw=2))
    return k.save_fig(fig, F("fig01_guess.png"))

def fig02_venn():
    """条件概率的几何解释：样本空间缩小 + 重新标准化"""
    fig, axes = plt.subplots(1, 2, figsize=(10.6, 4.4))
    for ax in axes:
        ax.set_xlim(0, 10); ax.set_ylim(0, 7.4); ax.axis("off")
    ax = axes[0]
    ax.add_patch(Rectangle((0.4, 0.4), 9.2, 5.6, fc="white", ec=k.M_INK, lw=1.8))
    ax.text(0.75, 5.35, "Ω", color=k.M_INK, style="italic")
    ax.add_patch(Ellipse((3.8, 3.2), 4.4, 3.8, fc=k.M_ACC2, alpha=0.28, ec=k.M_ACC2, lw=2.4))
    ax.add_patch(Ellipse((6.4, 3.2), 4.4, 3.8, fc=k.M_ACC, alpha=0.28, ec=k.M_ACC, lw=2.4))
    ax.text(2.3, 3.1, "A", color=k.M_ACC2, fontweight="bold", ha="center", va="center")
    ax.text(7.9, 3.1, "B", color=k.M_ACC, fontweight="bold", ha="center", va="center")
    ax.text(5.1, 3.1, "AB", color=k.M_INK, fontweight="bold", ha="center", va="center")
    ax.set_title("① 原样本空间 Ω", color=k.M_INK, pad=12)
    ax.text(5.0, 6.6, "P(A) = 面积A / 面积Ω", ha="center", va="center", color=k.M_SLATE)

    ax = axes[1]
    ax.add_patch(Rectangle((0.4, 0.4), 9.2, 5.6, fc="white", ec=k.M_RULE, lw=1.2, ls="--"))
    ax.add_patch(Ellipse((3.8, 3.2), 4.4, 3.8, fc="none", ec=k.M_RULE, lw=1.6, ls="--"))
    ax.add_patch(Ellipse((6.4, 3.2), 4.4, 3.8, fc=k.M_ACC, alpha=0.30, ec=k.M_ACC, lw=3))
    for x in np.linspace(3.4, 6.4, 260):
        d1 = 1 - ((x-3.8)/2.2)**2
        d2 = 1 - ((x-6.4)/2.2)**2
        if d1 > 0 and d2 > 0:
            h = 1.9*min(np.sqrt(d1), np.sqrt(d2))
            ax.plot([x, x], [3.2-h, 3.2+h], color=k.M_GRN, lw=1.2, alpha=0.6)
    ax.text(8.1, 3.1, "B", color=k.M_ACC, fontweight="bold", ha="center", va="center")
    ax.text(4.9, 1.0, "AB", color=k.M_GRN, fontweight="bold", ha="center", va="center")
    ax.annotate("", xy=(4.9, 2.2), xytext=(4.9, 1.4),
                arrowprops=dict(arrowstyle="->", color=k.M_GRN, lw=2))
    ax.set_title("② 已知 B 发生", color=k.M_RED, pad=12)
    ax.text(5.0, 6.6, "P(A|B) = 面积AB / 面积B", ha="center", va="center", color=k.M_RED)
    fig.subplots_adjust(left=0.02, right=0.98, top=0.80, bottom=0.03, wspace=0.10)
    fig.savefig(F("fig02_venn.png"), dpi=200)
    plt.close(fig)
    return F("fig02_venn.png")

def fig03_dice():
    """例1：骰子网格，已知点数和>8 求含6的概率"""
    fig, ax = k.new_fig(6.1, 4.0)
    ax.set_xlim(0.4, 6.6); ax.set_ylim(0.4, 6.6)
    for i in range(1, 7):
        for j in range(1, 7):
            s = i + j
            hit_b = s > 8
            hit_a = (i == 6 or j == 6)
            if hit_b and hit_a:
                c, tc = k.M_RED, "white"
            elif hit_b:
                c, tc = k.M_ACC, k.M_INK
            else:
                c, tc = "white", k.M_SLATE
            ax.add_patch(Rectangle((i-0.46, j-0.46), 0.92, 0.92, fc=c, ec=k.M_RULE, lw=1))
            ax.text(i, j, str(s), ha="center", va="center", color=tc)
    ax.set_xticks(range(1, 7)); ax.set_yticks(range(1, 7))
    for s in ax.spines.values():
        s.set_color(k.M_SLATE)
    ax.set_title("红 = AB（7格）　琥珀 = B 其余（3格）", color=k.M_INK, pad=10)
    ax.tick_params(colors=k.M_SLATE)
    return k.save_fig(fig, F("fig03_dice.png"))

def fig04_tree_basic():
    """乘法公式与树状图（分支带权）"""
    fig, ax = k.new_fig(11.3, 2.5)
    ax.set_xlim(0, 12); ax.set_ylim(0, 6); ax.axis("off")
    def node(x, y, t, c):
        ax.text(x, y, t, ha="center", va="center", color=c, fontweight="bold",
                bbox=dict(boxstyle="circle,pad=0.30", fc="white", ec=c, lw=2), zorder=4)
    def edge(x1, y1, x2, y2, w, c, dy=0.45):
        ax.annotate("", xy=(x2-0.55, y2), xytext=(x1+0.55, y1),
                    arrowprops=dict(arrowstyle="->", color=c, lw=2), zorder=2)
        ax.text((x1+x2)/2, (y1+y2)/2 + dy, w, ha="center", va="center",
                color=c, fontweight="bold", zorder=5,
                bbox=dict(boxstyle="round,pad=0.15", fc=k.M_PAPER, ec="none"))
    node(0.9, 3.0, "起", k.M_SLATE)
    node(5.2, 4.8, "B", k.M_ACC); node(5.2, 1.2, "B̄", k.M_ACC2)
    edge(0.9, 3.0, 5.2, 4.8, "P(B)", k.M_ACC)
    edge(0.9, 3.0, 5.2, 1.2, "P(B̄)", k.M_ACC2, -0.45)
    node(10.6, 5.6, "A", k.M_RED); node(10.6, 3.9, "Ā", k.M_SLATE)
    node(10.6, 2.1, "A", k.M_RED); node(10.6, 0.4, "Ā", k.M_SLATE)
    edge(5.2, 4.8, 10.6, 5.6, "P(A|B)", k.M_RED)
    edge(5.2, 4.8, 10.6, 3.9, "P(Ā|B)", k.M_SLATE, -0.45)
    edge(5.2, 1.2, 10.6, 2.1, "P(A|B̄)", k.M_RED)
    edge(5.2, 1.2, 10.6, 0.4, "P(Ā|B̄)", k.M_SLATE, -0.45)
    return k.save_fig(fig, F("fig04_tree.png"))

def fig05_ball():
    """例2：不放回摸球树（带具体权重）"""
    fig, ax = k.new_fig(10.4, 4.6)
    ax.set_xlim(0, 15.0); ax.set_ylim(-0.2, 6.4); ax.axis("off")
    def node(x, y, t, c):
        ax.text(x, y, t, ha="center", va="center", color="white", fontweight="bold",
                bbox=dict(boxstyle="circle,pad=0.30", fc=c, ec=c), zorder=4)
    def edge(x1, y1, x2, y2, w, c, dy=0.42):
        ax.annotate("", xy=(x2-0.55, y2), xytext=(x1+0.55, y1),
                    arrowprops=dict(arrowstyle="->", color=c, lw=2.2), zorder=2)
        ax.text((x1+x2)/2, (y1+y2)/2 + dy, w, ha="center", va="center",
                color=c, fontweight="bold", zorder=5,
                bbox=dict(boxstyle="round,pad=0.12", fc=k.M_PAPER, ec="none"))
    ax.text(0.8, 3.0, "袋", ha="center", va="center", color=k.M_SLATE, fontweight="bold",
            bbox=dict(boxstyle="circle,pad=0.30", fc="white", ec=k.M_SLATE, lw=2))
    ax.text(0.8, 1.9, "3红2白", ha="center", va="center", color=k.M_SLATE)
    node(4.6, 4.9, "红", k.M_RED); node(4.6, 1.1, "白", k.M_ACC2)
    edge(0.8, 3.0, 4.6, 4.9, "3/5", k.M_RED)
    edge(0.8, 3.0, 4.6, 1.1, "2/5", k.M_ACC2, -0.42)
    node(8.4, 5.9, "红", k.M_RED); node(8.4, 4.0, "白", k.M_ACC2)
    node(8.4, 2.1, "红", k.M_RED); node(8.4, 0.2, "白", k.M_ACC2)
    edge(4.6, 4.9, 8.4, 5.9, "2/4", k.M_RED)
    edge(4.6, 4.9, 8.4, 4.0, "2/4", k.M_ACC2, -0.42)
    edge(4.6, 1.1, 8.4, 2.1, "3/4", k.M_RED)
    edge(4.6, 1.1, 8.4, 0.2, "1/4", k.M_ACC2, -0.42)
    for y, t, c in [(5.9, "6/20", k.M_INK), (4.0, "6/20", k.M_INK),
                    (2.1, "6/20", k.M_INK), (0.2, "2/20", k.M_INK)]:
        ax.text(9.4, y, t, va="center", color=c)
    ax.annotate("", xy=(11.0, 2.1), xytext=(11.0, 5.9),
                arrowprops=dict(arrowstyle="<->", color=k.M_GRN, lw=2.2))
    ax.text(12.9, 4.0, "第二次红\n6/20+6/20\n= 3/5", va="center", ha="center",
            color=k.M_GRN, fontweight="bold", linespacing=1.5)
    fig.subplots_adjust(left=0.01, right=0.99, top=0.99, bottom=0.01)
    fig.savefig(F("fig05_ball.png"), dpi=200)
    plt.close(fig)
    return F("fig05_ball.png")

def fig06_partition():
    """样本空间的划分"""
    fig, ax = k.new_fig(6.2, 3.8)
    ax.set_xlim(0, 12); ax.set_ylim(0, 6.4); ax.axis("off")
    ax.add_patch(Rectangle((0.5, 0.5), 11, 5.4, fc="white", ec=k.M_INK, lw=1.8))
    ax.text(1.0, 5.3, "Ω", color=k.M_INK, style="italic", va="center")
    cols = [k.M_ACC, k.M_ACC2, k.M_GRN, k.M_RED]
    xs = [0.5, 3.25, 6.0, 8.75]
    for i, (x, c) in enumerate(zip(xs, cols)):
        ax.add_patch(Rectangle((x, 0.5), 2.75, 5.4, fc=c, alpha=0.22, ec=c, lw=1.6))
        ax.text(x + 1.37, 1.15, f"B{i+1}", ha="center", va="center",
                color=c, fontweight="bold")
    ax.add_patch(Ellipse((6.0, 3.5), 10.2, 3.0, fc=k.M_INK, alpha=0.32, ec=k.M_INK, lw=2.2))
    ax.text(6.0, 4.55, "A", ha="center", va="center", color="white", fontweight="bold")
    for i, x in enumerate(xs):
        ax.text(x + 1.37, 3.0, f"AB{i+1}", ha="center", va="center", color="white")
    ax.set_title("A 被划分切成互斥的 AB₁…ABₙ", color=k.M_INK, pad=8)
    return k.save_fig(fig, F("fig06_partition.png"))

def fig07_terms():
    """先验·似然·证据·后验 四术语示意"""
    fig, ax = k.new_fig(5.4, 2.0)
    ax.set_xlim(0, 12); ax.set_ylim(0, 3.4); ax.axis("off")
    items = [("先验\nP(B)", k.M_ACC2, 1.4), ("似然\nP(A|B)", k.M_ACC, 4.5),
             ("证据\nP(A)", k.M_SLATE, 7.5), ("后验\nP(B|A)", k.M_GRN, 10.6)]
    for t, c, x in items:
        ax.text(x, 1.9, t, ha="center", va="center", color=c, fontweight="bold",
                linespacing=1.4,
                bbox=dict(boxstyle="round,pad=0.35", fc="white", ec=c, lw=2.4))
    for x, op in ((2.95, "×"), (6.0, "÷"), (9.05, "=")):
        ax.text(x, 1.9, op, ha="center", va="center", color=k.M_INK, fontweight="bold")
    ax.text(6.0, 0.25, "后验 ∝ 先验 × 似然", ha="center", va="center",
            color=k.M_RED, fontweight="bold")
    return k.save_fig(fig, F("fig07_terms.png"))

def fig08_factory():
    """例3：三车间次品率，堆叠条形 + 后验"""
    fig, axes = plt.subplots(1, 2, figsize=(7.1, 4.6))
    prior = [0.45, 0.35, 0.20]
    dr = [0.04, 0.02, 0.05]
    names = ["甲", "乙", "丙"]
    cols = [k.M_ACC, k.M_ACC2, k.M_GRN]
    joint = [p*d for p, d in zip(prior, dr)]

    ax = axes[0]
    ax.bar(names, prior, color=cols, alpha=0.30, width=0.62, zorder=3)
    ax.bar(names, joint, color=cols, width=0.62, zorder=4)
    for i, j in enumerate(joint):
        ax.text(i, j + 0.02, f"{j:.3f}", ha="center", color=k.M_RED, fontweight="bold")
    ax.set_ylim(0, 0.62)
    ax.set_yticks([0, 0.2, 0.4])
    k.style_axes(ax, "", "", origin=False)
    ax.grid(True, axis="y", color=k.M_RULE, lw=0.8, zorder=0)
    ax.set_title("① 次品 P(BᵢA)\nΣ = 0.035", color=k.M_INK, pad=10)

    ax = axes[1]
    post = [j/sum(joint) for j in joint]
    ax.bar(names, post, color=cols, width=0.62, zorder=3)
    for i, v in enumerate(post):
        ax.text(i, v + 0.03, f"{v:.1%}", ha="center", color=k.M_INK, fontweight="bold")
    ax.set_ylim(0, 0.80)
    ax.set_yticks([0, 0.25, 0.5])
    k.style_axes(ax, "", "", origin=False)
    ax.grid(True, axis="y", color=k.M_RULE, lw=0.8, zorder=0)
    ax.set_title("② 后验 P(Bᵢ|A)\n已知是次品", color=k.M_RED, pad=10)
    fig.subplots_adjust(left=0.10, right=0.98, top=0.76, bottom=0.08, wspace=0.36)
    fig.savefig(F("fig08_factory.png"), dpi=200)
    plt.close(fig)
    return F("fig08_factory.png")

def fig09_grid():
    """★核心★ 自然频率图：左=10000人全景，右=阳性池放大（110人，可逐格数）"""
    from matplotlib.colors import to_rgb
    import matplotlib.patches as mp
    fig, (axL, axR) = plt.subplots(1, 2, figsize=(11.4, 5.6))
    GREY = "#E7E4DC"
    N = 100
    grid = np.zeros((N, N), dtype=int)
    grid[0, 0:10] = 1            # 真阳性 10（红）
    grid[1, 0:100] = 2           # 假阳性 100（琥珀）
    img = np.zeros((N, N, 3))
    for v, c in {0: GREY, 1: k.M_RED, 2: k.M_ACC}.items():
        img[grid == v] = to_rgb(c)
    axL.imshow(img, interpolation="nearest")
    axL.set_xticks([]); axL.set_yticks([])
    for s in axL.spines.values():
        s.set_color(k.M_SLATE)
    axL.add_patch(Rectangle((-0.5, -0.5), 100, 2, fill=False, ec=k.M_GRN, lw=2.6, zorder=5))
    axL.set_title("① 10 000 人（每格 1 人）", color=k.M_INK, fontweight="bold", pad=10)
    axL.text(50, 58, "健康且阴性\n9 890 人", ha="center", va="center", color=k.M_SLATE)
    axL.annotate("", xy=(112, 1.0), xytext=(101, 1.0),
                 arrowprops=dict(arrowstyle="->", color=k.M_GRN, lw=2.6),
                 annotation_clip=False)

    axR.set_xlim(-0.6, 15.0); axR.set_ylim(-1.0, 12.6)
    axR.invert_yaxis(); axR.axis("off")
    for i in range(110):
        r, c = divmod(i, 10)
        col = k.M_RED if i < 10 else k.M_ACC
        axR.add_patch(Rectangle((c, r), 0.86, 0.86, fc=col, ec="white", lw=1.2))
    axR.add_patch(Rectangle((-0.14, -0.14), 10.14, 1.14, fill=False, ec=k.M_INK, lw=2.4))
    axR.text(10.6, 0.4, "真患病 10", color=k.M_RED, fontweight="bold", va="center")
    axR.text(10.6, 6.0, "假阳性 100", color=k.M_ACC, fontweight="bold", va="center")
    axR.set_title("② 阳性的 110 人（逐格可数）", color=k.M_INK, fontweight="bold", pad=10)

    hs = [mp.Patch(fc=GREY, label="健康·阴性 9 890"),
          mp.Patch(fc=k.M_ACC, label="假阳性 100"),
          mp.Patch(fc=k.M_RED, label="真阳性 10")]
    fig.legend(handles=hs, loc="lower center", ncol=3, fontsize=20,
               frameon=False, bbox_to_anchor=(0.5, 0.085))
    fig.text(0.5, 0.012, "P(病|阳) = 10/110 ≈ 9.1%", ha="center",
             fontsize=24, color=k.M_RED, fontweight="bold")
    fig.subplots_adjust(left=0.03, right=0.99, top=0.90, bottom=0.20, wspace=0.10)
    fig.savefig(F("fig09_grid.png"), dpi=200)
    plt.close(fig)
    return F("fig09_grid.png")

def fig10_medtree():
    """疾病检验树状图（自然频率）"""
    fig, ax = k.new_fig(7.4, 4.6)
    ax.set_xlim(0, 11.8); ax.set_ylim(0, 6.6); ax.axis("off")
    def box(x, y, t, c):
        ax.text(x, y, t, ha="center", va="center", color=c, fontweight="bold",
                bbox=dict(boxstyle="round,pad=0.30", fc="white", ec=c, lw=2), zorder=4)
    def edge(x1, y1, x2, y2, w, c, dy=0.42):
        ax.annotate("", xy=(x2-1.25, y2), xytext=(x1+1.05, y1),
                    arrowprops=dict(arrowstyle="->", color=c, lw=2), zorder=2)
        ax.text((x1+x2)/2, (y1+y2)/2 + dy, w, ha="center", va="center", color=c,
                fontweight="bold", zorder=5,
                bbox=dict(boxstyle="round,pad=0.10", fc=k.M_PAPER, ec="none"))
    box(1.2, 3.2, "10 000", k.M_INK)
    box(4.9, 5.4, "病 10", k.M_RED)
    box(4.9, 1.0, "健 9 990", k.M_ACC2)
    edge(1.2, 3.2, 4.9, 5.4, "0.1%", k.M_RED)
    edge(1.2, 3.2, 4.9, 1.0, "99.9%", k.M_ACC2, -0.42)
    box(8.8, 6.2, "阳 10", k.M_RED)
    box(8.8, 4.5, "阴 0", k.M_SLATE)
    box(8.8, 1.9, "阳 100", k.M_ACC)
    box(8.8, 0.3, "阴 9 890", k.M_SLATE)
    edge(4.9, 5.4, 8.8, 6.2, "99%", k.M_RED)
    edge(4.9, 5.4, 8.8, 4.5, "1%", k.M_SLATE, -0.42)
    edge(4.9, 1.0, 8.8, 1.9, "1%", k.M_ACC)
    edge(4.9, 1.0, 8.8, 0.3, "99%", k.M_SLATE, -0.42)
    ax.annotate("", xy=(10.4, 1.9), xytext=(10.4, 6.2),
                arrowprops=dict(arrowstyle="<->", color=k.M_GRN, lw=2.4))
    ax.text(11.3, 4.05, "阳性池 110", color=k.M_GRN, fontweight="bold",
            va="center", ha="center", rotation=90)
    fig.subplots_adjust(left=0.01, right=0.99, top=0.99, bottom=0.01)
    fig.savefig(F("fig10_medtree.png"), dpi=200)
    plt.close(fig)
    return F("fig10_medtree.png")

def fig11_prior_curve():
    """后验随患病率（先验）变化的曲线"""
    fig, ax = k.new_fig(7.3, 4.6)
    p = np.linspace(0.0001, 0.30, 500)
    se = sp = 0.99
    post = se*p / (se*p + (1-sp)*(1-p))
    ax.plot(p*100, post*100, color=k.M_ACC, lw=3.4, zorder=3)
    ax.axhline(50, color=k.M_RULE, lw=1.2, ls="--")
    pts = [(0.001, "0.1% → 9%", 5.0, 6.0), (0.01, "1% → 50%", 8.0, 40.0),
           (0.10, "10% → 92%", 13.0, 74.0)]
    for px, lab, tx, ty in pts:
        py = se*px/(se*px + 0.01*(1-px))
        ax.plot(px*100, py*100, "o", ms=11, color=k.M_RED, zorder=5)
        ax.annotate(lab, xy=(px*100, py*100), xytext=(tx, ty), color=k.M_INK,
                    arrowprops=dict(arrowstyle="->", color=k.M_RED, lw=1.6))
    ax.set_xlim(0, 30); ax.set_ylim(0, 112)
    ax.set_yticks([0, 50, 100])
    k.style_axes(ax, "先验 %", "后验 %", origin=False)
    ax.grid(True, color=k.M_RULE, lw=0.8, zorder=0)
    return k.save_fig(fig, F("fig11_prior_curve.png"))

def fig12_update():
    """先验 → 似然 → 后验 三联柱状图"""
    fig, axes = plt.subplots(1, 3, figsize=(11.2, 4.0))
    lab = ["患病", "健康"]
    data = [([0.001, 0.999], "① 先验", k.M_ACC2, ["0.1%", "99.9%"]),
            ([0.99, 0.01], "② 似然（阳性）", k.M_ACC, ["99%", "1%"]),
            ([0.0909, 0.9091], "③ 后验", k.M_GRN, ["9.1%", "90.9%"])]
    for ax, (v, t, c, tx) in zip(axes, data):
        b = ax.bar(lab, v, color=[c, k.M_RULE], width=0.55, zorder=3)
        for r, s in zip(b, tx):
            ax.text(r.get_x()+r.get_width()/2, r.get_height()+0.05, s,
                    ha="center", color=k.M_INK, fontweight="bold")
        ax.set_ylim(0, 1.25)
        ax.set_yticks([0, 0.5, 1.0])
        ax.set_title(t, color=c, fontweight="bold", pad=10)
        k.style_axes(ax, "", "", origin=False, grid=False)
        ax.grid(True, axis="y", color=k.M_RULE, lw=0.7, zorder=0)
    fig.subplots_adjust(left=0.06, right=0.99, top=0.86, bottom=0.10, wspace=0.30)
    fig.savefig(F("fig12_update.png"), dpi=200)
    plt.close(fig)
    return F("fig12_update.png")

def fig13_seq():
    """贝叶斯连续更新：多帧静态动画（连续阳性）"""
    fig, axes = plt.subplots(1, 4, figsize=(11.4, 4.0))
    p = 0.001
    frames = [p]
    for _ in range(3):
        p = 0.99*p / (0.99*p + 0.01*(1-p))
        frames.append(p)
    titles = ["检验前", "1 次阳性", "2 次阳性", "3 次阳性"]
    for i, (ax, v, t) in enumerate(zip(axes, frames, titles)):
        ax.bar(["病"], [v], color=k.M_RED, width=0.5, zorder=3)
        ax.bar(["健"], [1-v], color=k.M_RULE, width=0.5, zorder=3)
        ax.text(0.5, 1.30, f"{v:.1%}", ha="center", color=k.M_RED, fontweight="bold")
        ax.set_ylim(0, 1.55)
        ax.set_yticks([0, 0.5, 1.0])
        ax.set_title(t, color=k.M_INK, fontweight="bold", pad=10)
        k.style_axes(ax, "", "", origin=False, grid=False)
        ax.grid(True, axis="y", color=k.M_RULE, lw=0.7, zorder=0)
        if i:
            ax.annotate("", xy=(-0.42, 0.60), xytext=(-0.78, 0.60),
                        xycoords="axes fraction", textcoords="axes fraction",
                        arrowprops=dict(arrowstyle="->", color=k.M_ACC, lw=2.6),
                        annotation_clip=False)
    fig.subplots_adjust(left=0.06, right=0.99, top=0.86, bottom=0.10, wspace=0.62)
    fig.savefig(F("fig13_seq.png"), dpi=200)
    plt.close(fig)
    return F("fig13_seq.png")

def fig14_monty():
    """蒙提霍尔：三门枚举图"""
    fig, axes = plt.subplots(1, 3, figsize=(11.0, 4.5))
    cases = [("车在 1 号", [1, 0, 0], "换 → 输", k.M_RED),
             ("车在 2 号", [0, 1, 0], "换 → 赢", k.M_GRN),
             ("车在 3 号", [0, 0, 1], "换 → 赢", k.M_GRN)]
    for ax, (t, car, note, c) in zip(axes, cases):
        ax.set_xlim(0, 3.3); ax.set_ylim(0, 3.6); ax.axis("off")
        for i in range(3):
            x = 0.2 + i*1.0
            ax.add_patch(Rectangle((x, 1.15), 0.8, 1.7, fc="white", ec=k.M_INK, lw=2))
            if car[i]:
                ax.add_patch(Rectangle((x+0.12, 1.78), 0.56, 0.26, fc=k.M_RED, ec=k.M_RED))
                ax.add_patch(Rectangle((x+0.24, 2.04), 0.32, 0.18, fc=k.M_RED, ec=k.M_RED))
                for cx in (x+0.24, x+0.56):
                    ax.add_patch(Circle((cx, 1.76), 0.075, fc=k.M_INK, ec=k.M_INK))
            else:
                ax.add_patch(Ellipse((x+0.40, 1.95), 0.50, 0.30, fc=k.M_SLATE, ec=k.M_SLATE))
                ax.add_patch(Circle((x+0.63, 2.12), 0.13, fc=k.M_SLATE, ec=k.M_SLATE))
                for hx in (x+0.60, x+0.68):
                    ax.plot([hx, hx+0.03], [2.22, 2.34], color=k.M_INK, lw=1.4)
                for lx in (x+0.26, x+0.52):
                    ax.plot([lx, lx], [1.80, 1.62], color=k.M_SLATE, lw=2)
            ax.text(x+0.4, 0.92, f"{i+1}", ha="center", va="top", color=k.M_SLATE)
        ax.text(0.6, 0.30, "你选", ha="center", va="center", color=k.M_ACC, fontweight="bold")
        ax.text(1.65, 3.35, t, ha="center", va="center", color=k.M_INK, fontweight="bold")
        ax.text(2.55, 0.30, note, ha="center", va="center", color=c, fontweight="bold")
    fig.subplots_adjust(left=0.01, right=0.99, top=0.97, bottom=0.02, wspace=0.05)
    fig.savefig(F("fig14_monty.png"), dpi=200)
    plt.close(fig)
    return F("fig14_monty.png")

def fig15_switch():
    """换/不换胜率对比 + 模拟收敛"""
    fig, axes = plt.subplots(1, 2, figsize=(7.3, 4.4))
    ax = axes[0]
    b = ax.bar(["不换", "换"], [1/3, 2/3], color=[k.M_SLATE, k.M_GRN], width=0.55, zorder=3)
    for r, v in zip(b, [1/3, 2/3]):
        ax.text(r.get_x()+r.get_width()/2, v+0.03, f"{v:.0%}", ha="center",
                color=k.M_INK, fontweight="bold")
    ax.set_ylim(0, 0.92)
    ax.set_yticks([0, 1/3, 2/3])
    ax.set_yticklabels(["0", "1/3", "2/3"])
    k.style_axes(ax, "", "", origin=False)
    ax.grid(True, axis="y", color=k.M_RULE, lw=0.8, zorder=0)
    ax.set_title("胜率：换 = 2 倍", color=k.M_INK, pad=10)

    ax = axes[1]
    rng = np.random.default_rng(2024)
    n = 3000
    car = rng.integers(0, 3, n)
    win_sw = (car != 0).astype(float)
    win_st = (car == 0).astype(float)
    t = np.arange(1, n+1)
    ax.plot(t, np.cumsum(win_sw)/t, color=k.M_GRN, lw=2.4, label="换")
    ax.plot(t, np.cumsum(win_st)/t, color=k.M_SLATE, lw=2.4, label="不换")
    ax.axhline(2/3, color=k.M_GRN, ls="--", lw=1)
    ax.axhline(1/3, color=k.M_SLATE, ls="--", lw=1)
    ax.set_ylim(0, 1.02); ax.set_xlim(1, n)
    ax.set_xticks([0, 1500])
    ax.set_yticks([0, 1/3, 2/3, 1])
    ax.set_yticklabels(["0", "1/3", "2/3", "1"])
    k.style_axes(ax, "", "", origin=False)
    ax.grid(True, color=k.M_RULE, lw=0.8, zorder=0)
    ax.legend(frameon=False, loc="center right")
    ax.set_title("3000 局模拟", color=k.M_INK, pad=10)
    fig.subplots_adjust(left=0.09, right=0.97, top=0.86, bottom=0.12, wspace=0.36)
    fig.savefig(F("fig15_switch.png"), dpi=200)
    plt.close(fig)
    return F("fig15_switch.png")

def fig16_spam():
    """垃圾邮件过滤：词条似然比 + 后验"""
    fig, axes = plt.subplots(1, 2, figsize=(7.4, 4.4))
    ax = axes[0]
    words = ["中奖", "免费", "发票"]
    ps = [0.40, 0.25, 0.30]     # P(词|垃圾)
    ph = [0.001, 0.03, 0.02]    # P(词|正常)
    x = np.arange(len(words)); w = 0.36
    ax.bar(x-w/2, ps, w, color=k.M_RED, label="垃圾", zorder=3)
    ax.bar(x+w/2, ph, w, color=k.M_ACC2, label="正常", zorder=3)
    ax.set_xticks(x); ax.set_xticklabels(words)
    ax.set_ylim(0, 0.60)
    ax.set_yticks([0, 0.2, 0.4])
    k.style_axes(ax, "", "", origin=False)
    ax.grid(True, axis="y", color=k.M_RULE, lw=0.8, zorder=0)
    ax.legend(frameon=False, loc="upper right")
    ax.set_title("词的似然", color=k.M_INK, pad=10)

    ax = axes[1]
    prior = 0.20
    seq = ["先验", "中奖", "免费", "发票"]
    v = prior; vals = [v]
    for s, h in [(0.40, 0.001), (0.25, 0.03), (0.30, 0.02)]:
        v = s*v / (s*v + h*(1-v)); vals.append(v)
    ax.plot(seq, vals, "o-", color=k.M_ACC, lw=3, ms=11, zorder=3)
    for i in (0, 3):
        ax.text(i, vals[i]+0.10, f"{vals[i]:.0%}", ha="center",
                color=k.M_INK, fontweight="bold")
    ax.set_ylim(0, 1.32); ax.set_xlim(-0.45, 3.45)
    ax.set_yticks([0, 0.5, 1.0])
    k.style_axes(ax, "", "", origin=False)
    ax.grid(True, color=k.M_RULE, lw=0.8, zorder=0)
    ax.set_title("P(垃圾) 逐词更新", color=k.M_GRN, pad=10)
    fig.subplots_adjust(left=0.09, right=0.97, top=0.86, bottom=0.14, wspace=0.34)
    fig.savefig(F("fig16_spam.png"), dpi=200)
    plt.close(fig)
    return F("fig16_spam.png")

def fig17_hundred():
    """极端化钥匙：100 扇门"""
    fig, axes = plt.subplots(2, 1, figsize=(6.4, 4.0))
    for ax, (title, opened, c) in zip(axes, [
            ("① 你选 1 号　P = 1/100", 0, k.M_SLATE),
            ("② 开掉 98 扇羊门", 98, k.M_GRN)]):
        ax.set_xlim(-1, 104); ax.set_ylim(0, 1.9); ax.axis("off")
        for i in range(100):
            keep = (i == 0) or (i == 36 and opened)
            if opened and not keep:
                fc, ec = "#EDEAE3", k.M_RULE
            elif i == 0:
                fc, ec = k.M_ACC, k.M_ACC
            else:
                fc, ec = "white", k.M_INK
            ax.add_patch(Rectangle((i + 0.12, 0.75), 0.76, 1.0, fc=fc, ec=ec, lw=0.8))
        ax.text(0, 0.45, "你选", ha="left", va="top", color=k.M_ACC, fontweight="bold")
        if opened:
            ax.add_patch(Rectangle((36.12, 0.75), 0.76, 1.0, fc=k.M_GRN, ec=k.M_GRN, lw=0.8))
            ax.annotate("这扇门 = 99/100", xy=(37.0, 0.75), xytext=(46, 0.30),
                        color=k.M_GRN, fontweight="bold", va="center",
                        arrowprops=dict(arrowstyle="->", color=k.M_GRN, lw=2))
        ax.set_title(title, color=c, fontweight="bold", pad=6)
    fig.subplots_adjust(left=0.02, right=0.98, top=0.88, bottom=0.03, hspace=0.62)
    fig.savefig(F("fig17_hundred.png"), dpi=200)
    plt.close(fig)
    return F("fig17_hundred.png")

def fig18_board():
    """板书布局示意图"""
    fig, ax = k.new_fig(10.2, 4.4)
    ax.set_xlim(0, 12); ax.set_ylim(0, 6.8); ax.axis("off")
    ax.add_patch(Rectangle((0.15, 0.15), 11.7, 6.5, fc="#1F2A38", ec=k.M_INK, lw=2))
    cols = [(0.45, "左 · 概念", k.M_ACC2, "文氏图 Ω→B\nP(A|B)=P(AB)/P(B)\n乘法公式\n骰子 7/10"),
            (4.30, "中 · 推演", k.M_ACC, "划分 B₁…Bₙ\n全概率公式\n贝叶斯（红框）\n先验·似然·后验"),
            (8.15, "右 · 案例", k.M_GRN, "体检 10000 人\n阳性 10+100\n= 110 → 9.1%\n蒙提霍尔 2/3")]
    for x, t, c, body in cols:
        ax.add_patch(Rectangle((x, 1.45), 3.4, 4.9, fc="none", ec=c, lw=1.6, ls="--"))
        ax.text(x + 1.7, 5.9, t, ha="center", va="center", color=c, fontweight="bold")
        ax.text(x + 1.7, 3.7, body, ha="center", va="center", color="white", linespacing=1.9)
    ax.add_patch(Rectangle((0.45, 0.40), 11.1, 0.80, fc="none", ec=k.M_RED, lw=1.8))
    ax.text(6.0, 0.80, "P(A|B) ≠ P(B|A)　│　勿忘先验", ha="center", va="center",
            color=k.M_RED, fontweight="bold")
    return k.save_fig(fig, F("fig18_board.png"))

figs = {}
for fn in [fig01_guess, fig02_venn, fig03_dice, fig04_tree_basic, fig05_ball,
           fig06_partition, fig07_terms, fig08_factory, fig09_grid, fig10_medtree,
           fig11_prior_curve, fig12_update, fig13_seq, fig14_monty, fig15_switch,
           fig16_spam, fig17_hundred, fig18_board]:
    figs[fn.__name__[:5]] = fn()
print("figures:", len(figs))

# ============================ PPT ============================
p = k.new_deck()
FQ = lambda n: F(f"eq{n}.png")

# 1 封面
k.title_slide(p, "条件概率与贝叶斯定理",
              "第 05 讲　│　从「样本空间的缩小」到「信念的更新」",
              "高中数学 · 概率统计模块", "60 分钟 · 名师课件包")

# 2 学习目标
s = k.content_slide(p, "本节学习目标", "导览")
k.bullets(s, [
    "理解条件概率 P(A|B) 的定义与几何本质：样本空间的重新标准化",
    "掌握乘法公式、全概率公式，会用树状图组织多阶段随机试验",
    "推导并运用贝叶斯定理，说清先验 / 似然 / 证据 / 后验四个术语",
    "能用自然频率破解「疾病检验悖论」，解释假阳性为何压倒真阳性",
    "会完整分析蒙提霍尔问题，识别常见的概率直觉错误",
    "了解贝叶斯思想在垃圾邮件过滤、机器学习中的现代意义",
], x=0.85, y=1.6, w=7.0, size=17)
k.picture(s, figs["fig07"], x=7.4, y=2.7, w=5.4)
k.callout(s, "重点：贝叶斯定理的推导与应用\n难点：小概率事件的直觉偏差（基础率谬误）",
          x=0.85, y=5.9, w=7.0, h=1.1, kind="key")

# 3 第1幕
k.section_slide(p, "第 1 幕 · 情境导入", "一道让 90% 的人答错的题", "8 min")

# 4 悬念投票
s = k.content_slide(p, "先猜一猜：体检报告上的那个「阳性」", "8 min")
k.bullets(s, [
    "某地某病的患病率为 0.1%（每 1000 人中 1 人）",
    "某检验试剂准确率 99%：病人检出阳性 99%，健康人误报阳性仅 1%",
    ("问：小明体检呈阳性，他真患病的概率是多少？", 1),
    "请全班举手投票：99% ？90% ？50% ？还是更低？",
], x=0.85, y=1.6, w=5.4, size=17)
k.picture(s, figs["fig01"], x=6.5, y=1.5, w=6.2)
k.callout(s, "先记下你的直觉答案。本节课结束时，我们会用一张 100×100 的方格图，\n"
             "把这个答案彻底改写。", x=0.85, y=5.6, w=5.4, h=1.4, kind="warn")

# 5 第2幕
k.section_slide(p, "第 2 幕 · 概念建构", "条件概率：把世界缩小", "10 min")

# 6 文氏图几何解释
s = k.content_slide(p, "条件概率的几何本质", "概念")
k.full_picture(s, figs["fig02"], y=1.45, w=10.6)
k.callout(s, "「已知 B 发生」＝ 把样本空间从 Ω 换成 B。概率不是消失了，是分母换人了。",
          x=1.4, y=6.35, w=10.6, h=0.85, kind="note")

# 7 定义公式
s = k.content_slide(p, "定义与公式", "公式")
k.formula(s, r"$P(A\mid B)=\dfrac{P(AB)}{P(B)}\quad\left(P(B)>0\right)$",
          x=1.6, y=1.55, w=10.1, out=FQ(1))
k.bullets(s, [
    "分子 P(AB)：A 与 B 同时发生（交集的面积）",
    "分母 P(B)：新的「全集」的面积——重新标准化因子",
    "验证它确实是概率：P(Ω|B)=P(B)/P(B)=1，非负性、可加性均成立",
    "独立的等价刻画：P(A|B)=P(A) ⟺ P(AB)=P(A)P(B)",
], x=0.9, y=3.3, w=11.5, size=17)
k.callout(s, "常见错误：把 P(A|B) 与 P(AB) 混为一谈。前者分母是 P(B)，后者分母是 1。",
          x=0.9, y=6.25, w=11.5, h=0.85, kind="warn")

# 8 例1
s = k.content_slide(p, "例 1　掷两颗骰子（板演）", "例题")
k.picture(s, figs["fig03"], x=6.6, y=1.5, w=6.1)
k.bullets(s, [
    "设 A = {至少出现一个 6}，B = {点数之和 > 8}",
    ("无条件：P(A) = 11/36 ≈ 30.6%", 1),
    ("B 含 10 个样本点：P(B) = 10/36", 1),
    ("AB 含 7 个样本点：P(AB) = 7/36", 1),
    ("P(A|B) = (7/36) / (10/36) = 7/10 = 70%", 1),
    "结论：知道「和大于 8」这一信息后，出现 6 的可能性从 30.6% 跃升到 70%",
], x=0.85, y=1.6, w=5.6, size=16)
k.callout(s, "板演要点：先数 B 的格子（新分母），再数 AB 的格子（新分子）——\n"
             "条件概率永远是「在小房间里数数」。",
          x=0.85, y=5.75, w=5.6, h=1.3, kind="key")

# 9 乘法公式
s = k.content_slide(p, "乘法公式：把定义反过来用", "公式")
k.formula(s, r"$P(AB)=P(B)\,P(A\mid B)=P(A)\,P(B\mid A)$", x=1.6, y=1.5, w=10.1, out=FQ(2))
k.formula(s, r"$P(A_1A_2\cdots A_n)=P(A_1)P(A_2\mid A_1)P(A_3\mid A_1A_2)\cdots P(A_n\mid A_1\cdots A_{n-1})$",
          x=0.9, y=3.0, w=11.5, size=0.72, out=FQ(3))
k.picture(s, figs["fig04"], x=1.0, y=4.15, w=11.3)
k.callout(s, "树状图口诀：沿路径「连乘」，同名叶子「相加」。",
          x=0.9, y=6.55, w=11.5, h=0.7, kind="note")

# 10 例2 树
s = k.content_slide(p, "例 2　不放回摸球（多阶段试验）", "例题")
k.full_picture(s, figs["fig05"], y=1.45, w=10.4)
k.callout(s, "P(第二次红) = 3/5·2/4 + 2/5·3/4 = 3/5 —— 与第一次的 3/5 相同！\n"
             "「抽签公平性」的严格证明，正是全概率公式。",
          x=1.5, y=6.25, w=10.4, h=1.0, kind="key")

# 11 第3幕
k.section_slide(p, "第 3 幕 · 定理推演", "全概率公式 与 贝叶斯定理", "12 min")

# 12 划分
s = k.content_slide(p, "第一步：样本空间的划分", "推演")
k.picture(s, figs["fig06"], x=6.6, y=1.9, w=6.2)
k.bullets(s, [
    "称 B₁, B₂, …, Bₙ 为 Ω 的一个划分，若：",
    ("① 两两互斥：BᵢBⱼ = ∅ (i≠j)", 1),
    ("② 完备：B₁ ∪ B₂ ∪ … ∪ Bₙ = Ω", 1),
    ("③ 每个 P(Bᵢ) > 0", 1),
    "直观：把 Ω 切成互不重叠、又恰好铺满的若干块",
    "于是任一事件 A 被自然切分：A = AB₁ ∪ … ∪ ABₙ，且各块互斥",
], x=0.85, y=1.6, w=5.6, size=16)
k.callout(s, "划分 = 「原因」的完备清单。每个 Bᵢ 是 A 发生的一种可能原因。",
          x=0.85, y=5.9, w=5.6, h=1.0, kind="note")

# 13 全概率
s = k.content_slide(p, "全概率公式：由因求果", "推演")
k.formula(s, r"$P(A)=\sum_{i=1}^{n}P(AB_i)=\sum_{i=1}^{n}P(B_i)\,P(A\mid B_i)$",
          x=1.3, y=1.55, w=10.7, out=FQ(4))
k.bullets(s, [
    "推导：A = ⋃ ABᵢ 互斥 ⇒ P(A) = Σ P(ABᵢ)（可加性）",
    "再对每一项用乘法公式 P(ABᵢ) = P(Bᵢ)·P(A|Bᵢ)　□",
    "读法：P(A) 是各条「原因通道」贡献的加权平均，权重就是先验 P(Bᵢ)",
    "在树状图上：走到 A 的所有路径，路径连乘后求和",
], x=0.9, y=3.5, w=11.5, size=17)
k.callout(s, "全概率公式是贝叶斯定理的分母。分母算错，后验必错。",
          x=0.9, y=6.3, w=11.5, h=0.8, kind="warn")

# 14 贝叶斯推导
s = k.content_slide(p, "贝叶斯定理：由果溯因", "推演")
k.formula(s, r"$P(B_k\mid A)=\frac{P(AB_k)}{P(A)}=\frac{P(B_k)\,P(A\mid B_k)}{\sum_{i=1}^{n}P(B_i)\,P(A\mid B_i)}$",
          x=1.3, y=1.5, w=10.7, size=0.95, out=FQ(5))
k.bullets(s, [
    "推导只有两步：条件概率定义（换分母）+ 全概率公式（拆分母）",
    "分子来自乘法公式，分母来自全概率公式——没有任何新假设",
    "方向的翻转：已知 P(A|Bᵢ)（易测），求 P(Bᵢ|A)（想知道）",
    "这正是科学推断的骨架：从观测到的现象，反推背后的原因",
], x=0.9, y=3.5, w=11.5, size=17)
k.callout(s, "P(A|B) ≠ P(B|A)。「下雨时地湿」的概率≈1，但「地湿是因为下雨」远小于 1。",
          x=0.9, y=6.3, w=11.5, h=0.8, kind="warn")

# 15 四术语
s = k.content_slide(p, "四个术语：先验 · 似然 · 证据 · 后验", "概念")
k.full_picture(s, figs["fig07"], y=1.5, w=10.6)
k.formula(s, r"$P(B\mid A)=\frac{P(B)\cdot P(A\mid B)}{P(A)}$",
          x=4.2, y=5.5, w=4.9, size=0.85, out=FQ(6))
k.callout(s, "记忆句：后验 ∝ 先验 × 似然。", x=4.3, y=6.6, w=4.7, h=0.7, kind="key")

# 16 例3
s = k.content_slide(p, "例 3　三车间的次品（板演）", "例题")
k.picture(s, figs["fig08"], x=5.7, y=1.9, w=7.1)
k.bullets(s, [
    "甲/乙/丙车间产量占 45%/35%/20%，次品率 4%/2%/5%",
    ("(1) 任取一件是次品的概率？", 1),
    ("P(A)=.45×.04+.35×.02+.20×.05 = .018+.007+.010 = 0.035", 1),
    ("(2) 已知是次品，来自甲的概率？", 1),
    ("P(甲|A)=0.018/0.035 ≈ 51.4%", 1),
    ("同理：乙 20.0%，丙 28.6%（三者之和为 1，可自检）", 1),
], x=0.85, y=1.6, w=4.9, size=15.5)
k.callout(s, "注意丙：次品率最高（5%），但产量小，「背锅」份额只有 28.6%。\n"
             "先验（产量）与似然（次品率）共同决定后验。",
          x=0.85, y=5.9, w=4.9, h=1.3, kind="key")

# 17 第4幕
k.section_slide(p, "第 4 幕 · 高光时刻", "疾病检验悖论：直觉的崩塌", "15 min")

# 18 揭晓前的形式化
s = k.content_slide(p, "把体检问题写成数学语言", "建模")
k.bullets(s, [
    "记 D = 患病，D̄ = 健康；+ = 检验阳性",
    ("先验：P(D) = 0.001，P(D̄) = 0.999", 1),
    ("似然（灵敏度）：P(+|D) = 0.99", 1),
    ("似然（假阳性率 = 1−特异度）：P(+|D̄) = 0.01", 1),
    "要求的是后验 P(D|+)，很多人却在心里算成了似然 P(+|D)=99%",
], x=0.85, y=1.6, w=11.5, size=17)
k.formula(s, r"$P(D\mid +)=\frac{0.001\times0.99}{0.001\times0.99+0.999\times0.01}=\frac{0.00099}{0.01098}\approx 9.0\%$",
          x=0.9, y=4.5, w=11.5, size=0.9, out=FQ(7))
k.callout(s, "★ 基础率谬误（base rate fallacy）：人们把 P(+|D) 误当成 P(D|+)，"
             "同时完全忽略了 0.1% 这个先验。",
          x=0.9, y=6.25, w=11.5, h=0.9, kind="warn")

# 19 ★方块图
s = k.content_slide(p, "★ 用自然频率打破直觉：10 000 人在此", "高光")
k.full_picture(s, figs["fig09"], y=1.4, w=11.4)

# 20 方块图解读
s = k.content_slide(p, "为什么假阳性会「压倒」真阳性", "剖析")
k.picture(s, figs["fig10"], x=5.5, y=1.7, w=7.4)
k.bullets(s, [
    "1% 的误报率，乘上 9 990 个健康人 = 约 100 个假阳性",
    "99% 的检出率，乘上仅 10 个病人 = 约 10 个真阳性",
    "阳性池共 110 人，真病人只占 10/110 ≈ 9.1%",
    "关键：健康人的基数太大，即使误报率很低，绝对人数也碾压真病人",
], x=0.85, y=1.7, w=4.6, size=15.5)
k.callout(s, "一句话：小概率 × 大基数 ＞ 大概率 × 小基数。",
          x=0.85, y=5.9, w=4.6, h=1.0, kind="key")

# 21 三联更新图
s = k.content_slide(p, "先验 → 似然 → 后验：一次信念更新", "可视化")
k.full_picture(s, figs["fig12"], y=1.6, w=11.2)
k.callout(s, "检验并非无用：它把患病概率放大了 91 倍（0.1% → 9.1%）。\n"
             "它只是没有把「可能」变成「确定」。",
          x=1.9, y=5.9, w=9.5, h=1.1, kind="note")

# 22 连续更新
s = k.content_slide(p, "换一台仪器再测一次：贝叶斯可迭代", "可视化")
k.full_picture(s, figs["fig13"], y=1.6, w=11.4)
k.callout(s, "上一次的后验 = 下一次的先验。两次独立阳性后达 90.8%，三次后 99.0%。\n"
             "这正是医学上「初筛 + 确诊」两步走的数学依据。",
          x=1.9, y=5.8, w=9.5, h=1.2, kind="key")

# 23 先验曲线
s = k.content_slide(p, "先验决定一切：同一台仪器，不同的结论", "延伸")
k.picture(s, figs["fig11"], x=5.6, y=1.7, w=7.3)
k.bullets(s, [
    "普查（先验 0.1%）：阳性 → 仅 9.0% 真患病",
    "高危人群（先验 1%）：阳性 → 50.0% 真患病",
    "已有症状就医（先验 10%）：阳性 → 91.7% 真患病",
    "医生的问诊、症状、家族史，本质上都是在「抬高先验」",
], x=0.85, y=1.8, w=4.6, size=15.5)
k.callout(s, "所以「不建议对低患病率疾病全民普查」——不是技术问题，是贝叶斯问题。",
          x=0.85, y=5.7, w=4.6, h=1.3, kind="warn")

# 24 第5幕
k.section_slide(p, "第 5 幕 · 经典与现代", "蒙提霍尔问题 与 贝叶斯思想", "12 min")

# 25 蒙提霍尔枚举
s = k.content_slide(p, "蒙提霍尔问题：完整枚举", "例题")
k.full_picture(s, figs["fig14"], y=1.45, w=11.0)
k.callout(s, "规则前提：主持人知道车在哪，且必定打开一扇有羊的、你没选的门。\n"
             "三种等可能情形中，换门赢两次 → P(换门赢) = 2/3。",
          x=1.4, y=6.1, w=10.6, h=1.1, kind="key")

# 26 蒙提霍尔贝叶斯
s = k.content_slide(p, "蒙提霍尔的贝叶斯算法", "推演")
k.picture(s, figs["fig15"], x=5.6, y=1.7, w=7.3)
k.bullets(s, [
    "设你选 1 号，主持人开了 3 号（记为事件 H₃）",
    ("先验：P(车在 i) = 1/3，i = 1,2,3", 1),
    ("似然：P(H₃|车在1) = 1/2（可开2或3，随机）", 1),
    ("似然：P(H₃|车在2) = 1（只能开3）", 1),
    ("似然：P(H₃|车在3) = 0（不会开有车的门）", 1),
    ("证据：P(H₃) = 1/3·1/2 + 1/3·1 + 0 = 1/2", 1),
], x=0.85, y=1.7, w=4.6, size=14.5)
k.formula(s, r"$P(C_2\mid H_3)=\frac{(1/3)\times 1}{1/2}=\frac{2}{3}$",
          x=0.7, y=5.6, w=4.9, size=0.62, out=FQ(8))
k.callout(s, "记 C₂ = 车在 2 号门。换门即选 C₂，胜率 2/3。",
          x=0.85, y=6.65, w=4.6, h=0.65, kind="key")

# 27 错误剖析
s = k.content_slide(p, "常见错误剖析：为什么「50-50」是错的", "辨析")
k.bullets(s, [
    "错误论证：剩两扇门，所以各 1/2。",
    ("破绽：主持人的开门不是随机事件，他知情且被规则约束。", 1),
    ("「不开 2 号」这件事本身就是证据。", 1),
    "验证的两把钥匙：",
    ("① 主持人若不知情、随机开门恰好开出羊 → 才真是 50-50", 1),
    ("② 极端化：100 扇门（见右图）", 1),
    "你最初选中的概率永远是 1/3，主持人不改变它；"
    "剩下的 2/3 被压缩到了另一扇门上。",
], x=0.85, y=1.6, w=5.3, size=15)
k.picture(s, figs["fig17"], x=6.4, y=1.7, w=6.4)
k.callout(s, "教学锚点：概率会因「谁在提供信息、按什么规则提供」而改变。"
             "条件概率的条件，必须写清是「哪一个事件」。",
          x=0.85, y=5.85, w=5.3, h=1.4, kind="warn")

# 28 变式
s = k.content_slide(p, "变式训练（限时 6 分钟）", "变式")
k.bullets(s, [
    "变式 1（假阳性再造）：某罕见病患病率 1/10 000，检验灵敏度 100%（绝不漏诊），"
    "假阳性率 0.5%。求阳性者真患病的概率。",
    ("提示：P = (1e-4×1) / (1e-4×1 + 0.9999×0.005) ≈ 1.96%。"
     "即便「绝不漏诊」，阳性仍有 98% 是虚惊。", 1),
    "变式 2（三门推广）：n 扇门（n ≥ 3），你选 1 扇，主持人开掉 n−2 扇羊门，"
    "只留你选的和另一扇。求换门的胜率。",
    ("答：不换 1/n；换 = (n−1)/n。n = 100 时换门胜率 99%——极端化让直觉复位。", 1),
], x=0.85, y=1.6, w=11.5, size=16)
k.callout(s, "解题范式（三步走）：① 定义事件与划分　② 写清先验与似然　③ 套贝叶斯，"
             "算完必查「后验之和 = 1」。",
          x=0.85, y=5.9, w=11.5, h=1.1, kind="key")

# 29 现代意义
s = k.content_slide(p, "贝叶斯的现代生命力", "延伸")
k.picture(s, figs["fig16"], x=5.5, y=1.8, w=7.4)
k.bullets(s, [
    "垃圾邮件过滤：朴素贝叶斯，每个词都是一条证据，逐词更新后验",
    "机器学习：先验 = 正则化项；后验 = 训练后的模型信念",
    "医学诊断、司法证据、A/B 测试、故障定位……都是「由果溯因」",
    "科学方法论：不追求「证明」，而追求「用证据不断修正信念」",
], x=0.85, y=1.9, w=4.5, size=15),
k.callout(s, "贝叶斯给出的不是真理，而是「在现有证据下最合理的信念」——\n"
             "以及一条把新证据吸收进来的、可计算的规则。",
          x=0.85, y=5.7, w=4.5, h=1.4, kind="note")

# 30 小结
s = k.content_slide(p, "课堂小结：一条主线，三个公式", "小结")
k.formula(s, r"$P(A\mid B)=\frac{P(AB)}{P(B)}\ \Rightarrow\ P(A)=\sum_i P(B_i)P(A\mid B_i)\ \Rightarrow\ P(B_k\mid A)=\frac{P(B_k)P(A\mid B_k)}{\sum_i P(B_i)P(A\mid B_i)}$",
          x=0.6, y=1.5, w=12.1, size=0.62, out=FQ(9))
k.bullets(s, [
    "条件概率 = 样本空间的重新标准化（分母换人）",
    "全概率公式：由因求果，把 A 拆到每条原因通道上",
    "贝叶斯定理：由果溯因，后验 ∝ 先验 × 似然",
    "基础率谬误：忽视先验，是概率直觉最常见的塌方点",
    "自然频率（讲人数，不讲百分数）是打通直觉的最有效工具",
], x=0.85, y=3.0, w=11.5, size=17)
k.callout(s, "带走一句话：P(A|B) 与 P(B|A) 不是一回事，中间隔着一个先验。",
          x=0.85, y=6.4, w=11.5, h=0.75, kind="key")

# 31 分层作业
s = k.content_slide(p, "分层作业", "作业")
k.bullets(s, [
    "A 层（必做 · 巩固）",
    ("1. 教材习题：条件概率定义题 3 道（骰子/扑克/摸球）", 1),
    ("2. 三车间例题改数：产量 50/30/20，次品率 3%/5%/2%，求 P(A) 与三个后验", 1),
    "B 层（选做 · 提升）",
    ("3. 某测谎仪对说谎者判「谎」的概率 0.88，对诚实者误判「谎」0.15。"
     "若受测者中 5% 说谎，被判「谎」者真说谎的概率是多少？（答：≈ 23.6%）", 1),
    ("4. 连续两次独立检验均为阳性，重算体检问题的后验（答：≈ 90.8%）", 1),
    "C 层（挑战 · 拓展）",
    ("5. 写 20 行 Python 蒙特卡洛验证蒙提霍尔 2/3，并把主持人改为「不知情随机开门」，"
     "解释胜率为何变回 1/2", 1),
    ("6. 查阅「朴素贝叶斯」，说明它「朴素」在哪里（独立性假设），并举一个失效例子", 1),
], x=0.85, y=1.55, w=11.5, size=14.5)

# 32 板书提纲
s = k.content_slide(p, "板书设计提纲", "板书")
k.full_picture(s, figs["fig18"], y=1.45, w=10.2)
k.callout(s, "书写节奏：左栏随第 2 幕落笔，中栏随第 3 幕推演生成（贝叶斯用红框圈定），"
             "右栏随第 4、5 幕的案例逐步填充；下沿警示条在小结时统一补写。",
          x=1.55, y=6.15, w=10.2, h=1.05, kind="key")

pptx = k.save(p, os.path.join(OUT, "05_条件概率与贝叶斯定理.pptx"))
print("slides:", len(p.slides.__iter__.__self__._sldIdLst))

# ============================ 教案 DOCX ============================
from docx import Document
from docx.shared import Pt as DPt, Cm, RGBColor as DRGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Noto Serif CJK SC"
st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)

def H(t, lv=1):
    h = doc.add_heading(t, level=lv)
    for r in h.runs:
        r.font.name = "Noto Serif CJK SC"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")
        r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
        r.font.size = DPt(15 if lv == 1 else 12.5)
    return h

def P(t, bold=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = DPt(4)
    r = p.add_run(t); r.bold = bold; r.font.size = DPt(size)
    r.font.name = "Noto Serif CJK SC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")
    return p

t = doc.add_heading("《条件概率与贝叶斯定理》教学设计", level=0)
for r in t.runs:
    r.font.name = "Noto Serif CJK SC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
P("第 05 讲　│　高中数学 · 概率统计模块　│　新授课　│　1 课时（60 分钟）", size=10)

info = doc.add_table(rows=3, cols=4); info.style = "Table Grid"
rows = [("课题", "条件概率与贝叶斯定理", "课时", "1 课时（60 分钟）"),
        ("课型", "新授课（概念建构 + 定理推演 + 案例剖析）", "授课对象", "高二年级"),
        ("配套课件", "05_条件概率与贝叶斯定理.pptx（32 页，18 张精绘配图）", "教具", "投影、粉笔三色、计算器")]
for i, r in enumerate(rows):
    for j, c in enumerate(r):
        cell = info.rows[i].cells[j]
        cell.text = c
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.font.size = DPt(10)
                rr.font.name = "Noto Serif CJK SC"
                rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")
                if j % 2 == 0: rr.bold = True

H("一、教材分析")
P("本讲位于概率统计模块的核心枢纽位置。前承古典概型与事件的独立性，后启随机变量的分布与统计推断。"
  "条件概率是概率论中第一个真正「有思想」的概念——它把静态的样本空间计算，升级为动态的信息更新。"
  "教材通常按「定义 → 乘法公式 → 全概率 → 贝叶斯」的逻辑链推进，本设计在此基础上强化两点："
  "其一，用文氏图与面积模型揭示条件概率的几何本质（样本空间的重新标准化），避免学生把公式当作机械记忆；"
  "其二，把贝叶斯定理放到「疾病检验悖论」这一真实情境中引爆认知冲突，让公式承担起纠正直觉的功能。"
  "蒙提霍尔问题作为经典反直觉案例，既检验学生对「条件」的精确刻画能力，也是概率教学中不可多得的思维体操。")

H("二、学情分析")
P("学生已掌握古典概型、事件的和/积运算与独立性判定，具备用列举法计算简单概率的能力。存在的困难有三："
  "（1）符号障碍——P(A|B)、P(AB)、P(A)P(B) 三者常被混淆；"
  "（2）直觉障碍——面对小概率事件时普遍犯「基础率谬误」，会把 P(阳性|患病) 直接当成 P(患病|阳性)；"
  "（3）建模障碍——多阶段试验中不会主动构造划分，导致全概率公式的分母写错。"
  "对策：以自然频率（讲人数而非百分数）为脚手架，以树状图为组织工具，以「先猜后算」的认知冲突为驱动力。")

H("三、教学目标（三维）")
P("【知识与技能】", True)
P("1. 理解条件概率的定义 P(A|B)=P(AB)/P(B)，能说明其几何意义；\n"
  "2. 掌握乘法公式、全概率公式与贝叶斯定理，能独立完成推导；\n"
  "3. 能准确辨识并使用先验、似然、证据、后验四个术语；\n"
  "4. 能用贝叶斯定理求解疾病检验、次品溯源、蒙提霍尔等典型问题。")
P("【过程与方法】", True)
P("1. 经历「猜测—冲突—可视化—推演—反思」的完整认知过程，体会数学对直觉的矫正作用；\n"
  "2. 学会用树状图与自然频率方块图组织多阶段随机试验，发展数学建模与直观想象素养；\n"
  "3. 通过蒙特卡洛模拟验证理论结果，体会数据与理论的相互印证。")
P("【情感态度与价值观】", True)
P("1. 在疾病检验悖论中体会「小概率 × 大基数」的震撼，形成对数据的审慎态度；\n"
  "2. 了解贝叶斯思想在人工智能、医学、司法中的现代价值，感受数学的现实力量；\n"
  "3. 养成「用证据修正信念」的理性思维习惯。")

H("四、教学重点与难点")
P("【重点】条件概率的本质理解；全概率公式与贝叶斯定理的推导及应用。", True)
P("【难点】", True)
P("1. 小概率事件的直觉偏差（基础率谬误）——学生会「算对公式但不信结果」，需用自然频率方块图彻底击穿；\n"
  "2. 蒙提霍尔问题中「主持人知情且受规则约束」这一条件的数学刻画——难在把「信息来源」纳入事件定义；\n"
  "3. P(A|B) 与 P(B|A) 的方向性区分。")

H("五、教法与学法")
P("【教法】情境驱动法（体检悬念贯穿全课）、认知冲突法（先投票再揭晓）、"
  "数形结合法（文氏图 / 方块图 / 树状图三线并行）、变式教学法。")
P("【学法】猜想—验证、自主推导、小组辨析、动手枚举、反思归纳。")

H("六、教学准备")
P("教师：32 页 PPT 课件（含 18 张 matplotlib 精绘配图）、10000 人方格图彩色投影、三色粉笔、蒙提霍尔道具（三个纸杯 + 一枚硬币）。\n"
  "学生：练习本、计算器；课前完成独立性与古典概型的复习卷。")

H("七、教学过程（60 分钟时间轴）")
tb = doc.add_table(rows=1, cols=6)
tb.style = "Table Grid"
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ["环节 / 时长", "PPT", "教师活动", "学生活动", "设计意图"]
hd = tb.rows[0].cells
heads = ["环节", "时长", "PPT 页", "教师活动", "学生活动", "设计意图"]
for c, t2 in zip(hd, heads):
    c.text = t2
    for pp in c.paragraphs:
        for rr in pp.runs:
            rr.bold = True; rr.font.size = DPt(9.5)
            rr.font.name = "Noto Serif CJK SC"
            rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")

TL = [
 ("第1幕 情境导入", "8 min", "P1–P4",
  "呈现体检情境：患病率0.1%、准确率99%，某人阳性。组织全班对「他真患病的概率」举手投票（99%/90%/50%/<20%），"
  "把票数写在黑板一角，暂不揭晓答案，只留下一句「记住你的答案」。宣布本课学习目标。",
  "独立思考并举手投票；多数学生选择 99% 或 90%；对「暂不揭晓」产生强烈期待。",
  "以认知冲突开场，制造 60 分钟的悬念张力。让错误的直觉先「公开出场」，后续的推翻才有力量。"),
 ("第2幕 概念建构", "10 min", "P5–P10",
  "用两幅文氏图对比讲解：已知 B 发生，样本空间由 Ω 缩为 B，概率是「面积之比」的重新标准化。"
  "板书定义 P(A|B)=P(AB)/P(B)。以骰子网格图（例1）现场板演：先数 B 的 10 格，再数 AB 的 7 格。"
  "反过来得乘法公式，引出树状图，讲解「沿路径连乘、同名叶子相加」。例2 不放回摸球，指出抽签公平性。",
  "跟随文氏图理解「分母换人」；在草稿纸上数骰子网格，独立算出 7/10；在树状图上标注分支权重，"
  "自行验证「第二次摸到红球」的概率仍是 3/5，产生惊讶。",
  "以几何直观为概念奠基，避免公式沦为记忆负担。骰子网格提供可数的、可触摸的样本空间；"
  "抽签公平性是学生熟悉的生活经验，用它奖励刚学的公式。"),
 ("第3幕 定理推演", "12 min", "P11–P16",
  "引入「划分」概念（互斥 + 完备），说明它是「原因的完备清单」。引导学生自行推导全概率公式（两步：可加性 + 乘法公式）。"
  "在此基础上，只需再用一次条件概率定义，即得贝叶斯定理——强调「没有引入任何新假设」。"
  "用四色方框图讲清先验/似然/证据/后验，提炼「后验 ∝ 先验 × 似然」。例3 三车间次品，完整板演。",
  "小组合作完成全概率公式的推导（教师只给出 A = ⋃ABᵢ 的提示）；"
  "上台板演例3 的两问；发现「丙车间次品率最高但背锅份额不是最大」，讨论原因。",
  "让学生亲手推出定理，而非被动接受，落实「过程与方法」目标。三车间案例中先验与似然的拉锯，"
  "为下一幕的疾病悖论做认知铺垫——学生已经预感到「先验很重要」。"),
 ("第4幕 高光时刻", "15 min", "P17–P23",
  "回到开课的体检问题，先形式化（写清先验 0.001、似然 0.99 与 0.01），代入贝叶斯得 9.0%，全班哗然。"
  "随即投影 10000 人方格图：10 个真阳性（红）vs 100 个假阳性（琥珀），让「假阳性压倒真阳性」肉眼可见。"
  "追问：检验是否无用？——展示三联柱状图（0.1%→9.1%，放大 91 倍）与连续更新序列（两次阳性 90.8%）。"
  "最后展示后验-先验曲线，解释为何医学上「初筛+确诊」两步走、为何不对低患病率疾病全民普查。",
  "对照自己开课时的投票，体验强烈的认知反差；数方格图中的色块，用「人数」而非百分数重新算一遍 10/110；"
  "讨论「为什么医生要先问诊」，自己给出「抬高先验」的解释。",
  "本课高光。先用公式给出反直觉答案（制造冲突），再用自然频率给出可见解释（化解冲突），"
  "最后用「检验并非无用」化解可能的虚无感，落到「贝叶斯是修正而非推翻」的价值观。"
  "自然频率是认知心理学证实的最有效表征，务必留足时间让学生「数格子」。"),
 ("第5幕 经典与现代", "12 min", "P24–P29",
  "拿出三个纸杯做蒙提霍尔现场演示（每组玩 10 轮，记录换/不换的胜负）。随后投影三情形枚举图，"
  "再用贝叶斯严格计算 P(车在2|H₃)=2/3。剖析「剩两扇门所以 50-50」的错误：主持人知情且受规则约束，"
  "他的「不开 2 号」本身携带信息；给出「100 扇门」的极端化钥匙。布置 6 分钟变式训练（变式1、变式2）。"
  "最后展示垃圾邮件过滤的朴素贝叶斯图，点出机器学习中的先验。",
  "分组用纸杯实做 10 轮，汇总全班数据（换门胜率接近 2/3）；在「50-50」的错误论证上展开辩论；"
  "限时完成两道变式，同桌互批；对「100 扇门」的极端化恍然大悟。",
  "动手实验先于理论，让数据说服学生。极端化（100 扇门）是数学思维中最锋利的直觉工具。"
  "变式1 与本课高光同构，检验迁移；变式2 推广到 n 门，训练一般化能力。现代应用回应「学这个有什么用」。"),
 ("小结与作业", "3 min", "P30–P32",
  "以一条主线三个公式收束：条件概率 → 全概率（由因求果）→ 贝叶斯（由果溯因）。"
  "回到黑板一角的投票记录，请学生对比开课的直觉与现在的理解。布置 A/B/C 三层作业，讲清各层要求。",
  "齐读小结句「P(A|B) 与 P(B|A) 不是一回事，中间隔着一个先验」；记录作业；"
  "对自己开课时的答案做一句话反思。",
  "首尾呼应，让 60 分钟成为一个完整的认知闭环。分层作业照顾差异：A 层保底，B 层迁移，C 层用编程与文献拓展。"),
]
for r in TL:
    cells = tb.add_row().cells
    for c, t2 in zip(cells, r):
        c.text = t2
        for pp in c.paragraphs:
            pp.paragraph_format.space_after = DPt(2)
            for rr in pp.runs:
                rr.font.size = DPt(9)
                rr.font.name = "Noto Serif CJK SC"
                rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")
widths = [Cm(2.3), Cm(1.4), Cm(1.6), Cm(5.6), Cm(4.0), Cm(4.6)]
for row in tb.rows:
    for i, cw in enumerate(widths):
        row.cells[i].width = cw
P("合计：8 + 10 + 12 + 15 + 12 + 3 = 60 分钟。", True)

H("八、板书设计")
bb = doc.add_table(rows=2, cols=3); bb.style = "Table Grid"
bd = [("【左栏 · 概念区】\n文氏图：Ω 缩为 B\nP(A|B)=P(AB)/P(B)\n乘法公式 P(AB)=P(B)P(A|B)\n骰子网格：7/10",
       "【中栏 · 推演区】\n划分 B₁…Bₙ（互斥 + 完备）\n全概率 P(A)=ΣP(Bᵢ)P(A|Bᵢ)\n贝叶斯（红框重点）：\nP(Bₖ|A)=P(Bₖ)P(A|Bₖ)/P(A)\n先验 · 似然 · 证据 · 后验",
       "【右栏 · 案例区】\n体检树：10000\n → 病10 / 健9990\n → 阳性 10 + 100 = 110\n → 9.1%（红笔圈出）\n蒙提霍尔：1/3 vs 2/3")]
for j, c in enumerate(bd[0]):
    bb.rows[0].cells[j].text = c
bb.rows[1].cells[0].merge(bb.rows[1].cells[2]).text = \
    "【下沿 · 易错警示条】　P(A|B) ≠ P(B|A)　│　勿忘先验（基础率谬误）　│　小概率 × 大基数 ＞ 大概率 × 小基数"
for row in bb.rows:
    for c in row.cells:
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.font.size = DPt(9.5)
                rr.font.name = "Noto Serif CJK SC"
                rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")

H("九、分层作业")
P("【A 层 · 必做（全体，约 20 分钟）】", True)
P("1. 教材条件概率定义题 3 道（骰子、扑克、摸球）。\n"
  "2. 三车间改数：产量 50%/30%/20%，次品率 3%/5%/2%。求 P(次品) 及三个后验，并自检三者之和为 1。")
P("【B 层 · 选做（学有余力，约 20 分钟）】", True)
P("3. 测谎仪：对说谎者判「谎」概率 0.88，对诚实者误判 0.15，受测者 5% 说谎。求被判「谎」者真说谎的概率。（≈23.6%）\n"
  "4. 连续两次独立检验均阳性，重算体检问题的后验。（≈90.8%）体会「上一次的后验是下一次的先验」。")
P("【C 层 · 挑战（拓展，不限时）】", True)
P("5. 用 20 行 Python 蒙特卡洛验证蒙提霍尔的 2/3；再把主持人改为「不知情随机开门（恰好开出羊）」，"
  "解释胜率为何回到 1/2。撰写 200 字说明。\n"
  "6. 查阅「朴素贝叶斯分类器」，说明它「朴素」在何处（条件独立性假设），并举一个该假设失效的实例。")

H("十、教学反思（课后填写）")
for q in ["1. 开课投票中选择「99%」的学生比例：______%。揭晓答案时的课堂反应：",
          "2. 自然频率方格图是否有效击穿了直觉障碍？仍有疑问的学生集中在哪一环节？",
          "3. 蒙提霍尔纸杯实验的全班汇总胜率：换门 ______ / 不换 ______。与理论值的偏差是否引发讨论？",
          "4. 学生在变式训练中的主要错误类型（先验遗漏 / 分母写错 / 方向颠倒）：",
          "5. 时间分配是否合理？哪一幕需要压缩或扩展？",
          "6. 下次改进措施："]:
    P(q)
    doc.add_paragraph("　")
    doc.add_paragraph("________________________________________________________________________")

docx_path = os.path.join(OUT, "教案_05_条件概率与贝叶斯定理.docx")
doc.save(docx_path)
print("docx:", docx_path)
print("pptx:", pptx)
