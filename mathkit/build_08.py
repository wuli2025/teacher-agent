# -*- coding: utf-8 -*-
"""第08讲：空间向量与立体几何的代数化 —— 60分钟课件包生成脚本。"""
import sys, os
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, Polygon, Arc

OUT = '/mnt/c/Users/mi/Desktop/数学名师课件包/08_空间向量与立体几何'
FIG = os.path.join(OUT, 'figures')
os.makedirs(FIG, exist_ok=True)

# ---------------------------------------------------------------- 等轴测投影
# x 轴向右，y 轴向"里"（右上 45°压缩），z 轴向上。斜二测风格：
#   X = x + 0.42y ,  Y = z + 0.35y
# 视线方向沿 y 增大方向"进入纸面"，故 y 最大且被前方遮挡的顶点为隐藏顶点。
KY, KZ = 0.42, 0.35


def P(p):
    x, y, z = p
    return (x + KY * y, z + KZ * y)


def seg(ax, p, q, hidden=False, c=k.M_INK, lw=1.8, z=3):
    (x1, y1), (x2, y2) = P(p), P(q)
    kw = dict(dashes=(4, 3)) if hidden else {}
    ax.plot([x1, x2], [y1, y2], ls='--' if hidden else '-',
            color=c, lw=lw * (0.8 if hidden else 1.0),
            zorder=z, solid_capstyle='round', **kw)


def arrow(ax, p, q, c=k.M_ACC, lw=2.2, z=6, style='-|>', ms=16):
    (x1, y1), (x2, y2) = P(p), P(q)
    ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle=style,
                                 mutation_scale=ms, color=c, lw=lw, zorder=z,
                                 shrinkA=0, shrinkB=0))


def dot(ax, p, c=k.M_INK, s=34, z=7):
    X, Y = P(p)
    ax.scatter([X], [Y], s=s, color=c, zorder=z)


def lab(ax, p, t, dx=0.0, dy=0.0, c=k.M_INK, fs=20, w='bold'):
    """在顶点 p 的投影位置外侧写标签（dx/dy 为投影平面上的外移量）。"""
    X, Y = P(p)
    ax.text(X + dx, Y + dy, t, color=c, fontsize=fs, fontweight=w,
            ha='center', va='center', zorder=9)


def txt(ax, x, y, t, c=k.M_INK, fs=20, w='normal', ha='left', va='center'):
    """直接在投影平面（数据坐标）上写字。"""
    return ax.text(x, y, t, color=c, fontsize=fs, fontweight=w,
                   ha=ha, va=va, zorder=9)


def face(ax, pts, c=k.M_ACC, a=0.16, z=1):
    ax.add_patch(Polygon([P(p) for p in pts], closed=True, facecolor=c,
                         edgecolor='none', alpha=a, zorder=z))


def clean(ax, xl, xr, yb, yt):
    ax.set_xlim(xl, xr); ax.set_ylim(yb, yt)
    ax.set_aspect('equal'); ax.axis('off')


# 正方体顶点：A 原点，x=AB，y=AD，z=AA1；隐藏顶点为 D(0,1,0)
def cube_pts(s=1.0):
    return dict(A=(0, 0, 0), B=(s, 0, 0), C=(s, s, 0), D=(0, s, 0),
                A1=(0, 0, s), B1=(s, 0, s), C1=(s, s, s), D1=(0, s, s))


# 顶点标签外移向量（投影平面，基准：2.0 英寸/单位；osc 用于换算其它比例的图）
# 说明：投影后 A,B,C,C1,D1,A1 在轮廓六边形上，B1 与 D 是"内部"顶点，
# 它们的标签只能落进面内，故取该顶点三条棱之间最空的扇区方向。
VOFF = {'A': (-.19, -.16), 'B': (.11, -.21), 'C': (.21, -.13), 'D': (-.24, .07),
        'A1': (-.21, .10), 'B1': (-.11, .21), 'C1': (.21, .14), 'D1': (-.06, .21)}
VNM = {'A1': '$A_1$', 'B1': '$B_1$', 'C1': '$C_1$', 'D1': '$D_1$'}


def draw_cube(ax, s=1.0, only=None, osc=1.0, lw=1.8, c=k.M_INK, fs=20, vfix=None):
    """画正方体线框；only=需要标注的顶点列表（None=全标，[]=不标）。"""
    v = cube_pts(s)
    vis = [('A', 'B'), ('B', 'C'), ('A', 'A1'), ('B', 'B1'), ('C', 'C1'),
           ('A1', 'B1'), ('B1', 'C1'), ('C1', 'D1'), ('D1', 'A1')]
    hid = [('D', 'A'), ('D', 'C'), ('D', 'D1')]
    for p, q in vis:
        seg(ax, v[p], v[q], False, c, lw)
    for p, q in hid:
        seg(ax, v[p], v[q], True, k.M_SLATE, lw)
    keys = list(v) if only is None else only
    for key in keys:
        if vfix and key in vfix:          # 绝对偏移（不随 osc 放大，避免压到棱线）
            dx, dy = vfix[key]
        else:
            dx, dy = [t * osc for t in VOFF[key]]
        lab(ax, v[key], VNM.get(key, key), dx, dy, k.M_INK, fs)
    return v


def axes3(ax, o=(0, 0, 0), L=1.55, names=('x', 'y', 'z'), osc=1.0):
    ox, oy, oz = o
    arrow(ax, o, (ox + L, oy, oz), k.M_ACC2, 1.6, 5, ms=13)
    arrow(ax, o, (ox, oy + L, oz), k.M_ACC2, 1.6, 5, ms=13)
    arrow(ax, o, (ox, oy, oz + L), k.M_ACC2, 1.6, 5, ms=13)
    lab(ax, (ox + L, oy, oz), names[0], .12 * osc, -.14 * osc, k.M_ACC2, 20)
    lab(ax, (ox, oy + L, oz), names[1], .16 * osc, .06 * osc, k.M_ACC2, 20)
    lab(ax, (ox, oy, oz + L), names[2], -.04 * osc, .17 * osc, k.M_ACC2, 20)


# ================================================================ 图 01 建系与坐标
def f01():
    fig, ax = k.new_fig(8.6, 4.3)          # 幻灯片 full_picture(w=8.6)
    v = draw_cube(ax, osc=1.10)
    axes3(ax, L=1.42, osc=1.10)
    face(ax, [v['A'], v['B'], v['C'], v['D']], k.M_ACC2, .10)
    txt(ax, -0.35, 1.90, '正方体建系：实线可见，虚线被遮挡', k.M_INK, 20, 'bold')
    ys = 1.35
    for s_ in ['A(0, 0, 0)', 'B(1, 0, 0)', 'C(1, 1, 0)', '$C_1$(1, 1, 1)']:
        txt(ax, 2.55, ys, s_, k.M_ACC, 20, 'bold'); ys -= 0.34
    txt(ax, 2.55, -0.08, '以 A 为原点', k.M_SLATE, 20)
    clean(ax, -0.50, 4.30, -0.38, 2.05)
    return k.save_fig(fig, f'{FIG}/f01_cube.png')


# ================================================================ 图 02 基底分解
def f02():
    fig, ax = k.new_fig(6.3, 5.2)          # 右栏 picture(w=6.3)
    v = draw_cube(ax, only=['A', 'B', 'C1'])
    arrow(ax, (0, 0, 0), (1, 1, 1), k.M_RED, 2.8, 8)
    arrow(ax, (0, 0, 0), (1, 0, 0), k.M_ACC, 2.2, 7)
    arrow(ax, (1, 0, 0), (1, 1, 0), k.M_ACC, 2.2, 7)
    arrow(ax, (1, 1, 0), (1, 1, 1), k.M_ACC, 2.2, 7)
    lab(ax, (.5, 0, 0), r'$x\vec{e_1}$', 0, -.24, k.M_ACC, 20)
    lab(ax, (1, .5, 0), r'$y\vec{e_2}$', .36, -.18, k.M_ACC, 20)
    lab(ax, (1, 1, .5), r'$z\vec{e_3}$', .40, .02, k.M_ACC, 20)
    lab(ax, (.55, .55, .55), r'$\vec{a}$', .26, -.16, k.M_RED, 22)
    txt(ax, -0.55, 1.82, r'$\vec a=x\vec{e_1}+y\vec{e_2}+z\vec{e_3}$',
        k.M_INK, 22, 'bold')
    clean(ax, -0.60, 2.45, -0.55, 1.95)
    return k.save_fig(fig, f'{FIG}/f02_basis.png')


# ================================================================ 图 03 数量积投影
def f03():
    fig, ax = k.new_fig(8.4, 4.3)          # full_picture(w=8.4)
    O = np.array([0, 0]); a = np.array([3.4, 0.0]); b = np.array([1.9, 2.0])
    ax.add_patch(FancyArrowPatch(O, a, arrowstyle='-|>', mutation_scale=18,
                                 color=k.M_ACC, lw=2.8))
    ax.add_patch(FancyArrowPatch(O, b, arrowstyle='-|>', mutation_scale=18,
                                 color=k.M_ACC2, lw=2.8))
    proj = np.array([b[0], 0])
    ax.plot([b[0], proj[0]], [b[1], proj[1]], ls='--', color=k.M_SLATE, lw=1.5)
    ax.add_patch(FancyArrowPatch(O, proj, arrowstyle='-|>', mutation_scale=16,
                                 color=k.M_RED, lw=3.4))
    ax.plot([b[0] - .16, b[0] - .16, b[0]], [0, .16, .16], color=k.M_SLATE, lw=1.2)
    ax.add_patch(Arc(O, 1.4, 1.4, theta1=0, theta2=46.4, color=k.M_INK, lw=1.5))
    txt(ax, .92, .30, r'$\theta$', k.M_INK, 22, 'bold')
    txt(ax, 3.55, .05, r'$\vec{a}$', k.M_ACC, 24, 'bold')
    txt(ax, 1.55, 2.22, r'$\vec{b}$', k.M_ACC2, 24, 'bold')
    txt(ax, 0.30, -.55, r'投影长 $|\vec b|\cos\theta$', k.M_RED, 20, 'bold')
    txt(ax, -0.85, 2.72,
        r'数量积：$\vec a\cdot\vec b=|\vec a||\vec b|\cos\theta'
        r'=x_1x_2+y_1y_2+z_1z_2$', k.M_INK, 21, 'bold')
    ax.set_xlim(-0.90, 6.60); ax.set_ylim(-1.00, 2.95)
    ax.set_aspect('equal'); ax.axis('off')
    return k.save_fig(fig, f'{FIG}/f03_dot.png')


# ================================================================ 图 04 建系三型
def f04():
    fig, axs = plt.subplots(1, 3, figsize=(10.4, 3.9))   # full_picture(w=10.4)
    for a in axs:
        a.set_facecolor(k.M_PAPER)
    # ① 正方体：现成三垂直
    ax = axs[0]
    draw_cube(ax, only=[])
    axes3(ax, L=1.35, osc=2.15)
    ax.set_title('① 现成三垂直\n原点取公共顶点', fontsize=20, color=k.M_INK, pad=8)
    clean(ax, -0.7, 2.9, -0.6, 2.09)
    # ② 直棱柱：侧棱作 z 轴
    ax = axs[1]
    A, B, C = (0, 0, 0), (1.4, 0, 0), (0, 1.2, 0)
    A1, B1, C1 = (0, 0, 1.4), (1.4, 0, 1.4), (0, 1.2, 1.4)
    for p, q in [(A, B), (B, C), (A, A1), (B, B1), (C, C1), (A1, B1), (B1, C1), (C1, A1)]:
        seg(ax, p, q)
    seg(ax, A, C, True, k.M_SLATE)
    axes3(ax, L=1.55, osc=2.15)
    lab(ax, A, 'A', -.36, -.32); lab(ax, B, 'B', .02, -.48); lab(ax, C, 'C', -.30, .40)
    ax.set_title('② 侧棱⊥底面\n侧棱作 z 轴', fontsize=20, color=k.M_INK, pad=8)
    clean(ax, -0.7, 2.9, -0.6, 2.09)
    # ③ 无现成直角：作垂线
    ax = axs[2]
    A, B, C, Pk = (0, 0, 0), (1.6, 0, 0), (0.55, 1.3, 0), (0.8, 0.43, 1.5)
    for p, q in [(A, B), (B, C), (C, A), (A, Pk), (B, Pk), (C, Pk)]:
        seg(ax, p, q)
    H = (0.8, 0.43, 0)
    seg(ax, Pk, H, True, k.M_RED, 1.8)
    dot(ax, H, k.M_RED)
    lab(ax, H, 'H', .06, -.34, k.M_RED, 20)
    lab(ax, A, 'A', -.34, -.28); lab(ax, B, 'B', .34, -.28)
    lab(ax, C, 'C', .46, .10); lab(ax, Pk, 'P', -.34, .22)
    ax.set_title('③ 无现成直角\n作垂线"造"坐标系', fontsize=20, color=k.M_INK, pad=8)
    clean(ax, -0.7, 2.9, -0.6, 2.09)
    fig.suptitle('建系三型：建得好不好，决定这道题算得快不快',
                 fontsize=21, color=k.M_INK, fontweight='bold', y=0.99)
    return k.save_fig(fig, f'{FIG}/f04_setup.png')


# ================================================================ 图 05 异面直线夹角
def f05():
    fig, ax = k.new_fig(6.2, 5.2)          # 右栏 picture(w=6.2)
    v = draw_cube(ax, only=['A', 'B', 'C', 'A1', 'B1'])
    arrow(ax, v['A1'], v['B'], k.M_RED, 2.8, 8)
    arrow(ax, v['B1'], v['C'], k.M_ACC2, 2.8, 8)
    lab(ax, (0.55, 0, 0.45), r'$\vec a$', .17, .13, k.M_RED, 22)
    lab(ax, (1, 0.5, 0.5), r'$\vec b$', .30, .06, k.M_ACC2, 22)
    txt(ax, -0.55, 1.86, r'$\vec a=\overrightarrow{A_1B}$（红）', k.M_RED, 20, 'bold')
    txt(ax, -0.55, 1.62, r'$\vec b=\overrightarrow{B_1C}$（蓝）', k.M_ACC2, 20, 'bold')
    clean(ax, -0.60, 2.40, -0.55, 1.98)
    return k.save_fig(fig, f'{FIG}/f05_skew.png')


# ================================================================ 图 06 线面角
def f06():
    fig, ax = k.new_fig(6.3, 5.2)          # 右栏 picture(w=6.3)
    v = draw_cube(ax, only=['A1', 'B', 'B1', 'D', 'D1'])
    face(ax, [v['B'], v['D'], v['D1'], v['B1']], k.M_ACC2, .20)
    seg(ax, v['B'], v['D'], True, k.M_ACC2, 1.8)
    seg(ax, v['B1'], v['D1'], False, k.M_ACC2, 1.8)
    arrow(ax, v['A1'], v['B'], k.M_RED, 2.8, 8)
    n0 = (0.5, 0.5, 0.5)
    arrow(ax, n0, (n0[0] + 0.62, n0[1] + 0.62, n0[2]), k.M_ACC, 2.4, 9)
    lab(ax, (1.12, 1.12, 0.5), r'$\vec n$', .26, .04, k.M_ACC, 22)
    lab(ax, (0.55, 0, 0.45), r'$\vec a$', -.22, .10, k.M_RED, 22)
    txt(ax, -0.60, 1.84, r'$\sin\theta=\dfrac{|\vec a\cdot\vec n|}{|\vec a||\vec n|}$',
        k.M_RED, 22, 'bold')
    txt(ax, 0.62, 1.86, '（是 sin，不是 cos）', k.M_SLATE, 20)
    clean(ax, -0.65, 2.50, -0.55, 2.05)
    return k.save_fig(fig, f'{FIG}/f06_lineplane.png')


# ================================================================ 图 07 二面角与法向量
def f07():
    fig, ax = k.new_fig(9.4, 4.3)          # full_picture(w=9.4)
    l0, l1 = (0, -0.2, 0), (0, 2.2, 0)
    seg(ax, l0, l1, False, k.M_INK, 2.6)
    a1, a2 = (1.5, -0.2, 0), (1.5, 2.2, 0)
    face(ax, [l0, l1, a2, a1], k.M_ACC2, .18)
    seg(ax, l0, a1, False, k.M_SLATE, 1.4); seg(ax, l1, a2, False, k.M_SLATE, 1.4)
    seg(ax, a1, a2, False, k.M_SLATE, 1.4)
    b1, b2 = (-0.85, -0.2, 1.25), (-0.85, 2.2, 1.25)
    face(ax, [l0, l1, b2, b1], k.M_ACC, .18)
    seg(ax, l0, b1, False, k.M_SLATE, 1.4); seg(ax, l1, b2, False, k.M_SLATE, 1.4)
    seg(ax, b1, b2, False, k.M_SLATE, 1.4)
    lab(ax, (1.25, 0.15, 0), r'$\alpha$', .10, -.04, k.M_ACC2, 24)
    lab(ax, (-0.72, 0.15, 1.05), r'$\beta$', -.10, .02, k.M_ACC, 24)
    lab(ax, (0, -0.2, 0), r'$l$', -.24, -.06, k.M_INK, 24)
    m = (0, 1.0, 0)
    arrow(ax, m, (0, 1.0, 0.95), k.M_ACC2, 2.4, 9)
    lab(ax, (0, 1.0, 0.95), r'$\vec{n_1}$', -.24, .14, k.M_ACC2, 21)
    arrow(ax, m, (0.78, 1.0, 0.53), k.M_ACC, 2.4, 9)
    lab(ax, (0.78, 1.0, 0.53), r'$\vec{n_2}$', .28, .04, k.M_ACC, 21)
    txt(ax, -1.85, 2.42, r'$|\cos\varphi|='
        r'\dfrac{|\vec{n_1}\cdot\vec{n_2}|}{|\vec{n_1}||\vec{n_2}|}$',
        k.M_RED, 22, 'bold')
    txt(ax, 0.75, 2.44, '公式只给"大小"，正负号要看图判断', k.M_INK, 20, 'bold')
    clean(ax, -1.95, 4.55, -0.45, 2.65)
    return k.save_fig(fig, f'{FIG}/f07_dihedral.png')


# ================================================================ 图 08 符号取舍
def f08():
    fig, axs = plt.subplots(1, 2, figsize=(9.0, 3.6))   # full_picture(w=9.0)
    for a in axs:
        a.set_facecolor(k.M_PAPER)
    for ax, ttl, sgn in [
            (axs[0], '一进一出：$\\langle n_1,n_2\\rangle=\\varphi$', +1),
            (axs[1], '同进同出：$\\langle n_1,n_2\\rangle=\\pi-\\varphi$', -1)]:
        l0, l1 = (0, -0.2, 0), (0, 1.9, 0)
        seg(ax, l0, l1, False, k.M_INK, 2.4)
        a1, a2 = (1.3, -0.2, 0), (1.3, 1.9, 0)
        face(ax, [l0, l1, a2, a1], k.M_ACC2, .16)
        seg(ax, l0, a1); seg(ax, l1, a2); seg(ax, a1, a2)
        b1, b2 = (-0.75, -0.2, 1.1), (-0.75, 1.9, 1.1)
        face(ax, [l0, l1, b2, b1], k.M_ACC, .16)
        seg(ax, l0, b1); seg(ax, l1, b2); seg(ax, b1, b2)
        m = (0, 0.9, 0)
        arrow(ax, m, (0, 0.9, 0.9), k.M_ACC2, 2.2, 9)
        lab(ax, (0, 0.9, 0.9), r'$\vec{n_1}$', -.30, .16, k.M_ACC2, 20)
        if sgn > 0:
            arrow(ax, m, (-0.62, 0.9, -0.42), k.M_ACC, 2.2, 9)
            lab(ax, (-0.62, 0.9, -0.42), r'$\vec{n_2}$', -.30, -.14, k.M_ACC, 20)
        else:
            arrow(ax, m, (0.62, 0.9, 0.42), k.M_ACC, 2.2, 9)
            lab(ax, (0.62, 0.9, 0.42), r'$\vec{n_2}$', .34, .04, k.M_ACC, 20)
        ax.set_title(ttl, fontsize=20, color=k.M_INK, pad=8)
        clean(ax, -2.05, 2.35, -0.95, 1.45)
    fig.suptitle('先算 |cos|，再看图判断锐钝',
                 fontsize=21, color=k.M_RED, fontweight='bold', y=0.99)
    return k.save_fig(fig, f'{FIG}/f08_sign.png')


# ================================================================ 图 09 点到平面距离
def f09():
    fig, ax = k.new_fig(6.4, 5.2)          # 右栏 picture(w=6.4)
    v = draw_cube(ax, only=['A', 'B', 'D', 'A1'])
    face(ax, [v['A1'], v['B'], v['D']], k.M_ACC2, .28)
    seg(ax, v['A1'], v['B'], False, k.M_ACC2, 2.2)
    seg(ax, v['B'], v['D'], True, k.M_ACC2, 2.2)
    seg(ax, v['D'], v['A1'], True, k.M_ACC2, 2.2)
    G = (1 / 3, 1 / 3, 1 / 3)
    arrow(ax, G, (G[0] + 0.55, G[1] + 0.55, G[2] + 0.55), k.M_ACC, 2.4, 9)
    lab(ax, (0.95, 0.95, 0.95), r'$\vec n$', .26, .08, k.M_ACC, 22)
    seg(ax, v['A'], G, True, k.M_RED, 1.8, z=8)
    dot(ax, G, k.M_RED)
    lab(ax, (1 / 3, 1 / 3, 1 / 3), r'$d$', -.20, -.34, k.M_RED, 22)
    txt(ax, -0.65, 1.86, r'$d=\dfrac{|\overrightarrow{AM}\cdot\vec n|}{|\vec n|}$',
        k.M_RED, 22, 'bold')
    txt(ax, 0.35, 1.88, '　（不用作垂足）', k.M_SLATE, 20)
    clean(ax, -0.70, 2.55, -0.55, 2.05)
    return k.save_fig(fig, f'{FIG}/f09_dist.png')


# ================================================================ 图 10 例1
def f10():
    fig, ax = k.new_fig(5.9, 4.9)          # 右栏 picture(w=5.9)
    v = draw_cube(ax, only=['A', 'B', 'C', 'A1', 'B1'], osc=1.05)
    axes3(ax, L=1.25, osc=1.05)
    arrow(ax, v['A1'], v['B'], k.M_RED, 2.8, 8)
    arrow(ax, v['B1'], v['C'], k.M_ACC2, 2.8, 8)
    txt(ax, -0.60, 1.86, r'$\overrightarrow{A_1B}=(1,0,-1)$', k.M_RED, 20, 'bold')
    txt(ax, -0.60, 1.62, r'$\overrightarrow{B_1C}=(0,1,-1)$', k.M_ACC2, 20, 'bold')
    clean(ax, -0.65, 2.40, -0.55, 1.98)
    return k.save_fig(fig, f'{FIG}/f10_ex1.png')


# ================================================================ 图 11 例2
def f11():
    fig, ax = k.new_fig(5.9, 4.9)          # 右栏 picture(w=5.9)
    v = draw_cube(ax, only=['A1', 'B', 'B1', 'D', 'D1'], osc=1.05,
                  vfix={'D': (-0.21, 0.23)})
    axes3(ax, L=1.25, osc=1.05)
    face(ax, [v['B'], v['D'], v['D1'], v['B1']], k.M_ACC2, .22)
    seg(ax, v['B'], v['D'], True, k.M_ACC2, 1.8)
    seg(ax, v['B1'], v['D1'], False, k.M_ACC2, 1.8)
    seg(ax, v['B'], v['B1'], False, k.M_ACC2, 1.8)
    seg(ax, v['D'], v['D1'], True, k.M_ACC2, 1.8)
    arrow(ax, v['A1'], v['B'], k.M_RED, 2.8, 8)
    txt(ax, -0.60, 1.86, r'平面 $BB_1D_1D$', k.M_ACC2, 20, 'bold')
    txt(ax, -0.60, 1.62, r'$\vec n=(1,1,0)$', k.M_ACC, 20, 'bold')
    clean(ax, -0.65, 2.40, -0.55, 1.98)
    return k.save_fig(fig, f'{FIG}/f11_ex2.png')


# ================================================================ 图 12 例3 棱锥
def pyr():
    return dict(D=(0, 0, 0), A=(2, 0, 0), B=(2, 2, 0), C=(0, 2, 0), P=(0, 0, 2))


def f12():
    fig, ax = k.new_fig(5.9, 5.0)          # 右栏 picture(w=5.9)
    v = pyr()
    D, A, B, C, Pp = v['D'], v['A'], v['B'], v['C'], v['P']
    for p, q in [(A, B), (B, C), (Pp, A), (Pp, B), (Pp, C)]:
        seg(ax, p, q)
    seg(ax, D, A, True, k.M_SLATE); seg(ax, D, C, True, k.M_SLATE)
    seg(ax, D, Pp, True, k.M_SLATE)
    face(ax, [Pp, A, C], k.M_ACC, .20)
    face(ax, [Pp, C, D], k.M_ACC2, .20)
    seg(ax, Pp, C, False, k.M_RED, 2.8, z=6)
    axes3(ax, o=(0, 0, 0), L=2.2, osc=1.45, names=('x', '', 'z'))
    lab(ax, (0, 1.15, 0), 'y', -.20, .26, k.M_ACC2, 20)
    lab(ax, D, 'D', -.30, -.26); lab(ax, A, 'A', -.10, -.44)
    lab(ax, B, 'B', .34, -.10); lab(ax, C, 'C', .34, .26)
    lab(ax, Pp, 'P', -.30, .20)
    txt(ax, -0.50, 2.82, '棱 PC（红）：求二面角 $A\\!-\\!PC\\!-\\!D$',
        k.M_RED, 20, 'bold')
    clean(ax, -0.55, 3.67, -0.60, 2.95)
    return k.save_fig(fig, f'{FIG}/f12_ex3.png')


# ================================================================ 图 13 综合法 vs 向量法
def f13():
    fig, axs = plt.subplots(2, 1, figsize=(6.2, 5.6))   # 右栏 picture(w=6.2)
    for a in axs:
        a.set_facecolor(k.M_PAPER)
    # 上：综合法
    ax = axs[0]
    v = draw_cube(ax, only=['A1', 'B'], osc=1.50)   # D 处五线交汇，不标以免压线
    seg(ax, v['A1'], v['B'], False, k.M_RED, 2.6, z=6)
    seg(ax, v['A1'], v['D'], True, k.M_GRN, 2.4, z=6)
    seg(ax, v['B'], v['D'], True, k.M_GRN, 2.4, z=6)
    seg(ax, v['B1'], v['C'], False, k.M_ACC2, 2.6, z=6)
    txt(ax, 2.10, 1.45, '综合法', k.M_SLATE, 21, 'bold')
    txt(ax, 2.10, 1.05, '平移 + 辅助线', k.M_SLATE, 20)
    txt(ax, 2.10, 0.65, '想不到就空白', k.M_RED, 20)
    clean(ax, -0.45, 4.60, -0.45, 1.87)
    # 下：向量法
    ax = axs[1]
    v = draw_cube(ax, only=['A', 'B', 'C', 'A1', 'B1'], osc=1.50)
    axes3(ax, L=1.4, osc=1.50)
    arrow(ax, v['A1'], v['B'], k.M_RED, 2.6, 8)
    arrow(ax, v['B1'], v['C'], k.M_ACC2, 2.6, 8)
    txt(ax, 2.10, 1.45, '向量法', k.M_ACC, 21, 'bold')
    txt(ax, 2.10, 1.05, '建系→坐标→公式', k.M_INK, 20)
    txt(ax, 2.10, 0.65, '四步机械执行', k.M_GRN, 20)
    clean(ax, -0.45, 4.60, -0.45, 1.87)
    return k.save_fig(fig, f'{FIG}/f13_compare.png')


# ================================================================ 图 14 变式
def f14():
    fig, axs = plt.subplots(1, 2, figsize=(10.2, 4.2))   # full_picture(w=10.2)
    for a in axs:
        a.set_facecolor(k.M_PAPER)
    ax = axs[0]
    v = draw_cube(ax, only=['A', 'B', 'D', 'A1'], osc=1.40,
                  vfix={'D': (-0.22, 0.26)})
    axes3(ax, L=1.3, osc=1.40)
    face(ax, [v['A1'], v['B'], v['D']], k.M_ACC, .25)
    seg(ax, v['A1'], v['B'], False, k.M_ACC, 2.4, z=6)
    seg(ax, v['B'], v['D'], True, k.M_ACC, 2.4, z=6)
    seg(ax, v['D'], v['A1'], True, k.M_ACC, 2.4, z=6)
    ax.set_title('变式1　二面角 $A_1\\!-\\!BD\\!-\\!A$', fontsize=20,
                 color=k.M_INK, pad=8)
    clean(ax, -0.75, 2.65, -0.55, 1.85)
    ax = axs[1]
    v = draw_cube(ax, only=['A', 'B', 'D', 'A1'], osc=1.40,
                  vfix={'D': (-0.22, 0.26)})
    axes3(ax, L=1.3, osc=1.40)
    face(ax, [v['A1'], v['B'], v['D']], k.M_ACC2, .25)
    seg(ax, v['A1'], v['B'], False, k.M_ACC2, 2.2, z=6)
    seg(ax, v['B'], v['D'], True, k.M_ACC2, 2.2, z=6)
    seg(ax, v['D'], v['A1'], True, k.M_ACC2, 2.2, z=6)
    G = (1 / 3, 1 / 3, 1 / 3)
    seg(ax, v['A'], G, True, k.M_RED, 1.8, z=8)
    dot(ax, G, k.M_RED)
    lab(ax, (1 / 3, 1 / 3, 1 / 3), r'$d$', -.22, -.42, k.M_RED, 22)
    ax.set_title('变式2　点 $A$ 到平面 $A_1BD$ 的距离', fontsize=20,
                 color=k.M_INK, pad=8)
    clean(ax, -0.75, 2.65, -0.55, 1.85)
    return k.save_fig(fig, f'{FIG}/f14_var.png')


# ================================================================ 图 15 四步算法
def f15():
    fig, ax = k.new_fig(5.9, 2.5)          # 板书页 picture(w=5.9) / 正文页放大贴
    steps = [('①建系', '找垂直', k.M_ACC2), ('②写坐标', '点变数', k.M_ACC),
             ('③算向量', '终减起', k.M_GRN), ('④套公式', '对号入座', k.M_RED)]
    for i, (t, s_, c) in enumerate(steps):
        x = 0.02 + i * 1.47
        ax.add_patch(plt.Rectangle((x, 0.62), 1.42, 1.24, facecolor='white',
                                   edgecolor=c, lw=2.0, zorder=2))
        ax.add_patch(plt.Rectangle((x, 0.62), 0.07, 1.24, facecolor=c,
                                   edgecolor='none', zorder=3))
        ax.text(x + 0.76, 1.52, t, fontsize=20, color=c, fontweight='bold',
                ha='center', va='center', zorder=4)
        ax.text(x + 0.76, 1.00, s_, fontsize=20, color=k.M_INK,
                ha='center', va='center', zorder=4)
    ax.text(2.95, 0.24, '会建系 + 会写坐标 + 记住公式 = 稳拿分',
            fontsize=20, color=k.M_SLATE, ha='center', va='center')
    ax.set_xlim(0, 5.9); ax.set_ylim(0, 2.05); ax.axis('off')
    return k.save_fig(fig, f'{FIG}/f15_flow.png')


# ================================================================ 图 16 知识地图
def f16():
    fig, ax = k.new_fig(6.0, 4.2)          # 右栏 picture(w=6.0)
    ax.add_patch(plt.Rectangle((0.15, 3.35), 5.7, 0.72, facecolor='white',
                               edgecolor=k.M_ACC, lw=2.2, zorder=2))
    ax.text(3.0, 3.71, '坐标化：数量积 + 法向量', fontsize=21, color=k.M_ACC,
            fontweight='bold', ha='center', va='center', zorder=3)
    tgt = [('异面角', r'$\cos\theta=|\vec a\cdot\vec b|/(|\vec a||\vec b|)$', 2.42, k.M_ACC),
           ('线面角', r'$\sin\theta=|\vec a\cdot\vec n|/(|\vec a||\vec n|)$', 1.66, k.M_GRN),
           ('二面角', r'$|\cos\varphi|=|\vec{n_1}\cdot\vec{n_2}|/(|\vec{n_1}||\vec{n_2}|)$', 0.90, k.M_RED),
           ('点面距', r'$d=|\overrightarrow{AM}\cdot\vec n|/|\vec n|$', 0.14, k.M_ACC2)]
    for name, f, y, c in tgt:
        ax.add_patch(plt.Rectangle((0.15, y), 5.7, 0.64, facecolor='white',
                                   edgecolor=c, lw=1.8, zorder=2))
        ax.add_patch(plt.Rectangle((0.15, y), 0.07, 0.64, facecolor=c,
                                   edgecolor='none', zorder=3))
        ax.text(0.36, y + 0.32, name, fontsize=20, color=c, fontweight='bold',
                va='center', zorder=3)
        ax.text(1.45, y + 0.32, f, fontsize=20, color=k.M_INK, va='center', zorder=3)
    ax.add_patch(FancyArrowPatch((3.0, 3.33), (3.0, 3.10), arrowstyle='-|>',
                                 mutation_scale=16, color=k.M_SLATE, lw=2))
    ax.set_xlim(0, 6.0); ax.set_ylim(0.05, 4.15); ax.axis('off')
    return k.save_fig(fig, f'{FIG}/f16_map.png')


FIGS = [f01(), f02(), f03(), f04(), f05(), f06(), f07(), f08(),
        f09(), f10(), f11(), f12(), f13(), f14(), f15(), f16()]
print('figures:', len(FIGS))

# ================================================================ PPT
prs = k.new_deck()
FML = 0


def fml(slide, tex, **kw):
    global FML
    FML += 1
    return k.formula(slide, tex, out=f'{FIG}/_fml{FML:02d}.png', **kw)


# 1 封面
k.title_slide(prs, '空间向量与立体几何的代数化',
              '把"添辅助线"的技巧，翻译成人人可执行的算法',
              '第 08 讲', '60 分钟 · 高中数学选择性必修一')

# 2 学习目标
s = k.content_slide(prs, '这节课，我们要拿到什么', '目标')
k.bullets(s, [
    '空间向量基本定理：几何 → 坐标',
    '建系三型 + 建系口诀',
    '数量积：夹角、垂直、长度',
    '法向量：线面角、二面角、点面距',
    ('二面角"符号取舍"：最大易错点', 1),
    '四步算法：建系→坐标→向量→公式',
], y=1.55, w=5.9, size=20)
k.picture(s, FIGS[15], x=6.75, y=1.9, w=6.0)
k.callout(s, '一句话主线：立体几何的难，难在"想不到辅助线"；\n向量法把"灵感"变成"流程"。',
          x=0.85, y=5.85, w=5.9, h=1.15, kind='key')

# 3 引入
s = k.content_slide(prs, '为什么立体几何总是"会了就简单，不会就死路"', '引入')
k.bullets(s, [
    '综合法：证平行、证垂直、作辅助线',
    ('每步都靠"看出来"，看不出就空白', 1),
    '建系之后，图形变成一张数表',
    ('点→数组；线→方向向量；面→法向量', 1),
    '"看出来"变成"算出来"：可检验、可拿分',
], y=1.6, w=5.9, size=20)
k.picture(s, FIGS[12], x=6.6, y=1.6, w=6.2)
k.callout(s, '笛卡尔：一切几何问题，都能化为代数方程。\n今天在立体几何里兑现一次。',
          x=0.85, y=5.6, w=5.9, h=1.2, kind='note')

# 4 第1幕
k.section_slide(prs, '第 1 幕', '空间的坐标语言：把图形写成数', '0–13 min')

# 5 建系与点坐标
s = k.content_slide(prs, '空间直角坐标系：给每个点一张"身份证"', '13 min')
k.full_picture(s, FIGS[0], y=1.45, w=8.6)
k.callout(s, '右手系：伸右手，四指从 x 轴转向 y 轴，拇指指向 z 轴正方向。',
          x=2.4, y=6.35, w=8.6, h=0.75, kind='note')

# 6 基底分解
s = k.content_slide(prs, '空间向量基本定理：唯一的"拆解"', '13 min')
k.picture(s, FIGS[1], x=6.5, y=1.6, w=6.3)
k.bullets(s, [
    '不共面的 e₁,e₂,e₃ 可作基底',
    'a 唯一写成 x e₁+y e₂+z e₃',
    ('唯一性 = 坐标不会有两个答案', 1),
    '取两两垂直的单位基底最省事',
], y=1.7, w=5.8, size=20)
fml(s, r'$\vec a=x\vec{e_1}+y\vec{e_2}+z\vec{e_3}\;\Longleftrightarrow\;\vec a=(x,y,z)$',
    x=0.85, y=4.5, w=5.6, size=0.62)
k.callout(s, '几何对象一旦有了坐标，就进入了代数的"管辖范围"。',
          x=0.85, y=5.9, w=5.6, h=0.9, kind='key')

# 7 坐标运算
s = k.content_slide(prs, '坐标运算：四则运算即几何操作', '13 min')
fml(s, r'$\vec a\pm\vec b=(x_1\pm x_2,\;y_1\pm y_2,\;z_1\pm z_2)\qquad '
       r'\lambda\vec a=(\lambda x,\lambda y,\lambda z)$', x=0.9, y=1.55, w=11.5, size=0.72)
fml(s, r'$\overrightarrow{AB}=\vec{OB}-\vec{OA}=(x_B-x_A,\;y_B-y_A,\;z_B-z_A)$',
    x=0.9, y=3.0, w=9.6, size=0.72)
k._tb(s, k.Inches(0.95), k.Inches(3.95), k.Inches(11.0), k.Inches(0.4),
      '——"终点减起点"，全场使用频率最高的一步', 15, k.SLATE)
fml(s, r'$|\vec a|=\sqrt{x^2+y^2+z^2}\qquad '
       r'\vec a\parallel\vec b\Leftrightarrow \vec a=\lambda\vec b\qquad '
       r'\vec a\perp\vec b\Leftrightarrow \vec a\cdot\vec b=0$',
    x=0.9, y=4.45, w=11.5, size=0.72)
k.callout(s, '"终点减起点"是全场使用频率最高的一步，写错坐标 = 全题崩塌。请务必逐点核对。',
          x=1.4, y=6.1, w=10.5, h=0.85, kind='warn')

# 8 数量积
s = k.content_slide(prs, '数量积：夹角、垂直、长度的总开关', '13 min')
k.full_picture(s, FIGS[2], y=1.4, w=8.4)
k.callout(s, 'a·b = 0 ⟺ a⊥b　　cos⟨a,b⟩ = (a·b)/(|a||b|)　　|a|² = a·a',
          x=1.9, y=6.35, w=9.5, h=0.75, kind='key')

# 9 建系口诀
s = k.content_slide(prs, '建系口诀：建得好，一步登天；建得差，寸步难行', '13 min')
k.full_picture(s, FIGS[3], y=1.4, w=10.4)
k.callout(s, '口诀：垂线找轴，交点定原；两两垂直，右手排序；先底后顶，逐点落标。',
          x=1.4, y=5.40, w=10.5, h=0.85, kind='key')
k.callout(s, '原则：让尽可能多的点落在坐标轴或坐标面上——0 越多，算得越快。',
          x=1.4, y=6.32, w=10.5, h=0.5, kind='note')

# 10 第2幕
k.section_slide(prs, '第 2 幕', '三大工具：夹角 · 距离 · 法向量', '13–26 min')

# 11 异面直线夹角
s = k.content_slide(prs, '工具一：异面直线所成角', '26 min')
k.picture(s, FIGS[4], x=6.6, y=1.65, w=6.2)
k.bullets(s, [
    '取两直线的方向向量 a、b',
    '公式必须加绝对值',
    ('θ ∈ (0°, 90°]，不可能是钝角', 1),
], y=1.7, w=5.8, size=20)
fml(s, r'$\cos\theta=\dfrac{|\vec a\cdot\vec b|}{|\vec a|\,|\vec b|}$',
    x=0.9, y=3.7, w=4.2, size=0.75)
k.callout(s, '易错：忘记绝对值，算出 cosθ = −1/2 就写 120°。异面直线夹角没有钝角！',
          x=0.85, y=5.75, w=5.8, h=1.05, kind='warn')

# 12 法向量求法
s = k.content_slide(prs, '工具二：平面的法向量怎么求', '26 min')
k.bullets(s, [
    '平面内取两个不共线向量 AB、AC',
    '设 n = (x, y, z)，解方程组',
    ('n·AB = 0 且 n·AC = 0', 1),
    '无穷多解 → 令某分量为 1，凑整取一个',
    ('凑整技巧：让分母消失，不带分数', 1),
], y=1.65, w=6.2, size=20)
fml(s, r'$\vec n\cdot\overrightarrow{AB}=0\;,\quad \vec n\cdot\overrightarrow{AC}=0$',
    x=7.2, y=2.2, w=5.2, size=0.62)
fml(s, r'$\vec n=(1,1,1)\;\sim\;(2,2,2)$', x=7.9, y=4.2, w=3.6, size=0.55)
k._tb(s, k.Inches(7.2), k.Inches(5.0), k.Inches(5.4), k.Inches(0.4),
      '（等价：法向量只看方向，不看长短）', 14, k.SLATE)
k.callout(s, '法向量不唯一，长度无所谓、方向可反——但"方向可反"正是二面角符号问题的根源。',
          x=0.85, y=5.9, w=11.6, h=0.9, kind='note')

# 13 线面角
s = k.content_slide(prs, '工具三：线面角（注意是 sin！）', '26 min')
k.picture(s, FIGS[5], x=6.5, y=1.6, w=6.3)
k.bullets(s, [
    '直线方向向量 a，平面法向量 n',
    'θ 与 ⟨a, n⟩ 互余，故用 sin',
    ('θ ∈ [0°, 90°]，同样加绝对值', 1),
], y=1.7, w=5.7, size=20)
fml(s, r'$\sin\theta=\dfrac{|\vec a\cdot\vec n|}{|\vec a|\,|\vec n|}$',
    x=0.9, y=3.6, w=4.4, size=0.75)
k.callout(s, '记忆法："线面角，正弦管"——因为向量夹角量的是"和法线的偏离"，正好差 90°。',
          x=0.85, y=5.75, w=5.7, h=1.1, kind='key')

# 14 二面角
s = k.content_slide(prs, '工具四：二面角（最难，也最能拿分）', '26 min')
k.full_picture(s, FIGS[6], y=1.4, w=9.4)
k.callout(s, '两个法向量的夹角，要么等于二面角，要么等于它的补角——公式只给"大小"，不给"符号"。',
          x=1.4, y=6.3, w=10.5, h=0.85, kind='warn')

# 15 符号取舍
s = k.content_slide(prs, '★ 二面角的符号取舍：本讲最大易错点', '易错')
k.full_picture(s, FIGS[7], y=1.35, w=9.0)
k.callout(s, '三条判断路线（考场首选第 ① 条）：\n'
             '① 看图：二面角是锐角还是钝角？据此给 |cos| 添正负号。\n'
             '② 在两个半平面内各取一条垂直于棱的向量，直接算夹角，无需取舍。\n'
             '③ 法向量定向：一进一出取同号，同进同出取补角。',
          x=0.85, y=5.15, w=11.6, h=1.85, kind='warn')

# 16 点到平面距离
s = k.content_slide(prs, '工具五：点到平面的距离（免作垂足）', '26 min')
k.picture(s, FIGS[8], x=6.4, y=1.6, w=6.4)
k.bullets(s, [
    '取平面内任一点 M，作向量 AM',
    'AM 在 n 上的投影长即距离',
    ('传统法要作垂足证垂直，这里全省', 1),
], y=1.75, w=5.6, size=20)
fml(s, r'$d=\dfrac{|\overrightarrow{AM}\cdot\vec n|}{|\vec n|}$',
    x=0.9, y=3.65, w=4.2, size=0.75)
k.callout(s, '推论：线面距、面面距、异面直线间距离，都可转化为"点到平面距离"这一个公式。',
          x=0.85, y=5.75, w=5.6, h=1.15, kind='key')

# 17 第3幕
k.section_slide(prs, '第 3 幕', '四步算法 · 三题实战板演', '26–48 min')

# 18 四步算法
s = k.content_slide(prs, '四步算法：把技巧变成流水线', '48 min')
k.full_picture(s, FIGS[14], y=1.40, w=9.0)
k.callout(s, '① 建系：找三条两两垂直的线，原点取公共顶点　② 写坐标：逐点落标\n'
             '③ 算向量：终点减起点；解方程组求法向量　④ 套公式：夹角/距离对号入座\n'
             '这四步没有一步需要"灵感"——把天才的直觉，降维成凡人的步骤。',
          x=1.4, y=5.50, w=10.5, h=0.9, kind='key')

# 19 例1 题面
s = k.content_slide(prs, '例1　异面直线夹角（正方体）', '例题')
k.picture(s, FIGS[9], x=6.9, y=1.6, w=5.9)
k.bullets(s, [
    '题：正方体 ABCD-A₁B₁C₁D₁ 棱长 1，',
    ('求异面直线 A₁B 与 B₁C 所成角。', 1),
    '① 建系：以 A 为原点，三棱为轴',
    '② 写坐标：',
    ('A₁(0,0,1)、B(1,0,0)、B₁(1,0,1)、C(1,1,0)', 1),
], y=1.7, w=6.0, size=20)
k.callout(s, '建系提示：正方体是"送分体"——三条棱直接就是三条轴，不用任何辅助线。',
          x=0.85, y=5.9, w=6.0, h=1.0, kind='note')

# 20 例1 板演
s = k.content_slide(prs, '例1　板演：③ 算向量 → ④ 套公式', '例题')
fml(s, r'$\overrightarrow{A_1B}=B-A_1=(1,0,-1),\qquad '
       r'\overrightarrow{B_1C}=C-B_1=(0,1,-1)$', x=0.9, y=1.6, w=11.4, size=0.72)
fml(s, r'$\overrightarrow{A_1B}\cdot\overrightarrow{B_1C}=1\times0+0\times1+(-1)\times(-1)=1$',
    x=0.9, y=2.95, w=11.4, size=0.72)
fml(s, r'$\cos\theta=\frac{|1|}{\sqrt{1^2+0^2+(-1)^2}\cdot\sqrt{0^2+1^2+(-1)^2}}'
       r'=\frac{1}{\sqrt{2}\cdot\sqrt{2}}=\frac{1}{2}\;\Longrightarrow\;\theta=60^\circ$',
    x=0.9, y=4.3, w=11.4, size=0.72)
k.callout(s, '全过程没有一条辅助线，没有一次"想到"。四步走完，答案自己跳出来。',
          x=1.4, y=6.05, w=10.5, h=0.85, kind='key')

# 21 例2 题面
s = k.content_slide(prs, '例2　线面角（正方体）', '例题')
k.picture(s, FIGS[10], x=6.9, y=1.6, w=5.9)
k.bullets(s, [
    '题：求 A₁B 与平面 BB₁D₁D 所成角',
    '①② 同例1，坐标已知',
    '③ 求平面 BB₁D₁D 的法向量 n：',
    ('BD=(−1,1,0)，BB₁=(0,0,1)', 1),
    ('−x+y=0 且 z=0 → 取 n=(1,1,0)', 1),
], y=1.7, w=6.0, size=20)
k.callout(s, '小技巧：BB₁DD₁ 是"竖直"平面，法向量必然水平（z=0），可先猜后验。',
          x=0.85, y=5.9, w=6.0, h=1.0, kind='note')

# 22 例2 板演
s = k.content_slide(prs, '例2　板演：④ 套公式（记住是 sin）', '例题')
fml(s, r'$\vec a=\overrightarrow{A_1B}=(1,0,-1),\qquad \vec n=(1,1,0)$',
    x=0.9, y=1.6, w=10.6, size=0.72)
fml(s, r'$\vec a\cdot\vec n=1\times1+0\times1+(-1)\times0=1,\qquad '
       r'|\vec a|=\sqrt{2},\quad |\vec n|=\sqrt{2}$', x=0.9, y=2.9, w=11.4, size=0.72)
fml(s, r'$\sin\theta=\frac{|\vec a\cdot\vec n|}{|\vec a||\vec n|}'
       r'=\frac{1}{\sqrt{2}\cdot\sqrt{2}}=\frac{1}{2}\;\Longrightarrow\;\theta=30^\circ$',
    x=0.9, y=4.2, w=11.4, size=0.72)
k.callout(s, '此处若写成 cosθ = 1/2 → 60°，就是把"线面角"和"向量夹角"混为一谈——它们互余。',
          x=1.4, y=5.95, w=10.5, h=0.95, kind='warn')

# 23 例3 题面
s = k.content_slide(prs, '例3　二面角（四棱锥·含符号取舍）', '例题')
k.picture(s, FIGS[11], x=6.9, y=1.5, w=5.9)
k.bullets(s, [
    '题：P-ABCD 中 PD⊥底面 ABCD，',
    ('底面正方形边长 2，PD=2；求二面角 A-PC-D 的余弦值。', 1),
    '① 建系：DP 作 z 轴，DA 作 x 轴，DC 作 y 轴',
    '② 写坐标：',
    ('D(0,0,0)、A(2,0,0)、C(0,2,0)、P(0,0,2)', 1),
], y=1.65, w=6.0, size=20)
k.callout(s, '建系信号词："PD⊥底面"——这五个字就是出题人递给你的 z 轴。',
          x=0.85, y=5.9, w=6.0, h=1.0, kind='key')

# 24 例3 板演
s = k.content_slide(prs, '例3　板演：两个法向量 + 符号取舍', '例题')
k._tb(s, k.Inches(0.9), k.Inches(1.5), k.Inches(2.2), k.Inches(0.5), '平面 PCD：', 17, k.INK, True)
fml(s, r'$x=0\ \Rightarrow\ \vec{n_1}=(1,0,0)$', x=3.0, y=1.45, w=4.6, size=0.62)
k._tb(s, k.Inches(0.9), k.Inches(2.7), k.Inches(2.2), k.Inches(0.5), '平面 PAC：', 17, k.INK, True)
fml(s, r'$\overrightarrow{PA}=(2,0,-2),\ \overrightarrow{PC}=(0,2,-2)\ \Rightarrow\ \vec{n_2}=(1,1,1)$',
    x=3.0, y=2.6, w=9.2, size=0.62)
fml(s, r'$\cos\langle \vec{n_1},\vec{n_2}\rangle=\frac{1}{1\times\sqrt{3}}=\frac{\sqrt{3}}{3}$',
    x=0.9, y=3.95, w=7.6, size=0.66)
k.callout(s, '★ 取舍：从图上看，二面角 A-PC-D 显然是锐角（A 与 D 在棱 PC 同侧偏近），\n'
             '故 cos = +√3/3。若图形显示为钝角，则应取 −√3/3。',
          x=0.85, y=5.00, w=11.6, h=1.2, kind='warn')
k.callout(s, '答：二面角 A-PC-D 的余弦值为 √3/3。',
          x=0.85, y=6.5, w=11.6, h=0.65, kind='key')

# 25 第4幕
k.section_slide(prs, '第 4 幕', '高光时刻：几何代数化的威力', '48–54 min')

# 26 对照
s = k.content_slide(prs, '同一道题：综合法 vs 向量法', '高光')
k.picture(s, FIGS[12], x=6.6, y=1.6, w=6.2)
k.bullets(s, [
    '综合法（上图）：',
    ('① 发现 B₁C ∥ A₁D——要想得到！', 1),
    ('② 连 BD，∠BA₁D 即所求角', 1),
    ('③ 证 △A₁BD 为等边三角形 → 60°', 1),
    '向量法（下图）：',
    ('① 建系 ② 写坐标 ③ 算向量 ④ 套公式', 1),
    ('cosθ = 1/(√2·√2) = 1/2 → 60°', 1),
], y=1.6, w=5.6, size=20)
k.callout(s, '综合法赢在优美，向量法赢在稳定。\n考场上，稳定比优美更值钱。',
          x=0.85, y=5.95, w=5.6, h=0.6, kind='key')

# 27 思想史
s = k.content_slide(prs, '注脚：笛卡尔的"几何代数化"', '思想')
k.bullets(s, [
    '1637 年，笛卡尔《几何学》：用坐标把曲线写成方程',
    ('从此"作图技巧"让位于"方程求解"', 1),
    '同样的革命，今天在你的立体几何卷子上重演：',
    ('点 → 数组；线 → 方向向量；面 → 法向量', 1),
    ('角与距离 → 数量积的四则运算', 1),
    '代价：失去几何直观的优雅；回报：得到可执行、可检验的算法',
], y=1.6, w=11.6, size=20)
k.callout(s, '"我思故我在"的那个笛卡尔，也是把几何交给代数的那个笛卡尔。\n数学的进步，常常来自一次成功的"翻译"。',
          x=0.85, y=5.85, w=11.6, h=1.1, kind='note')

# 28 第5幕
k.section_slide(prs, '第 5 幕', '变式训练与小结', '54–60 min')

# 29 变式
s = k.content_slide(prs, '变式训练（正方体棱长为 1）', '变式')
k.full_picture(s, FIGS[13], y=1.4, w=10.2)
k.callout(s, '变式1：n₁=(0,0,1)，n₂=(1,1,1)，cos = √3/3（锐角）。\n'
             '变式2：d = |AA₁·n|/|n| = 1/√3 = √3/3。（两题共用同一个 n）',
          x=1.4, y=5.75, w=10.5, h=1.25, kind='note')

# 30 小结
s = k.content_slide(prs, '小结：这节课的一条主线、四个公式、一个陷阱', '小结')
k.bullets(s, [
    '一条主线：坐标化 = 把"添辅助线的技巧"翻译成"可执行的算法"',
    '四个公式：异面角(cos，加绝对值)、线面角(sin)、二面角(|cos|+取舍)、点面距',
    '一个陷阱：二面角的符号——公式只算大小，正负必须看图判断',
    '一句叮嘱：建系决定成败，坐标决定生死',
    ('原点选在垂直关系最集中的顶点，让 0 尽量多', 1),
], y=1.65, w=11.6, size=20)
k.callout(s, '当你不再问"这题该添哪条辅助线"，而是问"这题该怎么建系"——\n你就已经站到了笛卡尔这一边。',
          x=0.85, y=5.5, w=11.6, h=1.2, kind='key')

# 31 分层作业
s = k.content_slide(prs, '分层作业', '作业')
k.bullets(s, [
    'A 层（必做·基础）：正方体中求 AC₁ 与 BD 所成角；',
    ('求 AC₁ 与底面 ABCD 所成角；教材坐标运算 6 题（限时 15 分钟）。', 1),
    'B 层（必做·提高）：四棱锥 P-ABCD（PD⊥底面，边长 2，PD=2）：',
    ('求二面角 B-PC-D 的余弦值；求点 D 到平面 PBC 的距离，写全四步。', 1),
    'C 层（选做·挑战）：自选一道曾用综合法做过的立体几何题，',
    ('改用向量法重做，并写 200 字比较两种方法的思维成本（解题日志）。', 1),
], y=1.60, w=11.8, size=20)
k.callout(s, '所有作业必须标注：① 建系图　② 各点坐标　③ 向量与法向量　④ 公式与取舍依据。',
          x=0.85, y=6.30, w=11.6, h=0.6, kind='warn')

# 32 板书提纲
s = k.content_slide(prs, '板书提纲', '板书')
k.bullets(s, [
    '【左栏·主线】几何对象 → 坐标 → 向量 → 数量积',
    '【中栏·四步算法】① 建系 ② 写坐标 ③ 算向量 ④ 套公式',
    ('口诀：垂线找轴，交点定原；两两垂直，右手排序', 1),
], y=1.55, w=6.1, size=20)
k.bullets(s, [
    '【右栏·公式墙】',
    ('异面角 cosθ=|a·b|/(|a||b|)', 1),
    ('线面角 sinθ=|a·n|/(|a||n|)', 1),
    ('二面角 |cosφ|=|n₁·n₂|/(|n₁||n₂|)', 1),
    ('点面距 d=|AM·n|/|n|', 1),
    '【副板】二面角符号：看图定锐钝',
], x=6.85, y=1.55, w=5.7, size=20)
k.picture(s, FIGS[14], x=0.85, y=4.75, w=5.9)
k.callout(s, '留白区：学生板演例3 的取舍判断。',
          x=6.85, y=5.80, w=5.7, h=0.8, kind='note')

path = k.save(prs, f'{OUT}/08_空间向量与立体几何.pptx')
print('slides:', len(prs.slides.__iter__.__self__._sldIdLst), 'ppt:', path)

# ================================================================ 教案 docx
from docx import Document
from docx.shared import Pt as DPt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Noto Serif CJK SC'
st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')


def H(t, lv=1):
    p = doc.add_heading(t, lv)
    for r in p.runs:
        r.font.name = 'Noto Serif CJK SC'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        r.font.color.rgb = None
    return p


def Pp(t, b=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = b
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    return p


tt = doc.add_heading('教案：第08讲　空间向量与立体几何的代数化', 0)
for r in tt.runs:
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')

H('一、课题与课时', 1)
Pp('课题：空间向量与立体几何的代数化（用坐标把立体几何翻译成代数运算）')
Pp('课时：1 课时（60 分钟）　　课型：新授 + 方法建构课')
Pp('教材：人教A版《选择性必修 第一册》第一章 空间向量与立体几何')

H('二、教材分析', 1)
Pp('本讲处于"空间向量基本定理—坐标运算—立体几何应用"的枢纽位置。教材前段完成了向量语言的建立，'
   '本讲的任务是完成一次"世界观切换"：把综合几何中依赖辅助线技巧的角、垂直、距离问题，'
   '统一翻译为数量积与法向量的代数运算。其思想源头是笛卡尔的几何代数化纲领——'
   '一切几何问题化为方程问题。掌握本讲，学生便获得一套可机械执行的四步算法'
   '（建系→写坐标→算向量→套公式），这既是高考立体几何大题的主流解法，'
   '也是"数学抽象、逻辑推理、数学运算"三大核心素养的集中训练点。')

H('三、学情分析', 1)
Pp('学生已掌握：平面向量数量积、空间点线面位置关系、平面向量坐标运算。')
Pp('主要困难：① 建系找不到三条两两垂直的直线，或原点选择不当导致坐标含分数根式；'
   '② 混淆异面直线夹角（cos，须加绝对值）与线面角（sin）；'
   '③ 二面角只会算 |cos|，不会判断正负号，这是失分最集中的地方；'
   '④ 心理上仍留恋综合法，不信任"算出来"的结果。')

H('四、教学目标（三维）', 1)
Pp('知识与技能：理解空间向量基本定理与坐标表示；掌握数量积、法向量的求法；'
   '能用向量法求异面直线夹角、线面角、二面角与点到平面的距离。', True)
Pp('过程与方法：经历"同一道题两种解法"的对比过程，体会代数化把技巧转化为算法的价值；'
   '形成建系→写坐标→算向量→套公式的四步操作流程。', True)
Pp('情感态度与价值观：通过笛卡尔几何代数化的思想史注脚，'
   '感受数学"翻译"的力量，建立"稳定胜过灵感"的解题信念。', True)

H('五、教学重点与难点', 1)
Pp('重点：空间直角坐标系的建立；数量积与法向量在夹角、距离问题中的应用；四步算法的固化。')
Pp('难点：① 无现成垂直关系时的建系（需作垂线"造"坐标系）；'
   '② 二面角的符号取舍（本讲最大易错点）。')
Pp('突破策略：建系给"三型 + 口诀"；符号取舍给"三条判断路线"，'
   '并在例3中现场板演一次完整取舍。')

H('六、教法与学法', 1)
Pp('教法：问题驱动 + 对比教学 + 算法化示范。以"为什么立体几何会了就简单、不会就死路"设问引入，'
   '用综合法与向量法并排对照制造认知冲突，再以四步算法收束。')
Pp('学法：观察—类比—模仿—变式迁移。学生在例题中随堂板演，在变式中独立跑通四步。')

H('七、教学准备', 1)
Pp('PPT 课件（32 页，含 16 幅等轴测立体图与公式贴图）；正方体教具或磁性立体模型；'
   '三角板；学案（含例题与变式的空白建系图）；实物投影仪用于展示学生板演。')

H('八、教学过程（分钟级时间轴）', 1)
rows = [
    ('0–2', 'PPT1–2　封面与学习目标', '出示课题与"知识地图"，点明本讲主线：把"添辅助线"翻译成算法。宣布本节课将用同一道题演示两种世界观。',
     '浏览目标，回忆上节课的向量坐标运算。', '给出全局路线图，降低认知负荷，让学生知道"算什么、往哪走"。'),
    ('2–6', 'PPT3　引入：立体几何为何"会与不会"两极分化',
     '追问：上次立体几何大题卡在哪一步？展示综合法与向量法并排图，制造冲突。',
     '举手回答"卡在想不到辅助线"；观察对照图，产生"原来可以不靠灵感"的期待。',
     '从学生真实痛点出发，建立学习动机；引出笛卡尔纲领。'),
    ('6–8', 'PPT4　第1幕过场', '宣布进入"空间的坐标语言"。', '整理笔记。', '结构清晰，节奏分明。'),
    ('8–11', 'PPT5　空间直角坐标系与点坐标',
     '在正方体等轴测图上现场建系，逐点报坐标；强调右手系与"虚实线区分遮挡"的画图规范。',
     '在学案的空白正方体上同步标注 8 个顶点坐标。', '手脑并用，把"图形→数表"的转换变成肌肉记忆。'),
    ('11–14', 'PPT6–7　空间向量基本定理与坐标运算',
     '讲解唯一分解；板书"终点减起点"；强调坐标写错则全题崩塌。',
     '完成 2 个口算：已知两点求向量、求模长。', '打牢代数运算的地基。'),
    ('14–17', 'PPT8　数量积及其投影意义',
     '用投影图讲清 a·b 的几何意义，导出夹角、垂直、长度三个推论。',
     '回答"a·b=0 意味着什么"。', '把数量积确立为"总开关"，后续四个公式皆由此派生。'),
    ('17–20', 'PPT9　建系口诀与建系三型',
     '出示三型图；教口诀"垂线找轴，交点定原；两两垂直，右手排序；先底后顶，逐点落标"；'
     '强调"让 0 尽量多"的选原点原则。',
     '齐读口诀；对三型各说一个见过的题目原型。', '把最关键的成败点（建系）显性化、口诀化。'),
    ('20–22', 'PPT10　第2幕过场 + 法向量求法引入', '宣布进入三大工具。', '整理笔记。', '承上启下。'),
    ('22–26', 'PPT11–12　异面直线夹角 · 法向量求法',
     '给出 cos 公式并重锤"必须加绝对值"；示范解方程组求法向量，讲取整技巧。',
     '当堂练：求正方体中平面 A₁BD 的一个法向量。', '先给工具，再进例题，避免例题中被工具卡住。'),
    ('26–31', 'PPT13–14　线面角 · 二面角',
     '对比 sin 与 cos 的适用场景；用二面角图说明"法向量夹角与二面角相等或互补"。',
     '在学案上填写"四个公式速查表"。', '厘清最容易混淆的两组公式。'),
    ('31–35', 'PPT15　★ 二面角符号取舍（难点攻坚）',
     '出示"一进一出/同进同出"对照图，给出三条判断路线；强调考场首选"看图定锐钝"。',
     '小组讨论：为什么公式只能给出大小？举手陈述理由。',
     '把最大失分点单独立页、单独训练，形成条件反射。'),
    ('35–37', 'PPT16　点到平面距离', '讲投影法求距离，指出线面距/面面距均可归约于此。',
     '记录公式。', '完成工具箱，形成闭环。'),
    ('37–39', 'PPT17–18　第3幕过场 + 四步算法流程',
     '出示四步流水线图，宣布"接下来三道题，每一题都严格走这四步"。',
     '齐读四步。', '把算法固化为课堂契约，便于后续自查。'),
    ('39–43', 'PPT19–20　例1 异面直线夹角（完整板演）',
     '教师板演①②步，请一名学生上台完成③④步；点评"绝对值"是否漏写。',
     '一人板演，其余同步在学案上计算，同桌互查。', '首题降低难度、保证成功体验，坐实四步流程。'),
    ('43–47', 'PPT21–22　例2 线面角（含法向量）',
     '引导学生自主求法向量，教师只点拨"竖直平面法向量必水平"；'
     '重锤"线面角用 sin"，故意示范一次错解 cos 再纠正。',
     '独立完成四步，2 分钟后核对；找出错解的错因。',
     '用"预设错误"制造深刻记忆，防止考场混淆。'),
    ('47–52', 'PPT23–24　例3 二面角（含符号取舍）',
     '完整板演建系（抓"PD⊥底面"这一信号词）、两个法向量、|cos| 计算；'
     '重点演示如何看图判断锐角并定号。',
     '学生上台陈述"为何是锐角"，其余同学质疑与补充。',
     '在真实题境中攻克难点，把判断路线用出来。'),
    ('52–56', 'PPT25–27　高光：综合法 vs 向量法 + 笛卡尔注脚',
     '并排复盘例1的两种解法，比较思维成本；讲述笛卡尔 1637 年的几何代数化革命。',
     '感受、发问；记录"稳定胜过优美"的结论。',
     '完成价值观层面的收束，让方法上升为思想。'),
    ('56–59', 'PPT28–29　变式训练',
     '出示变式1（二面角 A₁-BD-A）与变式2（点 A 到平面 A₁BD 的距离），'
     '提示"一图两吃、共用同一法向量"；巡视指导。',
     '限时 3 分钟独立完成，两名学生报出答案与取舍依据。',
     '即时迁移，检验四步算法是否真正内化。'),
    ('59–60', 'PPT30–32　小结 · 分层作业 · 板书回看',
     '收束"一条主线、四个公式、一个陷阱"；布置 A/B/C 三层作业与提交规范。',
     '记录作业，选择自己的层级。', '分层落实，兼顾保底与拔尖。'),
]
tb = doc.add_table(rows=1, cols=5)
tb.style = 'Table Grid'
hdr = ['时间(min)', '教学环节 / PPT', '教师活动', '学生活动', '设计意图']
for i, h in enumerate(hdr):
    c = tb.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h); r.bold = True
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for row in rows:
    cells = tb.add_row().cells
    for i, t in enumerate(row):
        cells[i].text = ''
        r = cells[i].paragraphs[0].add_run(t)
        r.font.name = 'Noto Serif CJK SC'
        r.font.size = DPt(9)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for i, w in enumerate([Cm(1.9), Cm(3.4), Cm(6.2), Cm(4.0), Cm(4.0)]):
    for c in tb.columns[i].cells:
        c.width = w
Pp('　')
Pp('合计：2+4+2+3+3+3+3+2+4+5+4+2+2+4+4+5+4+3+1 = 60 分钟（各环节时长已按上表时间轴精确对齐）。')

H('九、板书设计', 1)
Pp('【左栏·主线】立体几何代数化：几何对象 → 坐标 → 向量 → 数量积')
Pp('【中栏·四步算法】① 建系　② 写坐标　③ 算向量　④ 套公式')
Pp('　　建系口诀：垂线找轴，交点定原；两两垂直，右手排序；先底后顶，逐点落标。')
Pp('【右栏·公式墙】')
Pp('　　异面直线角：cosθ = |a·b| / (|a||b|)　　（加绝对值，θ∈(0°,90°]）')
Pp('　　线　面　角：sinθ = |a·n| / (|a||n|)　　（用 sin，勿用 cos）')
Pp('　　二　面　角：|cosφ| = |n₁·n₂| / (|n₁||n₂|)　+ 符号取舍（看图定锐钝）')
Pp('　　点面距离　：d = |AM·n| / |n|')
Pp('【副板·易错区】二面角符号：公式只给大小，正负必须看图。')
Pp('【留白区】学生板演例3 的取舍判断过程。')

H('十、分层作业', 1)
Pp('A 层（必做·基础）：正方体中求 AC₁ 与 BD 所成角；求 AC₁ 与底面 ABCD 所成角；'
   '教材空间向量坐标运算 6 题（限时 15 分钟）。')
Pp('B 层（必做·提高）：四棱锥 P-ABCD（PD⊥底面，正方形边长 2，PD=2）中，'
   '求二面角 B-PC-D 的余弦值；求点 D 到平面 PBC 的距离。要求写全四步。')
Pp('C 层（选做·挑战）：自选一道曾用综合法解过的立体几何题，改用向量法重做，'
   '并撰写 200 字的"两种方法思维成本比较"解题日志。')
Pp('提交规范：所有作业必须标注 ① 建系图　② 各点坐标　③ 向量与法向量　④ 公式与取舍依据。')

H('十一、教学反思（课后填写）', 1)
for t in ['1. 建系环节：学生是否能独立找出三条两两垂直的直线？"无现成直角"型的掌握率如何？',
          '2. 二面角符号取舍：三条判断路线中，学生实际最常用哪一条？错误率是否下降？',
          '3. 时间分配：例3 是否超时？高光对比环节是否被压缩？',
          '4. 分层作业：C 层解题日志的完成质量与思想深度。',
          '5. 改进设想：']:
    Pp(t)
    Pp('　')
    Pp('　')

dp = f'{OUT}/教案_08_空间向量与立体几何.docx'
doc.save(dp)
print('docx:', dp)
