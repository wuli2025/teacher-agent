# -*- coding: utf-8 -*-
"""第10讲：函数的极限与连续性 —— 课件包构建脚本"""
import sys, os
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, Wedge

OUT = '/mnt/c/Users/mi/Desktop/数学名师课件包/10_函数的极限与连续'
FIG = os.path.join(OUT, 'figures')
os.makedirs(FIG, exist_ok=True)
F = lambda n: os.path.join(FIG, n)

# ============================================================ 配图
# 画布宽度 = 该图在幻灯片上的展示宽度（不得更大，否则贴图被缩小 → 图内字 <20pt）
# 20pt 硬底线下，图内只留短符号标签；长句一律交给 PPT 的 bullets/callout。

# fig1 芝诺 阿基里斯追龟：等比级数收敛   展示 6.3in
def fig_zeno():
    fig, ax = k.new_fig(6.3, 3.9)
    n = np.arange(0, 11)
    terms = 100 * (0.1 ** n)
    part = np.cumsum(terms)
    ax.bar(n, terms, color=k.M_ACC2, alpha=.55, width=.55, label='$s_n$', zorder=3)
    ax.plot(n, part, 'o-', color=k.M_ACC, lw=2.6, ms=7, label='$S_n$', zorder=4)
    ax.axhline(1000 / 9, color=k.M_RED, ls='--', lw=2, zorder=2)
    ax.text(3.0, 122, '$S=111.1$', color=k.M_RED, fontsize=20)
    k.style_axes(ax, 'n', '米', origin=False)
    ax.set_xlim(-.7, 10.7); ax.set_ylim(0, 150)
    ax.set_xticks([0, 2, 4, 6, 8, 10])
    ax.set_yticks([0, 50, 100])
    ax.legend(loc='center right', framealpha=.9)
    return k.save_fig(fig, F('01_zeno.png'))

# fig2 直观逼近 f(x)=(x^2-1)/(x-1)   展示 5.2in（第5页放大到 7.0in）
def fig_intuitive():
    fig, ax = k.new_fig(5.2, 3.4)
    x1 = np.linspace(-.4, .97, 200); x2 = np.linspace(1.03, 2.5, 200)
    ax.plot(x1, x1 + 1, color=k.M_ACC2, lw=3)
    ax.plot(x2, x2 + 1, color=k.M_ACC2, lw=3)
    ax.plot([1], [2], 'o', mfc='white', mec=k.M_RED, mew=2.6, ms=13, zorder=5)
    ax.axhline(2, color=k.M_RED, ls='--', lw=1.6)
    ax.text(-0.62, 2.2, '$L=2$', color=k.M_RED, fontsize=20)
    ax.annotate('', xy=(0.95, 0.55), xytext=(0.35, 0.55),
                arrowprops=dict(arrowstyle='->', color=k.M_GRN, lw=2.4))
    ax.annotate('', xy=(1.05, 0.55), xytext=(1.65, 0.55),
                arrowprops=dict(arrowstyle='->', color=k.M_ACC, lw=2.4))
    ax.text(1.0, -0.55, '$a=1$', color=k.M_RED, fontsize=20, ha='center')
    k.style_axes(ax, '', '')
    ax.set_xlim(-.75, 2.7); ax.set_ylim(-.8, 4.0)
    ax.set_xticks([2]); ax.set_yticks([3])
    return k.save_fig(fig, F('02_intuitive.png'))

# fig3 lim 与 f(a) 无关：三种取值同一极限   展示 11.4in
def fig_indep():
    fig, axes = plt.subplots(1, 3, figsize=(11.4, 3.6))
    x1 = np.linspace(-.2, .95, 100); x2 = np.linspace(1.05, 2.2, 100)
    cases = [('① $f(1)$ 无定义', None), ('② $f(1)=2$', 2.0), ('③ $f(1)=3.2$', 3.2)]
    for ax, (t, fv) in zip(axes, cases):
        ax.plot(x1, x1 + 1, color=k.M_ACC2, lw=2.8)
        ax.plot(x2, x2 + 1, color=k.M_ACC2, lw=2.8)
        ax.plot([1], [2], 'o', mfc='white', mec=k.M_ACC2, mew=2.4, ms=12, zorder=4)
        if fv is not None:
            ax.plot([1], [fv], 'o', color=k.M_RED, ms=12, zorder=5)
        ax.axhline(2, color=k.M_GRN, ls='--', lw=1.6)
        ax.text(-.15, 2.25, '$L=2$', color=k.M_GRN, fontsize=20)
        ax.set_title(t, fontsize=20, color=k.M_INK, pad=8)
        k.style_axes(ax, '', '', grid=True)
        ax.set_xlim(-.35, 2.45); ax.set_ylim(-.25, 4.3)
        ax.set_xticks([2]); ax.set_yticks([4])
    return k.save_fig(fig, F('03_lim_indep.png'))

# fig4 ε-δ 双带图（核心）  展示 6.4in —— 图内只留短符号，解释句全在 PPT
def fig_epsdelta():
    fig, ax = k.new_fig(6.4, 4.4)
    a, L, eps, d = 1.0, 2.0, 0.6, 0.45
    x = np.linspace(-.3, 2.4, 400)
    ax.add_patch(Rectangle((-.4, L - eps), 4.0, 2 * eps, color=k.M_ACC, alpha=.20, zorder=1))
    ax.add_patch(Rectangle((a - d, -.7), 2 * d, 5.6, color=k.M_ACC2, alpha=.22, zorder=1))
    ax.plot(x, x + 1, color=k.M_INK, lw=3, zorder=4)
    ax.axhline(L, color=k.M_RED, ls='--', lw=1.4, zorder=3)
    ax.axvline(a, color=k.M_RED, ls='--', lw=1.4, zorder=3)
    ax.plot([a], [L], 'o', mfc='white', mec=k.M_RED, mew=2.4, ms=12, zorder=6)
    for s in (1, -1):
        ax.axhline(L + s * eps, color=k.M_ACC, lw=1.6, zorder=2)
        ax.axvline(a + s * d, color=k.M_ACC2, lw=1.6, zorder=2)
    ax.text(2.68, L + eps + .12, '$L+\\varepsilon$', color=k.M_ACC, fontsize=20)
    ax.text(2.68, L - eps - .42, '$L-\\varepsilon$', color=k.M_ACC, fontsize=20)
    ax.text(a - d - 0.10, 4.55, '$a-\\delta$', color=k.M_ACC2, fontsize=20, ha='right')
    ax.text(a + d + 0.10, 4.55, '$a+\\delta$', color=k.M_ACC2, fontsize=20, ha='left')
    ax.text(-0.28, L + .12, '$L$', color=k.M_RED, fontsize=20)
    ax.text(a + .08, -.50, '$a$', color=k.M_RED, fontsize=20)
    k.style_axes(ax, '', '')
    ax.set_xlim(-.60, 3.85); ax.set_ylim(-.65, 5.15)
    ax.set_xticks([]); ax.set_yticks([])
    return k.save_fig(fig, F('04_eps_delta.png'))

# fig5 ε 收紧 δ 随之收紧（三子图）   展示 12.2in
def fig_eps_shrink():
    fig, axes = plt.subplots(1, 3, figsize=(12.2, 4.0))
    a, L = 1.0, 2.0
    x = np.linspace(0, 2.0, 300)
    for ax, eps in zip(axes, [0.6, 0.3, 0.1]):
        d = eps
        ax.add_patch(Rectangle((0, L - eps), 2.0, 2 * eps, color=k.M_ACC, alpha=.22))
        ax.add_patch(Rectangle((a - d, 0.6), 2 * d, 2.9, color=k.M_ACC2, alpha=.25))
        ax.plot(x, x + 1, color=k.M_INK, lw=2.6)
        ax.plot([a], [L], 'o', mfc='white', mec=k.M_RED, mew=2.2, ms=10)
        ax.set_title('$\\varepsilon=%.1f\\Rightarrow\\delta=%.1f$' % (eps, d),
                     fontsize=20, color=k.M_ACC if eps > .2 else k.M_RED, pad=10)
        ax.set_xlim(.2, 1.8); ax.set_ylim(1.0, 3.0)
        ax.set_xticks([1]); ax.set_xticklabels(['$a$'])
        ax.set_yticks([2]); ax.set_yticklabels(['$L$'])
        ax.tick_params(colors=k.M_SLATE)
        for sp in ('right', 'top'):
            ax.spines[sp].set_visible(False)
    return k.save_fig(fig, F('05_eps_shrink.png'))

# fig6 左右极限不等：跳跃间断   展示 6.4in
def fig_jump():
    fig, ax = k.new_fig(6.4, 4.0)
    x1 = np.linspace(-2.2, 0, 200); x2 = np.linspace(0, 2.2, 200)
    ax.plot(x1, x1 + 1, color=k.M_ACC2, lw=3.2)
    ax.plot(x2, x2 - 1, color=k.M_ACC, lw=3.2)
    ax.plot([0], [1], 'o', color=k.M_ACC2, ms=12, zorder=5)
    ax.plot([0], [-1], 'o', mfc='white', mec=k.M_ACC, mew=2.6, ms=12, zorder=5)
    ax.annotate('', xy=(0.06, -0.92), xytext=(0.06, 0.92),
                arrowprops=dict(arrowstyle='<->', color=k.M_RED, lw=2.4))
    ax.text(0.30, 0.80, '跳跃 $=2$', color=k.M_RED, fontsize=20)
    ax.text(-2.75, 2.5, '$\\lim_{x\\to0^-}f=1$', color=k.M_ACC2, fontsize=20)
    ax.text(0.30, -3.05, '$\\lim_{x\\to0^+}f=-1$', color=k.M_ACC, fontsize=20)
    k.style_axes(ax, '', '')
    ax.set_xlim(-2.85, 2.9); ax.set_ylim(-3.5, 3.9)
    ax.set_xticks([-2, 2]); ax.set_yticks([-2, 2])
    return k.save_fig(fig, F('06_jump.png'))

# fig7 四类间断点 四宫格   展示 9.0in
def fig_four():
    fig, axes = plt.subplots(2, 2, figsize=(9.0, 4.65))
    for ax in axes.ravel():
        ax.grid(True, color=k.M_RULE, lw=.7)
        for s in ('right', 'top'):
            ax.spines[s].set_visible(False)
        ax.tick_params(colors=k.M_SLATE)
    (A, B), (C, D) = axes
    x1 = np.linspace(-2, -.05, 100); x2 = np.linspace(.05, 2, 100)
    A.plot(x1, np.sin(x1) / x1, color=k.M_ACC2, lw=2.6)
    A.plot(x2, np.sin(x2) / x2, color=k.M_ACC2, lw=2.6)
    A.plot([0], [1], 'o', mfc='white', mec=k.M_RED, mew=2.2, ms=11)
    A.set_title('① 可去  $\\sin x/x$', fontsize=20, color=k.M_INK, pad=6)
    A.set_ylim(.3, 1.35); A.set_xticks([-2, 0, 2]); A.set_yticks([0.5, 1.0])

    B.plot(np.linspace(-2, 0, 80), np.zeros(80), color=k.M_ACC, lw=3)
    B.plot(np.linspace(0, 2, 80), np.ones(80), color=k.M_ACC, lw=3)
    B.plot([0], [0], 'o', color=k.M_ACC, ms=11)
    B.plot([0], [1], 'o', mfc='white', mec=k.M_ACC, mew=2.2, ms=11)
    B.set_title('② 跳跃  符号型', fontsize=20, color=k.M_INK, pad=6)
    B.set_ylim(-.7, 1.8); B.set_xticks([-2, 0, 2]); B.set_yticks([0, 1])

    xa = np.linspace(-2, -.06, 200); xb = np.linspace(.06, 2, 200)
    C.plot(xa, 1 / xa, color=k.M_RED, lw=2.6); C.plot(xb, 1 / xb, color=k.M_RED, lw=2.6)
    C.axvline(0, color=k.M_SLATE, ls='--', lw=1.2)
    C.set_title('③ 无穷  $1/x$', fontsize=20, color=k.M_INK, pad=6)
    C.set_ylim(-12, 12); C.set_xticks([-2, 0, 2]); C.set_yticks([-10, 0, 10])

    xo = np.linspace(-.5, .5, 6000); xo = xo[np.abs(xo) > 1e-4]
    D.plot(xo, np.sin(1 / xo), color=k.M_GRN, lw=.9)
    D.set_title('④ 振荡  $\\sin(1/x)$', fontsize=20, color=k.M_INK, pad=6)
    D.set_ylim(-1.5, 1.5); D.set_xticks([-0.5, 0, 0.5]); D.set_yticks([-1, 0, 1])
    return k.save_fig(fig, F('07_four_disc.png'))

# fig8 sin(1/x) 振荡放大   展示 11.4in
def fig_osc():
    fig, axes = plt.subplots(1, 2, figsize=(11.4, 3.5))
    for ax, r, t in zip(axes, [0.6, 0.06], ['整体 $[-0.6,\\,0.6]$', '放大 $[-0.06,\\,0.06]$']):
        x = np.linspace(-r, r, 20000); x = x[np.abs(x) > 1e-6]
        ax.plot(x, np.sin(1 / x), color=k.M_ACC2, lw=.8)
        ax.axhline(1, color=k.M_RULE, lw=1); ax.axhline(-1, color=k.M_RULE, lw=1)
        ax.set_title(t, fontsize=20, color=k.M_INK, pad=6)
        ax.set_ylim(-1.45, 1.45)
        ax.set_yticks([-1, 0, 1])
        ax.set_xticks([-r, 0, r])
        ax.grid(True, color=k.M_RULE, lw=.6)
        for s in ('right', 'top'):
            ax.spines[s].set_visible(False)
        ax.tick_params(colors=k.M_SLATE)
    return k.save_fig(fig, F('08_osc.png'))

# fig9 单位圆几何证明   展示 6.6in（长推导链移到 PPT bullets）
def fig_unitcircle():
    fig, ax = k.new_fig(6.6, 4.9)
    th = np.deg2rad(38)
    t = np.linspace(0, np.pi / 2, 200)
    ax.plot(np.cos(t), np.sin(t), color=k.M_SLATE, lw=1.6)
    ax.add_patch(Wedge((0, 0), 1, 0, np.rad2deg(th), color=k.M_ACC2, alpha=.18))
    B = (np.cos(th), np.sin(th)); D = (1, np.tan(th))
    ax.plot([0, 1.3], [0, 0], color=k.M_INK, lw=1.6)
    ax.plot([0, 0], [0, 1.15], color=k.M_INK, lw=1.6)
    ax.plot([0, B[0]], [0, B[1]], color=k.M_INK, lw=1.8)
    ax.plot([0, D[0]], [0, D[1]], color=k.M_INK, lw=1.8)
    ax.plot([1, 1], [0, D[1]], color=k.M_RED, lw=4)
    ax.plot([B[0], B[0]], [0, B[1]], color=k.M_GRN, lw=4)
    tt = np.linspace(0, th, 60)
    ax.plot(np.cos(tt), np.sin(tt), color=k.M_ACC, lw=5)
    ax.plot([B[0]], [B[1]], 'o', color=k.M_INK, ms=7)
    ax.text(1.02, -0.15, 'A', fontsize=20)
    ax.text(B[0] - 0.11, B[1] + 0.04, 'B', fontsize=20)
    ax.text(B[0] - 0.03, -0.15, 'C', fontsize=20, ha='right')
    ax.text(1.04, D[1] + 0.02, 'D', fontsize=20)
    ax.text(1.10, D[1] / 2 - 0.06, '$\\tan x$', color=k.M_RED, fontsize=20)
    ax.text(B[0] - 0.05, 0.20, '$\\sin x$', color=k.M_GRN, fontsize=20, ha='right')
    ax.text(1.12, 0.02, '弧 $x$', color=k.M_ACC, fontsize=20)
    ax.set_aspect('equal'); ax.axis('off')
    ax.set_xlim(-0.36, 1.80); ax.set_ylim(-0.30, 1.30)
    return k.save_fig(fig, F('09_unit_circle.png'))

# fig10 夹逼定理三曲线   展示 7.4in
def fig_squeeze():
    fig, ax = k.new_fig(7.4, 4.4)
    x = np.linspace(-1.5, 1.5, 800); x = x[np.abs(x) > 1e-4]
    ax.plot(x, np.cos(x), color=k.M_GRN, lw=2.6, label='$\\cos x$')
    ax.plot(x, np.ones_like(x), color=k.M_RED, lw=2.6, label='$1$')
    ax.plot(x, np.sin(x) / x, color=k.M_ACC, lw=3.4, label='$\\sin x/x$')
    ax.fill_between(x, np.cos(x), 1, color=k.M_ACC2, alpha=.15)
    ax.plot([0], [1], 'o', mfc='white', mec=k.M_ACC, mew=2.6, ms=12, zorder=6)
    ax.annotate('挤到 $1$', xy=(0.05, 1.005), xytext=(0.40, 1.17),
                fontsize=20, color=k.M_INK,
                arrowprops=dict(arrowstyle='->', color=k.M_INK, lw=1.8))
    k.style_axes(ax, '', '', origin=False)
    ax.set_xlim(-1.6, 1.6); ax.set_ylim(0.45, 1.30)
    ax.set_xticks([-1, 0, 1]); ax.set_yticks([0.6, 0.8, 1.0])
    ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.06), ncol=3, frameon=False)
    return k.save_fig(fig, F('10_squeeze.png'))

# fig11 (sinx)/x 全景 + 数值表   展示 9.0in
def fig_sinx():
    fig, ax = k.new_fig(9.0, 4.6)
    x = np.linspace(-12, 12, 3000); x = x[np.abs(x) > 1e-3]
    ax.plot(x, np.sin(x) / x, color=k.M_ACC2, lw=2.4)
    ax.plot([0], [1], 'o', mfc='white', mec=k.M_RED, mew=2.6, ms=12, zorder=5)
    ax.axhline(1, color=k.M_RED, ls='--', lw=1.4)
    ax.annotate('$\\to 1$', xy=(0.3, 0.99), xytext=(3.2, 1.32),
                fontsize=20, color=k.M_RED,
                arrowprops=dict(arrowstyle='->', color=k.M_RED, lw=1.8))
    head = ['$x$', '1', '0.1', '0.01']
    body = ['$\\sin x/x$', '0.8415', '0.9983', '0.99998']
    tbl = ax.table(cellText=[body], colLabels=head, loc='center',
                   cellLoc='center', bbox=[0.06, 0.02, 0.88, 0.26])
    tbl.auto_set_font_size(False); tbl.set_fontsize(20)
    for c in tbl.get_celld().values():
        c.set_edgecolor(k.M_RULE); c.set_facecolor('white')
    k.style_axes(ax, '', '', origin=False)
    ax.set_xlim(-12.5, 12.5); ax.set_ylim(-1.35, 1.60)
    ax.set_xticks([-10, 0, 10]); ax.set_yticks([0, 1])
    return k.save_fig(fig, F('11_sinx_over_x.png'))

# fig12 (1+1/n)^n -> e   展示 7.3in
def fig_e():
    fig, ax = k.new_fig(7.3, 4.4)
    n = np.arange(1, 60)
    y = (1 + 1 / n) ** n
    z = (1 + 1 / n) ** (n + 1)
    ax.plot(n, y, 'o-', color=k.M_ACC, lw=2.0, ms=5, label='$a_n$')
    ax.plot(n, z, 's--', color=k.M_ACC2, lw=1.6, ms=4, label='$b_n$')
    ax.fill_between(n, y, z, color=k.M_GRN, alpha=.12)
    ax.axhline(np.e, color=k.M_RED, ls='--', lw=2)
    ax.text(22, 2.85, '$e\\approx 2.718$', color=k.M_RED, fontsize=20)
    k.style_axes(ax, 'n', '', origin=False)
    ax.set_xlim(0, 62); ax.set_ylim(1.8, 4.4)
    ax.set_xticks([0, 20, 40, 60]); ax.set_yticks([2, 3, 4])
    ax.legend(loc='upper right', ncol=2)
    return k.save_fig(fig, F('12_e_limit.png'))

# fig13 介值定理：必过横线   展示 5.4in（第25页放大到 7.2in）
def fig_ivt():
    fig, ax = k.new_fig(5.4, 3.6)
    x = np.linspace(0, 4, 400)
    f = lambda t: 0.45 * t ** 3 - 2.6 * t ** 2 + 3.2 * t + 0.6
    ax.plot(x, f(x), color=k.M_ACC2, lw=3.2)
    a, b = 0.25, 3.9
    ax.plot([a], [f(a)], 'o', color=k.M_GRN, ms=12)
    ax.plot([b], [f(b)], 'o', color=k.M_RED, ms=12)
    ax.text(0.45, 2.35, '$f(a)$', color=k.M_GRN, fontsize=20)
    ax.text(4.05, -0.55, '$f(b)$', color=k.M_RED, fontsize=20, ha='center')
    C = 1.0
    ax.axhline(C, color=k.M_ACC, ls='--', lw=2.2)
    ax.text(4.35, C + 0.18, '$C$', color=k.M_ACC, fontsize=20)
    xs = np.linspace(0, 4, 8000); vals = f(xs) - C
    roots = xs[:-1][np.sign(vals[:-1]) != np.sign(vals[1:])]
    for r in roots:
        ax.plot([r], [C], '*', color=k.M_RED, ms=20, zorder=6)
    k.style_axes(ax, '', '', origin=False)
    ax.set_xlim(0, 4.7); ax.set_ylim(-2.2, 3.3)
    ax.set_xticks([0, 2, 4]); ax.set_yticks([-2, 0, 2])
    return k.save_fig(fig, F('13_ivt.png'))

# fig14 对跖点同温   展示 11.2in
def fig_antipode():
    fig, axes = plt.subplots(1, 2, figsize=(11.2, 3.7))
    ax = axes[0]
    t = np.linspace(0, 2 * np.pi, 400)
    ax.plot(np.cos(t), np.sin(t), color=k.M_ACC2, lw=2.8)
    th = np.deg2rad(35)
    P = (np.cos(th), np.sin(th)); Q = (-P[0], -P[1])
    ax.plot([Q[0], P[0]], [Q[1], P[1]], '--', color=k.M_SLATE, lw=1.6)
    ax.plot(*P, 'o', color=k.M_RED, ms=13); ax.plot(*Q, 'o', color=k.M_GRN, ms=13)
    ax.text(P[0] + 0.12, P[1], '$P(\\theta)$', color=k.M_RED, fontsize=20, va='center')
    ax.text(Q[0] - 0.12, Q[1], '$P(\\theta+\\pi)$', color=k.M_GRN, fontsize=20,
            va='center', ha='right')
    ax.set_aspect('equal'); ax.axis('off')
    ax.set_xlim(-2.75, 2.35); ax.set_ylim(-1.3, 1.3)
    ax.set_title('赤道对跖点', fontsize=20, color=k.M_INK, pad=4)

    ax = axes[1]
    x = np.linspace(0, np.pi, 400)
    g = lambda t: (6 * np.sin(t + .7) + 2.5 * np.sin(3 * t)
                   - 6 * np.sin(t + .7 + np.pi) - 2.5 * np.sin(3 * (t + np.pi)))
    ax.plot(x, g(x), color=k.M_ACC, lw=3)
    ax.axhline(0, color=k.M_INK, lw=1.4)
    xs = np.linspace(0, np.pi, 4000); v = g(xs)
    for r in xs[:-1][np.sign(v[:-1]) != np.sign(v[1:])]:
        ax.plot([r], [0], '*', color=k.M_RED, ms=22, zorder=5)
    ax.grid(True, color=k.M_RULE, lw=.6)
    for s in ('right', 'top'):
        ax.spines[s].set_visible(False)
    ax.tick_params(colors=k.M_SLATE)
    ax.set_xlim(0, np.pi); ax.set_ylim(-20, 22)
    ax.set_xticks([0, np.pi / 2, np.pi]); ax.set_xticklabels(['$0$', '$\\pi/2$', '$\\pi$'])
    ax.set_yticks([-10, 0, 10])
    ax.set_title('$g(\\theta)=T(\\theta)-T(\\theta+\\pi)$', fontsize=20, color=k.M_INK, pad=4)
    return k.save_fig(fig, F('14_antipode.png'))

# fig15 知识网络   展示 5.4in（第29页放大到 7.0in）
def fig_map():
    fig, ax = k.new_fig(5.4, 4.2)
    ax.axis('off')
    nodes = [
        ('直观逼近', 2.3, 9.2, k.M_SLATE),
        ('$\\varepsilon$-$\\delta$ 定义', 7.7, 9.2, k.M_ACC),
        ('左右极限', 2.3, 6.5, k.M_ACC2),
        ('连续三条件', 7.7, 6.5, k.M_GRN),
        ('间断四类', 2.3, 3.8, k.M_RED),
        ('夹逼定理', 7.7, 3.8, k.M_ACC),
        ('两个重要极限', 2.5, 1.1, k.M_ACC2),
        ('介值定理', 7.9, 1.1, k.M_ACC),
    ]
    for t, x, y, c in nodes:
        ax.text(x, y, t, fontsize=20, color='white', ha='center', va='center',
                bbox=dict(boxstyle='round,pad=0.42', fc=c, ec='none'))
    arrows = [((4.2, 9.2), (5.7, 9.2)),
              ((7.7, 8.6), (7.7, 7.2)),
              ((6.0, 6.5), (3.8, 6.5)),
              ((2.3, 5.9), (2.3, 4.5)),
              ((7.7, 5.9), (7.7, 4.5)),
              ((7.0, 3.2), (4.6, 1.8))]
    for p, q in arrows:
        ax.annotate('', xy=q, xytext=p,
                    arrowprops=dict(arrowstyle='->', color=k.M_INK, lw=2))
    ax.annotate('', xy=(9.3, 1.5), xytext=(9.3, 6.1),
                arrowprops=dict(arrowstyle='->', color=k.M_INK, lw=2,
                                connectionstyle='arc3,rad=-0.35'))
    ax.set_xlim(-0.8, 11.6); ax.set_ylim(-0.1, 10.5)
    return k.save_fig(fig, F('15_map.png'))

# fig16 挑战—应战阶梯   展示 7.2in
def fig_ladder():
    fig, ax = k.new_fig(7.2, 2.6)
    eps = np.array([1e-1, 1e-2, 1e-3, 1e-4, 1e-5, 1e-6])
    dlt = eps / 2
    idx = np.arange(len(eps))
    ax.plot(idx, eps, 'o-', color=k.M_ACC, lw=2.6, ms=10, label='$\\varepsilon$')
    ax.plot(idx, dlt, 's-', color=k.M_ACC2, lw=2.6, ms=9, label='$\\delta=\\varepsilon/2$')
    for i in idx:
        ax.annotate('', xy=(i, dlt[i]), xytext=(i, eps[i]),
                    arrowprops=dict(arrowstyle='->', color=k.M_GRN, lw=1.6))
    ax.set_yscale('log')
    ax.set_xticks(idx); ax.set_xticklabels([str(i + 1) for i in idx])
    ax.set_xlabel('轮次', fontsize=20, color=k.M_INK)
    ax.set_yticks([1e-6, 1e-3, 1e0])
    ax.grid(True, color=k.M_RULE, lw=.7)
    for s in ('right', 'top'):
        ax.spines[s].set_visible(False)
    ax.tick_params(colors=k.M_SLATE)
    ax.set_xlim(-0.4, 7.4); ax.set_ylim(2e-7, 4e1)
    ax.legend(loc='upper right', ncol=2, framealpha=.9)
    return k.save_fig(fig, F('16_ladder.png'))

# fig17 板书分区示意   展示 12.0in
def fig_board():
    fig, ax = k.new_fig(12.0, 2.2)
    ax.axis('off')
    ax.add_patch(Rectangle((0, 0), 12, 2.2, fc='#22303F', ec=k.M_INK, lw=2))
    zones = [(0.15, '主板·左', k.M_ACC2, '概念主线\n芝诺→ε-δ→连续'),
             (3.10, '主板·中', k.M_ACC, '推演区\n单位圆夹逼链'),
             (6.05, '主板·右', k.M_GRN, '学生板演\n例2 · 例3 · 例4'),
             (9.00, '副板·角', k.M_RED, '易错墙\n✗代入 ✗弧度\n✗量词序')]
    for x, t, c, body in zones:
        ax.add_patch(Rectangle((x, 0.12), 2.85, 1.96, fc='none', ec=c, lw=2, ls='--'))
        ax.text(x + 1.425, 1.80, t, color=c, fontsize=20, ha='center', va='center', weight='bold')
        ax.text(x + 1.425, 0.82, body, color='#E8E4DA', fontsize=20, ha='center',
                va='center', linespacing=1.5)
    ax.set_xlim(-0.1, 12.1); ax.set_ylim(-0.1, 2.3)
    return k.save_fig(fig, F('17_board.png'))

FIGS = [fig_zeno(), fig_intuitive(), fig_indep(), fig_epsdelta(), fig_eps_shrink(),
        fig_jump(), fig_four(), fig_osc(), fig_unitcircle(), fig_squeeze(),
        fig_sinx(), fig_e(), fig_ivt(), fig_antipode(), fig_map(),
        fig_ladder(), fig_board()]
print('figures:', len(FIGS))

# ============================================================ PPT
prs = k.new_deck()
TEX = lambda n: F('tex_%s.png' % n)

# 20pt 字号下限下，默认行距会把长列表顶出页面：统一收紧行距（字号不动）
_bullets_raw = k.bullets
def _bullets(slide, items, ls=1.15, sa=4, **kw):
    tb = _bullets_raw(slide, items, **kw)
    for p in tb.text_frame.paragraphs:
        p.line_spacing = ls
        p.space_after = k.Pt(sa)
    return tb
k.bullets = _bullets

# 1 封面
k.title_slide(prs, '第10讲　函数的极限与连续性',
              '从“无限逼近”的直观，到 ε-δ 的严格——数学史上最重要的一次严格化',
              '数学名师课件包', '60 分钟 · 高中衔接/大学微积分先修')

# 2 学习目标
s = k.content_slide(prs, '学习目标 · 本节地图', 'Goal')
k.bullets(s, [
    '理解极限的直观含义：“无限逼近”，并明确 lim 与 f(a) 无关',
    ('极限只看逼近过程，不看该点取值——可去间断点是最好的反例', 1),
    '掌握 ε-δ 语言：把“无限接近”翻译成“要多近有多近”的博弈',
    ('对手任给 ε，我总能找到 δ 应战（本节核心难点）', 1),
    '掌握连续的三条件定义与间断点的四类分型',
    '会用夹逼定理证明 lim(sin x)/x = 1；理解 (1+1/n)ⁿ → e',
    '会用介值定理证明方程有根，感受“对跖点同温”的震撼推论',
], y=1.55, w=6.6, size=17)
k.picture(s, FIGS[14], x=7.5, y=1.7, w=5.4)
k.callout(s, '一句话主线：把“越来越接近”这句模糊的话，\n变成一条谁都无法抵赖的不等式。', x=0.85, y=6.15, w=6.4, h=0.95, kind='key')

# 3 第1幕
k.section_slide(prs, '第 1 幕 · 情境', '芝诺的挑衅：无限步，能走完吗？', '8 min')

# 4 芝诺
s = k.content_slide(prs, '阿基里斯追不上乌龟？', '导入')
k.bullets(s, [
    '乌龟先跑 100 米，阿基里斯速度是它的 10 倍',
    ('追到 100 米时，龟又前进 10 米；追到 10 米，龟又走 1 米……', 1),
    ('要追上，似乎要完成“无穷多”个步骤 → 永远追不上？', 1),
    '破绽：无穷多段，路程之和却是有限的',
], y=1.6, w=5.4, size=17)
k.picture(s, FIGS[0], x=6.5, y=1.45, w=6.3)
k.callout(s, '100+10+1+… = 100/(1-0.1) = 111.1 米。\n“无限个过程”可以有一个确定的归宿——这个归宿就叫极限。',
          x=0.85, y=5.0, w=5.4, h=1.6, kind='note')

# 5 直观极限
s = k.content_slide(prs, '直观的极限：无限逼近', '概念')
k.formula(s, '$\\lim_{x\\to a}f(x)=L$　：　$x$ 无限接近 $a\;(x\\neq a)$ 时，$f(x)$ 无限接近 $L$',
          x=0.85, y=1.45, w=11.6, size=0.72, out=TEX('lim_def'))
k.picture(s, FIGS[1], x=0.75, y=2.55, w=7.0)
k.bullets(s, [
    '注意 x ≠ a：极限研究的是“去心邻域”上的行为',
    '左右两侧都要逼近同一个 L',
    ('左极限 = 右极限 = L ⟺ 极限存在且为 L', 1),
    '“无限接近”还很模糊——模糊就无法证明',
], x=8.1, y=2.7, w=4.5, size=16)
k.callout(s, '一切严格化的起点：\n把“接近”量化。', x=8.1, y=5.85, w=4.5, h=1.0, kind='warn')

# 6 lim 与 f(a) 无关
s = k.content_slide(prs, '关键警示：lim 与 f(a) 毫无关系', '易错')
k.full_picture(s, FIGS[2], y=1.5, w=11.4)
k.callout(s, '求极限时，代入 x=a 只是“运气好”（连续时才成立）；一般情况必须考察逼近过程。',
          x=1.2, y=6.25, w=11.0, h=0.85, kind='warn')

# 7 第2幕
k.section_slide(prs, '第 2 幕 · 严格化', 'ε-δ：一场“挑战—应战”的博弈', '14 min')

# 8 博弈叙事
s = k.content_slide(prs, '柯西与魏尔斯特拉斯：把“接近”写成不等式', '思想')
k.bullets(s, [
    '【对手】“你说 f(x) 接近 L？我不信。我要求误差小于 ε = 0.1。”',
    '【我】“没问题，只要 x 落在 a 的 δ = 0.05 邻域内（且 x≠a），就能做到。”',
    '【对手】“那我加码，ε = 0.001！”',
    '【我】“我也加码，δ = 0.0005。”',
    '【对手】“ε = 10⁻¹⁰⁰⁰⁰！”',
    '【我】“……我依然找得到 δ。”',
    '对手无论怎么挑战，我总能应战 ⟹ 极限确实是 L',
], y=1.5, w=7.4, size=15)
k.callout(s, '逻辑顺序不可颠倒：\n∀ε 在前，∃δ 在后。\nδ 是被 ε 逼出来的，\n依赖于 ε（写作 δ(ε)）。',
          x=8.6, y=1.6, w=4.0, h=2.0, kind='key')
k.callout(s, '“无限接近”\n→“要多近，有多近”\n这就是 19 世纪的\n严格化革命。',
          x=8.6, y=3.8, w=4.0, h=1.7, kind='note')
k.picture(s, FIGS[15], x=0.7, y=4.75, w=7.2)

# 9 ε-δ 定义
s = k.content_slide(prs, 'ε-δ 定义（本讲核心）', '定义')
k.formula(s, '$\\lim_{x\\to a}f(x)=L\;\\Longleftrightarrow\;\\forall\\varepsilon>0,\\ \\exists\\delta>0$，使得 $0<|x-a|<\\delta\;\\Rightarrow\;|f(x)-L|<\\varepsilon$',
          x=0.6, y=1.5, w=12.1, size=0.63, out=TEX('epsdelta'))
k.picture(s, FIGS[3], x=0.75, y=2.7, w=6.4)
k.bullets(s, [
    '|f(x)−L| < ε：横带（对手划定的容错区）',
    '0 < |x−a| < δ：竖带（我划定的安全区，挖去 a）',
    '“竖带 ⊂ 横带的原像”即应战成功',
    ('图像在竖带内的部分，必须完全落进横带', 1),
], x=8.3, y=2.9, w=4.4, size=15)
k.callout(s, '0<|x−a| 中的 0，正是\n“与 f(a) 无关”的数学化身。',
          x=8.3, y=5.6, w=4.4, h=1.15, kind='warn')

# 10 ε 变小 δ 变小
s = k.content_slide(prs, 'ε 收紧，δ 随之收紧', '动态')
k.full_picture(s, FIGS[4], y=1.55, w=12.2)
k.callout(s, '看图说话：横带越窄，竖带就必须越窄；但只要每一次都找得到，极限就成立。',
          x=1.4, y=5.9, w=10.6, h=0.85, kind='key')

# 11 例1
s = k.content_slide(prs, '例1　用定义证明 lim(2x+1)=5 (x→2)', '例题')
k.formula(s, '分析：$|f(x)-L|=|2x+1-5|=2|x-2|<\\varepsilon\;\\Longleftrightarrow\;|x-2|<\\frac{\\varepsilon}{2}$',
          x=0.85, y=1.5, w=11.0, size=0.66, out=TEX('ex1a'))
k.bullets(s, [
    '【板演】证明：∀ε>0，取 δ = ε/2 > 0。',
    '当 0 < |x−2| < δ 时，',
    ('|f(x)−5| = 2|x−2| < 2δ = ε。', 1),
    '由定义知 lim(x→2)(2x+1) = 5。∎',
    '方法论：从结论 |f(x)−L|<ε 反解出 |x−a| 的范围，那就是 δ。',
], y=2.6, w=6.6, size=17)
k.callout(s, '“执果索因”：先把目标不等式化成 |x−a| 的形式，\nδ 就自动浮现。这叫“分析法找 δ，综合法写证明”。',
          x=7.6, y=2.7, w=5.0, h=1.8, kind='key')
k.callout(s, '变式（课堂口答）：证明 lim(x→3)(4x−1)=11，\n应取 δ = ______。（答：ε/4）',
          x=7.6, y=4.9, w=5.0, h=1.5, kind='note')

# 12 左右极限 / 跳跃
s = k.content_slide(prs, '左右极限：极限存在的“双向验票”', '概念')
k.picture(s, FIGS[5], x=6.4, y=1.5, w=6.4)
k.formula(s, r'$\lim_{x\to a}f(x)=L \Longleftrightarrow \lim_{x\to a^-}f(x)=\lim_{x\to a^+}f(x)=L$',
          x=0.7, y=1.5, w=5.6, size=0.35, out=TEX('lr'))
k.bullets(s, [
    '分段函数在分界点必须分左右讨论',
    '左右极限都存在但不相等 ⟹ 极限不存在（跳跃）',
    '典型：sgn(x)、取整函数 [x]、分段定价/分段计费模型',
], x=0.75, y=2.5, w=5.5, size=16)
k.callout(s, '“一票否决”：只要两侧走向不同的目的地，\n极限就不存在——哪怕两侧各自都很老实。',
          x=0.75, y=5.3, w=5.5, h=1.4, kind='warn')

# 13 第3幕
k.section_slide(prs, '第 3 幕 · 连续性', '连续 = 极限值恰好等于函数值', '12 min')

# 14 连续三条件
s = k.content_slide(prs, '连续的定义：三个条件缺一不可', '定义')
k.formula(s, '$f$ 在 $x=a$ 处连续 $\;\\Longleftrightarrow\;\\lim_{x\\to a}f(x)=f(a)$',
          x=0.85, y=1.45, w=10.0, size=0.68, out=TEX('cont'))
k.bullets(s, [
    '① f(a) 有定义（该点在定义域内）',
    '② lim(x→a) f(x) 存在（左右极限相等）',
    '③ 二者相等：极限值 = 函数值',
    '缺 ① → 可去/无穷间断；缺 ② → 跳跃/振荡；缺 ③ → 可去间断',
    '连续的直观：图像“一笔画成”，笔尖不离纸面',
], y=2.6, w=6.6, size=17)
k.picture(s, FIGS[1], x=7.5, y=2.5, w=5.2)
k.callout(s, '连续函数的最大红利：求极限可以“直接代入”。\n初等函数在其定义域上处处连续。',
          x=0.85, y=6.05, w=6.4, h=1.05, kind='key')

# 15 间断点四宫格
s = k.content_slide(prs, '间断点的四大类型', '分类')
k.full_picture(s, FIGS[6], y=1.25, w=9.0)
k.callout(s, '判型口诀：先看左右极限存不存在——都存在则第一类（相等=可去，不等=跳跃）；否则第二类（无穷 / 振荡）。',
          x=0.9, y=6.72, w=11.6, h=0.62, kind='key')

# 16 振荡
s = k.content_slide(prs, '振荡间断：sin(1/x) 的深渊', '反例')
k.full_picture(s, FIGS[7], y=1.4, w=11.4)
k.callout(s, '用 ε-δ 语言反驳：取 ε = 1/2，无论 δ 多小，竖带内总有点让 f = 1 与 f = −1 同时出现，\n不可能都落进宽度为 1 的横带 ⟹ 应战失败 ⟹ 极限不存在。',
          x=1.0, y=6.05, w=11.3, h=1.1, kind='warn')

# 17 例2
s = k.content_slide(prs, '例2　判断连续性并求间断点类型', '例题')
k.formula(s, r'$f(x)=\dfrac{x^2-3x+2}{x^2-1}$，求其间断点并分类',
          x=0.85, y=1.45, w=9.5, size=0.6, out=TEX('ex2'))
k.bullets(s, [
    '【板演】定义域：x ≠ ±1，故 x = 1 与 x = −1 是间断点。',
    'x = 1：f(x) = (x−1)(x−2)/[(x−1)(x+1)] = (x−2)/(x+1)（x≠1）',
    ('lim(x→1) f(x) = (1−2)/(1+1) = −1/2，极限存在 ⟹ 可去间断点', 1),
    ('补充定义 f(1) = −1/2 即可“修复”为连续', 1),
    'x = −1：分子 → (−1)²−3(−1)+2 = 6 ≠ 0，分母 → 0',
    ('lim(x→−1) f(x) = ∞ ⟹ 无穷间断点（第二类），无法修复', 1),
], y=2.5, w=7.6, size=16)
k.callout(s, '通法：因式分解 →\n约掉公因式看极限。\n\n约得掉 = 可去；\n约不掉且分母→0 = 无穷。',
          x=8.6, y=2.6, w=4.0, h=2.6, kind='key')
k.callout(s, '课堂变式1：求 f(x)=|x|/x\n在 x=0 的间断类型。\n（答：跳跃间断，左−1 右1）',
          x=8.6, y=5.5, w=4.0, h=1.5, kind='note')

# 18 第4幕
k.section_slide(prs, '第 4 幕 · 两个重要极限', '夹逼定理与自然常数 e', '14 min')

# 19 夹逼定理
s = k.content_slide(prs, '夹逼（三明治）定理', '工具')
k.formula(s, '$g(x)\\leq f(x)\\leq h(x)$ 且 $\\lim_{x\\to a}g(x)=\\lim_{x\\to a}h(x)=L\;\\Rightarrow\;\\lim_{x\\to a}f(x)=L$',
          x=0.6, y=1.45, w=12.1, size=0.6, out=TEX('squeeze'))
k.picture(s, FIGS[9], x=0.7, y=2.5, w=7.4)
k.bullets(s, [
    '上下两只手同时收拢，中间的人无处可逃',
    '关键：找到“好算”的上下界',
    '常与三角不等式、放缩法配合',
    'ε-δ 证明思路：g、h 各自的 δ 取较小者',
], x=8.4, y=2.8, w=4.3, size=16)
k.callout(s, '这是证明第一个重要极限的钥匙。',
          x=8.4, y=5.95, w=4.3, h=0.9, kind='key')

# 20 单位圆几何证明
s = k.content_slide(prs, '第一个重要极限：单位圆的几何证明', '推演')
k.picture(s, FIGS[8], x=0.55, y=1.45, w=6.6)
k.bullets(s, [
    '① 面积夹逼：½sin x < ½x < ½tan x (0<x<π/2)',
    '② 同除 ½sin x（正）：1 < x/sin x < 1/cos x',
    '③ 取倒数：cos x < (sin x)/x < 1',
    '④ x→0 时 cos x → 1，上界恒为 1',
    '⑤ 夹逼 ⟹ lim(x→0) (sin x)/x = 1',
    '⑥ 偶函数性质 ⟹ x<0 时同样成立',
], x=7.3, y=1.72, w=5.5, size=15)
k.callout(s, '前提：x 必须用弧度制！\n弧长 = 角度，几何才成立。',
          x=7.3, y=5.7, w=5.5, h=1.1, kind='warn')

# 21 sinx/x 全景
s = k.content_slide(prs, '第一个重要极限：图像与数值验证', '验证')
k.full_picture(s, FIGS[10], y=1.4, w=9.0)
k.formula(s, '$\\lim_{x\\to 0}\\frac{\\sin x}{x}=1$　　推广：$\\lim_{u\\to 0}\\frac{\\sin u}{u}=1$',
          x=3.2, y=6.42, w=7.0, size=0.52, out=TEX('sinx1'))

# 22 例3
s = k.content_slide(prs, '例3　0/0 型极限（板演）', '例题')
k.formula(s, r'$(1)\ \lim_{x\to0}\frac{\sin 3x}{5x}\qquad (2)\ \lim_{x\to0}\frac{1-\cos x}{x^2}\qquad(3)\ \lim_{x\to1}\frac{\sqrt{x+3}-2}{x-1}$',
          x=0.6, y=1.45, w=12.1, size=0.58, out=TEX('ex3'))
k.bullets(s, [
    '(1) 凑形：= lim (3/5)·(sin3x)/(3x) = 3/5 × 1 = 3/5',
    '(2) 半角：1−cos x = 2sin²(x/2)',
    ('= lim 2sin²(x/2)/x² = lim (1/2)·[sin(x/2)/(x/2)]² = 1/2', 1),
    '(3) 有理化：分子分母同乘 √(x+3)+2',
    ('= lim (x−1)/[(x−1)(√(x+3)+2)] = 1/(√4+2) = 1/4', 1),
], y=2.65, w=7.4, size=16)
k.callout(s, '0/0 型三板斧：\n① 因式分解约零因子\n② 根式有理化\n③ 凑重要极限的标准形',
          x=8.3, y=2.75, w=4.3, h=2.3, kind='key')
k.callout(s, '课堂变式2：\nlim(x→0) tan x / x = ?\n（答：1，拆成 sinx/x · 1/cosx）',
          x=8.3, y=5.4, w=4.3, h=1.5, kind='note')

# 23 第二个重要极限
s = k.content_slide(prs, '第二个重要极限：e 的诞生', '概念')
k.picture(s, FIGS[11], x=0.6, y=1.5, w=7.3)
k.formula(s, r'$\lim_{n\to\infty}\!\left(1+\frac{1}{n}\right)^{\!n}=e$',
          x=8.1, y=1.6, w=4.6, size=0.42, out=TEX('e'))
k.bullets(s, [
    '“1 的 ∞ 次方”是未定式：底趋于 1，指数趋于 ∞',
    '单调有界准则：递增且有上界 ⟹ 极限存在',
    '这个极限值定义为 e ≈ 2.71828…',
    '金融解读：年利率 100%、复利次数 → ∞ 的本息比',
], x=8.1, y=2.9, w=4.6, size=15)
k.callout(s, '推广式：lim(x→0)(1+x)^(1/x) = e',
          x=8.1, y=6.1, w=4.6, h=0.85, kind='key')

# 24 第5幕
k.section_slide(prs, '第 5 幕 · 整体性质', '介值定理：连续者不可瞬移', '10 min')

# 25 介值定理
s = k.content_slide(prs, '介值定理与最值定理', '定理')
k.picture(s, FIGS[12], x=0.6, y=1.5, w=7.2)
k.bullets(s, [
    '【介值定理】f 在 [a,b] 连续，C 介于 f(a)、f(b) 之间',
    ('⟹ ∃ξ∈(a,b)，使 f(ξ) = C', 1),
    '【零点定理】若 f(a)·f(b) < 0 ⟹ ∃ξ，f(ξ) = 0',
    '【最值定理】f 在闭区间 [a,b] 连续 ⟹ 必取到最大值与最小值',
    '三个前提缺一不可：闭区间 + 连续',
], x=7.9, y=1.75, w=4.9, size=15)
k.callout(s, '反例警醒：f(x)=1/x 在开区间 (0,1] 连续但无最大值；\n分段跳跃函数可以“跨过”C 而不取到它。',
          x=7.9, y=5.5, w=4.9, h=1.5, kind='warn')

# 26 高光：对跖点
s = k.content_slide(prs, '高光时刻：地球上必有一对对跖点同温', 'Wow')
k.full_picture(s, FIGS[13], y=1.35, w=11.2)
k.callout(s, '同理可证“椅子四脚着地”：地面连续起伏时，旋转椅子必存在一个角度使四脚同时落地。\n——一个纯粹的存在性定理，竟能断言现实世界的事实。',
          x=1.0, y=6.05, w=11.3, h=1.1, kind='key')

# 27 例4
s = k.content_slide(prs, '例4　用零点定理证明方程有根', '例题')
k.formula(s, '证明：方程 $x^3-4x^2+1=0$ 在 $(0,1)$ 内至少有一个实根',
          x=0.7, y=1.45, w=11.0, size=0.6, out=TEX('ex4'))
k.bullets(s, [
    '【板演】令 f(x) = x³ − 4x² + 1。',
    'f 是多项式函数 ⟹ 在 [0,1] 上连续。（前提①）',
    'f(0) = 1 > 0；f(1) = 1 − 4 + 1 = −2 < 0。（前提②异号）',
    '由零点定理：∃ξ ∈ (0,1)，使 f(ξ) = 0，',
    '即方程在 (0,1) 内至少有一个实根。∎',
    '追问：还能证明它在 (3,4) 内也有根吗？（f(3)=−8, f(4)=1 ⟹ 能）',
], y=2.55, w=7.6, size=16)
k.callout(s, '证根三步法：\n① 构造函数 f\n② 验证闭区间上连续\n③ 找两端异号的点\n\n三步齐了，结论必然。',
          x=8.6, y=2.6, w=4.0, h=2.9, kind='key')

# 28 变式演练
s = k.content_slide(prs, '变式演练（学生板演）', '变式')
k.bullets(s, [
    '变式1（连续性）：设 f(x) = { (sin 2x)/x , x<0 ; a , x=0 ; x+1 , x>0 }',
    ('求 a 使 f 在 x=0 连续。（提示：左极限 = 2，右极限 = 1 ⟹ 左右不等，', 1),
    ('无论 a 取何值都不连续，x=0 恒为跳跃间断点。这是个“陷阱题”！）', 1),
    '变式2（介值定理）：证明方程 x = cos x 在 (0, π/2) 内有唯一实根。',
    ('提示：令 g(x) = x − cos x；g(0) = −1 < 0，g(π/2) = π/2 > 0；', 1),
    ('又 g 单调递增（后续可用导数说明）⟹ 根唯一。', 1),
], y=1.65, w=6.9, size=16)
k.picture(s, FIGS[12], x=7.9, y=1.9, w=5.4)
k.callout(s, '变式1 的价值：提醒学生“连续”不是靠调一个点的值就能补出来的——\n必须先有极限，才谈得上让 f(a) 去等于它。',
          x=0.85, y=5.85, w=7.5, h=1.25, kind='warn')

# 29 小结
s = k.content_slide(prs, '课堂小结：从直观到严格，从局部到整体', '小结')
k.picture(s, FIGS[14], x=0.5, y=1.5, w=7.0)
k.bullets(s, [
    '极限：看逼近过程，与 f(a) 无关',
    'ε-δ：∀ε ∃δ，挑战—应战的博弈',
    '连续：三条件（有定义·有极限·相等）',
    '间断：可去/跳跃（一类）· 无穷/振荡（二类）',
    '两个重要极限：sinx/x → 1；(1+1/n)ⁿ → e',
    '介值/最值：闭区间 + 连续 = 强大存在性',
], x=7.8, y=1.9, w=5.0, size=15)
k.callout(s, '一句话：极限是微积分的地基，\n连续是这块地基上的第一块砖。',
          x=7.8, y=5.6, w=5.0, h=1.1, kind='key')

# 30 分层作业
s = k.content_slide(prs, '分层作业', '作业')
k.bullets(s, [
    'A 层 · 基础巩固（全体必做，约 20 min）',
    ('① 求 lim(x→0) sin5x/(2x)、lim(x→0)(1−cos2x)/x²', 1),
    ('② 求 f(x)=(x²−4)/(x−2) 的间断点并分类', 1),
    ('③ 用定义证明 lim(x→1)(3x−2)=1（写出 δ 的取法）', 1),
    'B 层 · 能力提升（选做，约 20 min）',
    ('④ 讨论 f(x)= x·sin(1/x) (x≠0), 0 (x=0) 在 x=0 的连续性（提示：夹逼）', 1),
    ('⑤ 证明 x·2ˣ = 1 在 (0,1) 内有实根', 1),
    ('⑥ 求 a、b 使 f 在 R 上连续（分段函数拼接题）', 1),
    'C 层 · 思维挑战（学有余力）',
    ('⑦ 用 ε-δ 严格证明 lim(x→2) x² = 4（提示：限定 |x−2|<1 先控制 |x+2|）', 1),
    ('⑧ 查阅“椅子四脚着地定理”，写 300 字说明它如何依赖介值定理', 1),
], y=1.6, w=11.8, size=15)

# 31 板书提纲
s = k.content_slide(prs, '板书设计提纲', '板书')
k.bullets(s, [
    '【主板 · 左】概念主线',
    ('1. 直观极限（芝诺图）→ 2. lim 与 f(a) 无关 → 3. ε-δ 定义（双带图）', 1),
    ('4. 连续三条件 → 5. 间断点四分类表', 1),
    '【主板 · 中】推演区',
    ('单位圆面积夹逼链：½sinx < ½x < ½tanx ⟹ cosx < sinx/x < 1 ⟹ 极限 = 1', 1),
    ('例1 δ = ε/2 的证明书写规范（分析→综合）', 1),
    '【主板 · 右】例题板演区（学生上台）',
    ('例2 间断点分类 · 例3 0/0 三板斧 · 例4 零点定理三步法', 1),
    '【副板 · 角】易错墙',
    ('✗ 代入求极限（未验连续） ✗ 忘记弧度制 ✗ ∀ε∃δ 顺序颠倒 ✗ 开区间用最值定理', 1),
], y=1.5, w=12.0, size=13, ls=1.0, sa=0)
k.full_picture(s, FIGS[16], y=5.25, w=12.0)

# 32 结语
s = k.content_slide(prs, '结语：一场持续两百年的严格化', 'End')
k.bullets(s, [
    '牛顿、莱布尼茨用“无穷小”算对了答案，却说不清它是什么',
    '贝克莱主教的嘲讽：“已死量的幽灵”',
    '柯西给出极限的描述，魏尔斯特拉斯给出 ε-δ 的铁律',
    '至此，微积分才从“灵光一现的技巧”变成“无懈可击的科学”',
    '你今天学的这两个希腊字母，是这场革命的最终答案',
], y=1.8, w=5.5, size=18)
k.picture(s, FIGS[3], x=6.6, y=1.7, w=6.4)
k.callout(s, '下一讲预告：导数——用极限，去捕捉“瞬间的变化率”。',
          x=0.85, y=5.9, w=5.5, h=0.95, kind='key')

pptx_path = k.save(prs, os.path.join(OUT, '10_函数的极限与连续.pptx'))
print('slides:', len(prs.slides.__iter__.__self__._sldIdLst), '->', pptx_path)

# ============================================================ 教案 docx
from docx import Document
from docx.shared import Pt as DPt, RGBColor as DRGB, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SERIF = 'Noto Serif CJK SC'
doc = Document()
st = doc.styles['Normal']
st.font.name = SERIF; st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)

def H(t, lv=1):
    p = doc.add_heading('', level=lv)
    r = p.add_run(t); r.font.name = SERIF; r.font.size = DPt(15 if lv == 1 else 12.5)
    r.font.bold = True; r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
    return p

def P(t, bold=False, size=10.5, ind=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = DPt(4)
    p.paragraph_format.line_spacing = 1.4
    if ind: p.paragraph_format.left_indent = Cm(ind)
    r = p.add_run(t); r.font.name = SERIF; r.font.size = DPt(size); r.bold = bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
    return p

def TABLE(head, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(head)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(head):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = DPt(9.5); r.font.name = SERIF
        r._element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            for j, ln in enumerate(str(v).split('\n')):
                pp = cells[i].paragraphs[0] if j == 0 else cells[i].add_paragraph()
                pp.paragraph_format.space_after = DPt(1)
                rr = pp.add_run(ln); rr.font.size = DPt(9); rr.font.name = SERIF
                rr._element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run('第10讲　函数的极限与连续性　教学设计')
r.font.name = SERIF; r.font.size = DPt(19); r.bold = True; r.font.color.rgb = DRGB(0x11, 0x1B, 0x2E)
r._element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('从“无限逼近”的直观到 ε-δ 的严格　|　1 课时（60 分钟）')
r.font.name = SERIF; r.font.size = DPt(11); r.font.color.rgb = DRGB(0x4A, 0x55, 0x68)
r._element.rPr.rFonts.set(qn('w:eastAsia'), SERIF)

H('一、基本信息', 1)
TABLE(['课题', '课时', '课型', '授课对象'],
      [['函数的极限与连续性', '1 课时 / 60 分钟', '新授课（概念建构 + 定理应用）', '高三衔接班 / 大学微积分先修']],
      [5.5, 3.8, 4.5, 3.5])

H('二、教材分析', 1)
P('本讲是微积分的“第一块地基”。在数学史上，牛顿与莱布尼茨凭“无穷小”算对了答案却无法解释它究竟是零还是非零，招致贝克莱主教“已死量的幽灵”的著名嘲讽；直到 19 世纪，柯西给出极限的描述性定义，魏尔斯特拉斯锻造出 ε-δ 语言，微积分才由“灵光一现的技巧”升格为“无懈可击的科学”。因此本讲不仅是知识点的传授，更是一次数学严格化精神的启蒙。')
P('内容结构上，本讲承担三重任务：其一，建立极限的直观（无限逼近）并立即指出其模糊性；其二，完成向 ε-δ 严格语言的跃迁，这是全讲的核心与制高点；其三，以极限为工具定义连续性，并推出介值定理、最值定理这类具有强烈“存在性”意味的整体性质。后继的导数、定积分、级数无一不以本讲为逻辑起点，故本讲的地位是奠基性的、不可替代的。')

H('三、学情分析', 1)
P('已有基础：学生已熟悉函数、分段函数、三角函数与不等式，能进行因式分解与根式有理化，具备基本的图像直观。')
P('可能困难：① ε-δ 定义中“∀ε ∃δ”的量词顺序是学生理解的最大障碍，学生极易将其误读为“δ 先定、ε 后配”；② 学生长期形成“求极限=代入”的操作惯性，难以接受 lim 与 f(a) 无关；③ 对“存在性证明”缺乏经验，看到介值定理只会背结论，不会用来证方程有根。')
P('应对策略：用“挑战—应战”的博弈叙事把抽象量词具象为一场对抗；用可去间断点的三联图强行冲击“代入”惯性；用“对跖点同温”这类反直觉推论点燃学生对存在性定理的敬畏与兴趣。')

H('四、教学目标（三维）', 1)
TABLE(['维度', '具体目标'], [
    ['知识与技能', '① 理解极限的直观含义与 ε-δ 严格定义，能用定义证明简单线性函数的极限；\n② 掌握左右极限、连续的三条件定义与间断点四类分型；\n③ 掌握夹逼定理，会证并会用两个重要极限；\n④ 会用零点/介值定理证明方程在指定区间内有根。'],
    ['过程与方法', '① 经历“直观→模糊→量化→严格”的概念精致化全过程，体会数学严格化的必要性；\n② 通过双带图的动态观察（ε 收紧 δ 随之收紧），把逻辑量词转化为可视图像；\n③ 通过单位圆的面积夹逼，体会“几何直观服务于代数证明”的方法论。'],
    ['情感态度与价值观', '① 通过芝诺悖论与贝克莱悖论的化解，感受数学在质疑中自我完善的力量；\n② 通过“对跖点同温”“椅子四脚着地”的震撼推论，体会纯粹数学对现实世界的断言力；\n③ 在“挑战—应战”的博弈中培养严谨、不含糊、不回避的理性精神。'],
], [3.2, 14.0])

H('五、教学重点与难点', 1)
P('教学重点：ε-δ 定义的理解与初步运用；连续的三条件；两个重要极限；介值（零点）定理的应用。', bold=True)
P('教学难点：① ε-δ 中量词的逻辑顺序（∀ε 在前、∃δ 在后，δ 依赖于 ε）；② 深刻理解 lim 与 f(a) 无关；③ 存在性定理的证明书写规范。', bold=True)
P('突破方法：博弈叙事 + 双带图 + ε 逐级收紧的多子图动态演示 + “三步法”证题模板。')

H('六、教法与学法', 1)
P('教法：情境导入法（芝诺悖论）、问题驱动法、数形结合法、博弈叙事法、变式训练法。')
P('学法：观察归纳、动手板演、同伴互评、错例辨析（变式 1 的“陷阱题”）。')

H('七、教学准备', 1)
P('教师：本讲 PPT（32 页，含 15 张 matplotlib 精确配图）、板书分区规划、例题分层作业单。')
P('学生：直尺、草稿纸；预习“函数的单调性与有界性”相关内容。')

H('八、教学过程（分钟级时间轴，总计 60 分钟）', 1)
TABLE(['时段\n(min)', 'PPT\n页码', '环节', '教师活动', '学生活动', '设计意图'], [
    ['0–2\n(2′)', 'P1–P2', '开课·目标', '出示课题与学习目标，抛出主线："把『越来越接近』这句模糊的话，变成一条谁都无法抵赖的不等式。"', '浏览学习目标，明确本课要跨越的那道坎（ε-δ）。', '以"模糊→严格"的张力开篇，先立靶子，让学生知道全课在攻什么。'],
    ['2–10\n(8′)', 'P3–P6', '第1幕\n情境导入', '讲述芝诺"阿基里斯追乌龟"悖论，制造认知冲突；出示等比级数收敛图，揭示"无穷多段、路程有限"；随即给出 (x²−1)/(x−1) 的逼近图；最后抛出杀手锏——可去间断三联图。', '口答"到底追不追得上"；观察级数收敛图，说出无穷过程的"归宿"；观察三联图，惊觉 f(1) 三种取值而极限相同。', '以千年悖论激发兴趣，用一张图化解它；三联图正面击碎"求极限=代入"的惯性，为 ε-δ 的必要性埋下伏笔。'],
    ['10–24\n(14′)', 'P7–P12', '第2幕\nε-δ 严格化', '用"挑战—应战"博弈剧本演绎：对手报 ε，我应 δ；对手加码，我再加码。随后给出双带图（横带=挑战，竖带=应战），再用四子图动态展示 ε 收紧则 δ 随之收紧。板演例1（δ=ε/2），强调"分析法找 δ、综合法写证明"。最后引出左右极限与跳跃间断。', '扮演"对手"随口报出 ε 值，全班合作找 δ；在双带图上指认横带与竖带；口答变式（lim(4x−1)=11 取 δ=ε/4）。', '把抽象量词转译为可对抗、可操作的博弈，是本讲的核心教学法。图像先行、定义后置，降低符号恐惧；例1 给出可迁移的证明模板。'],
    ['24–36\n(12′)', 'P13–P17', '第3幕\n连续与间断', '给出连续的三条件定义，逐条对应"缺什么→出什么间断"；出示四宫格厘清四类间断点，配 sin(1/x) 放大图强化"第二类"的深渊感；板演例2（有理分式的两个间断点分类）。', '归纳判型口诀（先看左右极限存不存在）；跟随板演完成例2；完成课堂变式1（|x|/x 在 0 处为跳跃间断）。', '连续性是极限的直接应用，"三条件"把前面所有铺垫收束成一个定义；四宫格为学生提供稳定的分类框架，可直接迁移到考题。'],
    ['36–50\n(14′)', 'P18–P23', '第4幕\n两个重要极限', '讲夹逼定理（"上下两只手收拢"）；用单位圆面积链 ½sin x < ½x < ½tan x 完成 (sin x)/x → 1 的几何证明（反复强调弧度制前提）；出示数值验证表；板演例3（0/0 型三板斧：约零因子、有理化、凑标准形）；最后以单调有界准则引出 e。', '跟随推导，独立完成 ③ 取倒数 这一步；上台板演例3(2) 的半角变形；完成变式2（tan x / x → 1）。', '让学生亲历一个"完整的、有几何血肉的"极限证明，而非死记结论；三板斧是解题的通用武器，必须在课堂上形成肌肉记忆。'],
    ['50–58\n(8′)', 'P24–P28', '第5幕\n整体性质', '给出介值定理、零点定理、最值定理，强调"闭区间+连续"缺一不可（举 1/x 在 (0,1] 的反例）；出示"必过横线"图；高光时刻：用零点定理证明地球上必存在一对同温的对跖点；板演例4（证 x³−4x²+1=0 在 (0,1) 有根），提炼"证根三步法"；布置变式演练。', '观察温差函数 g(θ)=T(θ)−T(θ+π) 的两端异号，齐声推出结论；跟随例4 完成三步法书写；小组讨论变式1 的"陷阱"（无论 a 取何值都不连续）。', '从局部性质跃升到整体性质，让学生看到"连续"二字的真正威力；对跖点推论制造全课的情感高潮，把抽象定理钉进学生的长期记忆。'],
    ['58–60\n(2′)', 'P29–P32', '小结·作业\n·结语', '用知识网络图串起全课；布置 A/B/C 三层作业；以"贝克莱的幽灵—柯西—魏尔斯特拉斯"的史话收束，预告下一讲"导数：用极限捕捉瞬间的变化率"。', '对照知识网络自查掌握情况；领取分层作业单。', '结构化收口，史话升华，并以悬念衔接下一讲，形成课程连贯性。'],
], [1.6, 1.5, 2.0, 5.6, 3.5, 4.0])

H('九、板书设计', 1)
P('【主板·左】概念主线：1. 直观极限（芝诺图）→ 2. lim 与 f(a) 无关 → 3. ε-δ 定义（双带示意）→ 4. 连续三条件 → 5. 间断点四分类表')
P('【主板·中】推演区：单位圆面积夹逼链 ½sin x < ½x < ½tan x ⟹ cos x < (sin x)/x < 1 ⟹ 极限 = 1；例1 中 δ = ε/2 的规范书写（分析→综合）')
P('【主板·右】例题板演区（留给学生上台）：例2 间断点分类 · 例3 0/0 三板斧 · 例4 零点定理三步法')
P('【副板·角】易错墙：✗ 未验连续就代入求极限　✗ 忘记弧度制前提　✗ ∀ε∃δ 顺序颠倒　✗ 开区间误用最值定理')

H('十、分层作业', 1)
TABLE(['层级', '题目', '目标与用时'], [
    ['A 层\n基础巩固\n（必做）',
     '① 求 lim(x→0) sin5x/(2x)；lim(x→0)(1−cos2x)/x²\n② 求 f(x)=(x²−4)/(x−2) 的间断点并分类\n③ 用 ε-δ 定义证明 lim(x→1)(3x−2)=1，写出 δ 的取法',
     '巩固两个重要极限的凑形、间断点判型与定义证明模板。约 20 分钟。'],
    ['B 层\n能力提升\n（选做）',
     '④ 讨论 f(x)=x·sin(1/x)(x≠0)、f(0)=0 在 x=0 的连续性（提示：夹逼）\n⑤ 证明 x·2ˣ=1 在 (0,1) 内有实根\n⑥ 求 a、b 使给定分段函数在 R 上连续',
     '训练夹逼定理的灵活运用、存在性证明与分段函数拼接。约 20 分钟。'],
    ['C 层\n思维挑战\n（学有余力）',
     '⑦ 用 ε-δ 严格证明 lim(x→2) x²=4（提示：先限定 |x−2|<1 以控制 |x+2|<5，再取 δ=min{1, ε/5}）\n⑧ 查阅"椅子四脚着地定理"，写 300 字说明它如何依赖介值定理',
     '突破非线性函数的 δ 取法（min 技巧），并将数学定理迁移到现实建模。约 30 分钟。'],
], [2.6, 10.0, 4.6])

H('十一、教学反思（课后填写）', 1)
for t in ['1. ε-δ 的"挑战—应战"叙事，学生的接受度如何？有多少人能独立复述量词顺序？',
          '2. 可去间断三联图是否真正撼动了"求极限=代入"的惯性？课后作业②的错误率可作检验。',
          '3. 单位圆几何证明中，"取倒数改变不等号方向"这一步是否有学生卡壳？',
          '4. "对跖点同温"是否达到了预期的情感高潮？课堂气氛的峰值出现在哪一分钟？',
          '5. 时间分配是否精确？第2幕（ε-δ，14 分钟）是否需要从第4幕挪用时间？',
          '6. 分层作业的 A/B/C 三层配比是否合适？C 层的 min 技巧是否需要下沉到课堂？']:
    P(t, ind=0.5)
for _ in range(6):
    doc.add_paragraph('　')

docx_path = os.path.join(OUT, '教案_10_函数的极限与连续.docx')
doc.save(docx_path)
print('docx ->', docx_path)
