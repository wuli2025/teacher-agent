# -*- coding: utf-8 -*-
"""第01讲：导数的概念与几何意义 —— 60分钟课件包生成脚本"""
import sys, os
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch

OUT = '/mnt/c/Users/mi/Desktop/数学名师课件包/01_导数的概念与几何意义'
FIG = os.path.join(OUT, 'figures')
os.makedirs(FIG, exist_ok=True)
P = lambda n: os.path.join(FIG, n)

# ==================== 配图 ====================
# 画布宽度 = 该图在幻灯片上的最小展示宽度（英寸），保证贴图不缩小、图内字 >= 20pt
from matplotlib.ticker import MaxNLocator

def fig_dive():
    """图1 高台跳水 h(t)=-4.9t^2+6.5t+10，平均速度=割线斜率（展示 6.3in）"""
    fig, ax = k.new_fig(6.3, 3.8)
    t = np.linspace(0, 2.05, 400)
    h = -4.9 * t**2 + 6.5 * t + 10
    ax.plot(t, h, color=k.M_ACC2, lw=2.6, zorder=3)
    t1, t2 = 0.5, 1.6
    h1, h2 = -4.9*t1**2+6.5*t1+10, -4.9*t2**2+6.5*t2+10
    ax.plot([t1, t2], [h1, h2], color=k.M_RED, lw=2.4, ls='--', zorder=4)
    ax.scatter([t1, t2], [h1, h2], s=80, color=k.M_RED, zorder=5)
    ax.plot([t1, t2], [h1, h1], color=k.M_SLATE, lw=1.2, ls=':')
    ax.plot([t2, t2], [h1, h2], color=k.M_SLATE, lw=1.2, ls=':')
    ax.text((t1+t2)/2, h1-4.6, r'$\Delta t$', ha='center', color=k.M_SLATE)
    ax.text(t2+0.06, (h1+h2)/2, r'$\Delta h$', color=k.M_SLATE, va='center')
    ax.text(1.05, 15.6, '割线 = 平均速度', color=k.M_RED, ha='center')
    ax.set_xlim(-0.15, 2.4); ax.set_ylim(0, 18)
    ax.set_xticks([0, 1, 2]); ax.set_yticks([5, 10, 15])
    k.style_axes(ax, 't/s', 'h/m')
    ax.set_title(r'$h(t)=-4.9t^2+6.5t+10$', color=k.M_INK, pad=10)
    return k.save_fig(fig, P('f01_dive.png'))

def fig_avg_rate():
    """图2 平均变化率 = 割线斜率（展示 6.1in）"""
    fig, ax = k.new_fig(6.1, 3.6)
    x = np.linspace(-0.3, 3.2, 400)
    f = lambda u: 0.45*u**2 + 0.3*u + 1
    ax.plot(x, f(x), color=k.M_ACC2, lw=2.6, zorder=3)
    x0, x1 = 0.8, 2.6
    sl = (f(x1)-f(x0))/(x1-x0)
    ax.plot([x0-0.5, x1+0.5], [f(x0)-sl*0.5, f(x1)+sl*0.5], color=k.M_RED, lw=2.2, zorder=4)
    ax.scatter([x0, x1], [f(x0), f(x1)], s=90, color=k.M_INK, zorder=6)
    ax.text(x0-0.15, f(x0)+0.25, '$A$', color=k.M_INK, ha='right')
    ax.text(x1-0.12, f(x1)+0.2, '$B$', color=k.M_INK, ha='right')
    ax.plot([x0, x1], [f(x0), f(x0)], color=k.M_GRN, lw=2.0)
    ax.plot([x1, x1], [f(x0), f(x1)], color=k.M_ACC, lw=2.0)
    ax.text((x0+x1)/2, f(x0)-0.85, r'$\Delta x$', ha='center', color=k.M_GRN)
    ax.text(x1+0.08, (f(x0)+f(x1))/2, r'$\Delta y$', color=k.M_ACC, va='center')
    ax.text(0.05, 6.0, r'$k_{AB}=\dfrac{\Delta y}{\Delta x}$', color=k.M_RED, va='top')
    ax.set_xlim(-0.5, 4.2); ax.set_ylim(-0.2, 7.6)
    ax.set_xticks([1, 2, 3]); ax.set_yticks([2, 4, 6])
    k.style_axes(ax)
    ax.set_title('割线 AB 的斜率', color=k.M_INK, pad=10)
    return k.save_fig(fig, P('f02_avgrate.png'))

def fig_secant_seq():
    """图3 割线→切线（整幅，展示 8.6in）"""
    fig, ax = k.new_fig(8.6, 4.6)
    f = lambda u: 0.5*u**2 + 1
    x0 = 1.0
    x = np.linspace(-0.3, 3.4, 400)
    ax.plot(x, f(x), color=k.M_ACC2, lw=2.8, zorder=3)
    cols = ['#D9B382', '#CE9A55', '#C87E2A']
    for dx, c in zip([2.0, 1.2, 0.6], cols):
        s = (f(x0+dx) - f(x0)) / dx
        xs = np.array([x0-0.6, x0+dx+0.5])
        ax.plot(xs, f(x0) + s*(xs-x0), color=c, lw=2.0, zorder=4,
                label=r'$\Delta x=%.1f$' % dx)
        ax.scatter([x0+dx], [f(x0+dx)], s=60, color=c, zorder=6)
    xs = np.array([x0-0.9, x0+2.2])
    ax.plot(xs, f(x0) + 1.0*(xs-x0), color=k.M_RED, lw=3.0, zorder=7, label='切线 $k=1$')
    ax.scatter([x0], [f(x0)], s=130, color=k.M_RED, zorder=8, ec='white', lw=1.5)
    ax.text(x0-0.15, f(x0)-0.9, '$P$', color=k.M_RED, ha='right')
    ax.legend(loc='center right', framealpha=0.95)
    ax.set_xlim(-0.6, 6.6); ax.set_ylim(-0.3, 6.6)
    ax.set_xticks([1, 2, 3]); ax.set_yticks([2, 4, 6])
    k.style_axes(ax)
    ax.set_title('割线绕 P 旋转，极限位置就是切线', color=k.M_INK, pad=10)
    return k.save_fig(fig, P('f03_secant_to_tangent.png'))

def fig_slope_table():
    """图4 割线斜率双侧逼近（最小展示 5.1in）"""
    fig, ax = k.new_fig(5.1, 3.4)
    dxs = np.array([1, 0.5, 0.1, 0.01, 0.001])
    f = lambda u: u**2
    ks_r = (f(1+dxs) - f(1)) / dxs
    ks_l = (f(1-dxs) - f(1)) / (-dxs)
    ax.plot(dxs, ks_r, 'o-', color=k.M_ACC, lw=2.2, ms=9)
    ax.plot(dxs, ks_l, 's-', color=k.M_ACC2, lw=2.2, ms=9)
    ax.axhline(2, color=k.M_RED, lw=2.4, ls='--')
    ax.text(0.0012, 2.75, r"$f'(1)=2$", color=k.M_RED)
    ax.text(0.05, 2.85, '右侧', color=k.M_ACC, ha='center')
    ax.text(0.05, 1.20, '左侧', color=k.M_ACC2, ha='center')
    ax.set_xscale('log')
    ax.set_xticks([1e-3, 1e-2, 1e-1, 1])
    ax.set_ylim(0.75, 3.5)
    ax.set_yticks([1.0, 2.0, 3.0])
    k.style_axes(ax, r'$|\Delta x|$', '斜率', origin=False)
    return k.save_fig(fig, P('f04_slope_table.png'))

def fig_zoom():
    """图5 局部线性化三级放大（整幅，展示 10.2in）"""
    fig, axes = plt.subplots(1, 3, figsize=(10.2, 3.9))
    fig.patch.set_facecolor(k.M_PAPER)
    f = lambda u: np.sin(u) + 0.3*u
    x0 = 1.0
    s = np.cos(x0) + 0.3
    for ax, r, ttl in zip(axes, [1.5, 0.4, 0.06], ['视野 ±1.5', '放大 ×4', '放大 ×25']):
        xx = np.linspace(x0-r, x0+r, 400)
        ax.plot(xx, f(xx), color=k.M_ACC2, lw=2.8, zorder=3, label='曲线')
        ax.plot(xx, f(x0) + s*(xx-x0), color=k.M_RED, lw=2.0, ls='--', zorder=4, label='切线')
        ax.scatter([x0], [f(x0)], s=70, color=k.M_INK, zorder=6)
        ax.set_title(ttl, color=k.M_INK)
        ax.set_facecolor(k.M_PAPER)
        ax.grid(True, color=k.M_RULE, lw=0.7)
        ax.xaxis.set_major_locator(MaxNLocator(3))
        ax.yaxis.set_major_locator(MaxNLocator(3))
        for sp in ('top', 'right'):
            ax.spines[sp].set_visible(False)
    axes[0].legend(loc='lower right', framealpha=0.95)
    axes[2].text(0.5, 0.08, '几乎重合', transform=axes[2].transAxes,
                 ha='center', color=k.M_GRN)
    fig.suptitle('不断放大，曲线“变直”', fontsize=20, color=k.M_INK)
    return k.save_fig(fig, P('f05_zoom.png'))

def fig_error():
    """图6 误差 o(h)（展示 6.3in）"""
    fig, ax = k.new_fig(6.3, 3.8)
    h = np.linspace(-0.9, 0.9, 400)
    f = lambda u: np.exp(u)
    err = np.abs(f(1+h) - (f(1) + f(1)*h))
    ax.plot(h, err, color=k.M_RED, lw=2.8, label='$|R(h)|=o(h)$')
    ax.plot(h, np.abs(h), color=k.M_SLATE, lw=2.0, ls='--', label='$|h|$')
    ax.fill_between(h, 0, err, color=k.M_RED, alpha=0.13)
    ax.set_xlim(-1.0, 1.0); ax.set_ylim(0, 1.05)
    ax.set_xticks([-0.5, 0, 0.5]); ax.set_yticks([0.25, 0.5, 0.75, 1.0])
    k.style_axes(ax, 'h', '误差', origin=False)
    ax.legend(loc='upper center', framealpha=0.95)
    return k.save_fig(fig, P('f06_error.png'))

def fig_best_linear():
    """图7 最佳线性逼近（最小展示 5.2in）"""
    fig, ax = k.new_fig(5.2, 3.4)
    x = np.linspace(0.2, 2.0, 400)
    f = lambda u: u**2
    x0 = 1.0
    ax.plot(x, f(x), color=k.M_ACC2, lw=3.0, zorder=4)
    ax.plot(x, f(x0) + 2*(x-x0), color=k.M_RED, lw=2.6, zorder=5)
    for kk in (1.2, 3.0):
        ax.plot(x, f(x0) + kk*(x-x0), color='#9AA6B5', lw=1.6, ls='--', zorder=3)
    ax.scatter([x0], [f(x0)], s=110, color=k.M_INK, zorder=8)
    ax.fill_between(x, f(x), f(x0)+2*(x-x0), color=k.M_GRN, alpha=0.15)
    ax.text(2.05, f(x0)+2*1.0, '切线', color=k.M_RED, va='center')
    ax.text(2.05, f(x0)+1.2*1.0, '$k=1.2$', color=k.M_SLATE, va='center', fontsize=20)
    ax.text(1.72, 4.55, '$k=3$', color=k.M_SLATE, ha='center')
    ax.text(0.25, 3.3, r'$y=x^2$', color=k.M_ACC2)
    ax.set_xlim(0, 3.4); ax.set_ylim(-1.0, 5.0)
    ax.set_xticks([1, 2]); ax.set_yticks([0, 2, 4])
    k.style_axes(ax)
    return k.save_fig(fig, P('f07_best_linear.png'))

def fig_tangent_ex():
    """图8 例题2：y=x^3-3x 在 x0=1 处的切线（展示 5.4in）"""
    fig, ax = k.new_fig(5.4, 3.4)
    x = np.linspace(-2.4, 2.4, 500)
    f = lambda u: u**3 - 3*u
    ax.plot(x, f(x), color=k.M_ACC2, lw=2.8, zorder=3)
    ax.plot(x, 0*x - 2, color=k.M_RED, lw=2.4, zorder=4)
    ax.scatter([1], [-2], s=110, color=k.M_RED, zorder=6, ec='white', lw=1.4)
    ax.text(2.55, -1.8, '$y=-2$', color=k.M_RED, va='bottom')
    ax.text(1.25, -3.6, r"$P(1,-2)$", color=k.M_RED)
    ax.plot(x, 0*x + 2, color=k.M_GRN, lw=1.8, ls='--', zorder=4)
    ax.scatter([-1], [2], s=80, color=k.M_GRN, zorder=6)
    ax.text(-2.7, 2.4, r"$f'(-1)=0$", color=k.M_GRN)
    ax.set_xlim(-2.9, 4.6); ax.set_ylim(-4.6, 4.4)
    ax.set_xticks([-2, -1, 1, 2]); ax.set_yticks([-4, -2, 2, 4])
    k.style_axes(ax, 'x', '')
    ax.set_title(r'$y=x^3-3x$', color=k.M_INK, pad=8)
    return k.save_fig(fig, P('f08_tangent_ex.png'))

def fig_abs():
    """图9 反例：y=|x|（展示 6.3in）"""
    fig, ax = k.new_fig(6.3, 3.8)
    x = np.linspace(-2, 2, 400)
    ax.plot(x, np.abs(x), color=k.M_ACC2, lw=3.2, zorder=3)
    ax.plot([-1.7, 0], [1.7, 0], color=k.M_RED, lw=2.2, ls='--', zorder=4)
    ax.plot([0, 1.7], [0, 1.7], color=k.M_GRN, lw=2.2, ls='--', zorder=4)
    ax.scatter([0], [0], s=130, color=k.M_INK, zorder=6, ec='white', lw=1.5)
    ax.text(-2.3, 2.15, r"$f'_-(0)=-1$", color=k.M_RED)
    ax.text(0.35, 2.15, r"$f'_+(0)=+1$", color=k.M_GRN)
    ax.annotate('尖点', xy=(0, 0), xytext=(0.75, 0.75), color=k.M_INK,
                arrowprops=dict(arrowstyle='->', color=k.M_INK, lw=1.8))
    ax.set_xlim(-2.6, 2.8); ax.set_ylim(-0.6, 3.0)
    ax.set_xticks([-2, -1, 1, 2]); ax.set_yticks([1, 2])
    k.style_axes(ax)
    return k.save_fig(fig, P('f09_abs.png'))

def fig_xsin():
    """图10 x·sin(1/x) vs x^2·sin(1/x)（整幅，展示 10.0in）"""
    fig, axes = plt.subplots(1, 2, figsize=(10.0, 4.3))
    fig.patch.set_facecolor(k.M_PAPER)
    x = np.linspace(-0.35, 0.35, 6000)
    x = x[np.abs(x) > 1e-6]
    specs = [(lambda u: u*np.sin(1/u), r'$g(x)=x\sin\frac{1}{x}$', '不可导', np.abs(x)),
             (lambda u: u**2*np.sin(1/u), r'$G(x)=x^2\sin\frac{1}{x}$', "可导，$G'(0)=0$", x**2)]
    for ax, (g, ttl, sub, env) in zip(axes, specs):
        ax.plot(x, g(x), color=k.M_ACC2, lw=1.2, zorder=3)
        ax.plot(x, env, color=k.M_RED, lw=1.4, ls='--', zorder=4)
        ax.plot(x, -env, color=k.M_RED, lw=1.4, ls='--', zorder=4)
        ax.scatter([0], [0], s=55, color=k.M_INK, zorder=6)
        ax.set_title(ttl + '　' + sub, color=k.M_INK, fontsize=20)
        ax.set_facecolor(k.M_PAPER)
        ax.grid(True, color=k.M_RULE, lw=0.7)
        ax.xaxis.set_major_locator(MaxNLocator(3))
        ax.yaxis.set_major_locator(MaxNLocator(4))
        for sp in ('top', 'right'):
            ax.spines[sp].set_visible(False)
    fig.suptitle('红虚线为包络：夹逼决定可导与否', fontsize=20, color=k.M_INK)
    return k.save_fig(fig, P('f10_xsin.png'))

def fig_cont_impl():
    """图11 可导 ⇒ 连续（最小展示 5.1in）"""
    fig, ax = k.new_fig(5.1, 3.2)
    from matplotlib.patches import Ellipse
    ax.add_patch(Ellipse((0.48, 0.56), 0.90, 0.62, fc='#E8ECF2', ec=k.M_ACC2, lw=2.2, zorder=2))
    ax.add_patch(Ellipse((0.36, 0.52), 0.48, 0.38, fc='#F6E3C8', ec=k.M_ACC, lw=2.2, zorder=3))
    ax.text(0.48, 0.86, '连续', color=k.M_ACC2, ha='center', zorder=5)
    ax.text(0.36, 0.52, '可导', color=k.M_ACC, ha='center', va='center', zorder=5)
    ax.scatter([0.70], [0.55], s=110, color=k.M_RED, zorder=6)
    ax.text(0.735, 0.55, r'$|x|$', color=k.M_RED, va='center')
    ax.scatter([0.70], [0.36], s=110, color=k.M_RED, zorder=6)
    ax.text(0.735, 0.36, r'$x\sin\frac{1}{x}$', color=k.M_RED, va='center')
    ax.text(0.5, 0.08, r'可导 $\Rightarrow$ 连续，反之不真', color=k.M_INK, ha='center')
    ax.set_xlim(0.0, 1.0); ax.set_ylim(0.0, 1.0); ax.axis('off')
    return k.save_fig(fig, P('f11_cont_diff.png'))

def fig_speed():
    """图12 位移-时间与瞬时速度（展示 6.2in）"""
    fig, axes = plt.subplots(2, 1, figsize=(6.2, 4.6), sharex=True,
                             gridspec_kw={'height_ratios': [1.3, 1]})
    fig.patch.set_facecolor(k.M_PAPER)
    t = np.linspace(0, 4, 400)
    s = 2*t**3 - 9*t**2 + 12*t
    v = 6*t**2 - 18*t + 12
    ax = axes[0]
    ax.plot(t, s, color=k.M_ACC2, lw=2.8, zorder=3)
    for t0 in (1, 2, 3):
        k0 = 6*t0**2 - 18*t0 + 12
        s0 = 2*t0**3 - 9*t0**2 + 12*t0
        tt = np.array([t0-0.6, t0+0.6])
        ax.plot(tt, s0 + k0*(tt-t0), color=k.M_RED, lw=2.0)
        ax.scatter([t0], [s0], s=70, color=k.M_INK, zorder=6)
    ax.text(1.0, 6.6, r"$s'=0$", color=k.M_RED, ha='center')
    ax.text(2.0, -1.0, r"$s'=0$", color=k.M_RED, ha='center')
    ax.annotate(r"$s'=12$", xy=(3, 9), xytext=(1.5, 15.0), color=k.M_RED,
                arrowprops=dict(arrowstyle='->', color=k.M_RED, lw=1.6))
    ax.set_ylim(-3, 20)
    ax.set_yticks([0, 10, 20])
    ax.set_ylabel('s/m', color=k.M_INK)
    ax = axes[1]
    ax.plot(t, v, color=k.M_GRN, lw=2.8)
    ax.axhline(0, color=k.M_INK, lw=1.2)
    ax.fill_between(t, 0, v, where=(v < 0), color=k.M_RED, alpha=0.18)
    ax.scatter([1, 2], [0, 0], s=80, color=k.M_ACC, zorder=6)
    ax.text(1.5, -12, '$v<0$', color=k.M_RED, ha='center', va='center')
    ax.text(1.7, 19, '换向点 $t=1,2$', color=k.M_ACC, ha='center')
    ax.set_ylim(-20, 30)
    ax.set_yticks([-10, 0, 20])
    ax.set_ylabel('v', color=k.M_INK)
    ax.set_xlabel('t/s', color=k.M_INK)
    ax.set_xticks([0, 1, 2, 3, 4])
    for a in axes:
        a.set_facecolor(k.M_PAPER)
        a.grid(True, color=k.M_RULE, lw=0.7)
        for sp in ('top', 'right'):
            a.spines[sp].set_visible(False)
    axes[0].set_title(r'$s(t)=2t^3-9t^2+12t$', color=k.M_INK, fontsize=20)
    return k.save_fig(fig, P('f12_speed.png'))

def fig_piecewise():
    """图13 分段函数光滑拼接（展示 5.2in）"""
    fig, ax = k.new_fig(5.2, 3.5)
    x1 = np.linspace(-1.6, 1, 200)
    x2 = np.linspace(1, 3, 200)
    ax.plot(x1, x1**2, color=k.M_ACC2, lw=3.0)
    ax.plot(x2, 2*x2 - 1, color=k.M_GRN, lw=3.0)
    ax.scatter([1], [1], s=120, color=k.M_RED, zorder=6, ec='white', lw=1.4)
    ax.text(-1.75, 4.6, '$y=x^2$', color=k.M_ACC2)
    ax.text(1.55, 5.0, '$ax+b$', color=k.M_GRN, ha='left')
    ax.annotate('接点 $(1,1)$', xy=(1, 1), xytext=(0.45, -2.8), color=k.M_INK, ha='center',
                arrowprops=dict(arrowstyle='->', color=k.M_INK, lw=1.8))
    ax.text(-1.95, -4.4, r'$a=2,\ b=-1$', color=k.M_RED)
    ax.set_xlim(-2.1, 3.8); ax.set_ylim(-5.2, 5.9)
    ax.set_xticks([-1, 1, 2, 3]); ax.set_yticks([2, 4])
    k.style_axes(ax)
    return k.save_fig(fig, P('f13_piecewise.png'))

def fig_weierstrass():
    """图14 Weierstrass 函数（最小展示 5.85in）"""
    fig, axes = plt.subplots(1, 3, figsize=(5.8, 2.5))
    fig.patch.set_facecolor(k.M_PAPER)
    def Wf(x, N=40):
        s = np.zeros_like(x)
        for n in range(N):
            s += 0.5**n * np.cos(7**n * np.pi * x)
        return s
    spans = [(-1, 1), (0.09, 0.13), (0.1095, 0.1105)]
    ttl = ['整体', '×50', '×2000']
    for ax, (a, b), t in zip(axes, spans, ttl):
        xx = np.linspace(a, b, 8000)
        ax.plot(xx, Wf(xx), color=k.M_INK, lw=0.8)
        ax.set_title(t, color=k.M_INK, fontsize=20, pad=6)
        ax.set_facecolor(k.M_PAPER)
        ax.set_xticks([]); ax.set_yticks([])
        for sp in ('top', 'right', 'left', 'bottom'):
            ax.spines[sp].set_color(k.M_RULE)
    return k.save_fig(fig, P('f14_weierstrass.png'))

def fig_mindmap():
    """图15 知识结构图（最小展示 5.5in）"""
    fig, ax = k.new_fig(5.5, 3.3)
    def node(x, y, txt, c):
        ax.text(x, y, txt, fontsize=20, color=c, ha='center', va='center', zorder=5,
                bbox=dict(boxstyle='round,pad=0.35', fc='white', ec=c, lw=1.8))
    node(0.50, 0.93, r'平均变化率 $\Delta y/\Delta x$', k.M_ACC2)
    node(0.50, 0.70, r"导数 $f'(x_0)$", k.M_RED)
    node(0.24, 0.44, '几何·切线斜率', k.M_GRN)
    node(0.76, 0.44, '物理·瞬时速率', k.M_GRN)
    node(0.50, 0.22, '分析·最佳线性逼近', k.M_GRN)
    node(0.50, 0.00, r'可导 $\Rightarrow$ 连续（逆不真）', k.M_SLATE)
    def arrow(p, q, c=k.M_SLATE):
        ax.add_patch(FancyArrowPatch(p, q, arrowstyle='-|>', mutation_scale=16,
                                     color=c, lw=1.6, shrinkA=22, shrinkB=24))
    arrow((0.50, 0.93), (0.50, 0.70), k.M_ACC)
    arrow((0.50, 0.70), (0.24, 0.44), k.M_GRN)
    arrow((0.50, 0.70), (0.76, 0.44), k.M_GRN)
    arrow((0.50, 0.70), (0.50, 0.22), k.M_GRN)
    ax.text(0.86, 0.82, r'$\Delta x\to0$', fontsize=20, color=k.M_ACC, ha='center', va='center')
    ax.set_xlim(0, 1); ax.set_ylim(-0.09, 1.03); ax.axis('off')
    return k.save_fig(fig, P('f15_mindmap.png'))

FIGS = [fig_dive(), fig_avg_rate(), fig_secant_seq(), fig_slope_table(), fig_zoom(),
        fig_error(), fig_best_linear(), fig_tangent_ex(), fig_abs(), fig_xsin(),
        fig_cont_impl(), fig_speed(), fig_piecewise(), fig_weierstrass(), fig_mindmap()]
print('figures:', len(FIGS))

# ==================== PPT ====================
prs = k.new_deck()

# 1 封面
k.title_slide(prs, '导数的概念与几何意义',
              '从平均变化率到瞬时变化率 —— 局部线性化的现代视角',
              '第 01 讲', '60 分钟 · 高中—大一衔接')

# 2 学习目标
s = k.content_slide(prs, '本课学习目标', '目标')
k.bullets(s, [
    '知识：理解平均变化率、瞬时变化率，掌握导数定义与几何意义',
    ('会用定义式求导数（三步法：作差—作商—取极限）', 1),
    '能力：能求曲线在某点的切线方程；能用导数刻画瞬时速率',
    ('掌握判定可导性的方法：左右导数、极限存在性', 1),
    '素养：建立“局部线性化 / 最佳线性逼近”的现代观点',
    ('体会“可导 ⇒ 连续，反之不真”，感受数学的严谨与深刻', 1),
], y=1.65, w=6.3, size=17)
k.callout(s, '一句话主线：把弯的东西，在极小的尺度上当成直的。', y=6.05, w=6.3, h=0.85)
k.picture(s, FIGS[14], x=7.25, y=1.7, w=5.5)

# 3 幕一
k.section_slide(prs, '第 1 幕 · 情境与冲突', '从“平均”到“瞬时”：一个说不清的速度', '8 min')

# 4 情境
s = k.content_slide(prs, '情境：高台跳水的速度', '2 min')
k.bullets(s, [
    '运动员离水面高度 $h(t)=-4.9t^2+6.5t+10$（m）',
    ('在 [0.5, 1.6] 内的平均速度 ≈ −4.79 m/s', 1),
    '追问：这个平均速度能刻画运动员“此刻”的快慢吗？',
    ('平均速度只给出一段的“总体效果”，抹平了过程', 1),
], y=1.65, w=5.4, size=17)
k.callout(s, '认知冲突①：t = 1 s 这一瞬间，速度是多少？\n可 Δt = 0 时，Δh/Δt = 0/0 —— 没有意义！', y=5.35, w=5.4, h=1.4, kind='warn')
k.picture(s, FIGS[0], x=6.4, y=1.6, w=6.3)

# 5 平均变化率
s = k.content_slide(prs, '概念一：平均变化率', '概念')
k.formula(s, r'$\bar{k}=\dfrac{\Delta y}{\Delta x}=\dfrac{f(x_0+\Delta x)-f(x_0)}{\Delta x}\quad(\Delta x\neq 0)$',
          x=0.85, y=1.55, w=11.6, out=P('q01_avgrate.png'))
k.bullets(s, [
    '几何意义：割线 AB 的斜率',
    '物理意义：一段时间内的平均速度',
    r'注意 $\Delta x$ 可正可负，但恒不为 0',
], y=3.05, w=5.5, size=17)
k.callout(s, '变化率是“比值”，是把变化“摊平”后的强度，而非变化量本身。', y=5.55, w=5.5, h=1.0, kind='note')
k.picture(s, FIGS[1], x=6.6, y=2.9, w=6.1)

# 6 冲突
s = k.content_slide(prs, '瞬时速度：让 Δx 无限逼近 0', '2 min')
k.bullets(s, [
    r'思路：$\Delta x$ 不取 0，但让它要多小有多小',
    ('计算一串 Δx = 1, 0.5, 0.1, 0.01, 0.001 … 的平均变化率', 1),
    ('观察这串数是否趋于一个确定的值', 1),
    '若左右两侧都趋于同一个数 —— 这个数就是我们要的“瞬时”',
], y=1.65, w=5.4, size=17)
k.callout(s, '数学化：这正是极限的思想。“无限逼近”不是“等于”。', y=5.5, w=5.4, h=1.0)
k.picture(s, FIGS[3], x=6.4, y=1.6, w=6.3)

# 7 幕二
k.section_slide(prs, '第 2 幕 · 概念建构', '导数的定义与几何意义', '14 min')

# 8 割线到切线
s = k.content_slide(prs, '割线 → 切线：极限位置', '几何')
k.full_picture(s, FIGS[2], y=1.45, w=8.6)
k.callout(s, '当 Δx → 0，割线 PQ 绕定点 P 旋转，趋于唯一的极限位置 —— 切线。\n注意：这是切线的现代定义，而非“与曲线只交于一点”。',
          x=1.9, y=6.35, w=9.5, h=1.0, kind='note')

# 9 导数定义
s = k.content_slide(prs, '概念二：导数的定义', '核心')
k.formula(s, r"$f'(x_0)=\lim_{\Delta x\to 0}\dfrac{\Delta y}{\Delta x}=\lim_{\Delta x\to 0}\dfrac{f(x_0+\Delta x)-f(x_0)}{\Delta x}$",
          x=0.85, y=1.5, w=11.6, out=P('q02_def.png'))
k._tb(s, k.Inches(2.7), k.Inches(3.05), k.Inches(2.2), k.Inches(0.5), "等价形式：", 18, k.SLATE)
k.formula(s, r"$f'(x_0)=\lim_{x\to x_0}\dfrac{f(x)-f(x_0)}{x-x_0}$",
          x=4.6, y=3.0, w=6.4, size=0.62, out=P('q03_def2.png'))
k.bullets(s, [
    '若该极限存在，称 f 在 x₀ 处可导，极限值记作 f′(x₀)',
    '若极限不存在，则 f 在 x₀ 处不可导',
    '导数是一个数；把 x₀ 换成变量 x，得到导函数 f′(x)',
], y=4.25, w=11.6, size=17)
k.callout(s, '几何意义：f′(x₀) 就是曲线 y = f(x) 在点 (x₀, f(x₀)) 处切线的斜率。',
          x=0.85, y=6.15, w=11.6, h=0.85)

# 10 三步法 + 例1
s = k.content_slide(prs, '例题1（硬算）：用定义求 f(x)=x² 在 x₀=1 处的导数', '例题')
k.bullets(s, [
    '第一步 作差：Δy = f(1+Δx) − f(1) = (1+Δx)² − 1 = 2Δx + (Δx)²',
    '第二步 作商：Δy/Δx = 2 + Δx',
    '第三步 取极限：f′(1) = lim(Δx→0)(2 + Δx) = 2',
], y=1.6, w=6.6, size=17)
k.formula(s, r"$f'(1)=\lim_{\Delta x\to0}\dfrac{(1+\Delta x)^2-1^2}{\Delta x}=\lim_{\Delta x\to0}(2+\Delta x)=2$",
          x=0.85, y=4.0, w=6.6, size=0.62, out=P('q04_ex1.png'))
k.callout(s, '关键：作商后必须先约去 Δx，把 0/0 型化为可直接代入的式子，再令 Δx → 0。',
          x=0.85, y=5.45, w=6.6, h=1.2, kind='warn')
k.picture(s, FIGS[3], x=7.7, y=1.6, w=5.1)

# 11 例1 推广
s = k.content_slide(prs, '例题1 · 推广：一般点 x₀ 与导函数', '板演')
k.formula(s, r"$f'(x_0)=\lim_{\Delta x\to0}\dfrac{(x_0+\Delta x)^2-x_0^2}{\Delta x}=\lim_{\Delta x\to0}(2x_0+\Delta x)=2x_0$",
          x=0.85, y=1.6, w=11.6, size=0.85, out=P('q05_ex1b.png'))
k.bullets(s, [
    r'把 $x_0$ 视作变量：得导函数 $f\,\!^{\prime}(x)=2x$',
    '同理可证：(x³)′ = 3x²，(1/x)′ = −1/x²，(√x)′ = 1/(2√x)',
    ('提示：√x 需分子有理化；1/x 需通分', 1),
    '导数是“函数 → 函数”的运算，这是第 02 讲的主题',
], y=3.3, w=11.6, size=17)
k.callout(s, '定义式是根，公式是叶。忘了公式可以推，忘了定义就寸步难行。',
          x=0.85, y=6.0, w=11.6, h=0.85)

# 12 幕三
k.section_slide(prs, '第 3 幕 · 现代视角', '局部线性化：把曲线“拉直”', '14 min')

# 13 放大
s = k.content_slide(prs, '实验：不断放大一条曲线', '直观')
k.full_picture(s, FIGS[4], y=1.5, w=10.2)
k.callout(s, '可导的本质：在足够小的邻域内，曲线与它的切线“看不出区别”。\n这就是“光滑”的精确含义 —— 局部线性化（local linearization）。',
          x=1.5, y=5.9, w=10.3, h=1.15, kind='note')

# 14 最佳线性逼近
s = k.content_slide(prs, '概念三：导数 = 最佳线性逼近的斜率', '核心')
k.formula(s, r"$f(x_0+h)=f(x_0)+f'(x_0)\,h+o(h)\qquad (h\to 0)$",
          x=0.85, y=1.5, w=11.6, out=P('q06_linapp.png'))
k.bullets(s, [
    r'$o(h)$ 读作“$h$ 的高阶无穷小”：$\lim_{h\to0} o(h)/h = 0$',
    '若存在常数 A 使 f(x₀+h) = f(x₀) + Ah + o(h)，则 A 唯一，且 A = f′(x₀)',
    '这是导数的“分析学定义”，可以原样推广到多元、向量、流形',
], y=3.0, w=6.4, size=16.5)
k.callout(s, '同学们记住：导数不是“斜率”，导数是“最好的那条直线”的斜率。',
          x=0.85, y=5.8, w=6.4, h=1.1)
k.picture(s, FIGS[6], x=7.5, y=2.1, w=5.3)

# 15 误差
s = k.content_slide(prs, '误差有多小？—— o(h) 的量化', '推演')
k.picture(s, FIGS[5], x=6.5, y=1.6, w=6.3)
k.bullets(s, [
    r'取 $f(x)=e^x,\ x_0=1$，切线 $L(h)=e+e\,h$',
    r'误差 $R(h)=f(1+h)-L(h)$',
    ('h = 0.1 → R ≈ 0.0143（约 h 的 14%）', 1),
    ('h = 0.01 → R ≈ 0.000136（约 h 的 1.4%）', 1),
    ('h 缩小 10 倍，误差缩小约 100 倍 —— 二阶！', 1),
], y=1.65, w=5.5, size=16.5)
k.callout(s, '线性近似的实用价值：e^1.01 ≈ e + 0.01e ≈ 2.7455（真值 2.7456）。\n工程师、物理学家每天都在这么算。',
          x=0.85, y=5.5, w=5.5, h=1.4, kind='note')

# 16 例2 切线
s = k.content_slide(prs, '例题2：求切线方程 y = x³ − 3x 在 x₀ = 1 处', '例题')
k.bullets(s, [
    '① 定义求导：Δy = (1+Δx)³ − 3(1+Δx) − (1−3) = 3Δx·0 + 3(Δx)² + (Δx)³',
    ('整理：Δy/Δx = 0 + 3Δx + (Δx)² → f′(1) = 0', 1),
    '② 求切点：f(1) = 1 − 3 = −2，切点 P(1, −2)',
    '③ 点斜式：y − (−2) = 0·(x − 1)，即 y = −2',
    '④ 检验：该点切线水平，正是极小值点',
], y=1.6, w=6.3, size=16)
k.callout(s, '易错点：“过点 P 的切线”与“在点 P 处的切线”不同！\n若 P 不在曲线上，须设切点 (t, f(t)) 反解 t。',
          x=0.85, y=5.55, w=6.3, h=1.35, kind='warn')
k.picture(s, FIGS[7], x=7.4, y=1.6, w=5.4)

# 17 变式1
s = k.content_slide(prs, '变式1：过曲线外一点作切线', '变式')
k.bullets(s, [
    r'求过点 $A(0,-16)$ 且与曲线 $y=x^3$ 相切的直线方程',
    '设切点 (t, t³)，由 (x³)′ = 3x² 得切线斜率 3t²',
    '切线：y − t³ = 3t²(x − t)。代入 A：−16 − t³ = 3t²(0 − t)',
    '即 −16 − t³ = −3t³ ⟹ 2t³ = 16 ⟹ t = 2',
    '切点 (2, 8)，斜率 12，切线：y = 12x − 16',
], y=1.6, w=7.0, size=16.5)
k.callout(s, '模型化：切线三要素 = 切点在曲线上 + 斜率是导数 + 直线过已知点。\n三个条件，恰好定出参数 t。',
          x=0.85, y=5.5, w=7.0, h=1.35)
k.picture(s, FIGS[6], x=7.95, y=1.9, w=5.2)

# 18 幕四
k.section_slide(prs, '第 4 幕 · 边界探究', '可导与连续：一个单向的箭头', '12 min')

# 19 可导⇒连续
s = k.content_slide(prs, '定理：可导 ⇒ 连续', '推演')
k.formula(s, r"$\lim_{h\to0}[f(x_0+h)-f(x_0)]=\lim_{h\to0}\dfrac{f(x_0+h)-f(x_0)}{h}\cdot h=f'(x_0)\cdot 0=0$",
          x=0.7, y=1.5, w=11.9, size=0.8, out=P('q07_thm.png'))
k.bullets(s, [
    '故 lim(h→0) f(x₀+h) = f(x₀)，即 f 在 x₀ 处连续',
    '也可由线性化式直接读出：f(x₀+h) = f(x₀) + f′(x₀)h + o(h) → f(x₀)',
], y=3.2, w=6.6, size=17)
k.callout(s, '逆命题为假！连续只保证“不断开”，可导还要求“不打折”。\n下面看两个经典反例。',
          x=0.85, y=5.4, w=6.6, h=1.35, kind='warn')
k.picture(s, FIGS[10], x=7.7, y=3.0, w=5.1)

# 20 反例1
s = k.content_slide(prs, '反例一：y = |x| 在 x = 0（尖点）', '反例')
k.picture(s, FIGS[8], x=6.5, y=1.6, w=6.3)
k.formula(s, r"$f'_{-}(0)=\lim_{h\to0^-}\dfrac{|h|}{h}=-1,\quad f'_{+}(0)=\lim_{h\to0^+}\dfrac{|h|}{h}=+1$",
          x=0.7, y=1.7, w=5.9, size=0.42, out=P('q08_abs.png'))
k.bullets(s, [
    '左右导数存在但不相等 ⟹ 极限不存在 ⟹ 不可导',
    '几何：尖点处没有唯一的切线（左右各有一条）',
    r'但 $\lim_{x\to0}|x|=0=|0|$，函数在 0 处连续',
], y=2.9, w=5.9, size=16.5)
k.callout(s, '判定口诀：可导 ⟺ 左导数 = 右导数 = 同一个有限数。',
          x=0.85, y=5.7, w=5.5, h=0.9, kind='note')

# 21 反例2
s = k.content_slide(prs, '反例二：x·sin(1/x) —— 震荡型不可导', '反例')
k.full_picture(s, FIGS[9], y=1.5, w=10.0)
k.callout(s, 'g(h)/h = sin(1/h) 在 h→0 时无休止地在 [−1,1] 间震荡，极限不存在。\n而 G(h)/h = h·sin(1/h)，被 |h| 夹逼归零 —— 多一个 x 因子，就“驯服”了震荡。',
          x=1.4, y=6.05, w=10.5, h=1.15, kind='note')

# 22 变式2
s = k.content_slide(prs, '变式2：分段函数的光滑拼接', '变式')
k.bullets(s, [
    '已知 f(x) = x²（x ≤ 1），= ax + b（x > 1），在 x = 1 处可导，求 a、b',
    '第一关（连续）：左极限 1 = 右极限 a + b ⟹ a + b = 1',
    '第二关（可导）：f′₋(1) = 2·1 = 2；f′₊(1) = a ⟹ a = 2',
    '联立得 a = 2，b = −1，此时 f 在 x = 1 处光滑',
], y=1.6, w=6.5, size=16.5)
k.callout(s, '顺序不可颠倒：先用连续定一个方程，再用可导定另一个。\n若只用可导而不查连续，会得到“断开却斜率相同”的假解。',
          x=0.85, y=5.5, w=6.5, h=1.35, kind='warn')
k.picture(s, FIGS[12], x=7.6, y=1.6, w=5.2)

# 23 幕五
k.section_slide(prs, '第 5 幕 · 应用与升华', '瞬时速率 · 认知冲突 · 数学之美', '12 min')

# 24 例3 实际
s = k.content_slide(prs, '例题3（实际背景）：质点的瞬时速度', '例题')
k.picture(s, FIGS[11], x=6.6, y=1.55, w=6.2)
k.bullets(s, [
    r'质点位移 $s(t)=2t^3-9t^2+12t$（m, s）',
    r'由定义可得 $v(t)=s\,\!^{\prime}(t)=6t^2-18t+12$',
    ('v(1) = 0：t = 1 s 时瞬时静止（换向点）', 1),
    ('v(2) = 0：t = 2 s 时再次换向', 1),
    ('1 < t < 2 时 v < 0：质点反向运动', 1),
    '追问：平均速度 v̄[0,3] = 9/3 = 3，能看出换向吗？不能。',
], y=1.6, w=5.6, size=15.5)
k.callout(s, '导数的威力：它把“整体的平均”升级为“逐点的精确刻画”。',
          x=0.85, y=6.15, w=5.6, h=0.85)

# 25 认知冲突
s = k.content_slide(prs, '认知冲突②：连续但处处不可导，可能吗？', '震撼')
k.bullets(s, [
    '|x| 只在一个点不可导；能否造出“到处都是尖点”的连续函数？',
    '1872 年，Weierstrass 给出肯定回答，震动整个数学界',
    ('当时主流认为：连续曲线至多在少数点不可导', 1),
], y=1.6, w=11.6, size=17)
k.picture(s, FIGS[13], x=1.87, y=3.1, w=9.6)

# 26 冲突反思
s = k.content_slide(prs, '这幅图告诉我们什么', '升华')
k.bullets(s, [
    '“连续”与“可导”是两种截然不同的正则性，不能靠直觉互推',
    '直觉会骗人 —— 这正是数学需要严格定义与证明的理由',
    'Weierstrass 函数并非病态玩具：它是布朗运动路径、股价、海岸线的数学原型（分形）',
    '可导 = 局部可拉直 = 光滑；不可导 = 无论怎么放大都粗糙',
], y=1.7, w=6.6, size=17)
k.callout(s, '数学的美，一半来自能被驯服的规律，一半来自拒绝被驯服的例外。',
          x=0.85, y=6.0, w=6.6, h=0.9)
k.picture(s, FIGS[13], x=7.35, y=2.9, w=5.85)

# 27 小结
s = k.content_slide(prs, '课堂小结：一条主线，三重身份', '小结')
k.full_picture(s, FIGS[14], y=1.45, w=8.0)
k.callout(s, '三重身份：几何 = 切线斜率；分析 = 最佳线性逼近；物理 = 瞬时变化率。',
          x=1.7, y=6.35, w=9.9, h=0.85)

# 28 方法清单
s = k.content_slide(prs, '方法清单（务必背下）', '小结')
k.bullets(s, [
    '定义求导三步法：作差 Δy → 作商 Δy/Δx（必须约去 Δx）→ 取极限',
    '求切线方程：先定切点，再求斜率 f′(x₀)，最后点斜式',
    ('若点不在曲线上：设切点 (t, f(t)) 反解', 1),
    '判可导：左导数 = 右导数 = 有限数；分段函数先查连续',
    '线性近似：f(x₀+h) ≈ f(x₀) + f′(x₀)h（h 很小时）',
], y=1.65, w=11.6, size=17)
k.callout(s, '常见错误 TOP3：① 忘约 Δx 直接代 0 得 0/0；② 切点未在曲线上就套点斜式；③ 只查斜率不查连续。',
          x=0.85, y=5.75, w=11.6, h=1.2, kind='warn')

# 29 作业
s = k.content_slide(prs, '分层作业', '作业')
k.bullets(s, [
    'A 层（必做·全体）',
    ('1. 用定义求 f(x)=3x²−1 在 x₀=2 处的导数', 1),
    ('2. 求 y=1/x 在点 (2, 1/2) 处的切线方程', 1),
    ('3. 判断 f(x)=|x−1| 在 x=1 处是否可导，说明理由', 1),
    'B 层（选做·进阶）',
    ('4. 求过点 (1, 0) 且与 y=x² 相切的所有直线', 1),
    ('5. 设 f(x)=x²(x≤2)，=ax+b(x>2) 在 x=2 可导，求 a, b', 1),
    'C 层（挑战·探究）',
    ('6. 证明：若 f 在 x₀ 可导且 f(x₀)=0，则 g(x)=|f(x)| 在 x₀ 可导 ⟺ f′(x₀)=0', 1),
    ('7. 用计算机作 W(x)=Σ2⁻ⁿcos(7ⁿπx) 前 5 项的图，观察“粗糙度”', 1),
], y=1.6, w=11.6, size=15)

# 30 板书提纲
s = k.content_slide(prs, '板书设计提纲', '板书')
k.bullets(s, [
    '【左板 · 主干】',
    ('课题：导数的概念与几何意义', 1),
    ('一、平均变化率 Δy/Δx（割线斜率）', 1),
    ('二、导数定义 f′(x₀)=lim(Δx→0) Δy/Δx  ★核心框住★', 1),
    ('三、几何意义：切线斜率；分析意义：f(x₀+h)=f(x₀)+f′(x₀)h+o(h)', 1),
    ('四、可导 ⇒ 连续（单向箭头，反例 |x|）', 1),
], y=1.6, w=6.2, size=15.5)
k.bullets(s, [
    '【右板 · 例题板演】',
    ('例1 定义求导（三步法完整展示，保留不擦）', 1),
    ('例2 切线方程（切点—斜率—点斜式，竖排）', 1),
    ('例3 瞬时速度（v(t) 表格 + 换向点）', 1),
    '【副板 · 随时擦写】',
    ('反例草图：|x| 尖点、x·sin(1/x) 震荡', 1),
    ('易错提醒：约 Δx！查连续！', 1),
], x=6.9, y=1.6, w=5.7, size=15.5)

path = k.save(prs, os.path.join(OUT, '01_导数的概念与几何意义.pptx'))
print('pptx:', path, '页数=', len(prs.slides.__iter__.__self__._sldIdLst))

# ==================== 教案 DOCX ====================
from docx import Document
from docx.shared import Pt as DPt, RGBColor as DRGB, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Noto Serif CJK SC'
st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Cm(2.2)

def H(t, lv=1):
    p = doc.add_heading('', level=lv)
    r = p.add_run(t)
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
    r.font.size = DPt(16 if lv == 1 else 13)
    r.font.bold = True
    return p

def PP(t, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    r.font.size = DPt(size); r.font.bold = bold
    p.paragraph_format.space_after = DPt(4)
    return p

t = doc.add_heading('', level=0)
r = t.add_run('教学设计（教案）  第 01 讲　导数的概念与几何意义')
r.font.name = 'Noto Serif CJK SC'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
r.font.size = DPt(20); r.font.bold = True
r.font.color.rgb = DRGB(0x11, 0x1B, 0x2E)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

H('一、基本信息', 1)
info = doc.add_table(rows=4, cols=4); info.style = 'Table Grid'
rows = [('课题', '导数的概念与几何意义', '课时', '1 课时（60 分钟）'),
        ('课型', '新授课（概念建构 + 探究）', '授课对象', '优秀高中 / 高中—大一衔接'),
        ('教材', '人教A版选择性必修二 5.1（拓展至线性逼近视角）', '配套课件', '01_导数的概念与几何意义.pptx（30 页）'),
        ('教具', '多媒体课件、几何画板/GeoGebra、板书', '配图', 'figures/ 共 15 张 matplotlib 数学图')]
for i, rw in enumerate(rows):
    for j, c in enumerate(rw):
        cell = info.cell(i, j); cell.text = ''
        p = cell.paragraphs[0]; run = p.add_run(c)
        run.font.name = 'Noto Serif CJK SC'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        run.font.size = DPt(9.5); run.font.bold = (j % 2 == 0)

H('二、教材分析', 1)
PP('导数是微积分的核心概念，是研究函数变化率的基本工具，也是高中数学由“常量数学”迈向“变量数学”的分水岭。本节课处于导数单元的起始位置，承担“从平均变化率到瞬时变化率”的概念奠基任务。教材以物理情境（高台跳水）引入，经由割线斜率的极限得到切线斜率，最终抽象出导数定义。')
PP('本设计在教材基础上作两点提升：（1）显式引入“局部线性化”这一现代视角，把导数定义式改写为 f(x₀+h)=f(x₀)+f′(x₀)h+o(h)，揭示导数即“最佳线性逼近的斜率”，为后续多元微分、微分方程乃至数值计算埋下伏笔；（2）通过 |x|、x·sin(1/x) 与 Weierstrass 函数三个反例，把“可导⇒连续，反之不真”讲透，培养学生对数学严谨性的敬畏。')

H('三、学情分析', 1)
PP('已有基础：学生已掌握函数、极限的初步概念，熟悉斜率、直线方程，具备一定的代数变形能力（因式分解、有理化）。')
PP('可能困难：① 对“无限逼近”与“等于”的区别理解不清，易把 Δx→0 直接当作 Δx=0；② 割线动态趋于切线的过程缺乏空间想象，需借助图形与动画；③ 难以接受“连续未必可导”，直觉与逻辑冲突；④ 作商后不知先约去 Δx，卡在 0/0。')
PP('应对策略：数值表格 + 动态割线序列图 + 三级放大实验，用“看得见”的方式支撑抽象；用反例制造认知冲突，再用严格定义化解。')

H('四、教学目标（三维）', 1)
PP('1. 知识与技能：理解平均变化率与瞬时变化率的含义；掌握导数的定义式及等价形式；能用定义（三步法）求简单函数在某点的导数；能求曲线在某点处的切线方程；掌握可导性的判定方法。', size=10.5)
PP('2. 过程与方法：经历“情境—冲突—数值探究—几何直观—抽象定义—反例检验”的完整概念建构过程；体会极限思想、逼近思想、以直代曲思想；学会用左右导数判定可导性。')
PP('3. 情感态度与核心素养：发展数学抽象、逻辑推理、直观想象三大核心素养；通过 Weierstrass 函数感受数学的深刻与美，体会“直觉需要证明来校正”的科学精神。')

H('五、教学重点与难点', 1)
PP('重点：导数的定义及其几何意义；用定义求导的三步法；切线方程的求法。', bold=True)
PP('难点：① 平均变化率到瞬时变化率的极限过程（Δx→0 的辩证理解）；② 导数作为“最佳线性逼近”的分析学理解；③ 可导与连续的关系及反例的构造思路。', bold=True)

H('六、教法与学法', 1)
PP('教法：情境驱动法（高台跳水）、问题链导学法、数形结合法、反例教学法、变式教学法。')
PP('学法：观察—猜想—验证—抽象的探究式学习；数值实验与几何直观并行；同伴互讲与板演展示。')

H('七、教学准备', 1)
PP('教师：PPT 课件 30 页、GeoGebra 割线动画、15 张精绘数学图、板书规划表。')
PP('学生：预习平均变化率；准备计算器（用于割线斜率数值逼近表）；带方格纸。')

H('八、教学过程（分钟级时间轴，总计 60 分钟）', 1)
COLS = ['时间', 'PPT页', '环节', '教师活动', '学生活动', '设计意图']
plan = [
 ('0–1′ (1)', 'P1', '导入', '出示课题与本讲主线“把弯的当成直的”，点明本课在微积分中的地位。', '静听，进入情境。', '开门见山，建立整体图式。'),
 ('1–3′ (2)', 'P2', '目标', '宣读三维学习目标，展示知识结构图（图15）预览全课。', '默读目标，圈出关键词“局部线性化”。', '目标前置，学习有靶心。'),
 ('3–4′ (1)', 'P3', '过场', '进入第 1 幕：从“平均”到“瞬时”。', '记录幕次。', '结构化提示，节奏感。'),
 ('4–6′ (2)', 'P4', '情境', '播放高台跳水视频/图1，计算 [0.5,1.6] 平均速度 ≈ −4.79 m/s。追问：这能代表 t=1 s 的快慢吗？', '计算平均速度；发现平均速度为负而运动员先上升后下降，产生疑惑。', '认知冲突①：平均抹平了过程。'),
 ('6–9′ (3)', 'P5', '概念一', '规范给出平均变化率定义（公式贴图），结合图2讲清“割线斜率”几何意义，强调 Δx≠0。', '在笔记本上默写定义式，标注 Δx、Δy 在图上的位置。', '把物理量数学化，形数对应。'),
 ('9–12′ (3)', 'P6', '探究', '组织数值实验：令 Δx=1,0.5,0.1,0.01,0.001 计算 f(x)=x² 在 x₀=1 的割线斜率（图4）。左右两侧同时算。', '分两组用计算器算左右两列数据，报出结果：2.5,2.25,2.1,2.01,2.001 与 1.5,1.75,1.9,1.99,1.999。', '数值证据先行，为极限提供实感；双侧逼近暗伏“左右导数”。'),
 ('12–13′ (1)', 'P7', '过场', '进入第 2 幕：概念建构。', '', '节奏切换。'),
 ('13–17′ (4)', 'P8', '几何', '用 GeoGebra 拖动 Q 点，配合图3展示割线绕 P 旋转趋于切线；强调切线的极限定义（而非“只交一点”）。', '观察动画，指出圆的切线定义在一般曲线上失效（如 y=x³ 在原点）。', '突破难点：切线的现代定义；直观想象素养。'),
 ('17–21′ (4)', 'P9', '概念二', '板书导数定义式与等价形式；说明“极限存在则可导”；辨析 f′(x₀) 是数、f′(x) 是函数。', '齐读定义；举手回答：极限不存在意味着什么？', '核心概念落地，语言精确化。'),
 ('21–25′ (4)', 'P10', '例题1', '板演：用定义求 f(x)=x² 在 x₀=1 的导数。逐字强调“作差—作商—取极限”，红笔圈出“约去 Δx”这一步。', '跟做；同桌互查是否先约分再取极限。', '突破易错点①：0/0 型必须先化简。'),
 ('25–27′ (2)', 'P11', '推广', '把 x₀ 一般化得 f′(x)=2x；布置口头推导 (x³)′、(1/x)′、(√x)′。', '口头说出 (x³)′ 的作差式；一名学生上台写 √x 的有理化。', '从“数”到“函数”，为第02讲铺路。'),
 ('27–28′ (1)', 'P12', '过场', '进入第 3 幕：现代视角。', '', '升华信号。'),
 ('28–31′ (3)', 'P13', '实验', '演示三级放大（图5）：视野 ±1.5 → ±0.4 → ±0.06，曲线逐渐与切线重合。', '观察并描述：越放大越像直线；说出“光滑”的直观含义。', '建立“局部线性化”的图像直觉。'),
 ('31–35′ (4)', 'P14', '概念三', '给出 f(x₀+h)=f(x₀)+f′(x₀)h+o(h)；解释 o(h) 的含义；指出满足此式的 A 唯一且等于 f′(x₀)（图7）。', '思考：为什么其他斜率的直线不行？误差是几阶？', '导数的分析学定义，衔接大学微积分，数学抽象素养。'),
 ('35–38′ (3)', 'P15', '量化', '用 f(x)=eˣ, x₀=1 算误差（图6）：h=0.1→R≈0.0143；h=0.01→R≈0.000136。h 缩 10 倍误差缩 100 倍。演示 e^1.01 的线性近似。', '用计算器验证 e+0.01e=2.7455 与真值 2.7456。', '让 o(h) 从符号变成可感的数；体会应用价值。'),
 ('38–42′ (4)', 'P16', '例题2', '板演求 y=x³−3x 在 x₀=1 处的切线：定义求导→f′(1)=0→切点(1,−2)→y=−2。强调“在点处”与“过点”的区别。', '独立完成，两名学生板演，全班纠错。', '规范切线求法；预警易错点②。'),
 ('42–45′ (3)', 'P17', '变式1', '变式：过 A(0,−16) 作 y=x³ 的切线。引导设切点 (t,t³) 反解 t=2。', '小组讨论 2 分钟，代表说思路；完成计算得 y=12x−16。', '模型化：切线三要素；提升思维层次。'),
 ('45–46′ (1)', 'P18', '过场', '进入第 4 幕：边界探究。', '', ''),
 ('46–48′ (2)', 'P19', '定理', '证明可导⇒连续（乘以 h 取极限）；引出逆命题真假之问。', '尝试举反例。', '逻辑推理素养；单向箭头意识。'),
 ('48–50′ (2)', 'P20', '反例一', '出示 y=|x|（图9）：左导数 −1、右导数 +1，不相等故不可导，但连续。给出判定口诀。', '齐答左右导数；在方格纸上画尖点。', '突破难点③；给出可操作判据。'),
 ('50–52′ (2)', 'P21', '反例二', '出示 x·sin(1/x) 与 x²·sin(1/x)（图10）：前者 sin(1/h) 震荡无极限，后者被 |h| 夹逼归零。', '观察包络线；说出“多一个 x 因子就驯服了震荡”。', '深化对极限存在性的理解；夹逼思想。'),
 ('52–54′ (2)', 'P22', '变式2', '分段函数 x²/(ax+b) 在 x=1 可导求 a,b。强调先连续后可导的顺序。', '独立完成得 a=2, b=−1；一人上台板演。', '综合运用；预警易错点③。'),
 ('54–55′ (1)', 'P23', '过场', '进入第 5 幕：应用与升华。', '', ''),
 ('55–56′ (1)', 'P24', '例题3', '实际背景：s(t)=2t³−9t²+12t，求 v(t) 与换向点 t=1、t=2（图12）。追问平均速度能否看出换向。', '读图回答 v<0 的时间段。', '回扣情境，凸显导数“逐点刻画”的威力。'),
 ('56–58′ (2)', 'P25–26', '认知冲突②', '出示 Weierstrass 函数三级放大图（图14）：处处连续、处处不可导，1872 年震动数学界。引申至分形、布朗运动。', '震撼、发问；讨论“直觉为什么会错”。', '情感态度高潮；数学美育；严谨性教育。'),
 ('58–59′ (1)', 'P27–28', '小结', '知识结构图（图15）串联主线；宣读方法清单与常见错误 TOP3。', '合书复述导数的三重身份。', '结构化沉淀，形成认知图式。'),
 ('59–60′ (1)', 'P29–30', '作业与板书', '布置 A/B/C 三层作业；回看板书主干，指出保留区。', '记录作业；拍照板书。', '分层落实，因材施教。'),
]
tb = doc.add_table(rows=1, cols=6)
tb.style = 'Table Grid'
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tb.rows[0].cells
for j, c in enumerate(COLS):
    p = hdr[j].paragraphs[0]; run = p.add_run(c)
    run.font.name = 'Noto Serif CJK SC'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    run.font.size = DPt(9.5); run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row in plan:
    cells = tb.add_row().cells
    for j, c in enumerate(row):
        p = cells[j].paragraphs[0]; run = p.add_run(c)
        run.font.name = 'Noto Serif CJK SC'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        run.font.size = DPt(8.5)
        p.paragraph_format.space_after = DPt(2)
widths = [Cm(2.0), Cm(1.6), Cm(1.6), Cm(5.4), Cm(4.2), Cm(3.4)]
for r_ in tb.rows:
    for j, w_ in enumerate(widths):
        r_.cells[j].width = w_
PP('')
PP('时间校验：1+2+1+2+3+3+1+4+4+4+2+1+3+4+3+4+3+1+2+2+2+2+1+1+2+1+1 = 60 分钟。', bold=True)

H('九、板书设计', 1)
bb = doc.add_table(rows=2, cols=3); bb.style = 'Table Grid'
bd = [('左板（主干·保留）', '右板（例题·板演）', '副板（草稿·随擦）'),
      ('课题：导数的概念与几何意义\n一、平均变化率 Δy/Δx = 割线斜率\n二、导数定义★\n　f′(x₀)=lim(Δx→0) [f(x₀+Δx)−f(x₀)]/Δx\n三、几何：切线斜率\n　分析：f(x₀+h)=f(x₀)+f′(x₀)h+o(h)\n四、可导 ⇒ 连续（反之不真）',
       '例1 定义求导（三步法完整保留）\n　作差 → 作商 → 取极限\n例2 切线方程\n　切点 → 斜率 f′(x₀) → 点斜式\n例3 瞬时速度\n　v(t)=6t²−18t+12，换向点 t=1,2',
       '反例草图：|x| 尖点\n　　　　　x·sin(1/x) 震荡\n数值表：2.5→2.25→2.1→2.01→2.001\n红字提醒：\n　① 先约 Δx！\n　② 查连续！')]
for i, rw in enumerate(bd):
    for j, c in enumerate(rw):
        cell = bb.cell(i, j)
        p = cell.paragraphs[0]; run = p.add_run(c)
        run.font.name = 'Noto Serif CJK SC'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        run.font.size = DPt(9); run.font.bold = (i == 0)

H('十、分层作业', 1)
PP('A 层（必做，全体，约 20 分钟）', bold=True)
for x_ in ['1. 用定义求 f(x)=3x²−1 在 x₀=2 处的导数。',
           '2. 求曲线 y=1/x 在点 (2, 1/2) 处的切线方程。',
           '3. 判断 f(x)=|x−1| 在 x=1 处是否可导，写出左右导数并说明理由。']:
    PP('　' + x_)
PP('B 层（选做，学有余力，约 15 分钟）', bold=True)
for x_ in ['4. 求过点 (1, 0) 且与曲线 y=x² 相切的所有直线方程。（提示：设切点，注意可能不止一条）',
           '5. 设 f(x)=x²（x≤2），=ax+b（x>2），若 f 在 x=2 处可导，求 a、b。']:
    PP('　' + x_)
PP('C 层（挑战，探究性，不限时）', bold=True)
for x_ in ['6. 证明：若 f 在 x₀ 处可导且 f(x₀)=0，则 g(x)=|f(x)| 在 x₀ 处可导的充要条件是 f′(x₀)=0。',
           '7. 用计算机（Python/GeoGebra）作 W(x)=Σ_{n=0}^{5} 2⁻ⁿcos(7ⁿπx) 的图像，逐级放大并撰写 200 字观察报告：它与 y=x² 的图像在放大过程中有何本质不同？']:
    PP('　' + x_)

H('十一、教学反思（课后填写）', 1)
for x_ in ['1. 数值探究环节学生是否顺利得出双侧逼近同一值？用时是否超出 3 分钟？',
           '2. “局部线性化”与 o(h) 的接受度如何？是否需要下节课再复习？',
           '3. Weierstrass 函数的震撼效果是否达到预期？学生提出了哪些问题？',
           '4. 例题板演中出现的典型错误：____________________________________',
           '5. 时间控制偏差：____ 分钟，需调整的环节：____________________',
           '6. 改进设想：__________________________________________________']:
    PP('　' + x_)
PP('\n\n\n')

dpath = os.path.join(OUT, '教案_01_导数的概念与几何意义.docx')
doc.save(dpath)
print('docx:', dpath)
