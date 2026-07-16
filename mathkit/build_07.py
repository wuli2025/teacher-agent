# -*- coding: utf-8 -*-
"""第07讲：复数、复平面与欧拉公式 —— 60分钟课件包生成脚本"""
import sys, os, math
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Arc, FancyArrowPatch, RegularPolygon

OUT = '/mnt/c/Users/mi/Desktop/数学名师课件包/07_复数与复平面'
FIG = os.path.join(OUT, 'figures')
TMP = os.path.join(FIG, '_tex')
os.makedirs(FIG, exist_ok=True); os.makedirs(TMP, exist_ok=True)
F = lambda n: os.path.join(FIG, n)
T = lambda n: os.path.join(TMP, n)

# ============================================================ 图 1..14
def arrow(ax, x0, y0, x1, y1, c, lw=2.2, ls='-', z=3):
    ax.add_patch(FancyArrowPatch((x0, y0), (x1, y1), arrowstyle='-|>',
                mutation_scale=16, color=c, lw=lw, linestyle=ls, zorder=z,
                shrinkA=0, shrinkB=0))

# fig1 卡尔达诺：x^3=15x+4 的图像，三个实根，但公式途经虚数
fig, ax = k.new_fig(5.6, 3.4)
x = np.linspace(-4.6, 4.6, 600)
ax.plot(x, x**3 - 15*x - 4, color=k.M_ACC2, lw=2.6, zorder=3,
        label='$y=x^{3}-15x-4$')
k.style_axes(ax, xlabel='x', ylabel='')
for r in [4.0, -0.26794919, -3.73205081]:
    ax.plot(r, 0, 'o', color=k.M_RED, ms=10, zorder=5)
ax.annotate('$x=4$', xy=(4, 0), xytext=(1.5, 26),
            color=k.M_RED, fontsize=20,
            arrowprops=dict(arrowstyle='->', color=k.M_RED, lw=2))
ax.set_ylim(-62, 96); ax.set_xlim(-5.2, 5.4)
ax.set_yticks([])
ax.legend(loc='upper center', frameon=False, fontsize=20, handlelength=1.4)
ax.set_title('三个实根，却绕不开虚数', fontsize=22, color=k.M_INK, pad=14)
k.save_fig(fig, F('01_cardano.png'))

# fig2 复平面：点 z=a+bi 的表示
fig, ax = k.new_fig(5.2, 3.2)
k.style_axes(ax, xlabel='Re', ylabel='Im')
a, b = 3, 2
arrow(ax, 0, 0, a, b, k.M_ACC)
ax.plot([a, a], [0, b], ls=':', color=k.M_SLATE, lw=1.6)
ax.plot([0, a], [b, b], ls=':', color=k.M_SLATE, lw=1.6)
ax.plot(a, b, 'o', color=k.M_ACC, ms=11, zorder=5)
ax.text(a + .18, b + .1, r'$z=a+bi$', fontsize=20, color=k.M_INK)
ax.text(a / 2 - .35, -.85, '$a$', fontsize=20, color=k.M_ACC2)
ax.text(a + .15, b / 2 - .2, '$b$', fontsize=20, color=k.M_ACC2)
ax.text(0.35, 2.95, r'$|z|=\sqrt{a^{2}+b^{2}}$', fontsize=20, color=k.M_ACC)
ax.set_xlim(-1.3, 5.6); ax.set_ylim(-1.3, 3.9)
ax.set_xticks([2, 4]); ax.set_yticks([2])
ax.set_title('复数 ↔ 点 ↔ 向量', fontsize=22, color=k.M_INK, pad=30)
k.save_fig(fig, F('02_plane.png'))

# fig3 向量加法（平行四边形）与减法
fig, axes = plt.subplots(1, 2, figsize=(9.6, 3.5))
z1 = (3, 1); z2 = (1, 2.5)
ax = axes[0]
k.style_axes(ax, xlabel='Re', ylabel='Im')
arrow(ax, 0, 0, *z1, k.M_ACC2); arrow(ax, 0, 0, *z2, k.M_GRN)
arrow(ax, 0, 0, z1[0] + z2[0], z1[1] + z2[1], k.M_RED, lw=2.8)
ax.plot([z1[0], z1[0] + z2[0]], [z1[1], z1[1] + z2[1]], ls='--', color=k.M_SLATE, lw=1.4)
ax.plot([z2[0], z1[0] + z2[0]], [z2[1], z1[1] + z2[1]], ls='--', color=k.M_SLATE, lw=1.4)
ax.text(3.2, .5, r'$z_1$', fontsize=20, color=k.M_ACC2)
ax.text(.35, 2.6, r'$z_2$', fontsize=20, color=k.M_GRN)
ax.text(3.6, 3.8, r'$z_1+z_2$', fontsize=20, color=k.M_RED)
ax.set_xlim(-1, 6.0); ax.set_ylim(-1, 5.0)
ax.set_xticks([2, 4]); ax.set_yticks([2, 4])
ax.set_title('加法 = 平行四边形', fontsize=21, color=k.M_INK, pad=8)
ax = axes[1]
k.style_axes(ax, xlabel='Re', ylabel='Im')
arrow(ax, 0, 0, *z1, k.M_ACC2); arrow(ax, 0, 0, *z2, k.M_GRN)
arrow(ax, z2[0], z2[1], z1[0], z1[1], k.M_RED, lw=2.8)
arrow(ax, 0, 0, z1[0] - z2[0], z1[1] - z2[1], k.M_RED, lw=2.0, ls='--')
ax.text(3.2, .5, r'$z_1$', fontsize=20, color=k.M_ACC2)
ax.text(.3, 2.7, r'$z_2$', fontsize=20, color=k.M_GRN)
ax.text(2.3, 2.3, r'$z_1-z_2$', fontsize=20, color=k.M_RED)
ax.set_xlim(-1, 6.0); ax.set_ylim(-2.1, 5.0)
ax.set_xticks([2, 4]); ax.set_yticks([2, 4])
ax.set_title('减法 = 两点连线', fontsize=21, color=k.M_INK, pad=8)
k.save_fig(fig, F('03_add_sub.png'))

# fig4 模与辐角
fig, ax = k.new_fig(5.7, 3.5)
k.style_axes(ax, xlabel='Re', ylabel='Im')
r, th = 3.2, math.radians(38)
a, b = r * math.cos(th), r * math.sin(th)
arrow(ax, 0, 0, a, b, k.M_ACC, lw=2.8)
ax.add_patch(Arc((0, 0), 2.0, 2.0, theta1=0, theta2=38, color=k.M_RED, lw=2.2))
ax.text(1.18, .32, r'$\theta$', fontsize=22, color=k.M_RED)
ax.text(0.95, 1.35, r'$r$', fontsize=22, color=k.M_ACC)   # 单字符，避免大字号横向压住向量
tt = np.linspace(0, 2 * np.pi, 300)
ax.plot(r * np.cos(tt), r * np.sin(tt), ls=':', color=k.M_RULE, lw=1.6)
ax.plot([a, a], [0, b], ls=':', color=k.M_SLATE)
ax.plot(a, b, 'o', color=k.M_ACC, ms=11, zorder=5)
ax.text(a + .18, b + .2, r'$z$', fontsize=22, color=k.M_INK)
ax.text(-4.6, 3.1, r'$z=r(\cos\theta+i\sin\theta)$', fontsize=20, color=k.M_INK)
ax.set_xlim(-4.8, 4.8); ax.set_ylim(-3.6, 4.0); ax.set_aspect('equal')
ax.set_xticks([-2, 2]); ax.set_yticks([-2, 2])
ax.set_title('模 $r$ 与辐角 $\\theta$', fontsize=22, color=k.M_INK, pad=30)
k.save_fig(fig, F('04_modarg.png'))

# fig5 乘以 i = 逆时针转 90°，连续四次
fig, ax = k.new_fig(5.2, 3.4)
k.style_axes(ax, xlabel='', ylabel='')
z0 = complex(2.6, 1.1)
cols = [k.M_ACC, k.M_ACC2, k.M_GRN, k.M_RED]
labs = [r'$z$', r'$iz$', r'$i^{2}z=-z$', r'$i^{3}z$']
pts = [z0 * (1j ** n) for n in range(4)]
for p, c, L in zip(pts, cols, labs):
    arrow(ax, 0, 0, p.real, p.imag, c, lw=2.6)
    ax.plot(p.real, p.imag, 'o', color=c, ms=9, zorder=5)
    ang = np.angle(p)
    R0 = abs(p) + 0.65
    ax.text(R0 * math.cos(ang), R0 * math.sin(ang), L, fontsize=20, color=c,
            ha='center', va='center')
R = abs(z0)
tt = np.linspace(0, 2 * np.pi, 400)
ax.plot(R * np.cos(tt), R * np.sin(tt), ls='--', color=k.M_RULE, lw=1.6)
for n in range(4):
    a0 = math.degrees(np.angle(pts[n])); a1 = a0 + 90
    ax.add_patch(Arc((0, 0), 1.5 + .25 * n, 1.5 + .25 * n, theta1=a0, theta2=a1,
                     color=cols[n], lw=1.8, ls=':'))
ax.set_xlim(-5.8, 5.8); ax.set_ylim(-3.8, 3.8); ax.set_aspect('equal')
ax.set_xticks([]); ax.set_yticks([])
ax.set_title('乘以 $i$ = 逆时针转 $90^{\\circ}$', fontsize=22, color=k.M_INK, pad=8)
k.save_fig(fig, F('05_times_i.png'))

# fig6 一般乘法 = 旋转 + 伸缩
fig, ax = k.new_fig(5.2, 3.4)
k.style_axes(ax, xlabel='', ylabel='')
z1 = 2.2 * np.exp(1j * math.radians(20))
z2 = 1.5 * np.exp(1j * math.radians(50))
z3 = z1 * z2
for z, c, L in [(z1, k.M_ACC2, r'$z_1=r_1(\cos\theta_1+i\sin\theta_1)$'),
                (z2, k.M_GRN, r'$z_2=r_2(\cos\theta_2+i\sin\theta_2)$'),
                (z3, k.M_RED, r'$z_1z_2=r_1r_2[\cos(\theta_1{+}\theta_2)+i\sin(\theta_1{+}\theta_2)]$')]:
    arrow(ax, 0, 0, z.real, z.imag, c, lw=2.6)
    ax.plot(z.real, z.imag, 'o', color=c, ms=8, zorder=5)
ax.text(2.35, .55, r'$z_1$', fontsize=20, color=k.M_ACC2)
ax.text(1.05, 1.35, r'$z_2$', fontsize=20, color=k.M_GRN)
ax.text(1.3, 3.35, r'$z_1z_2$', fontsize=20, color=k.M_RED)
ax.add_patch(Arc((0, 0), 1.2, 1.2, theta1=0, theta2=20, color=k.M_ACC2, lw=2))
ax.add_patch(Arc((0, 0), 1.9, 1.9, theta1=20, theta2=70, color=k.M_RED, lw=2, ls='--'))
ax.text(1.35, .05, r'$\theta_1$', fontsize=20, color=k.M_ACC2)
ax.text(3.0, 2.9, r'$\theta_1{+}\theta_2$', fontsize=20, color=k.M_RED)
ax.set_xlim(-2.0, 6.3); ax.set_ylim(-1.2, 4.2); ax.set_aspect('equal')
ax.set_xticks([2, 4]); ax.set_yticks([2])
ax.set_title('模相乘，辐角相加', fontsize=22, color=k.M_INK, pad=8)
k.save_fig(fig, F('06_mul_geo.png'))

# fig7 共轭对称
fig, ax = k.new_fig(6.0, 3.6)
k.style_axes(ax, xlabel='Re', ylabel='Im')
a, b = 2.8, 1.9
arrow(ax, 0, 0, a, b, k.M_ACC); arrow(ax, 0, 0, a, -b, k.M_ACC2)
ax.plot([a, a], [b, -b], ls=':', color=k.M_SLATE, lw=1.6)
ax.plot(a, b, 'o', color=k.M_ACC, ms=9, zorder=5)
ax.plot(a, -b, 'o', color=k.M_ACC2, ms=9, zorder=5)
ax.text(a + .2, b + .25, r'$z=a+bi$', fontsize=20, color=k.M_ACC)
ax.text(a + .2, -b - .6, r'$\bar z=a-bi$', fontsize=20, color=k.M_ACC2)
ax.axhline(0, color=k.M_RED, lw=1.0, ls='--', alpha=.6)
ax.text(-4.5, 2.0, r'$z\bar z=|z|^{2}$', fontsize=20, color=k.M_INK)
ax.set_xlim(-4.8, 4.8); ax.set_ylim(-2.85, 2.85); ax.set_aspect('equal')
ax.set_xticks([-2, 2]); ax.set_yticks([-2, 2])
ax.set_title('共轭：关于实轴对称', fontsize=22, color=k.M_INK, pad=30)
k.save_fig(fig, F('07_conj.png'))

# fig8 除法 = 反向旋转 + 缩短
fig, ax = k.new_fig(5.0, 3.3)
k.style_axes(ax, xlabel='', ylabel='')
z1 = 3.0 * np.exp(1j * math.radians(70)); z2 = 1.5 * np.exp(1j * math.radians(25))
z3 = z1 / z2
for z, c, L, dx in [(z1, k.M_ACC2, r'$z_1$', .2), (z2, k.M_GRN, r'$z_2$', .35),
                    (z3, k.M_RED, r'$z_1/z_2$', .3)]:
    arrow(ax, 0, 0, z.real, z.imag, c, lw=2.6)
    ax.plot(z.real, z.imag, 'o', color=c, ms=9, zorder=5)
    ax.text(z.real + dx, z.imag + .12, L, fontsize=20, color=c)
ax.add_patch(Arc((0, 0), 2.4, 2.4, theta1=45, theta2=70, color=k.M_RED, lw=2, ls='--'))
ax.set_xlim(-1.6, 6.0); ax.set_ylim(-1.0, 4.0); ax.set_aspect('equal')
ax.set_xticks([2, 4]); ax.set_yticks([2])
ax.set_title('除法：模相除，辐角相减', fontsize=22, color=k.M_INK, pad=8)
k.save_fig(fig, F('08_div.png'))

# fig9 棣莫弗：z^n 螺旋（r=1.15, θ=30°）
fig, ax = k.new_fig(5.8, 3.5)
k.style_axes(ax, xlabel='', ylabel='')
z = 1.18 * np.exp(1j * math.radians(30))
pts = [z ** n for n in range(0, 8)]
xs = [p.real for p in pts]; ys = [p.imag for p in pts]
ax.plot(xs, ys, '-', color=k.M_RULE, lw=2.0, zorder=2)
for n, p in enumerate(pts):
    arrow(ax, 0, 0, p.real, p.imag, k.M_ACC2 if n % 2 else k.M_ACC, lw=1.6)
    ax.plot(p.real, p.imag, 'o', color=k.M_RED, ms=8, zorder=6)
for n, dang in [(0, -22), (1, 0), (2, 0), (4, 0), (7, 0)]:   # z^0 沿径向下偏，避开实轴
    p = pts[n]
    ang = np.angle(p) + math.radians(dang); R0 = abs(p) + 0.7
    ax.text(R0 * math.cos(ang), R0 * math.sin(ang), f'$z^{{{n}}}$', fontsize=20,
            color=k.M_INK, ha='center', va='center')
ax.set_xlim(-4.6, 3.8); ax.set_ylim(-2.52, 2.52); ax.set_aspect('equal')
ax.set_xticks([]); ax.set_yticks([])
ax.set_title('乘方：模乘方，辐角倍增', fontsize=22, color=k.M_INK, pad=8)
k.save_fig(fig, F('09_demoivre.png'))

# fig10 n 次单位根：n=3,5,8 三子图正 n 边形
fig, axes = plt.subplots(1, 3, figsize=(5.9, 2.6))
for ax, n in zip(axes, [3, 5, 8]):
    tt = np.linspace(0, 2 * np.pi, 400)
    ax.plot(np.cos(tt), np.sin(tt), color=k.M_RULE, lw=1.8)
    roots = [np.exp(2j * np.pi * j / n) for j in range(n)]
    poly = roots + [roots[0]]
    ax.plot([p.real for p in poly], [p.imag for p in poly], color=k.M_ACC2, lw=2.0, zorder=3)
    show = range(n) if n <= 5 else (0, 1)      # n=8 顶点太密，只标 ω⁰、ω¹
    for j, p in enumerate(roots):
        arrow(ax, 0, 0, p.real, p.imag, k.M_ACC, lw=1.4)
        ax.plot(p.real, p.imag, 'o', color=k.M_RED, ms=9, zorder=6)
        if j in show:
            ax.text(p.real * 1.42, p.imag * 1.42, f'$\\omega^{{{j}}}$', fontsize=20,
                    color=k.M_INK, ha='center', va='center')
    ax.set_aspect('equal'); ax.set_xlim(-1.85, 1.85); ax.set_ylim(-1.85, 1.85)
    ax.set_xticks([]); ax.set_yticks([])
    for s in ax.spines.values():
        s.set_visible(False)
    ax.axhline(0, color=k.M_RULE, lw=.8); ax.axvline(0, color=k.M_RULE, lw=.8)
    ax.set_title(f'$n={n}$', fontsize=20, color=k.M_INK, pad=2)
fig.suptitle('单位根：把圆 $n$ 等分', fontsize=21, color=k.M_ACC)
k.save_fig(fig, F('10_roots.png'))

# fig11 z^4=1 详解
fig, ax = k.new_fig(5.1, 3.3)
tt = np.linspace(0, 2 * np.pi, 400)
ax.plot(np.cos(tt), np.sin(tt), color=k.M_RULE, lw=2.0)
k.style_axes(ax, xlabel='', ylabel='')
rs = [1, 1j, -1, -1j]; Ls = [r'$z_0=1$', r'$z_1=i$', r'$z_2=-1$', r'$z_3=-i$']
tx = [(1.95, 0.45), (0.75, 1.42), (-1.95, 0.45), (-0.85, -1.42)]   # 标签沿径向外移，避开实轴
poly = rs + [rs[0]]
ax.plot([p.real for p in poly], [p.imag for p in poly], color=k.M_ACC2, lw=2.2, zorder=3)
for p, L, (lx, ly) in zip(rs, Ls, tx):
    p = complex(p)
    arrow(ax, 0, 0, p.real, p.imag, k.M_ACC, lw=2.0)
    ax.plot(p.real, p.imag, 'o', color=k.M_RED, ms=11, zorder=6)
    ax.text(lx, ly, L, fontsize=20, color=k.M_INK, ha='center', va='center')
ax.set_aspect('equal'); ax.set_xlim(-2.85, 2.85); ax.set_ylim(-1.84, 1.84)
ax.set_xticks([]); ax.set_yticks([])
ax.set_title('$z^{4}=1$：四根成正方形', fontsize=22, color=k.M_INK, pad=6)
k.save_fig(fig, F('11_z4.png'))

# fig12 欧拉公式：e^{iθ} 沿单位圆运动 + 投影
fig, axes = plt.subplots(1, 2, figsize=(5.5, 3.5), gridspec_kw={'width_ratios': [1, 1.15]})
ax = axes[0]
tt = np.linspace(0, 2 * np.pi, 400)
ax.plot(np.cos(tt), np.sin(tt), color=k.M_RULE, lw=2)
for th, c in [(np.pi / 6, k.M_GRN), (2 * np.pi / 3, k.M_ACC2), (5 * np.pi / 4, k.M_RED)]:
    z = np.exp(1j * th)
    arrow(ax, 0, 0, z.real, z.imag, c, lw=2.2)
    ax.plot([z.real, z.real], [0, z.imag], ls=':', color=c, lw=1.4)
    ax.plot([0, z.real], [z.imag, z.imag], ls=':', color=c, lw=1.4)
    ax.plot(z.real, z.imag, 'o', color=c, ms=8, zorder=6)
ax.text(.75, 1.15, r'$e^{i\theta}$', fontsize=20, color=k.M_INK)
ax.axhline(0, color=k.M_INK, lw=1.1); ax.axvline(0, color=k.M_INK, lw=1.1)
ax.set_aspect('equal'); ax.set_xlim(-1.6, 1.6); ax.set_ylim(-1.6, 1.6)
ax.set_xticks([]); ax.set_yticks([])
ax.set_title('单位圆上匀速跑', fontsize=20, color=k.M_INK, pad=4)
for s in ax.spines.values():
    s.set_visible(False)
ax = axes[1]
t = np.linspace(0, 2 * np.pi, 500)
ax.plot(t, np.cos(t), color=k.M_ACC, lw=2.6, label=r'$\cos\theta$')
ax.plot(t, np.sin(t), color=k.M_ACC2, lw=2.6, label=r'$\sin\theta$')
ax.axhline(0, color=k.M_SLATE, lw=1)
ax.set_xticks([0, np.pi, 2 * np.pi])
ax.set_xticklabels(['0', r'$\pi$', r'$2\pi$'])
ax.set_yticks([-1, 1])
ax.set_ylim(-1.35, 1.35)
ax.legend(fontsize=20, frameon=False, loc='upper center',
          bbox_to_anchor=(0.5, -0.18), ncol=2, handlelength=1.2, columnspacing=1.2)
ax.grid(True, color=k.M_RULE, lw=.7)
for s in ('top', 'right'):
    ax.spines[s].set_visible(False)
ax.set_title('两个投影', fontsize=20, color=k.M_INK, pad=4)
# 公式已由 PPT 的 formula() 大字给出，此处不再重复 suptitle，把高度让给左图的圆
k.save_fig(fig, F('12_euler.png'))

# fig13 e^{iπ}+1=0 图解
fig, ax = k.new_fig(5.2, 3.4)
tt = np.linspace(0, np.pi, 300)
ax.plot(np.cos(tt), np.sin(tt), color=k.M_ACC, lw=3, zorder=3)
tt2 = np.linspace(np.pi, 2 * np.pi, 300)
ax.plot(np.cos(tt2), np.sin(tt2), color=k.M_RULE, lw=1.6, ls='--')
arrow(ax, 0, 0, 1, 0, k.M_GRN, lw=2.4)
arrow(ax, 0, 0, -1, 0, k.M_RED, lw=2.4)
ax.plot(1, 0, 'o', color=k.M_GRN, ms=11, zorder=6)
ax.plot(-1, 0, 'o', color=k.M_RED, ms=11, zorder=6)
ax.text(1.15, -.42, r'$1$', fontsize=20, color=k.M_GRN)
ax.text(-1.45, -.42, r'$-1$', fontsize=20, color=k.M_RED)
ax.annotate('', xy=(-0.72, 0.70), xytext=(0.72, 0.70),
            arrowprops=dict(arrowstyle='->', color=k.M_ACC, lw=2.4,
                            connectionstyle='arc3,rad=-0.35'))
ax.text(0.0, 1.28, r'转过 $\pi$', fontsize=20, color=k.M_ACC, ha='center')
ax.axhline(0, color=k.M_INK, lw=1.1); ax.axvline(0, color=k.M_INK, lw=1.1)
ax.text(0.0, -1.55, r'$e^{i\pi}+1=0$', fontsize=26, color=k.M_ACC, ha='center', va='center',
        bbox=dict(boxstyle='round,pad=0.35', fc='white', ec=k.M_ACC, lw=2))
ax.set_aspect('equal'); ax.set_xlim(-3.0, 3.0); ax.set_ylim(-2.05, 1.87)
ax.set_xticks([]); ax.set_yticks([])
for s in ax.spines.values():
    s.set_visible(False)
ax.set_title('走半圈：$e^{i\\pi}=-1$', fontsize=22, color=k.M_INK, pad=6)
k.save_fig(fig, F('13_eipi.png'))

# fig14 实系数多项式虚根成对
fig, axes = plt.subplots(1, 2, figsize=(6.5, 3.4))
ax = axes[0]
x = np.linspace(-2.4, 2.4, 400)
ax.plot(x, x**4 - 1, color=k.M_ACC2, lw=2.5)
ax.axhline(0, color=k.M_INK, lw=1.1); ax.axvline(0, color=k.M_INK, lw=1.1)
ax.plot([-1, 1], [0, 0], 'o', color=k.M_RED, ms=10, zorder=5)
ax.set_ylim(-3, 9); ax.set_yticks([0, 5]); ax.set_xticks([-2, 0, 2])
ax.grid(True, color=k.M_RULE, lw=.7)
for s in ('top', 'right'):
    ax.spines[s].set_visible(False)
ax.set_title('实轴：只见 2 根', fontsize=20, color=k.M_INK, pad=4)
ax = axes[1]
tt = np.linspace(0, 2 * np.pi, 300)
ax.plot(np.cos(tt), np.sin(tt), color=k.M_RULE, lw=1.6, ls=':')
for p, c, L, lx, ly in [(1, k.M_RED, '$1$', 1.35, -.38), (-1, k.M_RED, '$-1$', -1.42, -.38),
                        (1j, k.M_ACC, '$i$', .35, 1.35), (-1j, k.M_ACC, '$-i$', .45, -1.35)]:
    p = complex(p)
    ax.plot(p.real, p.imag, 'o', color=c, ms=12, zorder=6)
    ax.text(lx, ly, L, fontsize=20, color=c, ha='center', va='center')
ax.plot([0, 0], [1, -1], ls='--', color=k.M_ACC, lw=1.8)
ax.axhline(0, color=k.M_INK, lw=1.1); ax.axvline(0, color=k.M_INK, lw=1.1)
ax.set_aspect('equal'); ax.set_xlim(-1.85, 1.85); ax.set_ylim(-1.75, 1.75)
ax.set_xticks([]); ax.set_yticks([])
for s in ax.spines.values():
    s.set_visible(False)
ax.set_title('复平面：4 根俱在', fontsize=20, color=k.M_INK, pad=4)
fig.suptitle('虚根必共轭成对', fontsize=21, color=k.M_ACC)
k.save_fig(fig, F('14_conj_roots.png'))

print('figures done: 14')

# ============================================================ PPT
p = k.new_deck()
S = []

# 1 封面
S.append(k.title_slide(p, '第07讲　复数、复平面与欧拉公式',
    '从卡尔达诺的困境，到 $e^{i\\pi}+1=0$ 的震撼'.replace('$', '').replace('\\pi', 'π').replace('^{i}', ''),
    '数学名师课件包', '60 分钟 · 高中数学必修/选修衔接'))

# 2 学习目标
s = k.content_slide(p, '学习目标', '导入')
k.bullets(s, [
    '知道复数不是"硬造"的：三次方程求实根时必须途经虚数',
    '掌握复数四种形式：代数 / 几何(点·向量) / 三角 / 指数',
    ('a+bi ↔ 点(a,b) ↔ 向量 OZ ↔ r(cosθ+isinθ) ↔ re^{iθ}', 1),
    '理解复数乘法的几何意义：模相乘、辐角相加 = 旋转 + 伸缩（核心）',
    '会用棣莫弗定理求乘方，会求 n 次单位根并画出正 n 边形',
    '感受欧拉公式：复数是代数与几何、指数与三角的桥梁',
], y=1.7, w=7.0)
k.picture(s, F('13_eipi.png'), x=7.6, y=1.8, w=5.2)
k.callout(s, '本讲主线：一个「乘法」，串起旋转、单位根、欧拉公式。', x=0.85, y=6.05, w=6.6, h=0.85)
S.append(s)

# 3 幕一
S.append(k.section_slide(p, '第 1 幕 · 情境与动机', '虚数是被"逼"出来的', '0–8 min'))

# 4 卡尔达诺故事
s = k.content_slide(p, '一个历史真相：不是为了解 x²+1=0', '8 min')
k.bullets(s, [
    '流行说法："因为 x²+1=0 无解，所以造个 i"——这站不住脚。',
    ('无解就无解，古人从不为无解的方程发明新数。', 1),
    '真正的逼迫来自三次方程。卡尔达诺公式解 x³=15x+4：',
    ('公式给出 x = ∛(2+√-121) + ∛(2-√-121)', 1),
    ('可是画个图就看到：x=4 明明是实根！', 1),
    '要拿到明明白白的实根 4，中途必须穿过 √-121 这片"虚"的地带。',
], y=1.65, w=6.6)
k.picture(s, F('01_cardano.png'), x=7.2, y=1.6, w=5.6)
k.callout(s, '虚数是"通往实数答案的必经桥梁"——邦贝利(1572)硬着头皮算下去，才发现两个立方根加起来正好是 4。',
          x=0.85, y=5.75, w=6.4, h=1.35, kind='note')
S.append(s)

# 5 定义
s = k.content_slide(p, '复数的定义与代数形式', '概念')
k.formula(s, r'$i^{2}=-1,\qquad z=a+bi\ (a,b\in\mathbb{R})$', x=1.4, y=1.55, w=10.5, out=T('f01.png'))
k.bullets(s, [
    'a = Re z 实部，b = Im z 虚部',
    'b = 0 ⇒ 实数；b ≠ 0 ⇒ 虚数；a = 0 且 b ≠ 0 ⇒ 纯虚数',
    '相等：a+bi = c+di ⟺ a=c 且 b=d（一个复数方程 = 两个实方程）',
    'i 的周期性：i¹=i, i²=−1, i³=−i, i⁴=1（四次一循环）',
], y=3.05, w=7.0, size=18)
k.picture(s, F('02_plane.png'), x=7.6, y=3.0, w=5.2)
k.callout(s, '警示：复数无大小顺序，"z₁>z₂"没有意义（除非都是实数）。', x=0.85, y=6.3, w=6.6, h=0.75, kind='warn')
S.append(s)

# 6 幕二
S.append(k.section_slide(p, '第 2 幕 · 几何化', '复平面：让复数"看得见"', '8–18 min'))

# 7 复平面
s = k.content_slide(p, '复平面：三位一体', '几何形式')
k.bullets(s, [
    '横轴 = 实轴，纵轴 = 虚轴（除原点外，虚轴上的点都是纯虚数）',
    '复数 z = a+bi  ⟷  点 Z(a, b)  ⟷  向量 OZ',
    '模 |z| = √(a²+b²) = 向量长度 = 点到原点的距离',
    '这一步是"数"到"形"的翻译，复数的全部威力由此展开。',
], y=1.7, w=6.4)
k.picture(s, F('02_plane.png'), x=7.2, y=1.6, w=5.6)
k.callout(s, '记住：复数不是"一个怪数"，它是平面上的一个箭头。', x=0.85, y=5.6, w=6.2, h=0.8)
S.append(s)

# 8 加减法
s = k.content_slide(p, '加减法：向量的平行四边形法则', '运算')
k.formula(s, r'$(a+bi)\pm(c+di)=(a\pm c)+(b\pm d)i$', x=1.6, y=1.5, w=10.0, size=0.9, out=T('f02.png'))
k.full_picture(s, F('03_add_sub.png'), y=2.7, w=9.6)
k.callout(s, '|z₁−z₂| = 复平面上两点间的距离 —— 后面所有"轨迹问题"的钥匙。',
          x=1.9, y=6.55, w=9.5, h=0.7, kind='note')
S.append(s)

# 9 模与辐角
s = k.content_slide(p, '模与辐角 ⇒ 三角形式', '概念')
k.bullets(s, [
    '模 r = |z| ≥ 0；辐角 θ：向量与实轴正向的夹角',
    '辐角不唯一，相差 2kπ；主辐角 arg z ∈ [0, 2π)',
    'a = r cosθ，b = r sinθ',
    '三角形式：z = r(cosθ + i sinθ)',
    ('求 θ 必须先看 z 落在哪个象限，不能只用 arctan(b/a)。', 1),
], y=1.7, w=6.2)
k.picture(s, F('04_modarg.png'), x=7.2, y=1.6, w=5.7)
k.callout(s, '例：z = −1 + √3 i ⇒ r = 2，θ = 2π/3，z = 2(cos120° + i sin120°)',
          x=0.85, y=5.85, w=6.1, h=0.85)
S.append(s)

# 10 幕三
S.append(k.section_slide(p, '第 3 幕 · 本讲核心', '乘法 = 旋转 + 伸缩', '18–32 min'))

# 11 乘以 i
s = k.content_slide(p, '先看最简单的：乘以 i 会发生什么？', '探究')
k.picture(s, F('05_times_i.png'), x=6.6, y=1.5, w=6.3)
k.bullets(s, [
    '取 z = 2.6 + 1.1i，逐次乘 i：',
    ('iz = −1.1 + 2.6i；i²z = −z；i³z；i⁴z = z', 1),
    '观察：长度一点没变，每次整整转了 90°。',
    '转四次 = 转 360° = 回到原点 ⇒ i⁴ = 1。',
    '结论：乘以 i ≡ 绕原点逆时针旋转 90° 的一次操作。',
], y=1.75, w=6.0, size=17)
k.callout(s, '一个"数"竟然代表一个"动作"——这是复数最深刻的转身。', x=0.85, y=5.9, w=5.8, h=0.9, kind='key')
S.append(s)

# 12 一般乘法推演
s = k.content_slide(p, '推演：三角形式下的乘法', '推导')
k.formula(s, r'$z_1z_2=r_1r_2(\cos\theta_1+i\sin\theta_1)(\cos\theta_2+i\sin\theta_2)$',
          x=0.9, y=1.5, w=11.4, size=0.82, out=T('f03.png'))
k.formula(s, r'$=r_1r_2[(\cos\theta_1\cos\theta_2-\sin\theta_1\sin\theta_2)+i(\sin\theta_1\cos\theta_2+\cos\theta_1\sin\theta_2)]$',
          x=0.9, y=2.55, w=11.4, size=0.72, out=T('f04.png'))
k.formula(s, r'$\Rightarrow\ z_1z_2=r_1r_2[\cos(\theta_1+\theta_2)+i\sin(\theta_1+\theta_2)]$',
          x=0.9, y=3.65, w=11.4, size=0.85, out=T('f05.png'))
k.callout(s, '两角和公式，居然是复数乘法的"副产品"。\n模相乘，辐角相加 —— 八个字，定住本讲。',
          x=3.2, y=5.0, w=7.0, h=1.3, kind='key')
S.append(s)

# 13 乘法几何图
s = k.content_slide(p, '乘法的几何意义：一个变换', '核心')
k.full_picture(s, F('06_mul_geo.png'), y=1.45, w=7.0)
k.callout(s, '把 z₂ 看作"算子"：z ↦ z₂·z 就是把整个平面绕原点旋转 θ₂、再放大 r₂ 倍。'
             '这正是后续"旋转矩阵""相似变换"的源头。', x=1.9, y=6.35, w=9.5, h=0.9)
S.append(s)

# 14 共轭与除法
s = k.content_slide(p, '共轭与除法', '运算')
k.picture(s, F('07_conj.png'), x=0.7, y=1.55, w=6.0)
k.picture(s, F('08_div.png'), x=6.9, y=1.55, w=6.0)
k.callout(s, '除法口诀：分母实数化（上下同乘分母的共轭）；几何上是"反向旋转 + 按模缩放"。',
          x=1.4, y=5.95, w=10.5, h=0.85, kind='note')
S.append(s)

# 15 例题1
s = k.content_slide(p, '例题 1｜复数运算与化简', '例题')
k.bullets(s, [
    '计算 z = (1 + i)(2 − i) / (3 + i)，并求 |z| 与 arg z。',
    '【板演】分子：(1+i)(2−i) = 2 − i + 2i − i² = 3 + i',
    ('于是 z = (3+i)/(3+i) = 1', 1),
    '【模与辐角】|z| = 1，arg z = 0。',
    '【几何校验】用"模相乘/辐角相加"：',
    ('|1+i|=√2, |2−i|=√5, |3+i|=√10 ⇒ |z| = √2·√5/√10 = 1 ✓', 1),
    ('arg: 45° + (−26.57°) − 18.43° = 0° ✓', 1),
], y=1.65, w=7.4, size=17)
k.picture(s, F('08_div.png'), x=7.9, y=2.1, w=5.0)
k.callout(s, '好习惯：算完用"模与辐角"复核一遍，秒查错。', x=0.85, y=6.35, w=7.0, h=0.7)
S.append(s)

# 16 例题2 旋转
s = k.content_slide(p, '例题 2｜用复数解几何旋转', '例题')
k.bullets(s, [
    '题：把点 A(3, 1) 绕原点逆时针旋转 60°，求像点 A′。',
    '【复数化】A ↔ z = 3 + i；旋转 60° ↔ 乘以 w = cos60° + i sin60° = 1/2 + (√3/2)i',
    "【计算】z′ = z·w = (3+i)(1/2 + (√3/2)i)",
    ('= 3/2 + (3√3/2)i + (1/2)i + (√3/2)i² = (3−√3)/2 + (3√3+1)/2 · i', 1),
    "【答】A′( (3−√3)/2 , (3√3+1)/2 ) ≈ (0.634, 3.098)",
    '【推广】绕定点 z₀ 转 θ：z′ = z₀ + (z − z₀)·(cosθ + i sinθ)',
], y=1.65, w=7.3, size=17)
k.picture(s, F('06_mul_geo.png'), x=7.7, y=2.2, w=5.2)
k.callout(s, '几何题里最麻烦的"旋转"，在复数里退化成一次乘法。', x=0.85, y=6.35, w=7.0, h=0.7, kind='key')
S.append(s)

# 17 幕四
S.append(k.section_slide(p, '第 4 幕 · 乘方与开方', '棣莫弗定理与 n 次单位根', '32–46 min'))

# 18 棣莫弗
s = k.content_slide(p, '棣莫弗定理', '定理')
k.formula(s, r'$[r(\cos\theta+i\sin\theta)]^{n}=r^{n}(\cos n\theta+i\sin n\theta),\quad n\in\mathbb{Z}$',
          x=0.9, y=1.45, w=11.4, size=0.8, out=T('f06.png'))
k.bullets(s, [
    '证明：对乘法法则做 n 次归纳（模相乘 n 次 ⇒ rⁿ；辐角相加 n 次 ⇒ nθ）。',
    '副产品：令 r=1, n=2 即得二倍角公式 cos2θ = cos²θ − sin²θ。',
], y=2.7, w=6.4, size=17)
k.picture(s, F('09_demoivre.png'), x=7.0, y=2.5, w=5.8)
k.callout(s, '一句话：乘方就是"模乘方，辐角乘 n"。', x=0.85, y=5.9, w=5.9, h=0.75)
S.append(s)

# 19 单位根
s = k.content_slide(p, 'n 次单位根：把圆 n 等分', '核心')
k.formula(s, r'$z^{n}=1\ \Rightarrow\ z_k=\cos\frac{2k\pi}{n}+i\sin\frac{2k\pi}{n},\quad k=0,1,\dots,n-1$',
          x=0.9, y=1.4, w=11.4, size=0.72, out=T('f07.png'))
k.full_picture(s, F('10_roots.png'), y=2.5, w=8.4)
k.callout(s, 'n 个根全在单位圆上、辐角等差 2π/n ⇒ 它们恰是圆内接正 n 边形的顶点。'
             '"解方程"与"作正多边形"在此合流（高斯正十七边形的起点）。',
          x=1.4, y=6.35, w=10.5, h=0.9, kind='note')
S.append(s)

# 20 例题3
s = k.content_slide(p, '例题 3｜解方程 z⁴ = 1', '例题')
k.bullets(s, [
    '【三角化】1 = cos0 + i sin0，模为 1，辐角 0（+2kπ）。',
    '【开方】设 z = r(cosθ + i sinθ)，由棣莫弗：r⁴ = 1，4θ = 0 + 2kπ',
    ('⇒ r = 1，θ = kπ/2，k = 0,1,2,3', 1),
    '【四根】z₀ = 1，z₁ = i，z₂ = −1，z₃ = −i',
    '【验证】代数分解 z⁴−1 = (z²−1)(z²+1) = (z−1)(z+1)(z−i)(z+i) ✓',
    '【几何】四根构成单位圆内接正方形，两两共轭。',
], y=1.65, w=6.9, size=17)
k.picture(s, F('11_z4.png'), x=7.7, y=1.6, w=5.2)
k.callout(s, '通法：先化三角形式，再"模开 n 次方、辐角除以 n 并补 2kπ/n"。',
          x=0.85, y=6.2, w=6.6, h=0.8, kind='key')
S.append(s)

# 21 变式1
s = k.content_slide(p, '变式 1｜解 z³ = 8i', '变式')
k.bullets(s, [
    '【三角化】8i = 8(cos 90° + i sin 90°)',
    '【开方】r³ = 8 ⇒ r = 2；3θ = 90° + k·360° ⇒ θ = 30° + k·120°',
    '【三根】θ = 30°, 150°, 270°',
    ('z₀ = 2(cos30°+isin30°) = √3 + i', 1),
    ('z₁ = 2(cos150°+isin150°) = −√3 + i', 1),
    ('z₂ = 2(cos270°+isin270°) = −2i', 1),
    '【几何】三点在半径 2 的圆上，构成正三角形；三根之和为 0。',
], y=1.65, w=6.0, size=17)
k.picture(s, F('10_roots.png'), x=7.0, y=2.6, w=5.9)
k.callout(s, '通用结论：n 次方程 zⁿ = w 的 n 个根之和恒为 0（正 n 边形的重心在原点）。',
          x=0.85, y=6.3, w=6.6, h=0.75, kind='note')
S.append(s)

# 22 变式2
s = k.content_slide(p, '变式 2｜旋转的复合与 ω 的妙用', '变式')
k.bullets(s, [
    '记 ω = cos120° + i sin120° = −1/2 + (√3/2)i（三次单位根）。',
    '(1) 证明 1 + ω + ω² = 0。',
    ('法一：等比求和 (ω³−1)/(ω−1) = 0；法二：三向量首尾相接成正三角形。', 1),
    '(2) 已知正三角形 ABC 逆时针排列，A 对应 z₁、B 对应 z₂，求 C。',
    ('C 由 B 绕 A 逆时针转 60° 得到：z₃ = z₁ + (z₂ − z₁)(cos60° + i sin60°)', 1),
    '(3) 判定：z₁ + ω z₂ + ω² z₃ = 0 ⟺ △z₁z₂z₃ 为正三角形（逆时针）。',
], y=1.65, w=6.0, size=17)
k.picture(s, F('10_roots.png'), x=7.0, y=2.6, w=5.9)
k.callout(s, '单位根 ω 是"正三角形的代数身份证"。', x=0.85, y=6.3, w=6.6, h=0.75)
S.append(s)

# 23 共轭根
s = k.content_slide(p, '延伸：实系数方程的虚根必成对', '延伸')
k.bullets(s, [
    '设 f(x) 系数全为实数，若 f(z) = 0，则 f(z̄) = 0。',
    ('原因：共轭运算与加法、乘法可交换 ⇒ f(z̄) = conj(f(z)) = 0̄ = 0。', 1),
    '代数基本定理：n 次多项式在复数域内恰有 n 个根（含重数）。',
    '例：x⁴ − 1 在实轴上只见 2 个交点，在复平面上 4 个根一个不少。',
], y=1.65, w=5.9, size=17)
k.picture(s, F('14_conj_roots.png'), x=6.4, y=1.6, w=6.5)
k.callout(s, '"复数把方程的解补齐了"——这是引入复数最大的代数红利。',
          x=0.85, y=5.5, w=5.4, h=1.0, kind='key')
S.append(s)

# 24 幕五
S.append(k.section_slide(p, '第 5 幕 · 高光收尾', '欧拉公式：五大常数的会师', '46–56 min'))

# 25 欧拉公式
s = k.content_slide(p, '欧拉公式', '高光')
k.formula(s, r'$e^{i\theta}=\cos\theta+i\sin\theta$', x=2.6, y=1.4, w=8.2, size=1.15, out=T('f08.png'))
k.full_picture(s, F('12_euler.png'), y=2.85, w=5.6)
k.callout(s, '思路：把 e^x、cos x、sin x 的幂级数展开，代入 x = iθ，实部恰好拼成 cosθ，虚部恰好拼成 sinθ。',
          x=1.9, y=6.55, w=9.5, h=0.75, kind='note')
S.append(s)

# 26 指数形式
s = k.content_slide(p, '指数形式：所有规则一步归零', '统一')
k.formula(s, r'$z=re^{i\theta};\quad z_1z_2=r_1r_2e^{i(\theta_1+\theta_2)};\quad z^{n}=r^{n}e^{in\theta};\quad \bar z=re^{-i\theta}$',
          x=0.9, y=1.45, w=11.4, size=0.66, out=T('f09.png'))
k.bullets(s, [
    '"模相乘、辐角相加"⟵ 就是指数的 aᵐ·aⁿ = aᵐ⁺ⁿ！',
    '棣莫弗定理 ⟵ 就是幂的乘方 (aᵐ)ⁿ = aᵐⁿ！',
    'n 次单位根 ⟵ e^{2kπi/n}，k = 0..n−1，一眼看穿"n 等分圆"。',
    '前面所有几何结论，在指数形式下都退化成指数运算律。',
], y=2.7, w=6.6, size=17)
k.picture(s, F('12_euler.png'), x=7.4, y=2.5, w=5.5)
k.callout(s, '指数形式不是"新知识"，是把已学的一切压缩成一行。', x=0.85, y=5.9, w=6.2, h=0.8, kind='key')
S.append(s)

# 27 e^{iπ}+1=0
s = k.content_slide(p, '最美的公式：e^{iπ} + 1 = 0', '高光')
k.picture(s, F('13_eipi.png'), x=6.7, y=1.5, w=6.2)
k.bullets(s, [
    '在欧拉公式中令 θ = π：e^{iπ} = cosπ + i sinπ = −1',
    '移项即得 e^{iπ} + 1 = 0。',
    '一行式子里同时出现：',
    ('e（分析/增长）、i（代数/虚数）、π（几何/圆）', 1),
    ('1（乘法单位）、0（加法单位）', 1),
    '几何意义：从 1 出发，沿单位圆走过 π，正好走到 −1。',
], y=1.7, w=6.0, size=17)
k.callout(s, '复数是桥：代数 ↔ 几何、指数 ↔ 三角、离散 ↔ 连续。', x=0.85, y=6.1, w=5.8, h=0.85, kind='key')
S.append(s)

# 28 小结
s = k.content_slide(p, '课堂小结：一条主线', '小结')
k.bullets(s, [
    '① 动机：解三次方程求实根，虚数是必经之路（不是凭空硬造）。',
    '② 四位一体：a+bi ↔ 点(a,b)/向量 ↔ r(cosθ+isinθ) ↔ re^{iθ}。',
    '③ 加减 = 向量平移；|z₁−z₂| = 两点距离。',
    '④ 乘法 = 模相乘 + 辐角相加 = 旋转伸缩变换（本讲之魂）。',
    '⑤ 棣莫弗 ⇒ 乘方；反用 ⇒ n 次单位根 = 圆的 n 等分 = 正 n 边形。',
    '⑥ 欧拉公式把上述一切压缩为指数运算律，收束于 e^{iπ}+1=0。',
], y=1.7, w=7.3, size=17)
k.picture(s, F('05_times_i.png'), x=7.7, y=2.2, w=5.2)
k.callout(s, '若只带走一句话：复数乘法就是旋转。', x=0.85, y=6.4, w=7.0, h=0.7, kind='key')
S.append(s)

# 29 分层作业
s = k.content_slide(p, '分层作业', '作业')
k.bullets(s, [
    'A 基础（必做）',
    ('1. 化简 (2+3i)(1−2i)、(1+i)/(1−i)、i^{2026}。', 1),
    ('2. 把 z = −2 + 2i 写成三角形式与指数形式，求 |z|、arg z。', 1),
    'B 提高（必做）',
    ('3. 用棣莫弗求 (1+i)^{10}。', 1),
    ('4. 解 z³ = −27，画出三根在复平面上的位置。', 1),
    ('5. 点 P(2, 3) 绕点 Q(1, 1) 顺时针转 90°，求像点坐标。', 1),
    'C 挑战（选做）',
    ('6. 证明 1 + ω + ω² + … + ω^{n−1} = 0（ω 为 n 次单位根，ω ≠ 1）。', 1),
    ('7. 用欧拉公式导出 cos3θ = 4cos³θ − 3cosθ。', 1),
], y=1.6, w=6.8, size=16)
k.picture(s, F('11_z4.png'), x=7.8, y=2.4, w=5.1)
S.append(s)

# 30 板书提纲
s = k.content_slide(p, '板书提纲', '板书')
k.bullets(s, [
    '【左栏｜概念】',
    ('i²=−1；z=a+bi；复平面：点·向量', 1),
    ('|z|=√(a²+b²)；z=r(cosθ+isinθ)=re^{iθ}', 1),
    ('共轭：z·z̄=|z|²', 1),
    '【中栏｜核心推演】',
    ('z₁z₂ = r₁r₂[cos(θ₁+θ₂)+isin(θ₁+θ₂)]', 1),
    ('⇒ 模相乘、辐角相加 ⇒ 旋转+伸缩', 1),
    ('棣莫弗：zⁿ = rⁿ(cos nθ + i sin nθ)', 1),
    ('zⁿ=1 ⇒ z_k = e^{2kπi/n} ⇒ 正 n 边形', 1),
], y=1.6, w=6.6, size=16)
k.bullets(s, [
    '【右栏｜例题与高光】',
    ('例1 (1+i)(2−i)/(3+i) = 1', 1),
    ('例2 (3+i)·(cos60°+isin60°) 旋转', 1),
    ('例3 z⁴=1 ⇒ 1, i, −1, −i', 1),
    ('欧拉：e^{iθ}=cosθ+isinθ', 1),
    ('★ e^{iπ} + 1 = 0 ★', 1),
], x=7.4, y=1.6, w=5.4, size=16)
k.callout(s, '板书随堂生成：左栏定义写死不擦，中栏推演保留到下课，右栏例题可轮换。',
          x=0.85, y=6.35, w=11.6, h=0.7, kind='note')
S.append(s)

path_ppt = k.save(p, os.path.join(OUT, '07_复数与复平面.pptx'))
print('slides:', len(p.slides.__iter__.__self__._sldIdLst))

# ============================================================ 教案 docx
from docx import Document
from docx.shared import Pt as DPt, RGBColor as DRGB, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Noto Serif CJK SC'; st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')

def H(t, lv=1):
    h = doc.add_heading(t, level=lv)
    for r in h.runs:
        r.font.name = 'Noto Serif CJK SC'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
    return h

def P(t, b=False):
    p_ = doc.add_paragraph()
    r = p_.add_run(t); r.bold = b
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    return p_

t = doc.add_heading('教案：第07讲　复数、复平面与欧拉公式', 0)
for r in t.runs:
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

H('一、课题与课时', 1)
P('课题：复数、复平面与欧拉公式（数学名师课件包 第07讲）')
P('课时：1 课时，共 60 分钟（配套 PPT 30 页，插图 14 张）')
P('课型：概念建构 + 几何直观 + 方法综合课')

H('二、教材分析', 1)
P('复数是高中代数的收官内容，也是通往大学数学（线性代数、复变函数、傅里叶分析）的门槛。'
  '教材通常从"x²+1=0 无实数解"切入，逻辑上并不充分——历史上真正逼出虚数的是三次方程卡尔达诺公式：'
  '求纯实根时，中间步骤不得不出现负数的平方根。本课以此史实为动机起点，避免"人为造数"的生硬感。')
P('本课把复数的四种形式（代数、几何、三角、指数）统一为"一个对象的四张面孔"，'
  '并把全课的重心压在"乘法的几何意义"上：模相乘、辐角相加，等价于旋转 + 伸缩。'
  '由此顺势导出棣莫弗定理、n 次单位根（圆的 n 等分/正 n 边形），最后以欧拉公式收束，'
  '揭示复数作为"代数与几何、指数与三角之桥"的本质地位。')

H('三、学情分析', 1)
P('学生已掌握：向量的坐标与平行四边形法则、三角函数与两角和公式、多项式与方程根的概念、指数运算律。')
P('可能的障碍：① 把 i 当作"不存在的数"从而心理排斥；② 只会代数运算而无几何图像，'
  '导致辐角求解丢象限；③ 对"数即变换"的观念缺乏经验；④ 欧拉公式因涉及级数而只能"承认"不能"证明"，'
  '需用几何运动图景替代严格推导，保证可接受性。')

H('四、教学目标（三维）', 1)
P('知识与技能：', True)
P('1. 理解复数的概念与代数形式，掌握四则运算；2. 掌握复平面表示、模与辐角、三角形式与指数形式的互化；'
  '3. 掌握复数乘除的几何意义；4. 会用棣莫弗定理求乘方，会求 n 次单位根并作正 n 边形；'
  '5. 知道欧拉公式与 e^{iπ}+1=0。')
P('过程与方法：', True)
P('经历"历史困境 → 形式引入 → 几何化 → 运算的几何意义 → 统一形式"的完整建构过程；'
  '通过特例（乘以 i）归纳出一般规律，再用三角恒等变换严格推导，体会"归纳—演绎"的双轮驱动；'
  '经历数形结合、化归与转化的数学思想。')
P('情感态度与价值观：', True)
P('通过卡尔达诺—邦贝利的故事感受数学概念的诞生是"被问题逼出来的"；'
  '通过 e^{iπ}+1=0 感受数学的统一之美，激发对高等数学的向往。')

H('五、教学重点与难点', 1)
P('重点：复平面表示；复数乘法的几何意义（模相乘、辐角相加）；棣莫弗定理与 n 次单位根。', True)
P('难点：① 从"数"到"变换"的观念转变（乘法即旋转伸缩）；② n 次单位根的辐角参数化与正 n 边形的对应；'
  '③ 欧拉公式的合理性理解。', True)
P('突破策略：以"连续乘 i 四次"的动态图为认知锚点；以三角形式的严格展开完成演绎；'
  '以三个 n（3/5/8）的正多边形并列图揭示"n 等分圆"的共性；以单位圆上的匀速运动与其两个正交投影'
  '（余弦、正弦）解释欧拉公式的几何必然性。')

H('六、教法与学法', 1)
P('教法：史料情境导入法、探究发现法（特例→猜想→证明）、数形结合直观教学法、变式训练法。')
P('学法：观察—猜想—验证；动手在复平面上画向量；小组讨论"乘以 i 为何是旋转"；'
  '通过"代数算一遍、几何验一遍"的双通道自查。')

H('七、教学准备', 1)
P('教师：PPT（07_复数与复平面.pptx，30 页）、14 张 matplotlib 精确插图、几何画板/GeoGebra 备用动态演示、'
  '三色粉笔（黑：定义；蓝：推演；红：几何意义）。')
P('学生：直尺、量角器、坐标纸；复习向量运算与两角和公式。')

H('八、教学过程（分钟级时间轴）', 1)

rows = [
    ('0–2', '封面 · 学习目标\n(PPT 1–2)',
     '亮出 e^{iπ}+1=0，问："这行式子里五个符号，你认识几个？它们凭什么能站在一起？"宣布本课主线。',
     '被公式吸引，产生疑问；浏览学习目标。',
     '以最高光的结论开场，制造认知悬念，为 60 分钟设定"终点站"。'),
    ('2–8', '第 1 幕 · 卡尔达诺的困境\n(PPT 3–4)',
     '讲述史实：解 x³=15x+4，卡尔达诺公式给出 ∛(2+√−121)+∛(2−√−121)；投影函数图像，指出 x=4 是货真价实的实根。'
     '追问："答案是实数，为什么中途必须经过 √−121？"介绍邦贝利硬算下去得到 4 的故事。',
     '观察图像确认 x=4；意识到"无解方程需要新数"的说法站不住脚；接受"虚数是通往实答案的桥"。',
     '用真实历史替代教材的生硬动机，让新概念的引入具有必然性与说服力，落实数学文化。'),
    ('8–12', '复数的定义与代数形式\n(PPT 5)',
     '给出 i²=−1、z=a+bi；强调相等条件（一个复数方程 = 两个实方程）；带学生数 i 的四次周期；'
     '警示"复数不可比较大小"。',
     '记忆定义；口答 i^{2026} 的值；举反例理解"不可比大小"。',
     '把形式定义快速夯实，为几何化腾出时间；提前拦截高频错误。'),
    ('12–18', '第 2 幕 · 复平面与三角形式\n(PPT 6–9)',
     '建立复平面；展示"复数 ↔ 点 ↔ 向量"三位一体；引导学生自己说出加减法就是向量法则；'
     '强调 |z₁−z₂| 是两点距离；由 a=rcosθ、b=rsinθ 导出三角形式；强调辐角要看象限。',
     '在坐标纸上画出 z₁、z₂ 及 z₁+z₂ 的平行四边形；把 z=−1+√3i 化为三角形式并当堂报出 θ=2π/3。',
     '完成"数→形"的翻译，这是本课全部威力的地基；由学生自己发现加法法则，提高参与度。'),
    ('18–22', '第 3 幕 · 探究：乘以 i 是什么？\n(PPT 10–11)',
     '不给结论，先让学生算 z=2.6+1.1i 连续乘 i 的四个结果并描点；再投影四次旋转轨迹图，'
     '追问："长度变了吗？角度变了多少？为什么 i⁴=1？"',
     '动手计算并描点；发现模不变、每次转 90°；自己说出"i 是一个旋转 90° 的动作"。',
     '以最简特例制造"震撼"：数竟然代表动作。归纳先行，为随后的演绎提供动力。'),
    ('22–28', '推演：乘法的一般几何意义\n(PPT 12–13)',
     '板演三角形式相乘，用两角和公式收拢为 r₁r₂[cos(θ₁+θ₂)+isin(θ₁+θ₂)]；'
     '提炼八字口诀"模相乘、辐角相加"；把 z₂ 解释为作用在整个平面上的旋转伸缩算子。',
     '跟随板演，独立完成展开中的一步；齐读口诀；在图上指出 θ₁+θ₂。',
     '从特例猜想上升到严格证明，完成归纳—演绎闭环；确立本讲之魂。'),
    ('28–32', '共轭、除法与例题 1\n(PPT 14–15)',
     '给出共轭的镜面反射意义与 z·z̄=|z|²；讲分母实数化；板演例题 1 并用"模与辐角"复核。',
     '独立完成例 1 的分子化简；用模与辐角复核结果 |z|=1。',
     '把代数技巧与几何意义并行呈现，训练"双通道自查"的严谨习惯。'),
    ('32–36', '例题 2 · 旋转问题\n(PPT 16)',
     '提出几何题：A(3,1) 绕原点逆时针转 60°。引导学生把点复数化、把旋转乘法化；'
     '板演结果；给出绕定点旋转的通式 z′=z₀+(z−z₀)e^{iθ}。',
     '独立列式 z′=(3+i)(1/2+√3/2·i) 并完成计算；记录通式。',
     '让核心结论立刻变现为解题工具，形成"学了就能用"的正反馈。'),
    ('36–40', '第 4 幕 · 棣莫弗定理\n(PPT 17–18)',
     '由乘法法则归纳出 zⁿ；展示 zⁿ 的螺旋图；指出令 r=1,n=2 即得二倍角公式。',
     '口述归纳证明的关键步；观察螺旋图理解"模指数增长、辐角等差旋转"。',
     '定理不是空降的，而是核心结论的直接推论，减轻记忆负担。'),
    ('40–46', 'n 次单位根 · 例题 3\n(PPT 19–20)',
     '反用棣莫弗解 zⁿ=1；带学生做 z⁴=1 的完整板演（三角化→模开方→辐角除 n 补 2kπ/n）；'
     '投影 n=3/5/8 正多边形图，点明"解方程 = 把圆 n 等分"，提及高斯正十七边形。',
     '独立完成 z⁴=1 的四个根并画出正方形；用因式分解验证；观察三幅正多边形图归纳共性。',
     '代数与几何在此合流，把抽象的"方程的根"变成看得见的正多边形，是本课第二个高光。'),
    ('46–50', '变式 1、变式 2\n(PPT 21–22)',
     '布置变式 1（z³=8i）当堂限时 3 分钟，巡视并请一名学生板演；变式 2 讲 ω 的性质与正三角形判定。',
     '限时完成 z³=8i；一人上台板演；讨论 1+ω+ω²=0 的两种证法。',
     '及时迁移，检测通法掌握度；把单位根提升为"正多边形的代数身份证"。'),
    ('50–56', '第 5 幕 · 欧拉公式（高光）\n(PPT 23–27)',
     '先讲实系数虚根共轭成对与代数基本定理（复数把根补齐了）；再给出 e^{iθ}=cosθ+isinθ，'
     '用幂级数思路点到为止，重点用"单位圆匀速运动 + 两个投影"的图景解释；'
     '写出指数形式，指明"模相乘辐角相加"就是指数运算律；令 θ=π，写下 e^{iπ}+1=0，'
     '逐一点名 e、i、π、1、0 的出身。',
     '感受推导；齐答 e^{iπ}=−1；在笔记本上抄下最后一行式子。',
     '把全课所有结论压缩成一行指数律，实现认知的"降维压缩"；以数学之美收束情感目标。'),
    ('56–60', '小结 · 分层作业 · 板书回看\n(PPT 28–30)',
     '沿板书主线复盘六条；强调"若只带走一句话：复数乘法就是旋转"；布置 A/B/C 分层作业并解释选做要求。',
     '跟随复盘，补全笔记；记录作业；对 C 组挑战题提问。',
     '结构化收口，形成长时记忆的骨架；分层作业照顾差异，保证下限、抬高上限。'),
]
tb = doc.add_table(rows=1, cols=5)
tb.style = 'Table Grid'
hdr = ['时间', '环节 / 对应 PPT', '教师活动', '学生活动', '设计意图']
for i, hh in enumerate(hdr):
    c = tb.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(hh); r.bold = True; r.font.size = DPt(10)
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for row in rows:
    cells = tb.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = ''
        r = cells[i].paragraphs[0].add_run(v); r.font.size = DPt(9)
        r.font.name = 'Noto Serif CJK SC'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for w, col in zip([Cm(1.8), Cm(3.2), Cm(6.6), Cm(4.2), Cm(4.6)], tb.columns):
    for c in col.cells:
        c.width = w
P('合计：2+6+4+6+4+6+4+4+4+6+4+6+4 = 60 分钟。', True)

H('九、板书设计', 1)
P('黑板三栏式：')
P('左栏（概念，写死不擦）：i²=−1；z=a+bi；复平面：点·向量；|z|=√(a²+b²)；z=r(cosθ+isinθ)=re^{iθ}；z·z̄=|z|²')
P('中栏（核心推演，保留至下课）：z₁z₂=r₁r₂[cos(θ₁+θ₂)+isin(θ₁+θ₂)] ⇒ 模相乘、辐角相加 ⇒ 旋转+伸缩；'
  '棣莫弗 zⁿ=rⁿ(cos nθ+i sin nθ)；zⁿ=1 ⇒ z_k=e^{2kπi/n} ⇒ 正 n 边形')
P('右栏（例题与高光，可轮换擦写）：例1 (1+i)(2−i)/(3+i)=1；例2 (3+i)(cos60°+isin60°)；'
  '例3 z⁴=1 ⇒ 1, i, −1, −i；欧拉 e^{iθ}=cosθ+isinθ；★ e^{iπ}+1=0 ★（红框，下课不擦）')

H('十、分层作业', 1)
P('A 基础（全体必做，约 15 分钟）', True)
P('1. 化简：(2+3i)(1−2i)；(1+i)/(1−i)；i^{2026}。')
P('2. 把 z=−2+2i 写成三角形式与指数形式，求 |z| 与 arg z。')
P('B 提高（全体必做，约 20 分钟）', True)
P('3. 用棣莫弗定理求 (1+i)^{10}。')
P('4. 解 z³=−27，并在复平面上画出三个根的位置，指出它们构成的图形。')
P('5. 点 P(2,3) 绕点 Q(1,1) 顺时针旋转 90°，求像点坐标（要求用复数法）。')
P('C 挑战（学有余力选做）', True)
P('6. 证明：ω 为 n 次单位根且 ω≠1，则 1+ω+ω²+…+ω^{n−1}=0，并给出几何解释。')
P('7. 用欧拉公式（或棣莫弗定理）导出 cos3θ=4cos³θ−3cosθ 与 sin3θ=3sinθ−4sin³θ。')
P('8. 阅读拓展：查阅高斯用尺规作正十七边形的思路，说明它与 17 次单位根的关系。')

H('十一、教学反思（课后填写）', 1)
for t_ in ['1. 卡尔达诺情境的时长控制是否合适？学生是否真正被"实根却经虚数"打动？',
           '2. "连续乘 i"的探究环节，有多少学生自主归纳出了"旋转 90°"？',
           '3. 三角形式乘法的推导，展开步骤是否有学生跟丢？下次是否需要先复习两角和公式？',
           '4. z⁴=1 的板演，学生在"辐角补 2kπ/n"这一步的错误率如何？',
           '5. 欧拉公式的接受度：是"惊叹"还是"茫然"？几何图景的解释是否足够？',
           '6. 60 分钟是否超时？哪一幕可压缩？',
           '7. 其他：']:
    P(t_)
    doc.add_paragraph('　')

path_doc = os.path.join(OUT, '教案_07_复数与复平面.docx')
doc.save(path_doc)
print('docx saved')
print('PPT:', path_ppt)
print('DOC:', path_doc)
print('FIGS:', len([f for f in os.listdir(FIG) if f.endswith('.png')]))
