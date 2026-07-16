# -*- coding: utf-8 -*-
"""第03讲：圆锥曲线之椭圆——定义、方程与光学性质（60min 课件包）"""
import sys, os
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, Ellipse, Arc, FancyArrowPatch, Polygon

OUT = "/mnt/c/Users/mi/Desktop/数学名师课件包/03_椭圆的定义与性质"
FIG = os.path.join(OUT, "figures")
os.makedirs(FIG, exist_ok=True)
F = lambda n: os.path.join(FIG, n)

# ============================ 配图 ============================
def ell(ax, a, b, **kw):
    t = np.linspace(0, 2*np.pi, 400)
    ax.plot(a*np.cos(t), b*np.sin(t), **kw)

# fig01 绳圈画椭圆：两焦点定义
def fig01():
    fig, ax = k.new_fig(6.1, 4.2)
    a, b = 4.0, 3.0; c = np.sqrt(a*a-b*b)
    ell(ax, a, b, color=k.M_ACC2, lw=2.6, zorder=3)
    th = 2.1
    P = (a*np.cos(th), b*np.sin(th))
    for s, nm in ((-1, "$F_1$"), (1, "$F_2$")):
        ax.plot(s*c, 0, 'o', color=k.M_RED, ms=9, zorder=5)
        ax.annotate(nm, (s*c, 0), textcoords="offset points", xytext=(0, -34),
                    ha="center", color=k.M_RED)
    ax.plot(*P, 'o', color=k.M_ACC, ms=9, zorder=5)
    ax.annotate("$P$", P, textcoords="offset points", xytext=(-6, 12),
                ha="center", color=k.M_ACC)
    for s in (-1, 1):
        ax.plot([s*c, P[0]], [0, P[1]], color=k.M_ACC, lw=2.2, ls='-', zorder=4)
    ax.annotate("$r_1$", ((-c+P[0])/2, P[1]/2), xytext=(-30, 4),
                textcoords="offset points", color=k.M_INK)
    ax.annotate("$r_2$", ((c+P[0])/2, P[1]/2), xytext=(10, 8),
                textcoords="offset points", color=k.M_INK)
    ax.set_title("$r_1+r_2=2a$（绳长恒定）", color=k.M_INK, pad=30)
    k.style_axes(ax); ax.set_xlim(-6.4, 6.4); ax.set_ylim(-4.6, 4.6); ax.set_aspect('equal')
    return k.save_fig(fig, F("f01_definition.png"))

# fig02 绳圈实验三态：绳长/图钉间距变化
def fig02():
    fig, axes = plt.subplots(1, 3, figsize=(6.0, 2.6))
    cfgs = [(4.0, 1.0), (4.0, 3.0), (4.0, 3.9)]
    for ax, (a, c) in zip(axes, cfgs):
        b = np.sqrt(a*a-c*c)
        ell(ax, a, b, color=k.M_ACC2, lw=2.4)
        ax.plot([-c, c], [0, 0], 'o', color=k.M_RED, ms=7)
        th = 1.0; P = (a*np.cos(th), b*np.sin(th))
        for s in (-1, 1):
            ax.plot([s*c, P[0]], [0, P[1]], color=k.M_ACC, lw=1.6)
        ax.plot(*P, 'o', color=k.M_ACC, ms=6)
        ax.set_title(f"$e={c/a:.2f}$", color=k.M_INK, pad=8)
        ax.set_aspect('equal'); ax.set_xlim(-4.6, 4.6); ax.set_ylim(-4.4, 4.4)
        ax.set_xticks([]); ax.set_yticks([])
        for s in ax.spines.values(): s.set_color(k.M_RULE)
    return k.save_fig(fig, F("f02_string_experiment.png"))

# fig03 建系与推导
def fig03():
    fig, ax = k.new_fig(6.0, 4.2)
    a, b = 4.0, 3.0; c = np.sqrt(a*a-b*b)
    ell(ax, a, b, color=k.M_ACC2, lw=2.4)
    th = 1.9; P = (a*np.cos(th), b*np.sin(th))
    for s in (-1, 1):
        ax.plot(s*c, 0, 'o', color=k.M_RED, ms=8)
        ax.plot([s*c, P[0]], [0, P[1]], color=k.M_ACC, lw=2.0)
    ax.plot(*P, 'o', color=k.M_ACC, ms=8)
    ax.plot([P[0], P[0]], [0, P[1]], color=k.M_SLATE, ls=':', lw=1.3)
    ax.annotate("$P(x,y)$", P, xytext=(-12, 12), ha="right",
                textcoords="offset points", color=k.M_ACC)
    ax.annotate("$F_1$", (-c, 0), xytext=(-2, -36), ha="center",
                textcoords="offset points", color=k.M_RED)
    ax.annotate("$F_2$", (c, 0), xytext=(2, -36), ha="center",
                textcoords="offset points", color=k.M_RED)
    ax.set_title("对称建系：原点取 $F_1F_2$ 中点", color=k.M_INK, pad=30)
    k.style_axes(ax); ax.set_xlim(-6.6, 6.6); ax.set_ylim(-4.8, 4.8); ax.set_aspect('equal')
    return k.save_fig(fig, F("f03_setup.png"))

# fig04 a b c 直角三角形
def fig04():
    fig, ax = k.new_fig(5.9, 4.4)
    a, b = 4.0, 3.0; c = np.sqrt(a*a-b*b)
    ell(ax, a, b, color=k.M_ACC2, lw=2.4)
    B = (0, b)
    ax.plot([0, c], [0, 0], color=k.M_RED, lw=3)
    ax.plot([0, 0], [0, b], color=k.M_GRN, lw=3)
    ax.plot([c, 0], [0, b], color=k.M_ACC, lw=3)
    ax.plot([-c, c], [0, 0], 'o', color=k.M_RED, ms=8)
    ax.plot(0, b, 'o', color=k.M_GRN, ms=8)
    ax.plot([a, -a], [0, 0], 'o', color=k.M_INK, ms=6)
    ax.add_patch(Polygon([[0, 0], [0.35, 0], [0.35, 0.35], [0, 0.35]],
                         closed=True, fill=False, ec=k.M_SLATE, lw=1.2))
    ax.annotate("$c$", (c/2, 0), xytext=(0, -30), textcoords="offset points",
                ha="center", color=k.M_RED)
    ax.annotate("$b$", (0, b/2), xytext=(-62, -8), textcoords="offset points", color=k.M_GRN)
    ax.annotate("$a$", (c/2, b/2), xytext=(16, 10), textcoords="offset points", color=k.M_ACC)
    ax.annotate("$B(0,b)$", B, xytext=(10, 12), textcoords="offset points", color=k.M_GRN)
    ax.annotate("$A_2(a,0)$", (a, 0), xytext=(10, 14), ha="left",
                textcoords="offset points", color=k.M_INK)
    ax.set_title("$a^2=b^2+c^2$（$a$ 是斜边）", color=k.M_INK, pad=30)
    k.style_axes(ax); ax.set_xlim(-6.8, 8.0); ax.set_ylim(-3.8, 5.0); ax.set_aspect('equal')
    ax.set_xticks([-4, -2, 2, 4]); ax.set_yticks([-2, 2, 4])
    return k.save_fig(fig, F("f04_abc_triangle.png"))

# fig05 两种焦点位置对比
def fig05():
    fig, axes = plt.subplots(1, 2, figsize=(10.2, 4.2))
    for ax, foc in zip(axes, ("x", "y")):
        a, b = 4.0, 2.6; c = np.sqrt(a*a-b*b)
        if foc == "x":
            ell(ax, a, b, color=k.M_ACC2, lw=2.4); fs = [(-c, 0), (c, 0)]
            t = r"焦点在 $x$ 轴　$\dfrac{x^2}{a^2}+\dfrac{y^2}{b^2}=1$"
        else:
            ell(ax, b, a, color=k.M_ACC2, lw=2.4); fs = [(0, -c), (0, c)]
            t = r"焦点在 $y$ 轴　$\dfrac{y^2}{a^2}+\dfrac{x^2}{b^2}=1$"
        for p in fs:
            ax.plot(*p, 'o', color=k.M_RED, ms=8)
        ax.set_title(t, color=k.M_INK, pad=30)
        k.style_axes(ax); ax.set_xlim(-5.4, 5.4); ax.set_ylim(-5.2, 5.2); ax.set_aspect('equal')
        ax.set_xticks([-4, 4]); ax.set_yticks([-4, 4])
    return k.save_fig(fig, F("f05_two_forms.png"))

# fig06 离心率演变多子图
def fig06():
    es = [0.0, 0.4, 0.75, 0.9]
    fig, axes = plt.subplots(2, 2, figsize=(5.9, 5.0))
    a = 4.0
    for ax, e in zip(axes.ravel(), es):
        c = a*e; b = np.sqrt(a*a-c*c)
        ell(ax, a, b, color=k.M_ACC2, lw=2.6)
        ax.plot([-c, c], [0, 0], 'o', color=k.M_RED, ms=7)
        ax.axhline(0, color=k.M_RULE, lw=0.9); ax.axvline(0, color=k.M_RULE, lw=0.9)
        lab = "$e=0$（圆）" if e == 0 else f"$e={e}$"
        ax.set_title(lab, color=k.M_INK, pad=8)
        ax.set_aspect('equal'); ax.set_xlim(-4.6, 4.6); ax.set_ylim(-4.6, 4.6)
        ax.set_xticks([]); ax.set_yticks([])
        for s in ax.spines.values(): s.set_color(k.M_RULE)
    return k.save_fig(fig, F("f06_eccentricity.png"))

# fig07 焦点—准线统一定义 e=PF/PL
def fig07():
    fig, axes = plt.subplots(1, 3, figsize=(10.7, 3.6))
    for ax, (e, nm, col) in zip(axes, [(0.6, "$e<1$　椭圆", k.M_ACC2),
                                       (1.0, "$e=1$　抛物线", k.M_GRN),
                                       (1.6, "$e>1$　双曲线", k.M_RED)]):
        # 极坐标 r = e*d/(1+e cos θ)，焦点在原点，准线 x = d
        d = 2.2
        th = np.linspace(-np.pi, np.pi, 4000)
        den = 1 + e*np.cos(th)
        r = e*d/den
        m = (den > 0.05) & (r > 0) & (r < 14)
        ax.plot(r[m]*np.cos(th[m]), r[m]*np.sin(th[m]), '.', ms=2.0, color=col)
        if e > 1:
            m2 = (den < -0.05) & (r < 0) & (r > -14)
            ax.plot(r[m2]*np.cos(th[m2]), r[m2]*np.sin(th[m2]), '.', ms=2.0, color=col, alpha=0.55)
        ax.axvline(d, color=k.M_ACC, ls='--', lw=1.6)
        ax.text(d + 0.4, 4.6, "准线", color=k.M_ACC, ha="left")
        ax.plot(0, 0, 'o', color=k.M_RED, ms=8)
        ax.text(-0.6, -3.0, "$F$", color=k.M_RED, ha="right")
        t0 = 1.1; rr = e*d/(1+e*np.cos(t0)); P = (rr*np.cos(t0), rr*np.sin(t0))
        ax.plot(*P, 'o', color=k.M_INK, ms=7)
        ax.plot([0, P[0]], [0, P[1]], color=k.M_RED, lw=1.8)
        ax.plot([P[0], d], [P[1], P[1]], color=k.M_ACC, lw=1.8)
        ax.text(P[0]-0.4, P[1]+0.5, "$P$", color=k.M_INK, ha="right")
        ax.set_title(nm, color=k.M_INK, pad=8)
        ax.set_aspect('equal'); ax.set_xlim(-6.6, 9.6); ax.set_ylim(-6.2, 6.2)
        ax.set_xticks([]); ax.set_yticks([])
        for s in ax.spines.values(): s.set_color(k.M_RULE)
    return k.save_fig(fig, F("f07_unified.png"))

# fig08 椭圆的准线与第二定义
def fig08():
    fig, ax = k.new_fig(6.2, 4.6)
    a, b = 4.0, 3.0; c = np.sqrt(a*a-b*b); xd = a*a/c
    ell(ax, a, b, color=k.M_ACC2, lw=2.4)
    for s in (-1, 1):
        ax.axvline(s*xd, color=k.M_ACC, ls='--', lw=1.6)
        ax.plot(s*c, 0, 'o', color=k.M_RED, ms=8)
    th = 1.0; P = (a*np.cos(th), b*np.sin(th))
    ax.plot(*P, 'o', color=k.M_INK, ms=8)
    ax.plot([c, P[0]], [0, P[1]], color=k.M_RED, lw=2)
    ax.plot([P[0], xd], [P[1], P[1]], color=k.M_ACC, lw=2)
    ax.annotate("$P$", P, xytext=(-6, 12), ha="center",
                textcoords="offset points", color=k.M_INK)
    ax.annotate("$F_2$", (c, 0), xytext=(-14, -34), ha="center",
                textcoords="offset points", color=k.M_RED)
    ax.text(xd + 0.4, 5.2, "$x=a^2/c$", color=k.M_ACC, ha="center")
    ax.text(-xd - 0.4, 5.2, "$x=-a^2/c$", color=k.M_ACC, ha="center")
    ax.set_title("准线 $x=\\pm a^2/c$（在椭圆外）", color=k.M_INK, pad=30)
    k.style_axes(ax); ax.set_xlim(-10.4, 10.4); ax.set_ylim(-4.6, 7.0); ax.set_aspect('equal')
    ax.set_xticks([-5, 5]); ax.set_yticks([-3, 3])
    return k.save_fig(fig, F("f08_directrix.png"))

# fig09 圆锥截面 2D 投影
def fig09():
    fig, ax = k.new_fig(6.5, 4.4)
    H, R = 4.2, 3.0
    g = H/R                                   # 母线斜率（上半锥 y = g|x|）
    for s_, col, al in ((1, k.M_INK, 1.0), (-1, k.M_SLATE, 0.5)):
        ax.plot([-R, 0, R], [s_*H, 0, s_*H], color=col, lw=2.2, alpha=al)
        ax.add_patch(Ellipse((0, s_*H), 2*R, 0.95, fc='none', ec=col, lw=1.6, alpha=al))
    ax.plot(0, 0, 'o', color=k.M_INK, ms=6, zorder=5)
    ax.text(0.5, -0.9, "$V$", color=k.M_INK)
    ax.plot([0, 0], [-H, H], color=k.M_RULE, lw=1.0, ls=':')

    # 圆：水平截
    yc = 2.6; xc = yc/g
    ax.plot([-xc, xc], [yc, yc], color=k.M_GRN, lw=2.4, zorder=4)
    ax.add_patch(Ellipse((0, yc), 2*xc, 0.55, fc=k.M_GRN, ec=k.M_GRN, alpha=0.13, lw=1.4))
    ax.text(2.4, 2.5, "圆", color=k.M_GRN, ha="left", va="center")

    # 椭圆：斜截（|斜率| < 母线斜率），只交上半锥的两条母线
    m, q = -0.40, 2.80
    xL = -q/(g + m); yL = g*(-xL)          # 左母线 y=-g x
    xR = q/(g - m);  yR = g*xR             # 右母线 y=g x
    ax.plot([xL, xR], [yL, yR], color=k.M_ACC, lw=3.2, zorder=6)
    mid = ((xL+xR)/2, (yL+yR)/2)
    Lc = np.hypot(xR-xL, yR-yL)
    ax.add_patch(Ellipse(mid, Lc, 0.95, angle=np.rad2deg(np.arctan2(yR-yL, xR-xL)),
                         fc=k.M_ACC, ec=k.M_ACC, alpha=0.15, lw=1.6, zorder=3))
    ax.text(-3.2, 4.3, "椭圆", color=k.M_ACC, ha="right", va="center")

    # 抛物线：平行于右母线（斜率恰为 g）
    xs = np.array([-2.7, 1.9])
    ax.plot(xs, g*xs + 2.0, color=k.M_ACC2, lw=2.2, ls='--', zorder=4)
    ax.text(2.2, 4.8, "抛物线", color=k.M_ACC2, ha="left", va="center")

    # 双曲线：平行于轴，截到上下两支
    xv = -1.0
    ax.plot([xv, xv], [-H-0.3, H+0.3], color=k.M_RED, lw=2.2, ls=':', zorder=4)
    ax.text(xv, -5.2, "双曲线", color=k.M_RED, ha="center", va="center")

    ax.set_title("同一圆锥，切角不同", color=k.M_INK)
    ax.set_aspect('equal'); ax.set_xlim(-8.1, 8.1); ax.set_ylim(-6.0, 5.6)
    ax.axis('off')
    return k.save_fig(fig, F("f09_conic_sections.png"))

# fig10 Dandelin 双球（3D 配置的 2D 斜投影）
def fig10():
    fig, ax = k.new_fig(7.2, 5.6)
    half = np.deg2rad(30); t, sn = np.tan(half), np.sin(half)
    SH = 0.42                      # 斜投影：(X,Y,Z) -> (X, Z + SH*Y)
    pj = lambda X, Y, Z: (X, Z + SH*Y)

    # 圆锥的两条轮廓母线（Y=0）
    Zs = np.array([0.0, 7.4])
    ax.plot(t*Zs, Zs, color=k.M_INK, lw=2.4, zorder=2)
    ax.plot(-t*Zs, Zs, color=k.M_INK, lw=2.4, zorder=2)
    ax.plot([0, 0], [0, 7.4], color=k.M_SLATE, lw=0.9, ls=':', zorder=1)
    ax.plot(0, 0, 'o', color=k.M_INK, ms=6, zorder=5)
    ax.text(0.0, -1.1, "$V$", color=k.M_INK, ha="center")

    k1, k2 = 1.4, 5.0              # 球心高度（k2>=3k1 保证两球分离）
    r1, r2 = k1*sn, k2*sn
    th = np.linspace(0, 2*np.pi, 300)

    def sphere(kc, r, col, lab, side=1):
        ax.add_patch(Circle((0, kc), r, fc=col, ec=col, alpha=0.12, lw=0, zorder=3))
        ax.add_patch(Circle((0, kc), r, fc='none', ec=col, lw=2, zorder=4))
        ax.plot(0, kc, 'x', color=col, ms=6, zorder=5)
        ax.text(side*(r+0.3), kc, lab, color=col,
                ha="left" if side > 0 else "right", va="center", zorder=14)
        # 与圆锥的切圆：Z = kc*cos^2(half)，半径 rho = t*Z
        Zc = kc/(1+t*t); rho = t*Zc
        u, v = pj(rho*np.cos(th), rho*np.sin(th), Zc)
        ax.plot(u, v, color=col, lw=1.6, ls='--', alpha=0.85, zorder=6)
        return Zc, rho

    Zc1, _ = sphere(k1, r1, k.M_GRN, "$S_1$", side=-1)
    Zc2, _ = sphere(k2, r2, k.M_ACC2, "$S_2$", side=1)

    # 割平面 A*X + Z + C = 0（与两球内公切）
    n = (k2 - k1)/(r1 + r2); A = np.sqrt(n*n - 1); C = -r1*n - k1
    # 截线：X^2+Y^2 = t^2 Z^2 且 Z = -(A X + C)
    aa = 1 - t*t*A*A; bb = -2*t*t*A*C; cc = -t*t*C*C
    x0 = -bb/(2*aa); Yr2 = -(cc - bb*bb/(4*aa))
    Xh = np.sqrt(Yr2/aa); Yh = np.sqrt(Yr2)
    EX = x0 + Xh*np.cos(th); EY = Yh*np.sin(th); EZ = -(A*EX + C)
    ax.plot(*pj(EX, EY, EZ), color=k.M_ACC, lw=3.0, zorder=8)
    ax.annotate("截线", pj(x0+Xh, 0, -(A*(x0+Xh)+C)),
                xytext=(64, 40), textcoords="offset points", color=k.M_ACC,
                arrowprops=dict(arrowstyle="->", color=k.M_ACC))
    # 切点 F1,F2（在 Y=0 的轴平面内）
    def foot_cut(kc):
        d = (kc + C)/(A*A + 1)
        return np.array([-A*d, 0.0, kc - d])
    F1, F2 = foot_cut(k1), foot_cut(k2)

    # 截线上取一点 P（Y != 0，不在轴平面内）
    tp = 2.55
    P = np.array([x0 + Xh*np.cos(tp), Yh*np.sin(tp), 0.0]); P[2] = -(A*P[0] + C)
    # 过 P 的母线（V->P 方向）与两球的切点 T1,T2 = 球心到该直线的垂足
    u = P/np.linalg.norm(P)
    Ti = lambda kc: np.dot(np.array([0, 0, kc]), u)*u
    T1, T2 = Ti(k1), Ti(k2)
    Gend = 1.25*np.linalg.norm(T2)*u
    ax.plot(*zip(pj(0, 0, 0), pj(*Gend)), color=k.M_SLATE, lw=1.3, ls='-', alpha=0.8, zorder=8)
    ax.plot(*zip(pj(*T1), pj(*T2)), color=k.M_RED, lw=4.0, alpha=0.55, zorder=9)
    ax.plot(*zip(pj(*P), pj(*F1)), color=k.M_GRN, lw=2.2, zorder=10)
    ax.plot(*zip(pj(*P), pj(*F2)), color=k.M_ACC2, lw=2.2, zorder=10)

    for Q, nm, off in ((F1, "$F_1$", (10, -32)), (F2, "$F_2$", (18, 8))):
        ax.plot(*pj(*Q), 'o', color=k.M_RED, ms=10, zorder=12)
        ax.annotate(nm, pj(*Q), xytext=off, textcoords="offset points",
                    color=k.M_RED, zorder=14)
    for Q, nm, off in ((T1, "$T_1$", (16, 4)), (T2, "$T_2$", (16, 4))):
        ax.plot(*pj(*Q), 's', color=k.M_INK, ms=8, zorder=12)
        ax.annotate(nm, pj(*Q), xytext=off, textcoords="offset points",
                    color=k.M_INK, zorder=14)
    ax.plot(*pj(*P), 'o', color=k.M_ACC, ms=11, zorder=13)
    ax.annotate("$P$", pj(*P), xytext=(-38, -6), textcoords="offset points",
                color=k.M_ACC, zorder=14)
    ax.set_title("Dandelin 双球", color=k.M_INK)
    ax.set_aspect('equal'); ax.set_xlim(-4.6, 7.0); ax.set_ylim(-1.8, 8.2)
    ax.axis('off')
    return k.save_fig(fig, F("f10_dandelin.png"))

# fig11 Dandelin 关键引理：球外一点两条切线长相等
def fig11():
    fig, axes = plt.subplots(1, 2, figsize=(11.0, 5.2),
                             gridspec_kw=dict(width_ratios=[1.0, 1.5], wspace=0.05))
    ax = axes[0]
    O = (0, 0); R = 1.4
    ax.add_patch(Circle(O, R, fc=k.M_ACC2, ec=k.M_ACC2, alpha=0.15, lw=2))
    ax.add_patch(Circle(O, R, fc='none', ec=k.M_ACC2, lw=2))
    P = (3.4, 1.9); dP = np.hypot(*P)
    L = np.sqrt(dP**2 - R**2)
    ang = np.arctan2(P[1], P[0]); al = np.arccos(R/dP)
    for s in (1, -1):
        T = (R*np.cos(ang + s*al), R*np.sin(ang + s*al))
        ax.plot([P[0], T[0]], [P[1], T[1]], color=k.M_ACC, lw=2.2)
        ax.plot([0, T[0]], [0, T[1]], color=k.M_SLATE, lw=1.2, ls=':')
        ax.plot(*T, 'o', color=k.M_RED, ms=7)
        ax.annotate("$T_%d$" % (1 if s == 1 else 2), T,
                    xytext=(-6, 12) if s == 1 else (-4, -30),
                    ha="center", textcoords="offset points", color=k.M_RED)
    ax.plot(*P, 'o', color=k.M_INK, ms=8); ax.plot(*O, 'x', color=k.M_ACC2, ms=7)
    ax.plot([0, P[0]], [0, P[1]], color=k.M_SLATE, lw=1.0, ls='--')
    ax.annotate("$P$", P, xytext=(12, 0), textcoords="offset points", color=k.M_INK)
    ax.text(-0.25, -0.75, "$O$", color=k.M_ACC2, ha="center")
    ax.set_title("$|PT_1|=|PT_2|$", color=k.M_INK)
    ax.set_aspect('equal'); ax.set_xlim(-2.6, 5.6); ax.set_ylim(-2.8, 3.6); ax.axis('off')

    ax = axes[1]
    ax.axis('off')
    ax.text(0.0, 0.98, "把引理搬到圆锥上：", color=k.M_ACC, va="top", transform=ax.transAxes)
    steps = [
        "1. 母线切小球于 $T_1$，切大球于 $T_2$",
        "2. 割平面切小球于 $F_1$，切大球于 $F_2$",
        "3. $|PF_1|=|PT_1|$（同为切线段）",
        "4. $|PF_2|=|PT_2|$（同理）",
        "5. 相加：$|PF_1|+|PF_2|=|T_1T_2|$",
        "6. $|T_1T_2|$ 与 $P$ 无关 $\\Rightarrow$ 常数 $2a$",
        "$\\Rightarrow$ 截线是以 $F_1,F_2$ 为焦点的椭圆 $\\blacksquare$",
    ]
    for i, s in enumerate(steps):
        col = k.M_ACC if i == len(steps)-1 else k.M_INK
        ax.text(0.0, 0.83 - i*0.135, s, color=col, va="top", transform=ax.transAxes)
    return k.save_fig(fig, F("f11_dandelin_proof.png"))

# fig12 光学性质：反射光路
def fig12():
    fig, ax = k.new_fig(10.2, 5.0)
    a, b = 4.2, 3.0; c = np.sqrt(a*a-b*b)
    ell(ax, a, b, color=k.M_ACC2, lw=2.8)
    for s, nm in ((-1, "$F_1$"), (1, "$F_2$")):
        ax.plot(s*c, 0, 'o', color=k.M_RED, ms=10, zorder=6)
        ax.annotate(nm, (s*c, 0), xytext=(0, -36), ha="center",
                    textcoords="offset points", color=k.M_RED)
    for th in (0.55, 1.35, 2.25, 3.6, 4.7, 5.6):
        P = (a*np.cos(th), b*np.sin(th))
        ax.add_patch(FancyArrowPatch((-c, 0), P, arrowstyle='-|>', mutation_scale=13,
                                     color=k.M_ACC, lw=1.8, alpha=0.9))
        ax.add_patch(FancyArrowPatch(P, (c, 0), arrowstyle='-|>', mutation_scale=13,
                                     color=k.M_GRN, lw=1.8, alpha=0.9))
        # 法线
        n = np.array([P[0]/a**2, P[1]/b**2]); n = n/np.linalg.norm(n)
        ax.plot([P[0], P[0]+0.75*n[0]], [P[1], P[1]+0.75*n[1]], color=k.M_SLATE, ls=':', lw=1.2)
    ax.set_title("从 $F_1$ 射出的光，反射后必过 $F_2$", color=k.M_INK, pad=30)
    k.style_axes(ax); ax.set_xlim(-6.4, 6.4); ax.set_ylim(-4.6, 4.6); ax.set_aspect('equal')
    ax.set_xticks([-5, 5]); ax.set_yticks([-3, 3])
    return k.save_fig(fig, F("f12_optics.png"))

# fig13 光学性质证明（对称点/最短路径）
def fig13():
    fig, ax = k.new_fig(7.6, 5.4)
    a, b = 4.2, 3.0; c = np.sqrt(a*a-b*b)
    ell(ax, a, b, color=k.M_ACC2, lw=2.4, zorder=3)
    th = 0.85; P = np.array([a*np.cos(th), b*np.sin(th)])
    d = np.array([-a*np.sin(th), b*np.cos(th)]); d = d/np.linalg.norm(d)
    ts = np.array([-3.4, 3.4])
    ax.plot(P[0]+ts*d[0], P[1]+ts*d[1], color=k.M_ACC, lw=2.6, zorder=4)
    Ltip = P - 3.4*d                     # 切线的右下端，标签放这里，远离 Q
    ax.annotate("切线 $l$", Ltip, xytext=(12, -18), textcoords="offset points",
                color=k.M_ACC, ha="left")
    F1 = np.array([-c, 0.0]); F2 = np.array([c, 0.0])
    for Fp, nm in ((F1, "$F_1$"), (F2, "$F_2$")):
        ax.plot(*Fp, 'o', color=k.M_RED, ms=10, zorder=8)
        ax.annotate(nm, Fp, xytext=(0, -36), ha="center",
                    textcoords="offset points", color=k.M_RED)
    ax.plot([F1[0], P[0]], [F1[1], P[1]], color=k.M_RED, lw=2.4, zorder=6)
    ax.plot([F2[0], P[0]], [F2[1], P[1]], color=k.M_RED, lw=2.4, zorder=6)
    ax.plot(*P, 'o', color=k.M_INK, ms=10, zorder=9)
    ax.annotate("$P$", P, xytext=(16, -4), textcoords="offset points", color=k.M_INK)
    # F1 关于切线 l 的对称点
    v = F1 - P; F1r = P + (2*np.dot(v, d)*d - v)
    ax.plot(*F1r, 'o', color=k.M_GRN, ms=9, zorder=8)
    ax.annotate("$F_1'$", F1r, xytext=(10, 8), textcoords="offset points", color=k.M_GRN)
    ax.plot([F1[0], F1r[0]], [F1[1], F1r[1]], color=k.M_GRN, ls=':', lw=1.5, zorder=5)
    ax.plot([F1r[0], F2[0]], [F1r[1], F2[1]], color=k.M_GRN, ls='--', lw=2.0, zorder=5)
    # 切线上另取一点 Q，示意距离和更大
    Q = P + 2.4*d
    ax.plot(*Q, 'o', color=k.M_SLATE, ms=7, zorder=8)
    ax.annotate("$Q$", Q, xytext=(-8, 16), ha="center",
                textcoords="offset points", color=k.M_SLATE)
    for Fp in (F1, F2):
        ax.plot([Fp[0], Q[0]], [Fp[1], Q[1]], color=k.M_SLATE, lw=1.2, ls='-.', zorder=4)
    ax.set_title("$P$ 是 $l$ 上距离和最小的点", color=k.M_INK)
    ax.set_aspect('equal'); ax.set_xlim(-7.6, 8.4)
    ax.set_ylim(-4.0, max(7.6, F1r[1] + 1.6)); ax.axis('off')
    return k.save_fig(fig, F("f13_optics_proof.png"))

# fig14 应用：碎石机 / 耳语廊 / 椭圆桌球
def fig14():
    fig, axes = plt.subplots(1, 3, figsize=(11.2, 3.8))
    a, b = 3.6, 2.4; c = np.sqrt(a*a-b*b)
    # 碎石机（半椭圆）
    ax = axes[0]
    t = np.linspace(np.pi, 2*np.pi, 300)
    ax.plot(a*np.cos(t), b*np.sin(t), color=k.M_INK, lw=3)
    for th in np.linspace(np.pi+0.35, 2*np.pi-0.35, 7):
        P = (a*np.cos(th), b*np.sin(th))
        ax.add_patch(FancyArrowPatch((-c, 0), P, arrowstyle='-', color=k.M_ACC, lw=1.2, alpha=0.8))
        ax.add_patch(FancyArrowPatch(P, (c, 0), arrowstyle='-|>', mutation_scale=9,
                                     color=k.M_RED, lw=1.2, alpha=0.8))
    ax.plot(-c, 0, '*', color=k.M_ACC, ms=18); ax.plot(c, 0, 'o', color=k.M_RED, ms=10)
    ax.set_title("体外碎石机", color=k.M_INK, pad=8)
    # 耳语廊
    ax = axes[1]
    ell(ax, a, b, color=k.M_INK, lw=2.4)
    for th in np.linspace(0.3, 2*np.pi-0.3, 9):
        P = (a*np.cos(th), b*np.sin(th))
        ax.plot([-c, P[0]], [0, P[1]], color=k.M_ACC2, lw=0.9, alpha=0.75)
        ax.plot([P[0], c], [P[1], 0], color=k.M_ACC2, lw=0.9, alpha=0.75)
    ax.plot(-c, 0, 'o', color=k.M_ACC, ms=11); ax.plot(c, 0, 'o', color=k.M_GRN, ms=11)
    ax.set_title("耳语廊", color=k.M_INK, pad=8)
    # 椭圆桌球
    ax = axes[2]
    ell(ax, a, b, color=k.M_INK, lw=2.8)
    ax.add_patch(Circle((c, 0), 0.22, fc=k.M_INK, ec=k.M_INK))
    th = 2.0; P = (a*np.cos(th), b*np.sin(th))
    ax.add_patch(FancyArrowPatch((-c, 0), P, arrowstyle='-|>', mutation_scale=12, color=k.M_ACC, lw=2))
    ax.add_patch(FancyArrowPatch(P, (c, 0), arrowstyle='-|>', mutation_scale=12, color=k.M_RED, lw=2))
    ax.plot(-c, 0, 'o', color=k.M_ACC, ms=11)
    ax.set_title("椭圆桌球", color=k.M_INK, pad=8)
    for ax in axes:
        ax.set_aspect('equal'); ax.set_xlim(-4.2, 4.2); ax.set_ylim(-3.0, 3.0); ax.axis('off')
    return k.save_fig(fig, F("f14_applications.png"))

# fig15 行星轨道（开普勒第一定律）
def fig15():
    fig, ax = k.new_fig(5.4, 4.0)
    a, e = 4.0, 0.55
    c = a*e; b = np.sqrt(a*a-c*c)
    t = np.linspace(0, 2*np.pi, 400)
    ax.plot(a*np.cos(t)-c, b*np.sin(t), color=k.M_ACC2, lw=2.4)  # 太阳在原点
    ax.plot(0, 0, 'o', color=k.M_ACC, ms=18, zorder=6)
    ax.text(0.4, -1.1, "太阳", ha="left", va="center", color=k.M_ACC)
    ax.plot(-2*c, 0, 'x', color=k.M_SLATE, ms=10, mew=2)
    ax.plot(a-c, 0, 'o', color=k.M_RED, ms=8)
    ax.text(a-c+0.4, 0.9, "近日点", ha="center", va="center", color=k.M_RED)
    ax.plot(-a-c, 0, 'o', color=k.M_RED, ms=8)
    ax.text(-a-c+0.2, -1.2, "远日点", ha="center", va="center", color=k.M_RED)
    th = 2.3; P = (a*np.cos(th)-c, b*np.sin(th))
    ax.plot(*P, 'o', color=k.M_INK, ms=9)
    ax.plot([0, P[0]], [0, P[1]], color=k.M_INK, lw=1.6, ls='--')
    ax.text(P[0], P[1]+0.85, "行星", ha="center", va="center", color=k.M_INK)
    ax.set_title("太阳位于一个焦点", color=k.M_INK)
    ax.set_aspect('equal'); ax.set_xlim(-8.6, 5.8); ax.set_ylim(-4.4, 5.0); ax.axis('off')
    return k.save_fig(fig, F("f15_kepler.png"))

# fig16 例题2：焦点三角形 / 焦点弦
def fig16():
    fig, axes = plt.subplots(1, 2, figsize=(7.5, 3.8))
    ax = axes[0]
    a, b = 5.0, 3.0; c = 4.0
    ell(ax, a, b, color=k.M_ACC2, lw=2.2)
    th = 1.15; P = (a*np.cos(th), b*np.sin(th))
    for s in (-1, 1):
        ax.plot(s*c, 0, 'o', color=k.M_RED, ms=7)
        ax.plot([s*c, P[0]], [0, P[1]], color=k.M_ACC, lw=2)
    ax.set_xticks([-5, 5]); ax.set_yticks([-3, 3])
    ax.plot([-c, c], [0, 0], color=k.M_RED, lw=2)
    ax.plot(*P, 'o', color=k.M_INK, ms=8)
    ax.add_patch(Arc(P, 1.5, 1.5, theta1=np.rad2deg(np.arctan2(-P[1], -c-P[0])),
                     theta2=np.rad2deg(np.arctan2(-P[1], c-P[0])), color=k.M_GRN, lw=1.8))
    ax.text(P[0]-0.2, P[1]-1.55, "$\\theta$", color=k.M_GRN)
    ax.annotate("$P$", P, xytext=(-4, 12), ha="center",
                textcoords="offset points", color=k.M_INK)
    ax.set_title("焦点三角形", color=k.M_INK, pad=30)
    k.style_axes(ax); ax.set_aspect('equal'); ax.set_xlim(-7.0, 7.0); ax.set_ylim(-4.6, 5.2)
    ax.set_xticks([-5, 5]); ax.set_yticks([-3, 3])

    ax = axes[1]
    ell(ax, a, b, color=k.M_ACC2, lw=2.2)
    # 过 F2 的弦，斜率 1
    m = 1.0
    A_, B_, C_ = b*b + a*a*m*m, -2*a*a*m*m*c, a*a*m*m*c*c - a*a*b*b
    xs = np.roots([A_, B_, C_])
    pts = [(x, m*(x-c)) for x in xs]
    ax.plot([p[0] for p in pts], [p[1] for p in pts], color=k.M_ACC, lw=2.4)
    ax.set_xticks([-5, 5]); ax.set_yticks([-3, 3])
    for p, nm, off in zip(pts, ("$A$", "$B$"), ((10, 6), (14, -16))):
        ax.plot(*p, 'o', color=k.M_INK, ms=8)
        ax.annotate(nm, p, xytext=off, textcoords="offset points", color=k.M_INK)
    ax.plot(c, 0, 'o', color=k.M_RED, ms=7); ax.plot(-c, 0, 'o', color=k.M_RED, ms=7)
    for p in pts:
        ax.plot([-c, p[0]], [0, p[1]], color=k.M_GRN, lw=1.2, ls=':')
    ax.set_title("焦点弦 $AB$", color=k.M_INK, pad=30)
    k.style_axes(ax); ax.set_aspect('equal'); ax.set_xlim(-7.0, 7.0); ax.set_ylim(-4.6, 5.2)
    ax.set_xticks([-5, 5]); ax.set_yticks([-3, 3])
    return k.save_fig(fig, F("f16_focal.png"))

# fig17 焦半径公式图
def fig17():
    fig, ax = k.new_fig(6.0, 4.2)
    a, b = 5.0, 3.0; c = 4.0
    ell(ax, a, b, color=k.M_ACC2, lw=2.2)
    for th in (0.7, 2.4, 4.2):
        P = (a*np.cos(th), b*np.sin(th))
        ax.plot(*P, 'o', color=k.M_INK, ms=7)
        ax.plot([-c, P[0]], [0, P[1]], color=k.M_ACC, lw=1.5)
        ax.plot([c, P[0]], [0, P[1]], color=k.M_GRN, lw=1.5)
    ax.plot([-c, c], [0, 0], 'o', color=k.M_RED, ms=8)
    ax.set_title("$r=a\\pm e x_0$（左加右减）", color=k.M_INK, pad=30)
    k.style_axes(ax); ax.set_xlim(-7.6, 7.6); ax.set_ylim(-4.8, 5.2); ax.set_aspect('equal')
    return k.save_fig(fig, F("f17_focal_radius.png"))

# fig18 例1 图：求标准方程
def fig18():
    fig, ax = k.new_fig(6.0, 4.0)
    a, b = 4.0, 2.0; c = np.sqrt(a*a-b*b)
    ell(ax, a, b, color=k.M_ACC2, lw=2.6)
    ax.plot([-c, c], [0, 0], 'o', color=k.M_RED, ms=9)
    P = (2.0, np.sqrt(3))
    ax.plot(*P, 'o', color=k.M_ACC, ms=9)
    ax.annotate("$P(2,\\sqrt{3})$", P, xytext=(10, 8), textcoords="offset points",
                color=k.M_ACC)
    for s in (-1, 1):
        ax.plot([s*c, P[0]], [0, P[1]], color=k.M_ACC, lw=1.6, ls='--')
    ax.set_title("$x^2/16+y^2/4=1$", color=k.M_INK, pad=30)
    k.style_axes(ax); ax.set_xlim(-6.4, 7.6); ax.set_ylim(-3.4, 4.2); ax.set_aspect('equal')
    ax.set_xticks([-4, 4]); ax.set_yticks([-2, 2])
    return k.save_fig(fig, F("f18_ex1.png"))

# fig19 变式：e 的范围（斜率/角度约束）
def fig19():
    fig, ax = k.new_fig(9.5, 4.2)
    es = np.linspace(0.01, 0.99, 300)
    ax.plot(es, es/np.sqrt(1-es**2), color=k.M_ACC2, lw=2.8,
            label=r"$c/b=e/\sqrt{1-e^2}$")
    ax.axhline(1, color=k.M_RED, ls='--', lw=1.6)
    ax.axvline(np.sqrt(2)/2, color=k.M_ACC, ls='--', lw=1.6)
    ax.plot(np.sqrt(2)/2, 1, 'o', color=k.M_ACC, ms=10)
    ax.annotate(r"$e=\sqrt{2}/2$ 时 $b=c$",
                (np.sqrt(2)/2, 1), xytext=(-190, 55), textcoords="offset points",
                color=k.M_ACC, arrowprops=dict(arrowstyle="->", color=k.M_ACC))
    ax.set_xlim(0, 1); ax.set_ylim(0, 4)
    k.style_axes(ax, xlabel="$e$", ylabel="$c/b$", origin=False)
    ax.set_title("$e$ 一定，形状就定了", color=k.M_INK, pad=30)
    ax.legend(frameon=False, loc="upper left", bbox_to_anchor=(0.02, 0.99))
    return k.save_fig(fig, F("f19_e_range.png"))

# fig20 板书提纲图（结构导图）
def fig20():
    fig, ax = k.new_fig(5.6, 3.1)
    ax.axis('off')
    boxes = [
        (0.3, 4.3, "定义 $r_1+r_2=2a$", k.M_ACC),
        (5.3, 4.3, "标准方程", k.M_ACC2),
        (0.3, 3.0, "$a^2=b^2+c^2$", k.M_GRN),
        (5.3, 3.0, "离心率 $e=c/a$", k.M_ACC),
        (0.3, 1.7, "第二定义", k.M_ACC2),
        (5.3, 1.7, "Dandelin 双球", k.M_RED),
        (0.3, 0.4, "光学性质", k.M_ACC),
        (5.3, 0.4, "焦半径·焦点弦", k.M_GRN),
    ]
    for x, y, t, col in boxes:
        ax.add_patch(plt.Rectangle((x, y), 4.4, 0.95, fc='white', ec=col, lw=2))
        ax.text(x+2.2, y+0.48, t, ha="center", va="center", color=k.M_INK)
    for x0, y0, x1, y1 in [(4.7, 4.78, 5.3, 4.78),
                           (2.5, 4.3, 2.5, 3.95), (7.5, 4.3, 7.5, 3.95),
                           (2.5, 3.0, 2.5, 2.65), (7.5, 3.0, 7.5, 2.65),
                           (2.5, 1.7, 2.5, 1.35), (7.5, 1.7, 7.5, 1.35)]:
        ax.annotate("", (x1, y1), (x0, y0),
                    arrowprops=dict(arrowstyle="->", color=k.M_SLATE, lw=1.6))
    ax.set_xlim(-0.1, 10.0); ax.set_ylim(0.1, 5.5)
    ax.set_title("本讲知识地图", color=k.M_INK)
    return k.save_fig(fig, F("f20_map.png"))

FIGS = [f() for f in (fig01, fig02, fig03, fig04, fig05, fig06, fig07, fig08, fig09,
                      fig10, fig11, fig12, fig13, fig14, fig15, fig16, fig17, fig18,
                      fig19, fig20)]
print("figures:", len(FIGS))

# ============================ PPT ============================
prs = k.new_deck()
FM = lambda n: F("_tex_%02d.png" % n)
_n = [0]
def tex():
    _n[0] += 1
    return FM(_n[0])

# 1 封面
k.title_slide(prs, "第 03 讲　椭圆：定义、方程与光学性质",
              "从一根绳子到行星轨道 —— 圆锥曲线的第一站",
              "数学名师课件包", "高二 · 选择性必修一 · 60 分钟")

# 2 学习目标
s = k.content_slide(prs, "学习目标与本讲地图", "导入")
k.bullets(s, [
    "理解椭圆的第一定义，说清 2a>2c 这一存在性条件",
    ("会画：用绳圈作图，解释“绳长即 2a”", 1),
    "独立完成标准方程的推导（两次平方的技术细节）",
    ("说清“为什么令 b²=a²−c²”——不是凑，是几何", 1),
    "掌握离心率 e=c/a 的几何意义与统一的焦点—准线定义",
    "理解 Dandelin 双球：圆锥截面为何满足焦点定义",
    "会用光学性质解释碎石机、耳语廊、椭圆桌球",
], y=1.55, w=6.3, size=17)
k.picture(s, FIGS[19], x=7.2, y=1.5, w=5.6)
k.callout(s, "主线：一根绳（定义）→ 一个方程（代数）→ 一个圆锥（几何）→ 一束光（应用）",
          x=0.85, y=6.15, w=6.3, h=0.9)

# 3 幕1
k.section_slide(prs, "第 1 幕 · 定义与建系", "一根绳子能画出什么？", "0–12 min")

# 4 绳圈实验
s = k.content_slide(prs, "情境：两枚图钉 + 一根绳", "3 min")
k.bullets(s, [
    "两枚图钉固定在 F₁、F₂；套上长为 2a 的绳圈",
    "笔尖绷紧绳子沿纸移动 —— 画出的是什么？",
    "笔尖 P 满足：|PF₁| + |PF₂| = 常数 = 2a",
    ("追问：2a = 2c 退化为线段，2a < 2c 无轨迹", 1),
], y=1.5, w=5.6, size=16)
k.picture(s, FIGS[0], x=6.6, y=1.5, w=6.1)
k.callout(s, "存在性条件必须写：2a > 2c，即 a > c > 0", x=0.85, y=5.9, w=5.6, h=0.9, kind="warn")

# 5 定义
s = k.content_slide(prs, "椭圆的第一定义", "概念")
k.formula(s, r"$|PF_1|+|PF_2| = 2a \quad (2a > |F_1F_2| = 2c > 0)$",
          x=1.0, y=1.55, w=11.3, out=tex())
k.bullets(s, [
    "F₁、F₂ 称为焦点；|F₁F₂| = 2c 称为焦距",
    "常数记作 2a（先写 2a，是为了后面方程漂亮）",
    "三种情形一网打尽：",
    ("2a > 2c ⟹ 椭圆；　2a = 2c ⟹ 线段 F₁F₂；　2a < 2c ⟹ 空集", 1),
], y=3.1, w=6.0, size=17)
k.picture(s, FIGS[1], x=6.9, y=3.0, w=6.0)

# 7 幕2
k.section_slide(prs, "第 2 幕 · 方程的推导", "把几何条件翻译成代数", "12–26 min")

# 8 建系
s = k.content_slide(prs, "第一步：建系（怎么建最省力）", "推演")
k.bullets(s, [
    "原点取 F₁F₂ 的中点，x 轴取 F₁F₂ 所在直线",
    ("好处：F₁(−c,0)、F₂(c,0) 对称，方程不含一次项", 1),
    "设 P(x, y) 为椭圆上任意一点，代入定义：",
], y=1.6, w=5.7, size=17)
k.formula(s, r"$\sqrt{(x+c)^2+y^2}+\sqrt{(x-c)^2+y^2}=2a$",
          x=0.85, y=4.45, w=5.9, out=tex())
k.picture(s, FIGS[2], x=6.9, y=1.5, w=6.0)
k.callout(s, "难点：两个根号。策略——移项，让每次只平方掉一个根号。",
          x=0.85, y=5.5, w=5.9, h=0.85, kind="warn")

# 9 推导（一次平方）
s = k.content_slide(prs, "第二步：第一次平方（移项是关键）", "推演")
k.formula(s, r"$\sqrt{(x+c)^2+y^2}=2a-\sqrt{(x-c)^2+y^2}$", x=1.0, y=1.5, w=8.6, out=tex())
k.formula(s, r"$(x+c)^2+y^2=4a^2-4a\sqrt{(x-c)^2+y^2}+(x-c)^2+y^2$",
          x=1.0, y=2.75, w=10.6, out=tex())
k.formula(s, r"$4cx-4a^2=-4a\sqrt{(x-c)^2+y^2}\;\Longrightarrow\; a\sqrt{(x-c)^2+y^2}=a^2-cx$",
          x=1.0, y=4.1, w=11.0, out=tex())
k.callout(s, "注意：左边根号非负 ⟹ a² − cx ≥ 0 自动成立（因 |x| ≤ a、c < a），不必额外讨论。",
          x=1.0, y=5.7, w=11.0, h=0.9, kind="note")

# 10 推导（二次平方）
s = k.content_slide(prs, "第三步：第二次平方，化到最简", "推演")
k.formula(s, r"$a^2\left[(x-c)^2+y^2\right]=(a^2-cx)^2$", x=1.0, y=1.5, w=7.4, out=tex())
k.formula(s, r"$a^2x^2-2a^2cx+a^2c^2+a^2y^2=a^4-2a^2cx+c^2x^2$", x=1.0, y=2.7, w=10.4, out=tex())
k.formula(s, r"$(a^2-c^2)x^2+a^2y^2=a^2(a^2-c^2)$", x=1.0, y=4.0, w=7.8, out=tex())
k.formula(s, r"$\dfrac{x^2}{a^2}+\dfrac{y^2}{a^2-c^2}=1$", x=1.0, y=5.2, w=5.0, out=tex())
k.callout(s, "两边同除 a²(a²−c²)。合法性：a > c ⟹ a² − c² > 0，除数非零。",
          x=6.6, y=5.35, w=5.9, h=1.0)

# 11 为什么 b²=a²−c²
s = k.content_slide(prs, "为什么令 b² = a² − c²？——不是凑，是几何", "关键")
k.bullets(s, [
    "a² − c² > 0，它有资格当一个平方数",
    "它到底是什么？令 x = 0 代入方程：y² = a² − c²",
    ("即椭圆与 y 轴交点为 (0, ±√(a²−c²))，这正是短半轴长！", 1),
    "所以 b = 短半轴长，b² = a² − c² 是几何事实，不是记号游戏",
    "顺手得到黄金关系：a² = b² + c²",
], y=1.6, w=6.0, size=17)
k.picture(s, FIGS[3], x=7.0, y=1.5, w=5.9)
k.callout(s, "口诀：a 是斜边（最长），b、c 是两条直角边。见 a 想斜边，永不记错。",
          x=0.85, y=5.9, w=6.0, h=1.0)

# 12 标准方程
s = k.content_slide(prs, "椭圆的标准方程", "结论")
k.formula(s, r"$\dfrac{x^2}{a^2}+\dfrac{y^2}{b^2}=1\quad (a>b>0),\qquad a^2=b^2+c^2$",
          x=0.9, y=1.5, w=11.5, out=tex())
k.picture(s, FIGS[4], x=1.6, y=3.1, w=10.2)

# 13 性质表
s = k.content_slide(prs, "几何性质速查（焦点在 x 轴）", "结论")
k.bullets(s, [
    "范围：|x| ≤ a，|y| ≤ b —— 椭圆是有界闭曲线",
    "对称性：关于 x 轴、y 轴、原点都对称（中心对称图形）",
    "顶点：A₁(−a,0)、A₂(a,0)、B₁(0,−b)、B₂(0,b)",
    "长轴长 2a，短轴长 2b，焦距 2c",
    "离心率 e = c/a ∈ (0, 1)",
    ("圆是 e = 0 的极限情形（c = 0，两焦点重合于圆心）", 1),
], y=1.6, w=6.2, size=17)
k.picture(s, FIGS[5], x=7.0, y=1.5, w=5.9)

# 14 幕3
k.section_slide(prs, "第 3 幕 · 离心率与统一定义", "一个数，管住所有圆锥曲线", "26–36 min")

# 15 离心率
s = k.content_slide(prs, "离心率 e = c/a：形状的唯一参数", "概念")
k.formula(s, r"$e=\dfrac{c}{a}=\sqrt{1-\dfrac{b^2}{a^2}}\;\in(0,1)\qquad\Longleftrightarrow\qquad \dfrac{b}{a}=\sqrt{1-e^2}$",
          x=0.9, y=1.5, w=11.4, out=tex())
k.picture(s, FIGS[18], x=1.9, y=3.1, w=9.5)

# 17 统一定义
s = k.content_slide(prs, "第二定义：焦点—准线，统一三种曲线", "高光")
k.formula(s, r"$\dfrac{|PF|}{d(P,\,l)}=e\qquad\Longrightarrow\qquad r=\dfrac{ep}{1+e\cos\theta}$",
          x=0.9, y=1.45, w=10.4, out=tex())
k.picture(s, FIGS[6], x=1.3, y=2.95, w=10.7)
k.callout(s, "e < 1 椭圆　|　e = 1 抛物线　|　e > 1 双曲线　——同一个式子，只改一个数。",
          x=1.3, y=6.35, w=10.7, h=0.75, kind="key")

# 18 椭圆的准线
s = k.content_slide(prs, "椭圆的准线：x = ±a²/c", "推演")
k.bullets(s, [
    "由 |PF₂| = a − ex₀，而 d(P, l₂) = a²/c − x₀",
    ("比值 = (a − ex₀) / (a²/c − x₀) = c/a = e　（与 P 无关！）", 1),
    "准线在椭圆外：a²/c > a（因 c < a）",
    "两条准线关于 y 轴对称，与两个焦点一一配对",
], y=1.6, w=5.7, size=17)
k.picture(s, FIGS[7], x=6.6, y=1.45, w=6.2)
k.callout(s, "焦准距：a²/c − c = b²/c。半通径（过焦点垂直长轴的半弦）= b²/a = ep。",
          x=0.85, y=5.0, w=5.7, h=1.0, kind="note")

# 19 幕4
k.section_slide(prs, "第 4 幕 · Dandelin 双球", "为什么“圆锥的切面”就是“到两点距离和为定值”？", "36–46 min")

# 20 圆锥截面
s = k.content_slide(prs, "圆锥曲线之名从何而来", "背景")
k.bullets(s, [
    "阿波罗尼奥斯（约公元前 200 年）：用一个圆锥切出全部三种曲线",
    "但古人的定义是「截线」，我们的定义是「距离和」——凭什么是同一件事？",
    "1822 年，Dandelin 用两个内切球给出一个惊艳的证明",
], y=1.6, w=5.4, size=16)
k.picture(s, FIGS[8], x=6.3, y=1.4, w=6.5)
k.callout(s, "本讲高光：接下来 8 分钟，我们把 2000 年的鸿沟一步跨过去。",
          x=0.85, y=4.7, w=5.4, h=1.0, kind="key")

# 21 引理
s = k.content_slide(prs, "预备引理：球外一点的切线长相等", "证明")
k.full_picture(s, FIGS[10], y=1.5, w=11.0)

# 22 Dandelin 主图
s = k.content_slide(prs, "Dandelin 双球：把两个球塞进圆锥", "证明")
k.picture(s, FIGS[9], x=0.6, y=1.45, w=7.2)
k.bullets(s, [
    "在割平面两侧各放一球，同时与圆锥侧面、割平面相切",
    "割平面的两个切点，就是 F₁、F₂",
    "对截线上任一点 P，过 P 作母线，交两球于 T₁、T₂",
    "|PF₁| = |PT₁|，|PF₂| = |PT₂|（切线长相等）",
    "相加：|PF₁| + |PF₂| = |T₁T₂| = 常数 ✓",
], x=8.0, y=1.7, w=4.9, size=15)
k.callout(s, "|T₁T₂| 是两个切圆之间的母线段：由旋转对称性，它与 P 无关。",
          x=8.0, y=5.9, w=4.9, h=1.1, kind="key")

# 23 证明收束
s = k.content_slide(prs, "定理：圆锥的斜截线正是椭圆", "证明")
k.formula(s, r"$|PF_1|+|PF_2|=|PT_1|+|PT_2|=|T_1T_2|=\text{const}=2a$",
          x=0.9, y=1.5, w=11.4, out=tex())
k.bullets(s, [
    "两个定义完全等价：截线定义 ⟺ 焦点距离和定义",
    "推论：割平面倾角越大（越接近平行母线），双球越“瘦长”，e 越接近 1",
    "同一套方法可证抛物线（一个球 + 准线）与双曲线（两球分居两支）",
], y=3.0, w=11.4, size=17)
k.callout(s, "这就是数学的美：一个纯几何的构造，把代数定义和古典定义焊死在一起。",
          x=0.9, y=5.9, w=11.4, h=0.95)

# 24 幕5
k.section_slide(prs, "第 5 幕 · 光学性质与例题", "从一束光到三种工程", "46–60 min")

# 25 光学性质
s = k.content_slide(prs, "光学性质：F₁ 出发，必回 F₂", "性质")
k.full_picture(s, FIGS[11], y=1.45, w=10.2)

# 26 光学性质证明
s = k.content_slide(prs, "为什么？切线上距离和最小", "证明")
k.picture(s, FIGS[12], x=0.5, y=1.45, w=7.6)
k.bullets(s, [
    "椭圆上的点 = 使 |PF₁|+|PF₂| = 2a 的点",
    "切线 l 上除 P 外的点都在椭圆外 ⟹ 距离和 > 2a",
    "故 P 是 l 上使距离和最小的点",
    "由“将军饮马”：最小 ⟺ F₁ 的对称点 F₁′、P、F₂ 三点共线",
    "⟹ 入射角 = 反射角。证毕",
], x=8.3, y=1.7, w=4.6, size=15)

# 27 应用
s = k.content_slide(prs, "一条性质，三种工程", "应用")
k.full_picture(s, FIGS[13], y=1.5, w=11.2)
k.callout(s, "碎石机：声源置 F₁，结石置 F₂ | 耳语廊：北京天坛回音壁 | 椭圆桌球：F₁ 击球必进 F₂ 洞",
          x=1.05, y=5.75, w=11.2, h=0.85, kind="note")

# 28 行星轨道
s = k.content_slide(prs, "开普勒：天体也在用椭圆", "应用")
k.picture(s, FIGS[14], x=0.55, y=1.5, w=7.5)
k.bullets(s, [
    "开普勒第一定律：行星绕日作椭圆运动，太阳在一个焦点",
    "近日点 r = a − c，远日点 r = a + c",
    "地球 e ≈ 0.017（肉眼看是圆）",
    "哈雷彗星 e ≈ 0.967（极扁，76 年一遇）",
    "另一个焦点是空的 —— 这正是牛顿引力的几何签名",
], x=8.2, y=1.75, w=4.7, size=15)

# 29 例1
s = k.content_slide(prs, "例 1　求标准方程（待定系数）", "例题")
k.picture(s, FIGS[17], x=6.9, y=1.5, w=6.0)
k.bullets(s, [
    "题：焦点在 x 轴上，a = 2b，且过点 P(2, √3)，求标准方程与 e。",
    "设 x²/(4b²) + y²/b² = 1（a = 2b 直接代入，省一个未知数）",
    "代 P：4/(4b²) + 3/b² = 1 ⟹ 1/b² + 3/b² = 1 ⟹ b² = 4",
    "故 a² = 16，方程 x²/16 + y²/4 = 1",
    "c² = 16 − 4 = 12，c = 2√3，e = c/a = √3/2",
], y=1.65, w=6.1, size=15)
k.callout(s, "方法论：能用一个参数就别用两个。条件先化简，再代点。",
          x=0.85, y=5.85, w=6.1, h=0.9)

# 30 例2 + 例3
s = k.content_slide(prs, "例 2 焦点三角形　例 3 焦点弦", "例题")
k.picture(s, FIGS[15], x=0.5, y=1.45, w=7.5)
k.bullets(s, [
    "例2：x²/25 + y²/9 = 1，P 在椭圆上，∠F₁PF₂ = 60°，求 S△PF₁F₂。",
    ("r₁+r₂ = 10；余弦定理 r₁²+r₂²−r₁r₂ = 64", 1),
    ("(r₁+r₂)² − 3r₁r₂ = 64 ⟹ r₁r₂ = 12", 1),
    ("S = ½·12·sin60° = 3√3　【通法：b²tan(θ/2) = 9·tan30° = 3√3 ✓】", 1),
    "例3：过 F₂ 的弦 AB，求 △ABF₁ 的周长。",
    ("周长 = (|AF₁|+|AF₂|) + (|BF₁|+|BF₂|) = 2a + 2a = 4a = 20", 1),
], x=8.0, y=1.7, w=4.9, size=13)

# 31 变式
s = k.content_slide(prs, "变式演练（学生板演）", "变式")
k.picture(s, FIGS[16], x=7.0, y=1.5, w=6.0)
k.bullets(s, [
    "变式 1（求 e）：椭圆短轴端点 B 与两焦点构成 ∠F₁BF₂ = 90°，求 e。",
    ("提示：等腰直角 ⟹ b = c ⟹ a² = 2c² ⟹ e = √2/2", 1),
    "变式 2（焦半径）：x²/25 + y²/16 = 1 上一点 P 到 F₁(−3,0) 距离为 4，",
    ("求 P 的横坐标，并求 P 到右准线的距离。", 1),
    ("r₁ = a + ex₀ = 5 + (3/5)x₀ = 4 ⟹ x₀ = −5/3", 1),
    ("d = |PF₂| / e = (2a − 4) / (3/5) = 6 / 0.6 = 10", 1),
], y=1.65, w=6.1, size=14)

# 32 小结
s = k.content_slide(prs, "课堂小结：四个层次", "小结")
k.picture(s, FIGS[19], x=6.9, y=1.5, w=6.1)
k.bullets(s, [
    "定义层：|PF₁|+|PF₂| = 2a > 2c —— 一根绳",
    "代数层：x²/a² + y²/b² = 1，a² = b² + c² —— 两次平方",
    "统一层：e = |PF|/d，把三种曲线收进一个公式",
    "几何层：Dandelin 双球，圆锥截线 = 焦点定义",
    "应用层：光学性质，F₁ ⟶ F₂",
], y=1.7, w=6.0, size=16)
k.callout(s, "一句话带走：椭圆是「到两定点距离和恒定」的点集，e = c/a 决定它的胖瘦，光从一个焦点出发必回另一个焦点。",
          x=0.85, y=5.6, w=6.0, h=1.4, kind="key")

# 33 作业
s = k.content_slide(prs, "分层作业", "作业")
k.bullets(s, [
    "【必做 · A 层】",
    ("1. 求焦点(0,±4)、a=5 的椭圆标准方程", 1),
    ("2. x²/16+y²/7=1，求 a,b,c,e、顶点与准线", 1),
    ("3. 教材 P48 练习 1–4", 1),
    "【提高 · B 层】",
    ("4. 椭圆上一点 P 到两准线距离之比为 2:3，求 e 及 P 的横坐标", 1),
    ("5. 过 F₁ 的弦垂直于长轴，其长为 b²/a·2（通径），证明之", 1),
    "【拓展 · C 层】",
    ("6. 仿 Dandelin 双球，写出抛物线的双球（单球+准平面）证明思路", 1),
    ("7. 查资料：椭圆桌球「Alhazen 问题」为何不能尺规作图？", 1),
], y=1.55, w=6.6, size=14)
k.picture(s, FIGS[14], x=7.5, y=1.9, w=5.4)

# 34 板书
s = k.content_slide(prs, "板书设计", "板书")
k.full_picture(s, FIGS[19], y=1.5, w=10.4)

k.save(prs, os.path.join(OUT, "03_椭圆的定义与性质.pptx"))
print("slides:", len(prs.slides.__iter__.__self__._sldIdLst))

# ============================ 教案 DOCX ============================
from docx import Document
from docx.shared import Pt as DPt, Cm, RGBColor as DRGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Noto Serif CJK SC'; st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), 'Noto Serif CJK SC')

def H(t, lv=1):
    p = doc.add_heading(t, lv)
    for r in p.runs:
        r.font.name = 'Noto Serif CJK SC'
        r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
        r._element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), 'Noto Serif CJK SC')
    return p

def P(t, b=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = b
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), 'Noto Serif CJK SC')
    return p

t = doc.add_heading('教学设计　第 03 讲　椭圆：定义、方程与光学性质', 0)
for r in t.runs:
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(__import__('docx').oxml.ns.qn('w:eastAsia'), 'Noto Serif CJK SC')

H('一、基本信息', 1)
info = [('课题', '椭圆的定义、标准方程与光学性质'), ('课时', '1 课时（60 分钟）'),
        ('课型', '新授课（概念建构 + 推理论证 + 应用）'),
        ('教材', '人教 A 版 选择性必修第一册　第三章 圆锥曲线的方程 · 3.1 椭圆'),
        ('授课对象', '高二年级'), ('配套课件', '03_椭圆的定义与性质.pptx（32 页）')]
tb = doc.add_table(rows=0, cols=2); tb.style = 'Table Grid'
for a, b in info:
    c = tb.add_row().cells; c[0].text = a; c[1].text = b
    c[0].width = Cm(3.2)

H('二、教材分析', 1)
P('椭圆是圆锥曲线的第一课，处于「解析几何思想」由圆推广到一般二次曲线的关键节点。'
  '教材以「用绳圈作图」为起点，经由坐标法把几何条件 |PF₁|+|PF₂|=2a 翻译为代数方程，'
  '再由方程反推几何性质，完整演示「几何 → 代数 → 几何」的解析几何研究范式。'
  '本节的推导（两次平方去根号）是学生首次遇到的较复杂的方程化简，其技术处理（移项使每次只平方一个根号）'
  '和结构洞察（令 b²=a²−c² 的几何来源）是承前启后的关键。'
  '本设计在教材基础上补充两条高位内容：一是「焦点—准线」统一定义，为后续抛物线、双曲线搭好统一框架；'
  '二是 Dandelin 双球，回答「为什么圆锥的截线满足焦点定义」这一教材回避但学生必然会问的问题，'
  '是本讲的思维高光。')

H('三、学情分析', 1)
P('已有基础：学生已掌握圆的标准方程与坐标法的基本流程，具备两点间距离公式、'
  '根式化简、余弦定理等工具，对「用方程研究曲线」有初步经验。')
P('可能困难：(1) 两次平方的化简过程冗长，易在符号与移项上出错；'
  '(2) 对 b²=a²−c² 停留在「死记」层面，不理解 b 的几何身份；'
  '(3) 离心率被当作一个公式，缺乏「e 是形状的唯一参数」的直观；'
  '(4) 对圆锥曲线之「圆锥」二字毫无感觉，割裂了历史与代数。')
P('对策：绳圈实验建立直观 → 板书逐步推导并强调移项策略 → x=0 代入让 b 自己现身 → '
  '多子图看 e 的形变 → Dandelin 双球用 2D 剖面 + 切线长引理讲清。')

H('四、教学目标（三维）', 1)
P('【知识与技能】', True)
P('1. 理解椭圆的定义及存在性条件 2a>2c，能判别三种退化情形；\n'
  '2. 独立完成标准方程的推导，掌握两次平方的化简技术；\n'
  '3. 掌握 a²=b²+c² 与离心率 e=c/a，能由方程读出全部几何量；\n'
  '4. 掌握焦半径公式 r=a±ex₀ 与焦点弦、焦点三角形的常用结论。')
P('【过程与方法】', True)
P('1. 经历「实验—猜想—建系—推导—反思」的完整解析几何流程；\n'
  '2. 通过焦点—准线定义体会数学的统一性；\n'
  '3. 通过 Dandelin 双球体验「构造性证明」的力量，发展逻辑推理与直观想象素养。')
P('【情感态度与价值观】', True)
P('通过开普勒定律、体外碎石机、耳语廊等实例，感受数学模型对自然与工程的解释力；'
  '通过 Dandelin 的证明，体会数学之美与思维的震撼。')

H('五、重点与难点', 1)
P('教学重点：椭圆的定义、标准方程的推导与几何性质（离心率、光学性质）。', True)
P('教学难点：(1) 标准方程推导中两次平方的化简技术与 b²=a²−c² 的几何解释；'
  '(2) Dandelin 双球证明的空间想象与逻辑链条。', True)
P('突破策略：难点(1) 用「移项—平方—再移项—再平方」四步板书法，每步标注合法性；'
  '难点(2) 降维处理——只画轴剖面，先证平面几何引理（球外一点切线长相等），再一步升维。')

H('六、教法与学法', 1)
P('教法：实验探究法（绳圈作图）、启发讲授法（方程推导）、几何直观法（多子图与剖面图）、'
  '问题链驱动（为什么是 2a？为什么令 b²=a²−c²？为什么截线就是椭圆？）。')
P('学法：动手实验—观察归纳—代数演算—类比迁移—板演互评。')

H('七、教学准备', 1)
P('图钉、细绳、硬纸板（每组一套，供绳圈作图）；PPT 课件 32 页；椭圆桌球演示视频（可选）；'
  '几何画板/GeoGebra 动态演示 e 变化（可选）；学案与分层作业单。')

H('八、教学过程（60 分钟时间轴）', 1)
rows = [
    ("0–3", "情境导入", "PPT2–4",
     "发放图钉与绳圈，示范作图；提问：笔尖画出的是什么？绳子的长度扮演什么角色？",
     "两人一组动手作图，观察笔尖轨迹，记录「两段绳长之和不变」。",
     "以身体经验建立定义直观，让 2a 从「常数」变成「摸得到的绳长」。"),
    ("3–8", "形成定义", "PPT4–5",
     "追问：绳子恰好等于图钉间距会怎样？更短呢？板书三种情形。给出椭圆的严格定义。",
     "尝试拉紧绳子体会退化；口述定义，指出必须写 2a>2c。",
     "在「反例」中锤炼定义的严谨性，落实存在性条件这一易漏点。"),
    ("8–12", "感受形状", "PPT5",
     "固定绳长，移动图钉，投影三态图；引导用 c/a 描述胖瘦。",
     "观察并猜想：决定形状的是 c/a 而非 a 或 c 单独的值。",
     "为离心率埋伏笔，把「形状参数」的概念提前孕育。"),
    ("12–15", "建系与列式", "PPT6–7",
     "提问：怎样建系最省力？强调对称建系。写出含两个根号的方程，点明难点。",
     "自主完成建系并写出根式方程，同桌互查。",
     "落实坐标法第一步，体会「好的建系可以简化一半计算」。"),
    ("15–21", "推导：两次平方", "PPT8–9",
     "板书逐步演算：移项—平方—整理—再平方—再整理；每步标红「为什么可以这样做」。",
     "同步在草稿纸上演算，随教师节奏核对；指出教师故意留下的一处符号陷阱。",
     "难点突破。让学生亲历长演算，培养运算素养与耐心。"),
    ("21–24", "b²=a²−c² 的来历", "PPT10",
     "提问：a²−c² 是什么？引导令 x=0 代入，得 y²=a²−c²。揭示 b 就是短半轴。",
     "代入计算，恍然：b 不是凭空定义的记号，是椭圆与 y 轴交点的纵坐标。",
     "把「记忆」转化为「理解」，同时自然得到 a²=b²+c² 的直角三角形模型。"),
    ("24–26", "标准方程与性质", "PPT11–12",
     "给出两种焦点位置的标准形式；带学生速查范围、对称性、顶点、轴长。",
     "完成「看方程说性质」的口头快问快答（如 x²/25+y²/9=1）。",
     "形成结构化知识块，为解题建立检索路径。"),
    ("26–30", "离心率", "PPT13–14",
     "定义 e=c/a，推出 b/a=√(1−e²)；展示 e 从 0 到 0.9 的六联图。",
     "观察并归纳：e→0 越圆，e→1 越扁；e 与椭圆大小无关，只管形状。",
     "把前面孕育的直观正式命名，落实「离心率是形状的唯一参数」。"),
    ("30–36", "焦点—准线统一定义", "PPT15–16",
     "给出 |PF|/d=e；用极坐标统一式展示 e<1、=1、>1 三种曲线；推导椭圆准线 x=±a²/c。",
     "验证：由 r₂=a−ex₀ 与 d=a²/c−x₀ 相除，得到与 P 无关的常数 e。",
     "建立统一观，为后续抛物线、双曲线课时节省认知成本，凸显数学的简洁美。"),
    ("36–38", "圆锥曲线之名", "PPT17–18",
     "回到阿波罗尼奥斯：一个圆锥切出三种曲线。抛出核心问题：截线定义与距离和定义凭什么等价？",
     "产生认知冲突：这两件事看起来毫无关系。",
     "制造问题张力，为高光证明做情绪铺垫。"),
    ("38–41", "引理：切线长相等", "PPT19",
     "先降维：球外一点的两条切线长相等（勾股定理即得）。板书证明。",
     "独立在草稿纸上用勾股定理完成引理证明。",
     "把三维难点拆成一个二年级学生都能懂的平面引理，是突破的钥匙。"),
    ("41–46", "Dandelin 双球", "PPT20–21",
     "展示轴剖面图：割平面两侧各塞一球；标出 F₁、F₂、T₁、T₂；引导完成五步推理。",
     "跟随推理，独立说出 |PF₁|+|PF₂|=|T₁T₂| 为何是常数（旋转对称性）。",
     "本讲高光。让学生亲眼看到「两千年鸿沟一步跨过」，形成强记忆锚点。"),
    ("46–50", "光学性质及其证明", "PPT23–24",
     "演示反射光路图；用「切线上距离和最小 + 将军饮马」证明入射角=反射角。",
     "回忆将军饮马模型，补全共线论证；理解反证的逻辑。",
     "把新知挂到旧模型上（最短路径），实现知识迁移。"),
    ("50–53", "应用与开普勒", "PPT25–26",
     "介绍碎石机、耳语廊、椭圆桌球；讲开普勒第一定律与地球/哈雷彗星的 e。",
     "解释：为什么碎石机把结石放在另一个焦点上？",
     "回应「学这个有什么用」，体现数学建模与科学价值。"),
    ("53–58", "例题与变式板演", "PPT27–29",
     "例1 待定系数求方程；例2 焦点三角形面积（余弦定理 + b²tan(θ/2) 通法）；"
     "例3 焦点弦周长=4a。点两名学生板演变式 1、2。",
     "两人板演，其余在座位完成；同伴互评，指出板演中的规范性问题。",
     "把定义、方程、离心率、焦半径四类工具在真题中串联，形成解题范式。"),
    ("58–60", "小结与作业", "PPT30–32",
     "带学生用知识地图复盘四个层次；布置 A/B/C 三层作业；预告下一讲双曲线。",
     "用一句话概括本讲；对照地图自查掌握薄弱处。",
     "结构化收束，分层作业照顾差异，预告制造期待。"),
]
tb = doc.add_table(rows=1, cols=6); tb.style = 'Table Grid'; tb.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ["时间(min)", "环节", "对应PPT", "教师活动", "学生活动", "设计意图"]
for i, h in enumerate(hdr):
    c = tb.rows[0].cells[i]; c.text = h
    for p in c.paragraphs:
        for r in p.runs: r.bold = True; r.font.size = DPt(9.5)
widths = [Cm(1.7), Cm(2.0), Cm(1.7), Cm(6.2), Cm(4.8), Cm(5.0)]
for r_ in rows:
    cells = tb.add_row().cells
    for i, v in enumerate(r_):
        cells[i].text = v
        for p in cells[i].paragraphs:
            for run in p.runs: run.font.size = DPt(9)
for row in tb.rows:
    for i, c in enumerate(row.cells): c.width = widths[i]
P('合计：3+5+4+3+6+3+2+4+6+2+3+5+4+3+5+2 = 60 分钟。', True)

H('九、板书设计', 1)
P('主板（左）：椭圆定义 → 建系 → 根式方程 → 两次平方全过程（保留不擦，供学生回看）。')
P('主板（中）：标准方程 x²/a²+y²/b²=1、a²=b²+c² 与特征直角三角形图；离心率 e=c/a。')
P('副板（右）：Dandelin 五步推理框；光学性质示意；例题演算区（可擦）。')
P('板书结构图见课件第 32 页（figures/f20_map.png）。')

H('十、分层作业', 1)
P('A 层（必做·基础）：教材 P48 练习 1–4；求焦点 (0,±4)、a=5 的标准方程；'
  '由 x²/16+y²/7=1 求 a,b,c,e、顶点与准线。')
P('B 层（提高）：椭圆上一点到两准线距离之比为 2:3，求 e 与该点横坐标；证明通径长为 2b²/a；'
  '已知 ∠F₁PF₂=90° 求 e 的取值范围。')
P('C 层（拓展·选做）：仿 Dandelin 双球写出抛物线的证明思路；'
  '查阅椭圆桌球的 Alhazen 问题，说明其不可尺规作图的原因；'
  '用 GeoGebra 制作 e 连续变化的动画并写 200 字观察报告。')

H('十一、教学反思（课后填写）', 1)
for t_ in ['1. 绳圈实验的组织效率与学生参与度：',
           '2. 两次平方推导中学生的主要错误类型（符号 / 移项 / 漏条件）：',
           '3. Dandelin 双球的接受度：有多少学生能独立复述五步推理？',
           '4. 时间分配是否需要调整（哪一环节超时 / 提前）：',
           '5. 分层作业的完成质量与 C 层选做率：',
           '6. 下次改进点：']:
    P(t_)
    P('　')
    P('　')

doc.save(os.path.join(OUT, "教案_03_椭圆的定义与性质.docx"))
print("docx ok")
print("OUT:", OUT)
