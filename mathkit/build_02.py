# -*- coding: utf-8 -*-
"""第02讲：定积分与微积分基本定理 —— 60分钟课件包生成脚本"""
import sys, os
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon, FancyArrowPatch

OUT = '/mnt/c/Users/mi/Desktop/数学名师课件包/02_定积分与微积分基本定理'
FIG = os.path.join(OUT, 'figures')
TMP = os.path.join(FIG, '_tex')
os.makedirs(FIG, exist_ok=True); os.makedirs(TMP, exist_ok=True)
P = lambda n: os.path.join(FIG, n)
T = lambda n: os.path.join(TMP, n)

# ============================ 图表 ============================
# 画布宽度 == 该图在幻灯片上的展示宽度（deckkit 禁止缩小贴图，否则图内字 <20pt）
figs = []

def reg(path):
    figs.append(path); return path

def noticks(ax, xt=None, yt=None):
    ax.set_xticks(xt if xt is not None else [])
    ax.set_yticks(yt if yt is not None else [])

# --- fig01 两个实际背景：曲边梯形 + 变力做功（展示宽 5.2）---
fig, axes = plt.subplots(1, 2, figsize=(5.2, 3.4))
ax = axes[0]
x = np.linspace(0, 3.2, 400); f = 0.35 * x**2 + 0.6
ax.plot(x, f, color=k.M_ACC2, lw=2.6)
xs = np.linspace(0.5, 2.8, 200)
ax.fill_between(xs, 0, 0.35*xs**2+0.6, color=k.M_ACC, alpha=0.28)
ax.plot([0.5, 0.5], [0, 0.35*0.5**2+0.6], color=k.M_INK, lw=1.4, ls='--')
ax.plot([2.8, 2.8], [0, 0.35*2.8**2+0.6], color=k.M_INK, lw=1.4, ls='--')
ax.text(1.65, 1.05, '$S=?$', ha='center', color=k.M_INK)
ax.text(0.5, -1.05, '$a$', ha='center', color=k.M_RED)
ax.text(2.8, -1.05, '$b$', ha='center', color=k.M_RED)
ax.set_xlim(-0.4, 3.6); ax.set_ylim(-1.5, 5.2)
k.style_axes(ax, 'x', 'y'); noticks(ax)
ax.set_title('曲边梯形', color=k.M_INK, pad=16)
ax = axes[1]
s = np.linspace(0, 3.0, 400); F = 1.2 * s
ax.plot(s, F, color=k.M_GRN, lw=2.6)
ax.fill_between(s, 0, F, color=k.M_GRN, alpha=0.20)
ax.text(1.55, 1.2, '$W=?$', ha='center', color=k.M_INK)
ax.set_xlim(-0.4, 3.6); ax.set_ylim(-1.5, 5.2)
k.style_axes(ax, 'x', 'F'); noticks(ax)
ax.set_title('变力做功', color=k.M_GRN, pad=16)
k.save_fig(fig, reg(P('fig01_background.png')))

# --- fig02 四步法：分割—近似（展示宽 7.0）---
fig, ax = k.new_fig(7.0, 4.5)
f = lambda t: 0.35*t**2 + 0.6
x = np.linspace(0.2, 3.4, 400)
ax.plot(x, f(x), color=k.M_ACC2, lw=2.6, zorder=5)
a, b, n = 0.5, 2.9, 6
pts = np.linspace(a, b, n+1); dx = (b-a)/n
for i in range(n):
    ax.bar(pts[i], f(pts[i]), width=dx, align='edge', color=k.M_ACC, alpha=0.30,
           edgecolor=k.M_ACC, lw=1.2, zorder=2)
ax.plot(pts, [0]*len(pts), 'o', color=k.M_RED, ms=6, zorder=6)
ax.annotate('$\\Delta x$', xy=(pts[1]+dx/2, 0), xytext=(0.55, -1.35),
            color=k.M_RED, ha='center',
            arrowprops=dict(arrowstyle='->', color=k.M_RED, lw=1.4))
ax.annotate('以直代曲', xy=(pts[4]+dx/2, f(pts[4])), xytext=(3.05, 4.2),
            color=k.M_GRN, ha='center',
            arrowprops=dict(arrowstyle='->', color=k.M_GRN, lw=1.6))
ax.text(a, -0.75, '$a$', ha='center', color=k.M_RED)
ax.text(b, -0.75, '$b$', ha='center', color=k.M_RED)
ax.set_xlim(-0.25, 4.3); ax.set_ylim(-1.9, 5.6)
k.style_axes(ax, 'x', 'y'); noticks(ax)
ax.set_title('分割 · 以直代曲', color=k.M_INK, pad=10)
k.save_fig(fig, reg(P('fig02_partition.png')))

# --- fig03 黎曼和 n=4/8/16/64（展示宽 8.6）---
fig, axes = plt.subplots(2, 2, figsize=(8.0, 3.9))
f = lambda t: t**2
for ax, n in zip(axes.ravel(), [4, 8, 16, 64]):
    x = np.linspace(0, 1, 300)
    ax.plot(x, f(x), color=k.M_ACC2, lw=2.2, zorder=5)
    pts = np.linspace(0, 1, n+1); dx = 1.0/n
    ax.bar(pts[:-1], f(pts[1:]), width=dx, align='edge', color=k.M_ACC,
           alpha=0.32, edgecolor=k.M_ACC, lw=0.7 if n < 32 else 0.2, zorder=2)
    S = np.sum(f(pts[1:]) * dx)
    ax.set_title('$n=%d$　$S_n=%.3f$' % (n, S), color=k.M_INK, pad=6)
    ax.set_xlim(0, 1.02); ax.set_ylim(0, 1.05)
    k.style_axes(ax, '', '', origin=False); noticks(ax)
fig.suptitle('$S_n\\;\\longrightarrow\\;1/3$', color=k.M_GRN)
k.save_fig(fig, reg(P('fig03_riemann_seq.png')))

# --- fig04 上和下和夹逼（展示宽 9.6）---
fig, axes = plt.subplots(1, 2, figsize=(9.6, 4.3))
f = lambda t: t**2
n = 8; pts = np.linspace(0, 1, n+1); dx = 1.0/n
x = np.linspace(0, 1, 300)
for ax, mode in zip(axes, ['low', 'up']):
    ax.plot(x, f(x), color=k.M_ACC2, lw=2.4, zorder=5)
    h = f(pts[:-1]) if mode == 'low' else f(pts[1:])
    c = k.M_GRN if mode == 'low' else k.M_RED
    ax.bar(pts[:-1], h, width=dx, align='edge', color=c, alpha=0.28,
           edgecolor=c, lw=1.0, zorder=2)
    S = np.sum(h*dx)
    ax.set_title(('下和 $s_8=%.3f$' if mode == 'low' else '上和 $S_8=%.3f$') % S,
                 color=c, pad=8)
    ax.set_xlim(0, 1.03); ax.set_ylim(0, 1.08)
    k.style_axes(ax, 'x', 'y', origin=False); noticks(ax)
fig.suptitle('$s_n\\leq\\int_0^1\\!x^2dx\\leq S_n,\\quad S_n-s_n=1/n\\to 0$', color=k.M_INK)
k.save_fig(fig, reg(P('fig04_upper_lower.png')))

# --- fig05 上下和收敛曲线（展示宽 6.9）---
fig, ax = k.new_fig(6.9, 4.2)
ns = np.arange(2, 61)
low = np.array([np.sum((np.linspace(0,1,n+1)[:-1])**2)/n for n in ns])
up = np.array([np.sum((np.linspace(0,1,n+1)[1:])**2)/n for n in ns])
ax.plot(ns, up, '-', color=k.M_RED, lw=2.4, label='上和 $S_n$')
ax.plot(ns, low, '-', color=k.M_GRN, lw=2.4, label='下和 $s_n$')
ax.fill_between(ns, low, up, color=k.M_ACC, alpha=0.25)
ax.axhline(1/3, color=k.M_INK, ls='--', lw=1.6)
ax.text(62, 1/3, '$1/3$', color=k.M_INK, va='center', ha='left')
ax.set_xlim(0, 61); ax.set_ylim(0.12, 0.68)
k.style_axes(ax, 'n', '', origin=False)
ax.set_xticks([0, 20, 40, 60]); ax.set_yticks([0.2, 0.4, 0.6])
ax.legend(frameon=False, loc='upper right')
ax.set_title('两把钳子夹出同一个数', color=k.M_INK, pad=10)
k.save_fig(fig, reg(P('fig05_converge.png')))

# --- fig06 定义硬算 ∫₀¹x²dx（展示宽 9.4）---
fig, ax = k.new_fig(9.4, 4.6)
n = 10; pts = np.linspace(0, 1, n+1); dx = 1.0/n
x = np.linspace(0, 1.05, 300)
ax.plot(x, x**2, color=k.M_ACC2, lw=2.6, zorder=5)
ax.bar(pts[:-1], (pts[1:])**2, width=dx, align='edge', color=k.M_ACC, alpha=0.30,
       edgecolor=k.M_ACC, lw=1.0, zorder=2)
ax.annotate('$\\left(\\dfrac{i}{n}\\right)^{2}\\cdot\\dfrac{1}{n}$',
            xy=(pts[7]+dx/2, pts[8]**2), xytext=(0.30, 0.95),
            color=k.M_RED, ha='center',
            arrowprops=dict(arrowstyle='->', color=k.M_RED, lw=1.4))
ax.set_xlim(0, 1.30); ax.set_ylim(0, 1.35)
k.style_axes(ax, 'x', 'y', origin=False)
ax.set_xticks([0, 0.5, 1.0]); ax.set_yticks([0, 0.5, 1.0])
ax.set_title('例1：$\\int_0^1 x^2\\,dx$ 的和式画面（$n=10$）', color=k.M_INK, pad=10)
k.save_fig(fig, reg(P('fig06_x2_definition.png')))

# --- fig07 几何意义与性质（展示宽 7.4）---
fig, axes = plt.subplots(1, 2, figsize=(7.4, 4.4))
ax = axes[0]
x = np.linspace(-0.4, 3.6, 500); y = np.sin(x*1.6)*1.4
ax.plot(x, y, color=k.M_ACC2, lw=2.6, zorder=5)
ax.fill_between(x, 0, y, where=(y >= 0) & (x >= 0) & (x <= 3.3), color=k.M_GRN, alpha=0.30)
ax.fill_between(x, 0, y, where=(y < 0) & (x >= 0) & (x <= 3.3), color=k.M_RED, alpha=0.30)
ax.text(0.95, 0.45, '$+$', fontsize=26, color=k.M_GRN, ha='center')
ax.text(2.75, -0.75, '$-$', fontsize=26, color=k.M_RED, ha='center')
ax.set_xlim(-0.6, 4.0); ax.set_ylim(-2.4, 2.4)
k.style_axes(ax, 'x', 'y'); noticks(ax)
ax.set_title('下方计负', color=k.M_INK, pad=8)
ax = axes[1]
x = np.linspace(0, 4, 400); y = 0.25*x**2 + 0.5
ax.plot(x, y, color=k.M_ACC2, lw=2.6, zorder=5)
m1 = (x >= 0.4) & (x <= 2.0); m2 = (x >= 2.0) & (x <= 3.6)
ax.fill_between(x[m1], 0, y[m1], color=k.M_ACC, alpha=0.30)
ax.fill_between(x[m2], 0, y[m2], color=k.M_ACC2, alpha=0.25)
ax.plot([2.0, 2.0], [0, 0.25*4 + 0.5], color=k.M_RED, lw=1.8, ls='--', zorder=6)
for v, s in [(0.4, 'a'), (2.0, 'c'), (3.6, 'b')]:
    ax.text(v, -1.15, '$%s$' % s, ha='center', color=k.M_RED)
ax.text(1.2, 0.55, '$\\int_a^c$', color=k.M_INK, ha='center')
ax.text(2.85, 0.55, '$\\int_c^b$', color=k.M_INK, ha='center')
ax.set_xlim(-0.3, 4.6); ax.set_ylim(-1.7, 5.6)
k.style_axes(ax, 'x', 'y'); noticks(ax)
ax.set_title('区间可加', color=k.M_INK, pad=8)
k.save_fig(fig, reg(P('fig07_properties.png')))

# --- fig08 面积函数 A(x)（展示宽 9.6）---
fig, axes = plt.subplots(1, 2, figsize=(9.6, 3.5))
f = lambda t: 0.5*t + 0.8
ax = axes[0]
x = np.linspace(0, 4.2, 300)
ax.plot(x, f(x), color=k.M_ACC2, lw=2.6, zorder=5)
xs = np.linspace(0.5, 2.7, 200)
ax.fill_between(xs, 0, f(xs), color=k.M_ACC, alpha=0.32)
ax.axvline(2.7, ymax=0.52, color=k.M_RED, lw=2.0)
ax.text(1.6, 0.55, '$A(x)$', color=k.M_INK, ha='center')
ax.text(0.5, -1.05, '$a$', ha='center', color=k.M_RED)
ax.text(2.7, -1.05, '$x$', ha='center', color=k.M_RED)
ax.set_xlim(-0.3, 4.6); ax.set_ylim(-1.5, 4.4)
k.style_axes(ax, 't', 'y'); noticks(ax)
ax.set_title('$A(x)=\\int_a^x\\! f(t)\\,dt$', color=k.M_INK, pad=8)
ax = axes[1]
xx = np.linspace(0.5, 4.0, 200)
A = 0.25*(xx**2 - 0.5**2) + 0.8*(xx - 0.5)
ax.plot(xx, A, color=k.M_GRN, lw=3.0)
ax.plot([2.7], [0.25*(2.7**2-0.25)+0.8*2.2], 'o', color=k.M_RED, ms=10)
ax.set_xlim(0, 4.8); ax.set_ylim(-0.8, 7.2)
k.style_axes(ax, 'x', 'A'); noticks(ax)
ax.set_title('面积是 $x$ 的函数', color=k.M_GRN, pad=8)
k.save_fig(fig, reg(P('fig08_area_function.png')))

# --- fig09 增量 ΔA ≈ f(x)h（展示宽 9.5）---
fig, ax = k.new_fig(9.5, 4.9)
f = lambda t: 0.3*t**2 + 0.7
x = np.linspace(0, 4.0, 400)
ax.plot(x, f(x), color=k.M_ACC2, lw=2.8, zorder=6)
x0, h = 2.2, 0.75
xs = np.linspace(0.4, x0, 200)
ax.fill_between(xs, 0, f(xs), color=k.M_ACC, alpha=0.22)
xh = np.linspace(x0, x0+h, 100)
ax.fill_between(xh, 0, f(xh), color=k.M_GRN, alpha=0.45, zorder=3)
ax.add_patch(Polygon([[x0, 0], [x0+h, 0], [x0+h, f(x0)], [x0, f(x0)]], closed=True,
                     fill=False, edgecolor=k.M_RED, lw=2.2, ls='--', zorder=7))
ax.text(1.25, 0.6, '$A(x)$', color=k.M_INK, ha='center')
ax.annotate('$\\Delta A$', xy=(x0+h*0.75, f(x0)*0.6), xytext=(4.15, 1.1),
            color=k.M_GRN, ha='center',
            arrowprops=dict(arrowstyle='->', color=k.M_GRN, lw=1.6))
ax.annotate('矩形 $f(x)\\cdot h$', xy=(x0+h*0.3, f(x0)), xytext=(1.05, 4.6),
            color=k.M_RED, ha='center',
            arrowprops=dict(arrowstyle='->', color=k.M_RED, lw=1.6))
ax.text(x0, -0.95, '$x$', ha='center', color=k.M_INK)
ax.text(x0+h+0.12, -0.95, '$x+h$', ha='center', color=k.M_INK)
ax.set_xlim(-0.3, 5.6); ax.set_ylim(-1.6, 6.0)
k.style_axes(ax, 't', 'y'); noticks(ax)
ax.set_title('$\\Delta A/h\\approx f(x)\\;\\Rightarrow\\; A\'(x)=f(x)$', color=k.M_INK, pad=10)
k.save_fig(fig, reg(P('fig09_increment.png')))

# --- fig10 积分中值定理（展示宽 7.3）---
fig, ax = k.new_fig(7.3, 4.6)
f = lambda t: 0.3*t**2 + 0.7
x0, h = 2.2, 1.4
xh = np.linspace(x0, x0+h, 200)
x = np.linspace(1.6, 4.2, 300)
ax.plot(x, f(x), color=k.M_ACC2, lw=2.8, zorder=6)
ax.fill_between(xh, 0, f(xh), color=k.M_GRN, alpha=0.35, zorder=2)
xi = np.sqrt(((0.1*((x0+h)**3 - x0**3) + 0.7*h) / h - 0.7) / 0.3)
ax.add_patch(Polygon([[x0, 0], [x0+h, 0], [x0+h, f(xi)], [x0, f(xi)]], closed=True,
                     facecolor=k.M_ACC, alpha=0.30, edgecolor=k.M_RED, lw=2.2, zorder=4))
ax.plot([xi], [f(xi)], 'o', color=k.M_RED, ms=10, zorder=8)
ax.plot([xi, xi], [0, f(xi)], ls=':', color=k.M_RED, lw=1.6, zorder=5)
ax.text(xi, -0.95, '$\\xi$', ha='center', color=k.M_RED)
ax.text(x0 - 0.12, -0.95, '$x$', ha='center', color=k.M_INK)
ax.text(x0+h+0.2, -0.95, '$x+h$', ha='center', color=k.M_INK)
ax.set_xlim(1.3, 5.0); ax.set_ylim(-1.6, 5.6)
k.style_axes(ax, 't', 'y', origin=False); noticks(ax)
ax.set_title('$\\Delta A=f(\\xi)\\cdot h$', color=k.M_INK, pad=10)
k.save_fig(fig, reg(P('fig10_mvt.png')))

# --- fig11 牛顿-莱布尼茨：面积 ↔ 原函数之差（展示宽 5.2）---
fig, axes = plt.subplots(1, 2, figsize=(5.2, 3.2))
f = lambda t: t**2
Fn = lambda t: t**3/3
a, b = 0.6, 2.2
ax = axes[0]
x = np.linspace(0, 2.6, 300)
ax.plot(x, f(x), color=k.M_ACC2, lw=2.6, zorder=5)
xs = np.linspace(a, b, 200)
ax.fill_between(xs, 0, f(xs), color=k.M_ACC, alpha=0.35)
ax.text(1.45, 1.15, '$S$', color=k.M_INK, ha='center')
ax.text(a, -1.55, '$a$', ha='center', color=k.M_RED)
ax.text(b, -1.55, '$b$', ha='center', color=k.M_RED)
ax.set_xlim(-0.3, 2.9); ax.set_ylim(-2.3, 6.6)
k.style_axes(ax, 'x', 'y'); noticks(ax)
ax.set_title('一块面积', color=k.M_GRN, pad=8)
ax = axes[1]
ax.plot(x, Fn(x), color=k.M_GRN, lw=2.8, zorder=5)
ax.plot([b, b], [0, Fn(b)], ':', color=k.M_SLATE, lw=1.3)
ax.plot([0, b], [Fn(a), Fn(a)], ':', color=k.M_SLATE, lw=1.3)
ax.plot([0, b], [Fn(b), Fn(b)], ':', color=k.M_SLATE, lw=1.3)
ax.add_patch(FancyArrowPatch((b+0.14, Fn(a)), (b+0.14, Fn(b)),
                             arrowstyle='<->', color=k.M_RED, lw=2.2, mutation_scale=16))
ax.plot([a, b], [Fn(a), Fn(b)], 'o', color=k.M_RED, ms=8)
ax.text(-0.15, 5.6, '$F(b)-F(a)$', color=k.M_RED, ha='left', va='center')
ax.set_xlim(-0.3, 2.9); ax.set_ylim(-1.4, 6.6)
k.style_axes(ax, 'x', 'y'); noticks(ax)
ax.set_title('一个差', color=k.M_RED, pad=8)
k.save_fig(fig, reg(P('fig11_newton_leibniz.png')))

# --- fig12 例2 N-L 计算配图（展示宽 7.2）---
fig, ax = k.new_fig(7.2, 4.5)
x = np.linspace(0.05, 3.4, 400)
g = lambda t: t**2 - 2*t
ax.plot(x, g(x), color=k.M_ACC2, lw=2.8, zorder=5)
m_neg = (x >= 0) & (x <= 2); m_pos = (x >= 2) & (x <= 3)
ax.fill_between(x[m_neg], 0, g(x[m_neg]), color=k.M_RED, alpha=0.28)
ax.fill_between(x[m_pos], 0, g(x[m_pos]), color=k.M_GRN, alpha=0.30)
ax.text(1.0, -0.75, '$-$', fontsize=26, color=k.M_RED, ha='center')
ax.text(2.72, 0.75, '$+$', fontsize=26, color=k.M_GRN, ha='center')
for v in [2, 3]:
    ax.text(v, 0.35, '$%d$' % v, color=k.M_SLATE, ha='center')
ax.set_xlim(-0.3, 3.9); ax.set_ylim(-2.2, 4.2)
k.style_axes(ax, 'x', 'y'); noticks(ax)
ax.set_title('$\\int_0^3(x^2\\!-\\!2x)dx=0$，面积 $=8/3$', color=k.M_INK, pad=10)
k.save_fig(fig, reg(P('fig12_nl_example.png')))

# --- fig13 变速直线运动：位移 vs 路程（展示宽 9.4）---
fig, ax = k.new_fig(9.4, 4.6)
t = np.linspace(0, 5, 500)
v = lambda s: -s**2 + 4*s - 3
ax.plot(t, v(t), color=k.M_ACC2, lw=2.8, zorder=5)
for lo, hi, c, sg in [(0, 1, k.M_RED, '$-$'), (1, 3, k.M_GRN, '$+$'), (3, 4, k.M_RED, '$-$')]:
    m = (t >= lo) & (t <= hi)
    ax.fill_between(t[m], 0, v(t[m]), color=c, alpha=0.30)
    ax.text((lo+hi)/2, v((lo+hi)/2)*0.45, sg, fontsize=26, color=c, ha='center')
for xv in [1, 3, 4]:
    ax.plot([xv], [0], 'o', color=k.M_INK, ms=6)
    ax.text(xv, 0.32, '$%d$' % xv, color=k.M_SLATE, ha='center')
ax.set_xlim(-0.3, 5.4); ax.set_ylim(-4.4, 2.6)
k.style_axes(ax, 't/s', 'v'); noticks(ax)
ax.set_title('例3：$v(t)=-t^2+4t-3$，$[0,4]$ 的位移与路程', color=k.M_INK, pad=10)
k.save_fig(fig, reg(P('fig13_motion.png')))

# --- fig14 变力做功（弹簧）（展示宽 9.4）---
fig, ax = k.new_fig(9.4, 4.6)
x = np.linspace(0, 0.55, 300)
F = lambda s: 200*s
ax.plot(x, F(x), color=k.M_GRN, lw=3.0, zorder=5)
xs = np.linspace(0, 0.4, 200)
ax.fill_between(xs, 0, F(xs), color=k.M_ACC, alpha=0.34)
ax.plot([0.4, 0.4], [0, F(0.4)], '--', color=k.M_RED, lw=1.8)
ax.text(0.265, 20, '$W=16$ J', color=k.M_INK, ha='center')
ax.text(0.4, -14, '$0.4$ m', ha='center', color=k.M_RED)
ax.text(0.47, 105, '$F=200x$', color=k.M_GRN, ha='center')
ax.set_xlim(-0.04, 0.70); ax.set_ylim(-32, 145)
k.style_axes(ax, 'x/m', 'F/N'); noticks(ax)
ax.set_title('例4：$W=\\int_0^{0.4}\\!200x\\,dx$ —— 图线下的面积', color=k.M_INK, pad=10)
k.save_fig(fig, reg(P('fig14_work.png')))

# --- fig15 变式2：两曲线围成面积（展示宽 9.4）---
fig, ax = k.new_fig(9.4, 4.7)
x = np.linspace(-0.15, 1.35, 400)
ax.plot(x, np.sqrt(np.clip(x, 0, None)), color=k.M_ACC2, lw=2.8, label='$y=\\sqrt{x}$')
ax.plot(x, x**2, color=k.M_GRN, lw=2.8, label='$y=x^2$')
xs = np.linspace(0, 1, 300)
ax.fill_between(xs, xs**2, np.sqrt(xs), color=k.M_ACC, alpha=0.34)
ax.plot([0, 1], [0, 1], 'o', color=k.M_RED, ms=8, zorder=8)
ax.text(0.52, 0.42, '$S=1/3$', color=k.M_INK, ha='center')
ax.text(1.06, 1.0, '$(1,1)$', color=k.M_RED, va='center')
ax.set_xlim(-0.2, 1.75); ax.set_ylim(-0.35, 1.45)
k.style_axes(ax, 'x', 'y'); noticks(ax)
ax.legend(frameon=False, loc='upper left')
ax.set_title('$S=\\int_0^1(\\sqrt{x}-x^2)\\,dx$', color=k.M_INK, pad=10)
k.save_fig(fig, reg(P('fig15_between_curves.png')))

# --- fig16 Gabriel 号角（展示宽 6.3）---
fig, axes = plt.subplots(1, 2, figsize=(6.3, 3.1))
ax = axes[0]
x = np.linspace(1, 9, 400)
ax.plot(x, 1/x, color=k.M_ACC2, lw=2.6)
ax.plot(x, -1/x, color=k.M_ACC2, lw=2.6)
ax.fill_between(x, -1/x, 1/x, color=k.M_ACC, alpha=0.30)
for xv in np.linspace(1.3, 8.6, 10):
    ax.plot([xv, xv], [-1/xv, 1/xv], color=k.M_RULE, lw=0.9, zorder=1)
ax.set_xlim(0.4, 9.6); ax.set_ylim(-1.4, 1.4)
k.style_axes(ax, 'x', 'y', origin=False); noticks(ax)
ax.set_title('号角（侧视）', color=k.M_INK, pad=8)
ax = axes[1]
b = np.linspace(1, 60, 400)
ax.plot(b, 2*np.pi*np.log(b), color=k.M_RED, lw=2.8, label='$A$')
ax.plot(b, np.pi * (1 - 1/b), color=k.M_GRN, lw=2.8, label='$V$')
ax.axhline(np.pi, color=k.M_GRN, ls='--', lw=1.4)
ax.set_xlim(0, 62); ax.set_ylim(0, 30)
k.style_axes(ax, 'b', '', origin=False); noticks(ax)
ax.legend(frameon=False, loc='center right')
ax.set_title('$V\\to\\pi$，$A\\to\\infty$', color=k.M_INK, pad=8)
k.save_fig(fig, reg(P('fig16_gabriel.png')))

# --- fig17 ∫dx/x = ln x（展示宽 10.0）---
fig, axes = plt.subplots(1, 2, figsize=(10.0, 4.4))
ax = axes[0]
x = np.linspace(0.35, 8, 500)
ax.plot(x, 1/x, color=k.M_ACC2, lw=2.8)
for lo, hi, c in [(1, 2, k.M_ACC), (2, 4, k.M_GRN), (4, 8, k.M_ACC2)]:
    xs = np.linspace(lo, hi, 200)
    ax.fill_between(xs, 0, 1/xs, color=c, alpha=0.32)
    ax.text((lo*hi)**0.5, 1.32, '$\\ln 2$', color=k.M_INK, ha='center')
for v in [1, 2, 4, 8]:
    ax.text(v, -0.16, '$%d$' % v, color=k.M_SLATE, ha='center')
ax.set_xlim(0, 9.2); ax.set_ylim(-0.35, 1.75)
k.style_axes(ax, 'x', 'y', origin=False); noticks(ax)
ax.set_title('三块面积相等', color=k.M_INK, pad=8)
ax = axes[1]
b = np.linspace(1, 8, 300)
ax.plot(b, np.log(b), color=k.M_GRN, lw=3.0)
for v in [2, 4, 8]:
    ax.plot([v], [np.log(v)], 'o', color=k.M_RED, ms=8)
    ax.plot([v, v], [0, np.log(v)], ':', color=k.M_SLATE, lw=1.1)
ax.text(0.5, 2.25, '$A(xy)=A(x)+A(y)$', color=k.M_INK, ha='left', va='center')
ax.set_xlim(0, 9.2); ax.set_ylim(-0.2, 2.7)
k.style_axes(ax, 'b', 'A', origin=False); noticks(ax)
ax.set_title('$A(b)=\\ln b$', color=k.M_GRN, pad=8)
k.save_fig(fig, reg(P('fig17_log_discovery.png')))

print('figures:', len(figs))

# ============================ PPT ============================
prs = k.new_deck()

# 1 封面
k.title_slide(prs, '定积分与微积分基本定理',
              '第 02 讲　分割·近似·求和·取极限 —— 从和式的极限到牛顿-莱布尼茨公式',
              '数学名师课件包', '高中数学 选修 / 60 分钟')

# 2 学习目标
s = k.content_slide(prs, '本课学习目标', '导引')
k.bullets(s, [
    '理解定积分的定义：它是"和式的极限"，不是"原函数之差"',
    ('分割 → 近似代替 → 求和 → 取极限，四步缺一不可', 1),
    '会用定义（黎曼和）计算 ∫₀¹x²dx，体会"以直代曲"的力量',
    '理解微积分基本定理的直观证明：面积函数求导 + 积分中值定理',
    ('A(x)=∫ₐˣf(t)dt ⟹ A′(x)=f(x)：微分与积分互逆', 1),
    '熟练运用牛顿-莱布尼茨公式求定积分、面积、位移与功',
    '感受"无穷"的震撼：Gabriel 号角与对数的诞生',
], y=1.6, w=7.1, size=17)
k.picture(s, P('fig01_background.png'), x=7.6, y=2.3, w=5.2)
k.callout(s, '一句话主线：定积分把"无限细分的和"变成一个确定的数；\n基本定理告诉我们，这个数可以用求导的逆运算轻松拿到。',
          x=0.85, y=6.05, w=7.0, h=1.05, kind='key')

# 3 幕1
k.section_slide(prs, '第 1 幕 · 情境与建构', '从曲边梯形到"和式的极限"', '0–10 min')

# 4 两个背景
s = k.content_slide(prs, '两个古老而崭新的问题', '10 min')
k.full_picture(s, P('fig01_background.png'), y=1.55, w=7.0)
k.callout(s, '共同点：都要把"变化的量"在一个区间上累积起来。矩形面积、恒力做功我们会算；'
             '一旦"曲了""变了"，初等方法就失灵 —— 除非，我们把它切碎。',
          x=1.6, y=6.15, w=10.1, h=1.0, kind='note')

# 5 四步法（1）分割与近似
s = k.content_slide(prs, '四步法（一）：分割 · 近似代替', '10 min')
k.picture(s, P('fig02_partition.png'), x=5.9, y=1.5, w=7.0)
k.bullets(s, [
    '分割：把 [a,b] 等分成 n 份，Δx=(b−a)/n',
    ('小区间 [xᵢ₋₁, xᵢ]，i = 1,2,…,n', 1),
    '近似代替：在每个小区间上取一点 ξᵢ',
    ('用矩形 f(ξᵢ)·Δx 代替曲边小条 ΔSᵢ', 1),
    ('"以直代曲"——只要区间够窄，误差可以任意小', 1),
], y=1.6, w=5.1, size=17)
k.callout(s, '关键直觉：曲线在极小范围内"看起来是直的"。\n这与第 01 讲导数的"局部线性化"是同一个思想。',
          x=0.85, y=5.5, w=5.0, h=1.4, kind='key')

# 6 四步法（2）求和取极限
s = k.content_slide(prs, '四步法（二）：求和 · 取极限', '10 min')
k.formula(s, r'$S_n=\sum_{i=1}^{n} f(\xi_i)\,\Delta x\quad\longrightarrow\quad'
             r'\lim_{n\to\infty}\sum_{i=1}^{n} f(\xi_i)\,\Delta x$',
          x=1.2, y=1.5, w=10.9, size=0.95, out=T('t06.png'))
k.full_picture(s, P('fig03_riemann_seq.png'), y=2.55, w=8.0)
k.callout(s, 'n 从 4 到 64，误差从 0.22 掉到 0.008 —— 和式在"收敛"。', x=3.4, y=6.55, w=6.6, h=0.65, kind='note')

# 7 夹逼：极限确实存在
s = k.content_slide(prs, '它真的收敛吗？—— 上和与下和的夹逼', '10 min')
k.full_picture(s, P('fig04_upper_lower.png'), y=1.5, w=9.6)
k.callout(s, '不论 ξᵢ 取左端、右端还是中点，都被夹在下和与上和之间；'
             '而 Sₙ − sₙ = 1/n → 0。所以极限存在且唯一，与取点方式无关 —— 这才配叫"定"积分。',
          x=1.4, y=6.05, w=10.5, h=1.05, kind='key')

# 8 收敛可视化
s = k.content_slide(prs, '两把钳子，夹出唯一的那个数', '10 min')
k.picture(s, P('fig05_converge.png'), x=6.0, y=1.7, w=6.9)
k.bullets(s, [
    '下和 sₙ：单调递增（内接矩形，永远偏小）',
    '上和 Sₙ：单调递减（外接矩形，永远偏大）',
    'sₙ ≤ 真值 ≤ Sₙ，且 Sₙ − sₙ → 0',
    ('⟹ 存在唯一实数被夹住，就是 ∫ₐᵇ f(x)dx', 1),
    '这正是实数连续性（确界原理）的威力',
], y=1.9, w=5.2, size=17)
k.callout(s, '严谨性提醒：闭区间上的连续函数一定可积。\n本课默认 f(x) 在 [a,b] 上连续。',
          x=0.85, y=5.4, w=5.1, h=1.3, kind='warn')

# 9 定积分定义
s = k.content_slide(prs, '定积分的定义（务必背下"极限"二字）', '定义')
k.formula(s, r'$\int_a^b f(x)\,dx\;=\;\lim_{n\to\infty}\sum_{i=1}^{n} '
             r'f\!\left(\xi_i\right)\cdot\frac{b-a}{n}$',
          x=1.0, y=1.6, w=11.3, size=1.05, out=T('t09.png'))
k.bullets(s, [
    'a 积分下限，b 积分上限；f(x) 被积函数；x 积分变量（哑变量）',
    '∫ 是拉长的 S（Summa，"和"）—— 符号本身就在提醒你：这是个和！',
    '∫ₐᵇ f(x)dx = ∫ₐᵇ f(t)dt：结果是一个数，与字母无关',
], y=3.5, w=11.6, size=17)
k.callout(s, '常见误区（现在就掐死）：\n×  "定积分就是 F(b) − F(a)"　—— 那是牛顿-莱布尼茨【定理】的结论，不是【定义】。\n'
             '√  定积分是和式的极限；它先于原函数而存在，哪怕原函数写不出来（如 e^(−x²)）。',
          x=0.85, y=5.25, w=11.6, h=1.75, kind='warn')

# 10 幕2
k.section_slide(prs, '第 2 幕 · 定义的实战', '例 1：不用任何公式，硬算 ∫₀¹x²dx', '10–20 min')

# 11 例1 推演
s = k.content_slide(prs, '例1　用定义计算 ∫₀¹ x² dx', '例题')
k.bullets(s, [
    '① 分割：[0,1] 等分 n 份，Δx = 1/n，xᵢ = i/n',
    '② 近似：取右端点 ξᵢ = i/n，小矩形面积 (i/n)²·(1/n)',
], y=1.5, w=11.8, size=17)
k.formula(s, r'$S_n=\sum_{i=1}^{n}\frac{i^2}{n^2}\cdot\frac{1}{n}=\frac{1}{n^3}\sum_{i=1}^{n}i^2'
             r'=\frac{1}{n^3}\cdot\frac{n(n+1)(2n+1)}{6}=\frac{1}{6}\left(1+\frac{1}{n}\right)\left(2+\frac{1}{n}\right)$',
          x=0.7, y=2.65, w=11.9, size=0.86, out=T('t11.png'))
k.formula(s, r'$\int_0^1 x^2\,dx=\lim_{n\to\infty}S_n=\frac{1}{6}\cdot 1\cdot 2=\frac{1}{3}$',
          x=2.6, y=4.35, w=8.1, size=1.0, out=T('t11b.png'))
k.callout(s, '④ 取极限，答案 1/3。整个过程只用到了平方和公式与极限 —— 没有原函数，没有求导。\n'
             '这就是定积分的"出厂设置"。',
          x=0.85, y=5.75, w=11.6, h=1.2, kind='key')

# 12 例1 配图
s = k.content_slide(prs, '例1 的几何画面：一叠越来越薄的矩形', '例题')
k.full_picture(s, P('fig06_x2_definition.png'), y=1.5, w=9.4)
k.callout(s, '思考：若取左端点 ξᵢ=(i−1)/n，Sₙ = (1/6)(1−1/n)(2−1/n)，极限仍是 1/3。'
             '取点方式不影响结果 —— 与上一页夹逼结论完全吻合。',
          x=1.6, y=6.35, w=10.1, h=0.85, kind='note')

# 13 几何意义与性质
s = k.content_slide(prs, '几何意义与三条基本性质', '20 min')
k.picture(s, P('fig07_properties.png'), x=5.5, y=1.55, w=7.4)
k.bullets(s, [
    '几何意义：x 轴上方计正、下方计负的"代数面积"',
    '线性：∫(kf + mg) = k∫f + m∫g',
    '区间可加：∫ₐᵇ = ∫ₐᶜ + ∫꜀ᵇ（c 可在区间外）',
    '保号性：f ≥ 0 ⟹ ∫ₐᵇ f ≥ 0',
], y=1.7, w=5.0, size=16)
k.callout(s, '面积 ≠ 积分值！\n求面积必须先找零点、分段取绝对值。',
          x=0.85, y=5.3, w=4.9, h=1.2, kind='warn')

# 14 幕3
k.section_slide(prs, '第 3 幕 · 定理的诞生', '微积分基本定理：微分与积分互为逆运算', '20–35 min')

# 15 面积函数
s = k.content_slide(prs, '关键一步：把上限放开，造一个"面积函数"', '35 min')
k.formula(s, r'$A(x)=\int_a^x f(t)\,dt\qquad (a\leq x\leq b)$',
          x=2.9, y=1.45, w=7.5, size=0.95, out=T('t15.png'))
k.full_picture(s, P('fig08_area_function.png'), y=2.75, w=9.6)
k.callout(s, '注意 t 与 x 的分工：t 在区间内跑（哑变量），x 是自变量。'
             '牛顿的天才之处，就在于敢让积分上限"动起来"。',
          x=1.6, y=6.35, w=10.1, h=0.85, kind='note')

# 16 增量分析
s = k.content_slide(prs, '给 A(x) 求导：ΔA ≈ f(x)·h', '35 min')
k.full_picture(s, P('fig09_increment.png'), y=1.5, w=9.5)
k.callout(s, '细长条几乎就是矩形：底 h，高 ≈ f(x)。h 越小，"≈"越准。',
          x=3.2, y=6.55, w=7.0, h=0.65, kind='key')

# 17 积分中值定理
s = k.content_slide(prs, '把"≈"变成"="：积分中值定理', '35 min')
k.picture(s, P('fig10_mvt.png'), x=5.6, y=1.5, w=7.3)
k.bullets(s, [
    'f 在 [x, x+h] 上连续 ⟹ 有最小值 m、最大值 M',
    ('m·h ≤ ΔA ≤ M·h，即 m ≤ ΔA/h ≤ M', 1),
    '介值定理：f 取遍 [m, M] 的每个值',
    ('⟹ ∃ξ ∈ [x, x+h]，使 f(ξ) = ΔA/h', 1),
    'h → 0 时 ξ → x，由 f 连续得 f(ξ) → f(x)',
], y=1.7, w=4.9, size=15.5)
k.callout(s, '于是　A′(x) = lim(h→0) ΔA/h = f(x)　∎',
          x=0.85, y=5.55, w=4.9, h=0.95, kind='key')

# 18 FTC 第一部分
s = k.content_slide(prs, '微积分基本定理（第一部分）', '定理')
k.formula(s, r'$\frac{d}{dx}\int_a^x f(t)\,dt \;=\; f(x)$',
          x=2.4, y=1.55, w=8.5, size=1.15, out=T('t18.png'))
k.bullets(s, [
    '文字：先积分、再求导，回到原点 —— 二者互为逆运算。',
    '推论：连续函数必有原函数，A(x) 就是其中一个（还顺手证明了原函数的存在性）。',
    '这一步把"求面积"（几何、静态、无限过程）与"求导"（代数、动态、有限步骤）焊死在一起。',
], y=3.15, w=11.8, size=17)
k.callout(s, '数学史上最漂亮的桥梁之一：\n阿基米德算了一辈子面积，牛顿与莱布尼茨只用一个求导的逆运算就全接管了。',
          x=0.85, y=5.6, w=11.6, h=1.25, kind='note')

# 19 N-L 公式
s = k.content_slide(prs, '牛顿-莱布尼茨公式（第二部分）', '定理')
k.formula(s, r'$\int_a^b f(x)\,dx=F(b)-F(a)\;=\;\left[F(x)\right]_a^b,'
             r'\qquad F\,\!'"'"r'(x)=f(x)$',
          x=0.9, y=1.5, w=11.5, size=0.90, out=T('t19.png'))
k.full_picture(s, P('fig11_newton_leibniz.png'), y=2.6, w=6.0)
k.callout(s, '证明一行：A(x) 与 F(x) 都是 f 的原函数 ⟹ F(x)=A(x)+C；'
             '令 x=a 得 C=F(a)，再令 x=b 得 ∫ₐᵇf = A(b) = F(b)−F(a)。',
          x=1.0, y=6.35, w=11.3, h=0.9, kind='key')

# 20 幕4
k.section_slide(prs, '第 4 幕 · 例题与变式', '算得又快又对：N-L 公式的四种用法', '35–52 min')

# 21 常用原函数表
s = k.content_slide(prs, '工具箱：常用原函数（求导的逆）', '52 min')
k.formula(s, r'$\int x^{n}dx=\frac{x^{n+1}}{n+1}\ (n\neq-1),\quad'
             r'\int \frac{1}{x}dx=\ln|x|,\quad \int e^{x}dx=e^{x}$',
          x=0.7, y=1.6, w=11.9, size=0.82, out=T('t21.png'))
k.formula(s, r'$\int \sin x\,dx=-\cos x,\quad \int \cos x\,dx=\sin x,\quad'
             r'\int a^{x}dx=\frac{a^{x}}{\ln a}$',
          x=0.9, y=3.0, w=11.5, size=0.82, out=T('t21b.png'))
k.bullets(s, [
    '口诀：把导数表倒着背一遍就是原函数表。',
    '验算法宝：求出的 F(x) 求一次导，必须变回 f(x)。三秒自检，杜绝低级错。',
], y=4.4, w=11.8, size=17)
k.callout(s, 'n = −1 是幂函数积分公式唯一的"漏网之鱼"——它的原函数是 ln|x|。\n'
             '这条裂缝里，藏着本课最后的彩蛋。',
          x=0.85, y=5.75, w=11.6, h=1.2, kind='warn')

# 22 例2
s = k.content_slide(prs, '例2　∫₀³(x²−2x)dx，并求曲线与 x 轴围成的面积', '例题')
k.picture(s, P('fig12_nl_example.png'), x=5.7, y=1.5, w=7.2)
k.bullets(s, [
    '积分值：[x³/3 − x²]₀³ = 9 − 9 = 0',
    '面积：先解 x² − 2x = 0 ⟹ x = 0, 2',
    ('[0,2] 上 f ≤ 0，[2,3] 上 f ≥ 0', 1),
    ('S = −∫₀²f + ∫₂³f = 4/3 + 4/3 = 8/3', 1),
], y=1.7, w=5.0, size=16)
k.callout(s, '积分值 0 ≠ 面积 8/3。\n"正负抵消"是定积分的本性，不是错误。',
          x=0.85, y=5.2, w=4.9, h=1.25, kind='warn')

# 23 变式1
s = k.content_slide(prs, '变式1　含绝对值与分段函数', '变式')
k.formula(s, r'$\int_0^{2}|x^{2}-1|\,dx=\int_0^{1}(1-x^{2})dx+\int_1^{2}(x^{2}-1)dx$',
          x=1.0, y=1.55, w=11.3, size=0.85, out=T('t23.png'))
k.formula(s, r'$=\left[x-\frac{x^{3}}{3}\right]_0^{1}+\left[\frac{x^{3}}{3}-x\right]_1^{2}'
             r'=\frac{2}{3}+\left(\frac{2}{3}+\frac{2}{3}\right)=2$',
          x=1.6, y=3.05, w=10.1, size=0.85, out=T('t23b.png'))
k.bullets(s, [
    '三步走：① 去绝对值找零点　② 按符号分段　③ 逐段 N-L 后相加',
    '同理：分段函数在分界点处拆开积分（区间可加性正是为此而生）',
], y=4.5, w=11.8, size=17)
k.callout(s, '课堂速练（1 分钟）：∫₋₁¹ |x| dx = ?　（答案：1）',
          x=0.85, y=5.95, w=11.6, h=0.9, kind='note')

# 24 例3 运动
s = k.content_slide(prs, '例3　变速直线运动：位移与路程', '例题')
k.full_picture(s, P('fig13_motion.png'), y=1.5, w=9.4)
k.callout(s, '位移 = ∫v dt（带符号）；路程 = ∫|v| dt（先找 v 的零点 t=1, 3，再分段）。\n'
             '物理直觉：v–t 图线下的"代数面积"就是位移 —— 这正是定积分的原始出身。',
          x=1.2, y=6.25, w=10.9, h=1.05, kind='key')

# 25 例4 做功
s = k.content_slide(prs, '例4　变力做功：弹簧问题', '例题')
k.full_picture(s, P('fig14_work.png'), y=1.5, w=9.4)
k.callout(s, '思路统一：把 [0, 0.4] 分割，每小段位移内力近似为常力 → 微功 dW = F(x)dx → 求和取极限 = ∫F(x)dx。\n'
             '一切"变量的累积"，都是定积分。',
          x=1.2, y=6.25, w=10.9, h=1.05, kind='key')

# 26 变式2
s = k.content_slide(prs, '变式2　两曲线围成的面积', '变式')
k.full_picture(s, P('fig15_between_curves.png'), y=1.5, w=9.4)
k.callout(s, '通法：① 联立求交点定上下限　② 判断谁在上　③ ∫(上 − 下)dx。'
             '不需要讨论正负 —— 差值天然非负。',
          x=1.4, y=6.35, w=10.5, h=0.85, kind='note')

# 27 幕5
k.section_slide(prs, '第 5 幕 · 拔高与震撼', '当积分遇上"无穷"：两个改变数学史的瞬间', '52–58 min')

# 28 Gabriel 号角（震撼 + 哲学）
s = k.content_slide(prs, '震撼一：Gabriel 号角 —— 装得满，刷不完', '拔高')
k.picture(s, P('fig16_gabriel.png'), x=6.6, y=1.55, w=6.3)
k.formula(s, r'$V=\pi\!\int_1^{\infty}\!\frac{dx}{x^{2}}=\pi$',
          x=6.6, y=4.80, w=6.3, size=0.60, out=T('t28.png'))
k.formula(s, r'$A>2\pi\!\int_1^{\infty}\!\frac{dx}{x}=+\infty$',
          x=6.6, y=5.65, w=6.3, size=0.60, out=T('t28b.png'))
k.bullets(s, [
    '把 π 立方单位的油漆灌进号角 —— 它满了。',
    '可油漆此刻正贴着无限大的内壁 —— 它"刷"完了？',
    '出口：现实刷漆要求漆层有厚度，数学曲面厚度为零。',
    ('号角越往后越细，任何有厚度的漆刷都伸不进去。', 1),
    '教益：1/x² 收敛而 1/x 发散 —— "趋于 0" ≠ "加起来有限"。',
], y=1.65, w=6.0, size=16)
k.callout(s, '无穷不是很大的数，它是一种过程。\n直觉在此失效，唯有极限的定义可靠。',
          x=0.85, y=5.5, w=5.7, h=1.2, kind='key')

# 30 对数的诞生
s = k.content_slide(prs, '震撼二：∫dx/x 里长出了对数', '拔高')
k.full_picture(s, P('fig17_log_discovery.png'), y=1.5, w=10.0)
k.callout(s, '面积 A(b)=∫₁ᵇ dt/t 天然满足 A(xy)=A(x)+A(y) —— 把乘法压成加法。'
             '17 世纪的人先看见这条性质，才认出它就是对数。数学之美：定义从现象里"长"出来。',
          x=1.0, y=6.25, w=11.3, h=1.0, kind='note')

# 31 小结
s = k.content_slide(prs, '课堂小结：一条主线，两个身份', '小结')
k.bullets(s, [
    '定积分的【定义】：分割—近似—求和—取极限，∫ₐᵇf = lim Σ f(ξᵢ)Δx',
    '定积分的【算法】：牛顿-莱布尼茨公式 ∫ₐᵇf = F(b) − F(a)',
    '连接二者的【定理】：A(x)=∫ₐˣf(t)dt ⟹ A′(x)=f(x)（微分积分互逆）',
    '三类应用：面积（分段去绝对值）、位移与路程、变力做功',
    '一条警戒线：积分值 ≠ 面积；n = −1 时原函数是 ln|x|',
], y=1.7, w=6.9, size=17)
k.picture(s, P('fig11_newton_leibniz.png'), x=7.9, y=2.4, w=5.3)
k.callout(s, '记住这句：定义告诉你它"是什么"，定理告诉你它"怎么算"。\n考场上用定理，思想上守定义。',
          x=0.85, y=5.9, w=7.4, h=1.15, kind='key')

# 32 分层作业
s = k.content_slide(prs, '分层作业', '作业')
k.bullets(s, [
    'A 基础（人人必做）',
    ('① 用定义计算 ∫₀¹ x dx；② ∫₁² (x² + 1/x²)dx；③ ∫₀^π sin x dx', 1),
    'B 提高（多数完成）',
    ('④ ∫₀³ |x² − 4| dx；⑤ 求 y = x² 与 y = 2x 围成的面积', 1),
    ('⑥ v(t) = t² − 4t + 3 (m/s)，求 [0,3] 内的位移与路程', 1),
    'C 挑战（学有余力）',
    ('⑦ 证明：∫₀¹ xⁿ dx = 1/(n+1)，并用定义（不用 N-L）证 n = 3 的情形', 1),
    ('⑧ 弹簧劲度 k = 300 N/m，从伸长 0.1 m 拉到 0.3 m，求外力做的功', 1),
    ('⑨ 探究：∫₁^b dx/x^p 何时收敛？找出 Gabriel 号角悖论的临界指数', 1),
], y=1.6, w=11.9, size=15.5)

# 33 板书提纲
s = k.content_slide(prs, '板书提纲', '板书')
k.bullets(s, [
    '【左板 · 定义区】曲边梯形图 → 四步：分割/近似/求和/取极限 → ∫ₐᵇf = lim Σf(ξᵢ)Δx',
    '【中板 · 定理区】面积函数 A(x) 图 → ΔA ≈ f(x)h → 中值定理 → A′(x)=f(x) → N-L 公式',
    '【右板 · 例题区】例1 定义算 x²（保留全过程不擦）｜例2 分段求面积｜例3 位移/路程｜例4 做功',
    '【副板 · 警戒区】积分值 ≠ 面积　·　∫dx/x = ln|x|　·　常用原函数表',
], y=1.8, w=11.9, size=16)
k.callout(s, '板书节奏：例1 的黎曼和推导必须完整板演并全程保留 —— 它是本课"定义先于定理"的物证。',
          x=0.85, y=5.6, w=11.6, h=0.95, kind='key')

pptx_path = k.save(prs, os.path.join(OUT, '02_定积分与微积分基本定理.pptx'))
print('slides:', len(prs.slides.__iter__.__self__._sldIdLst))

# ============================ 教案 DOCX ============================
from docx import Document
from docx.shared import Pt as DPt, RGBColor as DRGB, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Noto Serif CJK SC'; st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Cm(2.2)

def H(t, lv=1):
    p = doc.add_heading('', level=lv)
    r = p.add_run(t); r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
    r.font.size = DPt(16 if lv == 1 else 13)
    r.font.bold = True
    return p

def Pp(t, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.font.size = DPt(size)
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    p.paragraph_format.line_spacing = 1.4
    return p

ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run('教案：第02讲　定积分与微积分基本定理')
r.font.size = DPt(20); r.font.bold = True; r.font.color.rgb = DRGB(0x11, 0x1B, 0x2E)
r.font.name = 'Noto Serif CJK SC'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')

info = doc.add_table(rows=2, cols=4); info.style = 'Table Grid'
info.alignment = WD_TABLE_ALIGNMENT.CENTER
kv = [('课　题', '定积分与微积分基本定理'), ('课　时', '1 课时（60 分钟）'),
      ('课　型', '新授课（概念建构 + 定理探究 + 应用）'), ('教　具', 'PPT 32 页、17 幅精确数学图、几何画板、彩色粉笔')]
for i, (a, b) in enumerate(kv):
    c = info.rows[i // 2].cells[(i % 2) * 2]
    c.text = a; c.paragraphs[0].runs[0].bold = True
    info.rows[i // 2].cells[(i % 2) * 2 + 1].text = b
doc.add_paragraph()

H('一、教材分析')
Pp('本课位于导数及其应用之后，是微积分模块的收官与升华。定积分以"曲边梯形面积、变速运动位移、变力做功"三大原型为背景，'
   '通过"分割—近似代替—求和—取极限"建构概念，其思想方法（无限逼近、以直代曲）贯穿整个高等数学。'
   '微积分基本定理则把看似彼此无关的两大分支——微分与积分——统一起来，是数学史上最深刻的定理之一。'
   '教材编排上先讲定义、后给定理，正是为了防止学生把"F(b)−F(a)"误当成定积分的本质。本课必须守住这一编排意图。')

H('二、学情分析')
Pp('学生已掌握导数的定义与运算、极限的初步概念，具备"局部线性化"的经验，为"以直代曲"提供了心理准备。'
   '但存在三个典型障碍：(1) 对"无限求和"心存疑虑，难以接受一个无穷过程能得到确定的数；'
   '(2) 极易把定积分等同于原函数之差，跳过定义直接套公式，导致遇到无初等原函数的情形（如 e^(−x²)）时无所适从；'
   '(3) 混淆"积分值"与"面积"，忽视 x 轴下方计负。教学中需以夹逼可视化破除疑虑，以例1（用定义硬算）建立敬畏，以反例强化辨析。')

H('三、教学目标')
Pp('【知识与技能】', True)
Pp('1. 理解定积分的定义，能复述"分割—近似—求和—取极限"四个步骤及其数学表达；\n'
   '2. 能用定义（黎曼和）计算 ∫₀¹x²dx 这类简单定积分；\n'
   '3. 理解并能陈述微积分基本定理，熟练运用牛顿-莱布尼茨公式求定积分；\n'
   '4. 会用定积分求平面图形面积、变速直线运动的位移与路程、变力做功。')
Pp('【过程与方法】', True)
Pp('1. 经历从实际问题抽象出定积分模型的完整过程，发展数学抽象与数学建模素养；\n'
   '2. 通过上和、下和的夹逼与数值实验，体会极限思想与实数连续性，发展逻辑推理素养；\n'
   '3. 通过面积函数增量分析 + 积分中值定理的直观证明，经历"发现—猜想—论证"的科学探究路径。')
Pp('【情感态度与价值观】', True)
Pp('1. 感受微积分基本定理沟通微分与积分的统一之美，体会数学的深刻与简洁；\n'
   '2. 通过 Gabriel 号角悖论与对数的积分起源，感受"无穷"带来的认知冲击，培养理性质疑精神与数学审美。')

H('四、教学重点与难点')
Pp('【重点】定积分的定义（和式的极限）；微积分基本定理与牛顿-莱布尼茨公式的运用。', True)
Pp('【难点】① 理解定积分是和式的极限而非"原函数之差"，把握定义与定理的逻辑先后；'
   '② 微积分基本定理的直观证明（面积函数求导 + 积分中值定理）；③ 区分积分值与面积。', True)
Pp('【突破策略】难点①用"n=4/8/16/64 逼近序列 + 上下和夹逼图"的双重可视化，并用例1"不许用公式"的硬算建立体验；'
   '难点②用细长条与矩形的几何对照图（ΔA ≈ f(x)h）铺垫，再以积分中值定理把"≈"升级为"="；'
   '难点③用 ∫₀³(x²−2x)dx = 0 而面积 = 8/3 的强反差反例击穿。')

H('五、教法与学法')
Pp('教法：问题驱动 + 直观几何 + 数值实验 + 讲练结合。以两个物理/几何原型引入，用 matplotlib 精确图与几何画板动态演示支撑抽象。')
Pp('学法：观察—猜想—验证—归纳。学生在"分割变细"的动态中自主发现收敛，在夹逼中确信极限存在，在例题中巩固算法。')

H('六、教学准备')
Pp('PPT 课件 32 页（含 17 幅精确数学插图）；几何画板黎曼和动态文件；分层作业单；三色粉笔（黑—主线，红—警戒，蓝—几何）。')

H('七、教学过程（总计 60 分钟）')

rows = [
    ('环节', '时间', '教师活动', '学生活动', '设计意图'),
    ('导入\n(P1–P4)', '0–4′\n(4)',
     '出示两个问题：① 抛物线下曲边梯形的面积；② 弹簧拉伸中变力做的功。提问："矩形面积、恒力做功你们都会算，'
     '为什么这里就卡住了？"板书两图。',
     '尝试用初等方法解决，发现"曲了""变了"就失效；意识到需要新工具。',
     '以真实问题制造认知失衡，激发内在需求；两个原型分别指向几何与物理，暗示定积分的普适性。'),
    ('建构一\n分割·近似\n(P5)', '4–8′\n(4)',
     '引导："如果把区间切成很多小段，每一小段上曲线还'"'"'弯'"'"'吗？"演示以直代曲，写出 ΔSᵢ ≈ f(ξᵢ)Δx。'
     '联系第01讲导数的局部线性化。',
     '观察分割图，说出"小区间上曲线近似是直的"；写出小矩形面积表达式。',
     '把新知锚定在已有的"局部线性化"经验上，降低抽象台阶；建立"以直代曲"这一核心策略。'),
    ('建构二\n求和·取极限\n(P6–P8)', '8–14′\n(6)',
     '几何画板演示 n = 4→8→16→64，同步报出和值与误差；再展示上和与下和的夹逼图与收敛折线图。'
     '追问："凭什么断定它一定收敛到某个数？"',
     '读出误差从 0.22 降至 0.008；发现上和递减、下和递增、差为 1/n → 0；得出"被唯一夹住"的结论。',
     '用数值实验代替 ε-δ，使"极限存在"可感可信；夹逼结构同时说明结果与 ξᵢ 取法无关，为"定"字正名。'),
    ('概念形成\n(P9)', '14–18′\n(4)',
     '板书定积分定义式，逐项讲解上下限、被积函数、哑变量。红笔重点标注："定积分是和式的极限"，'
     '并当场否定"定积分 = F(b) − F(a)"这一常见误解。',
     '齐读定义；辨析 ∫f(x)dx 与 ∫f(t)dt 的关系；记录警戒线笔记。',
     '在学生尚未接触 N-L 公式时先立定义，从时间顺序上杜绝"先入为主"的误解，守住教材的编排意图。'),
    ('例1\n定义硬算\n(P10–P12)', '18–28′\n(10)',
     '布置例1：用定义计算 ∫₀¹x²dx，明确"禁止使用任何积分公式"。板演四步，用到平方和公式。'
     '追问："若取左端点，答案会变吗？"',
     '独立完成分割与求和（3′），一名学生上台板演；计算左端点情形，验证极限仍为 1/3。',
     '让学生亲历定义的"笨重"，从而在下一节感受 N-L 公式的"轻盈"——为定理的震撼做情绪铺垫。'),
    ('性质梳理\n(P13)', '28–31′\n(3)',
     '结合正负面积图与区间可加图，讲解几何意义与线性、可加、保号三条性质。红笔写下"积分值 ≠ 面积"。',
     '观察图形，说出 x 轴下方部分的符号；记录三条性质。',
     '为后续例2的反例埋伏笔；性质是运算的依据，须在定理之前落地。'),
    ('定理探究\n面积函数\n(P14–P16)', '31–37′\n(6)',
     '关键一问："如果把积分上限放开、让它动起来，会得到什么？"引出 A(x) = ∫ₐˣ f(t)dt。'
     '几何画板拖动 x，展示面积随之变化。再画细长条与矩形，引导写出 ΔA ≈ f(x)·h。',
     '发现面积是 x 的函数；观察细长条"几乎是矩形"；写出 ΔA/h ≈ f(x) 并大胆猜想 A′(x) = f(x)。',
     '"让上限动起来"是全课的思维跳跃点，必须由学生在教师追问下自己迈出；猜想先于证明，符合科学发现的真实顺序。'),
    ('定理证明\n中值定理\n(P17–P18)', '37–42′\n(5)',
     '追问："≈ 不够严谨，能变成 = 吗？"引入积分中值定理：由最值定理与介值定理得 ∃ξ 使 ΔA = f(ξ)h。'
     '令 h → 0，由 f 连续得 A′(x) = f(x)。板书完整证明链。',
     '回忆最值定理与介值定理；跟随推导，理解 ξ → x 的极限过程；完成证明笔记。',
     '把直观猜想提升为严格论证，示范"数学的严谨如何为直觉兜底"；同时复习连续函数的两大定理。'),
    ('N-L 公式\n(P19–P21)', '42–46′\n(4)',
     '由 A(x) 与 F(x) 同为原函数、相差常数，一行推出 ∫ₐᵇf = F(b) − F(a)。展示"面积↔纵坐标之差"对照图。'
     '梳理常用原函数表，强调求导验算法。',
     '完成一行推导；对照两图说出"图形世界的一块面积 = 代数世界的一个差"；背诵原函数表。',
     '让定理的"神奇"落在可见的图形对照上；求导验算是最经济的自纠错手段，须当堂养成习惯。'),
    ('例2·变式1\n(P22–P23)', '46–50′\n(4)',
     '例2：∫₀³(x²−2x)dx 与其面积。故意先让学生算积分值得 0，再问"面积也是 0 吗？"'
     '变式1：∫₀²|x²−1|dx，示范"找零点—分段—相加"三步。',
     '计算得积分值 0，产生困惑；看图恍然，重算面积 8/3；独立完成变式1并互批。',
     '用"0 与 8/3"的强反差把易错点钉死在记忆里；变式1把方法固化为可迁移的三步流程。'),
    ('例3·例4\n应用\n(P24–P25)', '50–54′\n(4)',
     '例3：v(t) = −t²+4t−3，求 [0,4] 的位移与路程，强调 ∫v 与 ∫|v| 的区别。'
     '例4：弹簧 F = 200x 拉伸 0.4 m 的功，回到导入问题②。',
     '找 v 的零点分段；计算位移与路程；独立完成做功计算得 16 J，与导入呼应。',
     '完成"提出问题—解决问题"的教学闭环；两道应用题分别落在物理的运动学与功能关系，凸显定积分的建模价值。'),
    ('变式2\n(P26)', '54–55′\n(1)',
     '快速给出两曲线围成面积的通法：定交点—判上下—∫(上−下)。',
     '口答 y=√x 与 y=x² 围成面积为 1/3。',
     '一分钟补齐面积问题的最后一类，形成完整的方法谱系。'),
    ('拔高·震撼\n(P27–P29)', '55–58′\n(3)',
     'Gabriel 号角：V = π 有限，A = ∞ 无限——"灌得满，刷不完"。追问悖论的出口。'
     '再展示 ∫₁ᵇ dt/t 的等面积现象，指出对数正是这样被"积"出来的。',
     '被悖论震住；讨论并说出"数学曲面没有厚度"；观察 [1,2]、[2,4]、[4,8] 面积相等，惊叹 A(xy)=A(x)+A(y)。',
     '在认知高点收尾：用悖论击碎"直觉可靠"的幻觉，用对数的诞生展示"定义从现象中生长"的数学之美，把课堂推向情感与哲思的顶点。'),
    ('小结·作业\n(P30–P32)', '58–60′\n(2)',
     '归纳主线："定义告诉你它是什么，定理告诉你它怎么算。"布置 A/B/C 三层作业，说明 C 层第⑨题与号角悖论的关联。',
     '复述定义式与 N-L 公式；按自身水平选择作业层级。',
     '一句话锚定全课逻辑；分层作业照顾差异，C 层第⑨题把课堂的震撼延伸为课后的探究。'),
]
tb = doc.add_table(rows=len(rows), cols=5); tb.style = 'Table Grid'
widths = [Cm(2.3), Cm(1.8), Cm(5.6), Cm(3.6), Cm(4.4)]
for i, row in enumerate(rows):
    for j, txt in enumerate(row):
        c = tb.rows[i].cells[j]; c.width = widths[j]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(txt)
        r.font.size = DPt(8.5); r.font.bold = (i == 0)
        r.font.name = 'Noto Serif CJK SC'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        if i == 0:
            r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
        p.paragraph_format.line_spacing = 1.2
Pp('')
Pp('时间核验：4+4+6+4+10+3+6+5+4+4+4+1+3+2 = 60 分钟。', True)

H('八、板书设计')
Pp('【左板 · 定义区】曲边梯形图 →（分割 / 近似代替 / 求和 / 取极限）→ ∫ₐᵇ f(x)dx = lim(n→∞) Σᵢ₌₁ⁿ f(ξᵢ)·(b−a)/n')
Pp('【中板 · 定理区】面积函数 A(x)=∫ₐˣf(t)dt 图 → ΔA ≈ f(x)h → 积分中值定理 ΔA = f(ξ)h → A′(x)=f(x) → ∫ₐᵇf = F(b)−F(a)')
Pp('【右板 · 例题区】例1 用定义算 ∫₀¹x²dx（全过程保留不擦）｜例2 分段求面积 8/3｜例3 位移与路程｜例4 做功 16 J')
Pp('【副板 · 警戒区】红笔：积分值 ≠ 面积　·　∫dx/x = ln|x| (n = −1 例外)　·　常用原函数表')
Pp('说明：例1 的黎曼和推导必须完整板演并全程保留，它是本课"定义先于定理"的物证；'
   '课末指着左板与中板收束——"左边是定义，中间是定理"。')

H('九、分层作业')
Pp('A 层 · 基础（全体必做，约 15 分钟）', True)
Pp('① 用定义（黎曼和）计算 ∫₀¹ x dx；　② ∫₁² (x² + 1/x²) dx；　③ ∫₀^π sin x dx。')
Pp('B 层 · 提高（多数完成，约 20 分钟）', True)
Pp('④ ∫₀³ |x² − 4| dx；　⑤ 求 y = x² 与 y = 2x 所围图形的面积；'
   '　⑥ 已知 v(t) = t² − 4t + 3 (m/s)，求 t ∈ [0,3] 内质点的位移与路程。')
Pp('C 层 · 挑战（学有余力，选做）', True)
Pp('⑦ 证明 ∫₀¹ xⁿ dx = 1/(n+1)，并对 n = 3 用定义（不得使用 N-L 公式）给出证明；'
   '　⑧ 弹簧劲度系数 k = 300 N/m，将其从伸长 0.1 m 拉到伸长 0.3 m，求外力做的功；'
   '　⑨ 探究 ∫₁^{+∞} dx/x^p 何时收敛、何时发散，并据此解释 Gabriel 号角"体积有限而表面积无限"的临界指数。')

H('十、教学反思（课后填写）')
for t in ['1. 目标达成度：学生能否准确复述"定积分是和式的极限"？例1 用定义硬算的正确率约为　　　%。',
          '2. 难点突破：夹逼可视化是否真正消除了"无限求和能否收敛"的疑虑？还有哪些学生停留在死记公式？',
          '3. 定理证明：面积函数 + 中值定理的直观证明，学生跟得上吗？是否需要再放慢一档？',
          '4. 易错点：积分值 ≠ 面积的反例是否达到了预期的"震撼—纠错"效果？',
          '5. 拔高环节：Gabriel 号角引发的讨论质量如何？是否有学生主动追问 p 级数的收敛性？',
          '6. 时间调控：实际用时　　　分钟，超时/提前的环节是　　　　，下次调整方案：',
          '7. 改进设想：']:
    p = Pp(t)
    p.paragraph_format.space_after = DPt(14)
doc.add_paragraph('　\n　\n　')

docx_path = os.path.join(OUT, '教案_02_定积分与微积分基本定理.docx')
doc.save(docx_path)
print('docx:', docx_path)
print('OK')
