# -*- coding: utf-8 -*-
"""第06讲：三角函数的图像变换与简谐运动 —— 60分钟课件包生成脚本"""
import sys, os
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch

OUT = '/mnt/c/Users/mi/Desktop/数学名师课件包/06_三角函数图像与变换'
FIG = os.path.join(OUT, 'figures')
os.makedirs(FIG, exist_ok=True)
F = lambda n: os.path.join(FIG, n)
PI = np.pi

# ============================================================ 配图
figs = []

def reg(p):
    figs.append(p); return p

# --- fig01 单位圆展开成正弦曲线（右栏 5.9in 画布，最小标注） ---
def fig01():
    fig, axes = plt.subplots(1, 2, figsize=(5.9, 3.3),
                             gridspec_kw={'width_ratios': [1, 1.35]})
    ax = axes[0]
    t = np.linspace(0, 2 * PI, 400)
    ax.plot(np.cos(t), np.sin(t), color=k.M_ACC2, lw=2)
    a = 2 * PI / 3
    ax.plot([0, np.cos(a)], [0, np.sin(a)], color=k.M_ACC, lw=2.2)
    ax.plot([np.cos(a), np.cos(a)], [0, np.sin(a)], color=k.M_RED, lw=2.2, ls='--')
    ax.plot([np.cos(a)], [np.sin(a)], 'o', color=k.M_ACC, ms=9)
    ax.text(np.cos(a) - 0.10, np.sin(a) + 0.22, '$P$', color=k.M_ACC, ha='center')
    ax.set_xlim(-1.45, 1.45); ax.set_ylim(-1.35, 1.75); ax.set_aspect('equal')
    ax.set_xticks([]); ax.set_yticks([])
    k.style_axes(ax, '', '', grid=False)
    ax.set_title('单位圆', color=k.M_INK)

    ax = axes[1]
    x = np.linspace(0, 2 * PI, 600)
    ax.plot(x, np.sin(x), color=k.M_ACC2, lw=2.4)
    ax.plot([a], [np.sin(a)], 'o', color=k.M_ACC, ms=9)
    ax.plot([0, a], [np.sin(a), np.sin(a)], color=k.M_RED, lw=1.4, ls=':')
    ax.plot([a, a], [0, np.sin(a)], color=k.M_RED, lw=2.0, ls='--')
    ax.set_xticks([0, PI, 2 * PI])
    ax.set_xticklabels(['$O$', r'$\pi$', r'$2\pi$'])
    ax.set_yticks([])
    ax.set_ylim(-1.35, 1.75); ax.set_xlim(-0.35, 2 * PI + 0.35)
    k.style_axes(ax, '', '')
    ax.set_title(r'$y=\sin x$', color=k.M_INK)
    return k.save_fig(fig, reg(F('01_unit_circle_unroll.png')))

# --- fig02 周期性（右栏 6.1in） ---
def fig02():
    fig, ax = k.new_fig(6.1, 3.8)
    x = np.linspace(-2 * PI, 4 * PI, 1400)
    ax.plot(x, np.sin(x), color=k.M_ACC2, lw=2.4)
    ax.axvspan(0, 2 * PI, color=k.M_GRN, alpha=0.13)
    ax.annotate('', xy=(2 * PI, 1.30), xytext=(0, 1.30),
                arrowprops=dict(arrowstyle='<->', color=k.M_GRN, lw=1.8))
    ax.text(PI, 1.55, r'$T=2\pi$', color=k.M_GRN, ha='center')
    ax.set_xticks([-2 * PI, 0, 2 * PI, 4 * PI])
    ax.set_xticklabels([r'$-2\pi$', '$O$', r'$2\pi$', r'$4\pi$'])
    ax.set_yticks([-1, 1])
    ax.set_ylim(-1.6, 2.6)
    k.style_axes(ax, 'x', 'y')
    ax.set_title('周期性', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('02_period.png')))

# --- fig03 五点作图法（整幅 9.8in） ---
def fig03():
    fig, ax = k.new_fig(9.8, 4.6)
    x = np.linspace(-0.3, 2 * PI + 0.3, 600)
    ax.plot(x, np.sin(x), color=k.M_ACC2, lw=2.6, zorder=2)
    px = np.array([0, PI / 2, PI, 3 * PI / 2, 2 * PI]); py = np.sin(px)
    ax.plot(px, py, 'o', color=k.M_ACC, ms=12, zorder=3)
    labs = ['$(0,0)$', r'$(\frac{\pi}{2},1)$', r'$(\pi,0)$',
            r'$(\frac{3\pi}{2},-1)$', r'$(2\pi,0)$']
    offs = [(-46, 20), (0, 20), (0, 20), (0, -56), (-14, 22)]
    for i, (a, b) in enumerate(zip(px, py)):
        ax.annotate(labs[i], (a, b), textcoords='offset points',
                    xytext=offs[i], ha='center', color=k.M_ACC)
        ax.plot([a, a], [0, b], color=k.M_RULE, lw=1.2, ls='--', zorder=1)
    ax.set_xticks(px)
    ax.set_xticklabels(['$0$', r'$\frac{\pi}{2}$', r'$\pi$', r'$\frac{3\pi}{2}$', r'$2\pi$'])
    ax.set_yticks([-1, 1])
    ax.set_ylim(-2.7, 2.4)
    k.style_axes(ax, 'x', 'y')
    ax.set_title('五点作图法：零—峰—零—谷—零', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('03_five_points.png')))

# --- fig04 A 振幅（整幅 9.4in） ---
def fig04():
    fig, ax = k.new_fig(9.4, 4.6)
    x = np.linspace(0, 2 * PI, 600)
    for A, c, ls in [(1, k.M_RULE, '--'), (2, k.M_ACC, '-'), (0.5, k.M_ACC2, '-')]:
        ax.plot(x, A * np.sin(x), color=c, lw=2.6, ls=ls, label=rf'$y={A}\sin x$')
    ax.axhline(2, color=k.M_ACC, lw=1, ls=':'); ax.axhline(-2, color=k.M_ACC, lw=1, ls=':')
    ax.annotate('', xy=(PI / 2, 2), xytext=(PI / 2, 0),
                arrowprops=dict(arrowstyle='<->', color=k.M_RED, lw=1.8))
    ax.text(PI / 2 + 0.22, 1.55, r'$A=2$', color=k.M_RED)
    ax.set_xticks([0, PI, 2 * PI])
    ax.set_xticklabels(['$O$', r'$\pi$', r'$2\pi$'])
    ax.set_yticks([-2, 2])
    ax.set_ylim(-2.8, 5.2)
    k.style_axes(ax, 'x', 'y')
    ax.legend(loc='upper center', ncol=3, frameon=False)
    ax.set_title(r'参数 $A$：纵向拉伸，值域 $[-A,A]$', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('04_param_A.png')))

# --- fig05 ω 角频率（整幅 9.4in） ---
def fig05():
    fig, ax = k.new_fig(9.4, 4.6)
    x = np.linspace(0, 2 * PI, 900)
    ax.plot(x, np.sin(x), color=k.M_RULE, lw=2.6, ls='--', label=r'$\sin x$')
    ax.plot(x, np.sin(2 * x), color=k.M_ACC, lw=2.6, label=r'$\sin 2x$')
    ax.plot(x, np.sin(0.5 * x), color=k.M_ACC2, lw=2.6, label=r'$\sin\frac{x}{2}$')
    ax.annotate('', xy=(PI, -1.45), xytext=(0, -1.45),
                arrowprops=dict(arrowstyle='<->', color=k.M_ACC, lw=1.8))
    ax.text(PI / 2, -2.45, r'$T=\pi$', color=k.M_ACC, ha='center')
    ax.set_xticks([0, PI, 2 * PI])
    ax.set_xticklabels(['$O$', r'$\pi$', r'$2\pi$'])
    ax.set_yticks([-1, 1])
    ax.set_ylim(-2.9, 3.6)
    k.style_axes(ax, 'x', 'y')
    ax.legend(loc='upper center', ncol=3, frameon=False)
    ax.set_title(r'参数 $\omega$：横向压缩，$T=\dfrac{2\pi}{|\omega|}$', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('05_param_omega.png')))

# --- fig06 φ 初相（整幅 9.4in） ---
def fig06():
    fig, ax = k.new_fig(9.4, 4.6)
    x = np.linspace(-PI / 2, 2 * PI, 900)
    ax.plot(x, np.sin(x), color=k.M_RULE, lw=2.6, ls='--', label=r'$\sin x$')
    ax.plot(x, np.sin(x + PI / 3), color=k.M_ACC, lw=2.6, label=r'左移 $\frac{\pi}{3}$')
    ax.plot(x, np.sin(x - PI / 3), color=k.M_ACC2, lw=2.6, label=r'右移 $\frac{\pi}{3}$')
    ax.annotate('', xy=(-PI / 3, 0.12), xytext=(0, 0.12),
                arrowprops=dict(arrowstyle='->', color=k.M_ACC, lw=2.0))
    ax.text(-PI / 3 - 0.10, 0.30, '左加', color=k.M_ACC, ha='center')
    ax.annotate('', xy=(PI / 3, -0.12), xytext=(0, -0.12),
                arrowprops=dict(arrowstyle='->', color=k.M_ACC2, lw=2.0))
    ax.text(PI / 3 + 0.10, -0.70, '右减', color=k.M_ACC2, ha='center')
    ax.set_xticks([0, PI, 2 * PI])
    ax.set_xticklabels(['$O$', r'$\pi$', r'$2\pi$'])
    ax.set_yticks([-1, 1])
    ax.set_ylim(-2.0, 3.5)
    k.style_axes(ax, 'x', 'y')
    ax.legend(loc='upper center', ncol=3, frameon=False)
    ax.set_title(r'参数 $\varphi$：左加右减（$\omega=1$）', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('06_param_phi.png')))

# --- fig07 k 中心线（右栏 6.1in） ---
def fig07():
    fig, ax = k.new_fig(6.1, 3.9)
    x = np.linspace(0, 2 * PI, 600)
    ax.plot(x, np.sin(x), color=k.M_RULE, lw=2.4, ls='--', label=r'$\sin x$')
    ax.plot(x, np.sin(x) + 1.5, color=k.M_ACC, lw=2.6, label=r'$\sin x+1.5$')
    ax.axhline(1.5, color=k.M_GRN, lw=1.8, ls=':')
    ax.text(2.55, 0.55, r'$y=k=1.5$', color=k.M_GRN)
    ax.set_xticks([0, PI, 2 * PI]); ax.set_xticklabels(['$O$', r'$\pi$', r'$2\pi$'])
    ax.set_yticks([1.5])
    ax.set_ylim(-1.7, 5.4)
    k.style_axes(ax, 'x', 'y')
    ax.legend(loc='upper center', ncol=2, frameon=False)
    ax.set_title(r'参数 $k$：上下平移', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('07_param_k.png')))

# --- fig08 易错核心：先平移后伸缩 vs 先伸缩后平移（整幅 10.6in，4 子图） ---
def fig08():
    fig, axes = plt.subplots(2, 2, figsize=(10.6, 6.0))
    x = np.linspace(-PI, 3 * PI / 2, 900)
    tgt = lambda t: np.sin(2 * t + PI / 3)

    ax = axes[0, 0]
    ax.plot(x, np.sin(x), color=k.M_RULE, lw=2, ls='--')
    ax.plot(x, np.sin(x + PI / 3), color=k.M_ACC, lw=2.8)
    ax.set_title(r'甲① 先左移 $\frac{\pi}{3}$', color=k.M_ACC)

    ax = axes[0, 1]
    ax.plot(x, np.sin(x + PI / 3), color=k.M_RULE, lw=2, ls='--')
    ax.plot(x, tgt(x), color=k.M_GRN, lw=3.0)
    ax.set_title(r'甲② 再横缩 $\frac{1}{2}$【对】', color=k.M_GRN)

    ax = axes[1, 0]
    ax.plot(x, np.sin(x), color=k.M_RULE, lw=2, ls='--')
    ax.plot(x, np.sin(2 * x), color=k.M_ACC2, lw=2.8)
    ax.set_title(r'乙① 先横缩 $\frac{1}{2}$', color=k.M_ACC2)

    ax = axes[1, 1]
    ax.plot(x, tgt(x), color=k.M_GRN, lw=3.0)
    ax.plot(x, np.sin(2 * (x + PI / 3)), color=k.M_RED, lw=2.6)
    ax.axvspan(-PI / 3, -PI / 6, color=k.M_RED, alpha=0.14)
    ax.annotate('对', xy=(-PI / 6, 0), xytext=(0.55, -1.72),
                color=k.M_GRN,
                arrowprops=dict(arrowstyle='->', color=k.M_GRN, lw=1.6))
    ax.annotate('错', xy=(-PI / 3, 0), xytext=(-PI - 0.05, -1.72),
                color=k.M_RED,
                arrowprops=dict(arrowstyle='->', color=k.M_RED, lw=1.6))
    ax.set_title(r'乙② 只能左移 $\frac{\pi}{6}$', color=k.M_RED)

    for ax in axes.ravel():
        ax.set_xticks([-PI / 3, 0, PI])
        ax.set_xticklabels([r'$-\frac{\pi}{3}$', '$O$', r'$\pi$'])
        ax.set_yticks([])
        ax.set_ylim(-2.1, 3.3)
        k.style_axes(ax, '', '')
    return k.save_fig(fig, reg(F('08_order_trap.png')))

# --- fig09 同图叠加：正确 vs 错误（整幅 9.6in） ---
def fig09():
    fig, ax = k.new_fig(9.6, 4.7)
    x = np.linspace(-PI / 2, 3 * PI / 2, 1000)
    ax.plot(x, np.sin(2 * x), color=k.M_RULE, lw=2.0, ls=':', label=r'$\sin 2x$')
    ax.plot(x, np.sin(2 * x + PI / 3), color=k.M_GRN, lw=3.2,
            label=r'对：左移 $\frac{\pi}{6}$')
    ax.plot(x, np.sin(2 * x + 2 * PI / 3), color=k.M_RED, lw=2.8, ls='--',
            label=r'错：左移 $\frac{\pi}{3}$')
    for xv, c in [(-PI / 6, k.M_GRN), (-PI / 3, k.M_RED)]:
        ax.plot([xv], [0], 'o', color=c, ms=12)
    ax.set_xticks([-PI / 3, 0, PI / 2, PI])
    ax.set_xticklabels([r'$-\frac{\pi}{3}$', '$O$', r'$\frac{\pi}{2}$', r'$\pi$'])
    ax.set_yticks([-1, 1])
    ax.set_ylim(-2.0, 3.8)
    k.style_axes(ax, 'x', 'y')
    ax.legend(loc='upper center', ncol=3, frameon=False)
    ax.set_title(r'平移量 $=\dfrac{\varphi}{\omega}$，不是 $\varphi$', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('09_shift_amount.png')))

# --- fig10 例1 由图像求解析式（右栏 6.1in） ---
def fig10():
    fig, ax = k.new_fig(6.1, 4.0)
    A, w, ph = 2.0, 2.0, PI / 6
    x = np.linspace(-PI / 2, PI, 800)
    ax.plot(x, A * np.sin(w * x + ph), color=k.M_ACC2, lw=2.8)
    ax.plot([PI / 6], [2], 'o', color=k.M_ACC, ms=11)
    ax.annotate(r'$(\frac{\pi}{6},2)$', (PI / 6, 2), textcoords='offset points',
                xytext=(-8, 14), ha='right', color=k.M_ACC)
    ax.plot([-PI / 12], [0], 'o', color=k.M_GRN, ms=11)
    ax.annotate(r'$(-\frac{\pi}{12},0)$', (-PI / 12, 0), textcoords='offset points',
                xytext=(-6, -62), ha='center', color=k.M_GRN)
    ax.plot([2 * PI / 3], [-2], 'o', color=k.M_RED, ms=11)
    ax.annotate(r'$(\frac{2\pi}{3},-2)$', (2 * PI / 3, -2), textcoords='offset points',
                xytext=(4, -58), ha='center', color=k.M_RED)
    ax.annotate('', xy=(2 * PI / 3, 3.5), xytext=(PI / 6, 3.5),
                arrowprops=dict(arrowstyle='<->', color=k.M_SLATE, lw=1.8))
    ax.text(5 * PI / 12, 3.75, r'$\frac{T}{2}$', color=k.M_SLATE, ha='center')
    ax.set_xticks([0, PI / 2]); ax.set_xticklabels(['$O$', r'$\frac{\pi}{2}$'])
    ax.set_yticks([])
    ax.set_ylim(-4.8, 5.6)
    k.style_axes(ax, 'x', '')
    ax.set_title('例1：由图求解析式', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('10_ex1_read_graph.png')))

# --- fig11 例2 单调区间与对称（右栏 6.1in） ---
def fig11():
    fig, ax = k.new_fig(6.1, 4.0)
    x = np.linspace(-PI / 2, PI, 900)
    y = 2 * np.sin(2 * x + PI / 6)
    ax.axvspan(-PI / 3, PI / 6, color=k.M_GRN, alpha=0.15)
    ax.axvspan(PI / 6, 2 * PI / 3, color=k.M_RED, alpha=0.11)
    ax.plot(x, y, color=k.M_ACC2, lw=2.8)
    ax.axvline(PI / 6, color=k.M_ACC, lw=1.6, ls='--')
    ax.text(-PI / 12, -3.9, '增区间', color=k.M_GRN, ha='center')
    ax.text(5 * PI / 12, -3.9, '减区间', color=k.M_RED, ha='center')
    ax.text(PI / 6, 2.7, '对称轴', color=k.M_ACC, ha='center')
    ax.plot([5 * PI / 12], [0], 'o', color=k.M_INK, ms=10)
    ax.annotate('对称中心', (5 * PI / 12, 0), textcoords='offset points',
                xytext=(-4, 16), ha='center', color=k.M_INK)
    ax.set_xticks([-PI / 3, 0, 2 * PI / 3])
    ax.set_xticklabels([r'$-\frac{\pi}{3}$', '$O$', r'$\frac{2\pi}{3}$'])
    ax.set_yticks([-2, 2])
    ax.set_ylim(-4.8, 4.4)
    k.style_axes(ax, 'x', '')
    ax.set_title(r'例2：$y=2\sin(2x+\frac{\pi}{6})$', fontsize=20, color=k.M_INK)
    return k.save_fig(fig, reg(F('11_ex2_monotone.png')))

# --- fig12 简谐运动（整幅 10.2in） ---
def fig12():
    fig, axes = plt.subplots(1, 2, figsize=(10.2, 4.0),
                             gridspec_kw={'width_ratios': [1, 1.7]})
    ax = axes[0]
    ax.plot([-0.9, 0.9], [1.05, 1.05], color=k.M_INK, lw=3)
    for i in range(9):
        ax.plot([-0.9 + i * 0.225, -0.9 + i * 0.225], [1.0, 1.05], color=k.M_SLATE, lw=1)
    ys = np.linspace(1.0, -0.15, 300)
    ax.plot(0.18 * np.sin(np.linspace(0, 16 * PI, 300)), ys, color=k.M_ACC2, lw=1.8)
    ax.add_patch(plt.Rectangle((-0.28, -0.62), 0.56, 0.47, color=k.M_ACC))
    ax.axhline(-0.385, color=k.M_GRN, lw=1.4, ls='--', xmin=0.06, xmax=0.60)
    ax.text(-1.20, -1.35, '平衡位置', color=k.M_GRN)
    ax.annotate('', xy=(0.72, 0.30), xytext=(0.72, -1.05),
                arrowprops=dict(arrowstyle='<->', color=k.M_RED, lw=1.8))
    ax.text(0.85, -0.48, r'$2A$', color=k.M_RED)
    ax.set_xlim(-1.30, 1.75); ax.set_ylim(-1.75, 1.35); ax.axis('off')
    ax.set_title('弹簧振子', fontsize=22, color=k.M_INK)

    ax = axes[1]
    t = np.linspace(0, 4, 800)
    ax.plot(t, 3 * np.sin(PI * t + PI / 4), color=k.M_ACC2, lw=2.8)
    ax.axhline(3, color=k.M_RULE, lw=1, ls=':'); ax.axhline(-3, color=k.M_RULE, lw=1, ls=':')
    ax.annotate('', xy=(0.25, 3.0), xytext=(0.25, 0),
                arrowprops=dict(arrowstyle='<->', color=k.M_RED, lw=1.8))
    ax.text(0.40, 1.1, r'$A=3$', color=k.M_RED)
    ax.annotate('', xy=(2.25, -4.0), xytext=(0.25, -4.0),
                arrowprops=dict(arrowstyle='<->', color=k.M_GRN, lw=1.8))
    ax.text(1.25, -6.3, r'$T=2\,\mathrm{s}$', color=k.M_GRN, ha='center')
    ax.set_xticks([0, 2, 4]); ax.set_yticks([-3, 3])
    ax.set_ylim(-7.0, 4.4)
    k.style_axes(ax, 't', 'y')
    ax.set_title(r'$y=3\sin(\pi t+\frac{\pi}{4})$', fontsize=20, color=k.M_INK)
    return k.save_fig(fig, reg(F('12_shm_spring.png')))

# --- fig13 交流电（整幅 9.6in） ---
def fig13():
    fig, ax = k.new_fig(9.6, 4.4)
    t = np.linspace(0, 0.05, 1000)
    ax.plot(t * 1000, 311 * np.sin(100 * PI * t), color=k.M_RED, lw=2.8)
    ax.axhline(220, color=k.M_GRN, lw=1.6, ls='--')
    ax.annotate('有效值 220 V', xy=(46, 220), xytext=(50, -560), ha='right',
                color=k.M_GRN,
                arrowprops=dict(arrowstyle='->', color=k.M_GRN, lw=1.6))
    ax.axhline(311, color=k.M_ACC, lw=1.2, ls=':')
    ax.text(0.5, 360, '峰值 311 V', color=k.M_ACC)
    ax.annotate('', xy=(20, -390), xytext=(0, -390),
                arrowprops=dict(arrowstyle='<->', color=k.M_ACC2, lw=1.8))
    ax.text(10, -620, r'$T=0.02\,\mathrm{s}$', color=k.M_ACC2, ha='center')
    ax.set_xticks([0, 20, 40]); ax.set_yticks([-311, 311])
    ax.set_ylim(-760, 600)
    k.style_axes(ax, 't / ms', 'u / V')
    ax.set_title(r'家用交流电 $u=311\sin(100\pi t)$', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('13_ac_current.png')))

# --- fig14 傅里叶逼近方波（整幅 10.6in，精简为 4 格） ---
def fig14():
    fig, axes = plt.subplots(2, 2, figsize=(10.6, 6.0))
    x = np.linspace(-PI, 3 * PI, 2000)
    sq = np.sign(np.sin(x))
    for ax, n in zip(axes.ravel(), [1, 3, 5, 40]):
        s = np.zeros_like(x)
        for i in range(1, n + 1):
            m = 2 * i - 1
            s += np.sin(m * x) / m
        s *= 4 / PI
        ax.plot(x, sq, color=k.M_RULE, lw=2.2, ls='--')
        ax.plot(x, s, color=k.M_ACC if n < 40 else k.M_GRN, lw=2.6)
        ax.set_ylim(-1.9, 1.9)
        ax.set_xticks([0, PI, 2 * PI])
        ax.set_xticklabels(['$O$', r'$\pi$', r'$2\pi$'])
        ax.set_yticks([])
        k.style_axes(ax, '', '')
        ax.set_title('$n=%d$ 项' % n if n < 40 else '$n=40$ 项 → 方波',
                     color=k.M_GRN if n == 40 else k.M_INK)
    return k.save_fig(fig, reg(F('14_fourier_square.png')))

# --- fig15 叠加成任意波（整幅 10.4in） ---
def fig15():
    fig, axes = plt.subplots(1, 2, figsize=(10.4, 4.6))
    x = np.linspace(0, 4 * PI, 1200)
    c = [np.sin(x), 0.5 * np.sin(2 * x + 1.0), 0.3 * np.sin(5 * x + 2.0)]
    ax = axes[0]
    for i, (cc, col, nm) in enumerate(zip(c, [k.M_ACC2, k.M_ACC, k.M_RED],
                                          ['基频', '2 次谐波', '5 次谐波'])):
        ax.plot(x, cc + (1 - i) * 2.8, color=col, lw=2.2)
        ax.axhline((1 - i) * 2.8, color=k.M_RULE, lw=0.8, ls=':')
        ax.text(4 * PI + 0.8, (1 - i) * 2.8, nm, color=col, va='center')
    ax.set_xlim(0, 4 * PI + 9.0); ax.set_ylim(-4.6, 4.6)
    ax.axis('off')
    ax.set_title('三个正弦分量', fontsize=22, color=k.M_INK)

    ax = axes[1]
    ax.plot(x, sum(c), color=k.M_GRN, lw=3.0)
    ax.set_xticks([0, 2 * PI, 4 * PI])
    ax.set_xticklabels(['$O$', r'$2\pi$', r'$4\pi$'])
    ax.set_yticks([-1, 1])
    ax.set_ylim(-2.6, 2.8)
    k.style_axes(ax, 'x', 'y')
    ax.set_title('叠加 = 复杂波形', fontsize=22, color=k.M_GRN)
    return k.save_fig(fig, reg(F('15_superposition.png')))

# --- fig16 变换链（整幅 9.6in） ---
def fig16():
    fig, ax = k.new_fig(9.6, 4.7)
    x = np.linspace(-PI / 2, 2 * PI, 1000)
    ax.plot(x, np.sin(x), color=k.M_RULE, lw=2.2, ls=':', label=r'$\sin x$')
    ax.plot(x, np.sin(x - PI / 4), color=k.M_ACC2, lw=2.2, ls='--',
            label=r'①右移 $\frac{\pi}{4}$')
    ax.plot(x, np.sin(0.5 * x - PI / 4), color=k.M_ACC, lw=2.6,
            label=r'②横伸 2 倍')
    ax.plot(x, 3 * np.sin(0.5 * x - PI / 4), color=k.M_GRN, lw=3.0,
            label=r'③纵伸 3 倍【对】')
    ax.set_xticks([0, PI, 2 * PI]); ax.set_xticklabels(['$O$', r'$\pi$', r'$2\pi$'])
    ax.set_yticks([-3, 3])
    ax.set_ylim(-4.4, 10.0)
    k.style_axes(ax, 'x', 'y')
    ax.legend(loc='upper center', ncol=2, frameon=False)
    ax.set_title(r'三步到 $y=3\sin(\frac{x}{2}-\frac{\pi}{4})$', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('16_variant_chain.png')))

# --- fig17 摩天轮（右栏 6.1in） ---
def fig17():
    fig, axes = plt.subplots(1, 2, figsize=(6.1, 4.0),
                             gridspec_kw={'width_ratios': [1, 1.3]})
    ax = axes[0]
    th = np.linspace(0, 2 * PI, 400)
    ax.plot(50 * np.cos(th), 60 + 50 * np.sin(th), color=k.M_ACC2, lw=2.2)
    ax.plot([0], [60], 'o', color=k.M_INK, ms=6)
    ax.plot([0, -30], [60, 0], color=k.M_SLATE, lw=2)
    ax.plot([0, 30], [60, 0], color=k.M_SLATE, lw=2)
    ax.axhline(0, color=k.M_INK, lw=2)
    a0 = -PI / 2 + 1.1
    ax.plot([50 * np.cos(a0)], [60 + 50 * np.sin(a0)], 'o', color=k.M_ACC, ms=11)
    ax.annotate('', xy=(0, 60), xytext=(0, 0),
                arrowprops=dict(arrowstyle='<->', color=k.M_GRN, lw=1.6))
    ax.text(-5, -30, r'$k=60$', color=k.M_GRN, ha='center', va='top')
    ax.annotate('', xy=(-64, 60), xytext=(-64, 110),
                arrowprops=dict(arrowstyle='<->', color=k.M_RED, lw=1.6))
    ax.text(-55, 120, r'$A=50$', color=k.M_RED, ha='center')
    ax.set_xlim(-105, 95); ax.set_ylim(-72, 152); ax.set_aspect('equal'); ax.axis('off')
    ax.set_title('摩天轮', fontsize=22, color=k.M_INK)

    ax = axes[1]
    t = np.linspace(0, 24, 800)
    ax.plot(t, 60 - 50 * np.cos(PI * t / 6), color=k.M_ACC, lw=3.0)
    ax.axhline(60, color=k.M_GRN, lw=1.5, ls='--')
    ax.plot([0], [10], 'o', color=k.M_RED, ms=9)
    ax.plot([6], [110], 'o', color=k.M_ACC2, ms=9)
    ax.annotate('', xy=(12, -42), xytext=(0, -42),
                arrowprops=dict(arrowstyle='<->', color=k.M_SLATE, lw=1.6))
    ax.text(6, -92, r'$T=12$', color=k.M_SLATE, ha='center')
    ax.set_xticks([0, 12, 24]); ax.set_yticks([10, 60, 110])
    ax.set_ylim(-112, 142)
    k.style_axes(ax, 't / min', 'h / m')
    ax.set_title('高度—时间', fontsize=22, color=k.M_INK)
    return k.save_fig(fig, reg(F('17_ferris_wheel.png')))

# --- fig18 知识网络（整幅 9.8in） ---
def fig18():
    fig, ax = k.new_fig(9.8, 4.7)
    ax.axis('off')
    ax.set_xlim(0, 10); ax.set_ylim(0, 5.6)
    def box(cx, cy, w, h, t, c, fs=20):
        ax.add_patch(plt.Rectangle((cx - w / 2, cy - h / 2), w, h, facecolor='white',
                                   edgecolor=c, lw=2.0, zorder=2))
        ax.text(cx, cy, t, ha='center', va='center', fontsize=fs, color=k.M_INK, zorder=3)
    def arw(p, q, c):
        ax.add_patch(FancyArrowPatch(p, q, arrowstyle='-|>', mutation_scale=16,
                                     color=c, lw=1.8, zorder=1))
    box(5, 5.05, 3.4, 0.8, '单位圆 · 周期性', k.M_ACC2)
    box(5, 3.55, 4.6, 0.85, r'$y=A\sin(\omega x+\varphi)+k$', k.M_ACC, 22)
    arw((5, 4.62), (5, 4.02), k.M_ACC2)
    for cx, t, c in [(1.35, r'$A$ 振幅', k.M_RED),
                     (3.8, r'$\omega$ 周期', k.M_GRN),
                     (6.25, r'$\varphi$ 初相', k.M_ACC),
                     (8.7, r'$k$ 中心线', k.M_ACC2)]:
        box(cx, 2.1, 2.3, 0.8, t, c)
        arw((5, 3.1), (cx, 2.55), c)
    box(5, 0.6, 8.4, 0.85, '一切周期现象 = 正弦的叠加', k.M_GRN)
    arw((5, 1.68), (5, 1.06), k.M_GRN)
    return k.save_fig(fig, reg(F('18_summary_map.png')))

for fn in [fig01, fig02, fig03, fig04, fig05, fig06, fig07, fig08, fig09,
           fig10, fig11, fig12, fig13, fig14, fig15, fig16, fig17, fig18]:
    fn()
print('figures:', len(figs))

# ============================================================ PPT
prs = k.new_deck()
FM = lambda n: F('fx_%02d.png' % n)
_fc = [0]
def fx(slide, tex, **kw):
    _fc[0] += 1
    return k.formula(slide, tex, out=F('fx_%02d.png' % _fc[0]), **kw)

# 1 封面
k.title_slide(prs, '第06讲　三角函数的图像变换与简谐运动',
              '从单位圆到傅里叶级数——一切周期现象的基石',
              '高中数学名师课堂', '60 分钟 · 必修第一册 第五章')

# 2 学习目标
s = k.content_slide(prs, '本讲学习目标', '导入')
k.bullets(s, [
    '理解正弦曲线由单位圆"展开"而来，从旋转把握周期性',
    '掌握五点作图法，规范画出 y=Asin(ωx+φ)+k 的图像',
    ('列表 → 描点 → 光滑连线，令 ωx+φ 依次取 0, π/2, π, 3π/2, 2π', 1),
    '理解图像变换的本质是"自变量的替换"',
    ('★ 平移量是 φ/ω 而不是 φ —— 全章最大易错点', 1),
    '会由图像求解析式、求单调区间与对称轴',
    '能用简谐运动模型解决实际问题（振幅/角频率/初相/中心线）',
    '感受傅里叶级数：任意周期波都可由正弦叠加而成',
], y=1.55, w=6.2, size=17)
k.picture(s, F('01_unit_circle_unroll.png'), x=7.2, y=2.1, w=5.9)
k.callout(s, '一句话主线：圆的旋转 → 正弦的波动 → 世界的周期。', x=0.85, y=6.35, w=6.2, h=0.8)

# 3 第1幕
k.section_slide(prs, '第 1 幕 · 溯源', '正弦曲线是怎么"长"出来的', '0–8 min')

# 4 单位圆展开
s = k.content_slide(prs, '一、单位圆：正弦的出生地', '8 min')
k.full_picture(s, F('01_unit_circle_unroll.png'), y=1.5, w=8.0)
k.callout(s, '动点 P 绕单位圆匀速旋转，它的纵坐标随角 x 的变化规律，就是 y = sin x。'
             '正弦曲线是圆周运动投在墙上的"影子"。', x=1.5, y=6.35, w=10.3, h=0.85, kind='note')

# 5 周期性
s = k.content_slide(prs, '二、周期性：转一圈，回到原处')
k.picture(s, F('02_period.png'), x=6.9, y=1.8, w=6.1)
k.bullets(s, [
    '旋转 2π 后动点回到同一位置 ⇒ 纵坐标重复',
    '定义：若存在非零常数 T，使 f(x+T)=f(x) 恒成立，则 f 为周期函数',
    'y=sin x 与 y=cos x 的最小正周期均为 2π',
    'y=tan x 的最小正周期为 π',
], y=1.7, w=5.8, size=17)
fx(s, r'$\sin(x+2k\pi)=\sin x,\quad k\in\mathbb{Z}$', x=0.85, y=4.5, w=5.6, size=0.62)
k.callout(s, '周期性是三角函数区别于一切多项式函数的根本特征——\n这也是它能描述"重复发生的事"的原因。',
          x=0.85, y=5.6, w=5.6, h=1.2)

# 6 五点法
s = k.content_slide(prs, '三、五点作图法：规范流程', '必考')
k.full_picture(s, F('03_five_points.png'), y=1.5, w=9.8)
k.callout(s, '流程：①令整体 u=ωx+φ 依次取 0, π/2, π, 3π/2, 2π　②反解 x　'
             '③算 y　④列表描点　⑤光滑连线。切忌"看着像"随手画。',
          x=1.4, y=6.35, w=10.5, h=0.85)

# 7 五点法表格演示（文字）
s = k.content_slide(prs, '五点法示范：y = 2sin(2x + π/6)')
fx(s, r'$u=2x+\frac{\pi}{6},\qquad x=\frac{u}{2}-\frac{\pi}{12}$', x=0.85, y=1.55, w=6.0, size=0.6)
k.bullets(s, [
    'u = 0    →  x = −π/12,  y = 0　（上升零点）',
    'u = π/2  →  x = π/6,    y = 2　（最高点）',
    'u = π    →  x = 5π/12,  y = 0　（对称中心）',
    'u = 3π/2 →  x = 2π/3,   y = −2（最低点）',
    'u = 2π   →  x = 11π/12, y = 0　（一周期末）',
], y=2.6, w=5.8, size=16)
k.picture(s, F('11_ex2_monotone.png'), x=6.9, y=1.8, w=6.1)
k.callout(s, '五个 x 值等差，公差恰为 T/4 = π/4。发现这一点，作图速度翻倍。',
          x=0.85, y=5.9, w=5.8, h=0.9, kind='note')

# 8 第2幕
k.section_slide(prs, '第 2 幕 · 探究', '四个参数，四种变换', '8–22 min')

# 9 A
s = k.content_slide(prs, '参数 A：振幅（纵向伸缩）')
k.full_picture(s, F('04_param_A.png'), y=1.45, w=9.4)
k.callout(s, '纵坐标变为原来的 A 倍（横坐标不变）：值域 [−|A|, |A|]，周期不变。'
             '物理上 A 就是振幅——振动能达到的最大偏离。', x=1.9, y=6.3, w=9.5, h=0.85)

# 10 ω
s = k.content_slide(prs, '参数 ω：角频率（横向伸缩）')
k.full_picture(s, F('05_param_omega.png'), y=1.45, w=9.4)
k.callout(s, 'ω>1 波被压密（周期变小），0<ω<1 波被拉疏。切记横坐标"缩为原来的 1/ω 倍"，'
             '和 ω 的变化方向相反。', x=1.9, y=6.3, w=9.5, h=0.85, kind='warn')

# 11 φ
s = k.content_slide(prs, '参数 φ：初相（左右平移）')
k.full_picture(s, F('06_param_phi.png'), y=1.45, w=9.4)
k.callout(s, '"左加右减"是对自变量 x 而言的。当 ω=1 时平移量恰为 |φ|；'
             'ω≠1 时——请立刻翻到下一页。', x=1.9, y=6.3, w=9.5, h=0.85, kind='warn')

# 12 k
s = k.content_slide(prs, '参数 k：中心线（上下平移）')
k.picture(s, F('07_param_k.png'), x=6.9, y=1.9, w=6.1)
k.bullets(s, [
    '整体上移 k 个单位（上加下减）',
    '中心线（平衡位置）从 y=0 移到 y=k',
    '值域变为 [k−|A|, k+|A|]',
    '实际含义：潮汐的平均水位、摩天轮的中心高度、直流偏置',
], y=1.9, w=5.8, size=17)
fx(s, r'$y=A\sin(\omega x+\varphi)+k$', x=0.85, y=4.6, w=5.6, size=0.68)
k.callout(s, '完整四参数模型：A 管高低，ω 管快慢，φ 管早晚，k 管基准。',
          x=0.85, y=5.9, w=5.8, h=0.95)

# 13 第3幕
k.section_slide(prs, '第 3 幕 · 攻坚', '变换的本质：自变量的替换', '22–32 min')

# 14 本质
s = k.content_slide(prs, '本质：把 x 换成什么？', '核心')
fx(s, r'$y=\sin(\omega x+\varphi)=\sin\left[\,\omega\left(x+\dfrac{\varphi}{\omega}\right)\right]$',
   x=1.3, y=1.6, w=10.6, size=0.95)
k.bullets(s, [
    '平移看的是"x 本身被加/减了多少"，不是括号里最后那个常数',
    ('把 ω 提出来，括号内 x 后面跟的 φ/ω 才是真正的平移量', 1),
    '所以：由 y=sin ωx 得到 y=sin(ωx+φ)，需左移 φ/ω 个单位',
    '而由 y=sin x 得到 y=sin(x+φ)，才是左移 φ 个单位',
], y=3.5, w=11.5, size=18)
k.callout(s, '口诀：先提系数，再看平移。ω 一旦不是 1，φ 就要"打折"。',
          x=1.3, y=6.2, w=10.6, h=0.9, kind='warn')

# 15 易错对比 4 子图
s = k.content_slide(prs, '易错攻坚：先平移后伸缩 vs 先伸缩后平移', '易错')
k.full_picture(s, F('08_order_trap.png'), y=1.4, w=10.6)

# 16 叠加对比
s = k.content_slide(prs, '两条曲线，一目了然', '易错')
k.full_picture(s, F('09_shift_amount.png'), y=1.45, w=9.6)
k.callout(s, '记住：顺序不同，平移量必须跟着改。先平移(φ)后伸缩 = 先伸缩后平移(φ/ω)，殊途同归。',
          x=1.9, y=6.35, w=9.5, h=0.8, kind='warn')

# 17 变换链
s = k.content_slide(prs, '变换链演练：三步到位')
k.full_picture(s, F('16_variant_chain.png'), y=1.45, w=9.6)
k.callout(s, '注意 ω=1/2：此处先平移 π/4，再横向伸为 2 倍——因为 φ/ω = (−π/4)/(1/2) = −π/2 与之等价。'
             '两条路，同一终点。', x=1.9, y=6.35, w=9.5, h=0.85, kind='note')

# 18 第4幕
k.section_slide(prs, '第 4 幕 · 应用', '例题精讲与实际建模', '32–48 min')

# 19 例1
s = k.content_slide(prs, '例1　由图像求解析式', '例题')
k.picture(s, F('10_ex1_read_graph.png'), x=6.9, y=1.6, w=6.1)
k.bullets(s, [
    '题：图像最高点 (π/6, 2)，相邻最低点 (2π/3, −2)，求 y=Asin(ωx+φ) (A>0, ω>0, |φ|<π/2)',
    '① 由最值得 A = 2',
    '② 峰谷相距半周期：T/2 = 2π/3 − π/6 = π/2 ⇒ T = π ⇒ ω = 2π/T = 2',
    '③ 代最高点：2·(π/6) + φ = π/2 ⇒ φ = π/6（满足 |φ|<π/2）',
], y=1.6, w=5.8, size=15)
fx(s, r'$y=2\sin\left(2x+\dfrac{\pi}{6}\right)$', x=1.0, y=5.0, w=5.2, size=0.72)
k.callout(s, '定序：A 看幅度 → ω 看周期 → φ 代特殊点（优先代最高点，别代零点，避免增根）。',
          x=0.85, y=6.15, w=6.0, h=0.95)

# 20 例2
s = k.content_slide(prs, '例2　单调区间 · 对称轴 · 对称中心', '例题')
k.picture(s, F('11_ex2_monotone.png'), x=6.9, y=1.6, w=6.1)
fx(s, r'$y=2\sin\left(2x+\dfrac{\pi}{6}\right)$', x=0.85, y=1.5, w=4.4, size=0.5)
k.bullets(s, [
    '整体代换 u = 2x + π/6，套 y=sin u 的性质',
    '增区间：−π/2+2kπ ≤ u ≤ π/2+2kπ ⇒ x ∈ [−π/3+kπ, π/6+kπ]',
    '对称轴：u = π/2 + kπ ⇒ x = π/6 + kπ/2',
    '对称中心：u = kπ ⇒ x = −π/12 + kπ/2，中心为 (−π/12+kπ/2, 0)',
], y=2.6, w=5.8, size=15)
k.callout(s, '整体代换是三角函数性质题的万能钥匙：把 ωx+φ 当作一个新的 u。'
             '注意 ω<0 时要先用诱导公式化正。', x=0.85, y=6.1, w=5.8, h=1.0, kind='note')

# 21 例3 简谐运动
s = k.content_slide(prs, '例3　简谐运动应用题', '例题')
k.full_picture(s, F('12_shm_spring.png'), y=1.4, w=10.2)
k.bullets(s, [
    '题：某弹簧振子位移 y=3sin(πt + π/4) (cm)。求振幅、周期、频率、初相；'
    '并求 t=0 时的位移与首次到达最高点的时刻。',
], x=0.9, y=5.55, w=11.6, size=15)
k.callout(s, 'A=3 cm；T=2π/π=2 s；f=1/T=0.5 Hz；初相 π/4。'
             't=0：y=3sin(π/4)=3√2/2 ≈ 2.12 cm。首次最高：πt+π/4=π/2 ⇒ t=1/4 s。',
          x=0.9, y=6.25, w=11.6, h=1.0)

# 22 交流电
s = k.content_slide(prs, '现实：插座里的正弦波')
k.full_picture(s, F('13_ac_current.png'), y=1.45, w=9.6)
k.callout(s, '我国民用电 u = 311 sin(100πt) V：峰值 311 V，ω=100π，T=0.02 s，f=50 Hz，'
             '有效值 311/√2 ≈ 220 V。', x=1.9, y=6.3, w=9.5, h=0.85, kind='note')

# 23 变式1
s = k.content_slide(prs, '变式1　摩天轮建模', '变式')
k.picture(s, F('17_ferris_wheel.png'), x=6.9, y=1.6, w=6.1)
k.bullets(s, [
    '题：摩天轮半径 50 m，中心离地 60 m，12 min 转一圈，从最低点开始计时。'
    '求高度 h 关于时间 t 的函数，并求首次到 85 m 的时刻。',
    '① A=50，k=60，T=12 ⇒ ω = 2π/12 = π/6',
    '② t=0 在最低点：h(0)=10 ⇒ h = 60 − 50cos(πt/6)',
    '③ 令 60 − 50cos(πt/6) = 85 ⇒ cos(πt/6) = −1/2 ⇒ πt/6 = 2π/3 ⇒ t = 4 min',
], y=1.6, w=5.8, size=14)
k.callout(s, '实际建模四问：最大最小（定 A、k）→ 转一圈多久（定 ω）→ 起点在哪（定 φ）→ 代入求解。',
          x=0.85, y=6.15, w=5.8, h=0.95)

# 24 变式2
s = k.content_slide(prs, '变式2　变换顺序辨析', '变式')
k.bullets(s, [
    '题：把 y=sin x 的图像上所有点横坐标缩短为原来的 1/2（纵坐标不变），'
    '再向右平移 π/6 个单位，所得图像的解析式是（　）',
    'A. y=sin(2x − π/6)　　B. y=sin(2x − π/3)　　C. y=sin(2x + π/3)　　D. y=sin(x/2 − π/6)',
], x=0.9, y=1.6, w=11.6, size=17)
fx(s, r'$y=\sin 2x\ \ \longrightarrow\ (x\to x-\frac{\pi}{6})\ \longrightarrow\ \ y=\sin 2\left(x-\frac{\pi}{6}\right)=\sin\left(2x-\frac{\pi}{3}\right)$',
   x=0.9, y=3.5, w=11.4, size=0.78)
k.callout(s, '答案 B。陷阱选项 A 正是"把平移量直接塞进括号当常数"的典型错误——'
             '平移作用在 x 上，要连同 ω 一起分配。', x=0.9, y=5.6, w=11.6, h=1.1, kind='warn')

# 25 第5幕
k.section_slide(prs, '第 5 幕 · 升华', '傅里叶级数：世界是正弦做的', '48–56 min')

# 26 叠加
s = k.content_slide(prs, '两个正弦相加，会得到什么？', '高光')
k.full_picture(s, F('15_superposition.png'), y=1.5, w=10.4)
k.callout(s, '几个频率成整数倍的正弦叠加，就能造出复杂的波形——乐器的音色、心电图、语音，皆是如此。',
          x=1.5, y=6.35, w=10.3, h=0.85, kind='note')

# 27 傅里叶（高光）
s = k.content_slide(prs, '★ 傅里叶级数：正弦如何"拼"出方波', '高光')
k.full_picture(s, F('14_fourier_square.png'), y=1.35, w=10.6)

# 28 傅里叶公式
s = k.content_slide(prs, '一个惊人的结论')
fx(s, r'$\frac{4}{\pi}\left(\sin x+\frac{\sin 3x}{3}+\frac{\sin 5x}{5}+\frac{\sin 7x}{7}+\cdots\right)$' + '  =  方波',
   x=0.7, y=1.6, w=11.9, size=0.82)
k.bullets(s, [
    '光滑的正弦，居然能拼出带尖角的方波——项数越多，越像',
    '傅里叶（Fourier, 1807）：任何"足够正常"的周期函数，都可展开为正弦与余弦的无穷级数',
    '这就是 MP3 压缩、JPEG 图像、5G 通信、地震分析、CT 成像共同的数学底座',
    '你今天画的每一条 y=Asin(ωx+φ)，都是这台机器上的一颗零件',
], y=3.4, w=11.7, size=17)
k.callout(s, '三角函数不是"考试用的曲线"，而是人类描述一切周期现象的通用语言。',
          x=0.9, y=6.3, w=11.6, h=0.85)

# 29 小结
s = k.content_slide(prs, '课堂小结：一张图收束全课', '56–60 min')
k.full_picture(s, F('18_summary_map.png'), y=1.45, w=9.8)
k.callout(s, '三句话带走：①正弦来自圆的旋转；②变换的本质是自变量替换，平移量 φ/ω；'
             '③一切周期，皆可正弦。', x=1.7, y=6.35, w=9.9, h=0.85)

# 30 分层作业
s = k.content_slide(prs, '分层作业', '作业')
k.bullets(s, [
    'A 基础层（人人过关）',
    ('用五点法作 y=2sin(x/2 + π/3) 一个周期内的图像；写出 A、T、φ、值域', 1),
    ('教材 P243 练习 1、2、3', 1),
    'B 提高层（多数完成）',
    ('由图像求解析式 2 题（含 |φ|<π/2 的取舍讨论）', 1),
    ('求 y=3sin(π/3 − 2x) 的单调递增区间与对称中心（注意 ω<0 先化正）', 1),
    ('潮汐问题：某港口水深 h=5+2sin(πt/6)，求何时可通行吃水 6 m 的货轮', 1),
    'C 挑战层（学有余力）',
    ('用计算器/Desmos 画出 sin x + sin3x/3 + … 的前 20 项和，观察吉布斯现象并写 200 字说明', 1),
    ('查阅：为什么钢琴的中央 C 和小提琴的中央 C 听起来不同？用谐波解释', 1),
], y=1.55, w=12.0, size=15)

# 31 板书提纲
s = k.content_slide(prs, '板书提纲', '板书')
k.bullets(s, [
    '【左栏 · 概念区】',
    ('单位圆 → sin x；周期 T=2π/|ω|', 1),
    ('y = A sin(ωx + φ) + k　A振幅 ω角频率 φ初相 k中心线', 1),
    ('★ y=sin(ωx+φ)=sin[ω(x+φ/ω)]　平移量 = φ/ω', 1),
    '【中栏 · 演算区】',
    ('五点法表格：u = 0, π/2, π, 3π/2, 2π ↔ x, y', 1),
    ('例1 由图求式：A→ω→φ 三步', 1),
    ('例2 整体代换 u=ωx+φ 求单调/对称', 1),
    ('例3 简谐运动：A、T、f、初相', 1),
], x=0.85, y=1.55, w=6.2, size=15)
k.bullets(s, [
    '【右栏 · 图像区】',
    ('sin x 标准图（保留全课）', 1),
    ('三参数变化对比草图', 1),
    ('易错对比：两条错开 π/6 的曲线（红笔）', 1),
    ('傅里叶逼近方波示意（收尾画）', 1),
    '【角落 · 易错本】',
    ('平移量 ≠ φ　　②ω<0 先化正　　③五点是 u 的五点', 1),
], x=7.0, y=1.55, w=5.6, size=15)

path_ppt = k.save(prs, os.path.join(OUT, '06_三角函数图像与变换.pptx'))
n_slides = len(prs.slides._sldIdLst)
print('slides:', n_slides, path_ppt)

# ============================================================ 教案 docx
from docx import Document
from docx.shared import Pt as DPt, RGBColor as DRGB, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Noto Serif CJK SC'
st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Cm(2.2)

def H(t, lv=1):
    p = doc.add_heading(t, level=lv)
    for r in p.runs:
        r.font.name = 'Noto Serif CJK SC'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
        r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
    return p

def P(t, bold=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(t); r.bold = bold; r.font.size = DPt(size)
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    return p

ttl = doc.add_paragraph()
ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run('第06讲　三角函数的图像变换与简谐运动　教学设计')
r.bold = True; r.font.size = DPt(18)
r.font.name = 'Noto Serif CJK SC'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('人教A版必修第一册 第五章 5.6　　课时：1 课时（60 分钟）')
r.font.size = DPt(11); r.font.name = 'Noto Serif CJK SC'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')

H('一、课题与课时', 1)
P('课题：三角函数的图像变换与简谐运动（y = A sin(ωx + φ) + k）')
P('课时：第 06 讲，1 课时，60 分钟；配套课件 31 页，配图 18 张。')

H('二、教材分析', 1)
P('本节位于人教A版必修第一册第五章末，是三角函数由"定义与性质"走向"模型与应用"的枢纽。'
  '教材先由单位圆上动点的旋转引出正弦线，再通过参数 A、ω、φ 的逐一变化，'
  '归纳出 y=A sin(ωx+φ) 与 y=sin x 的图像关系，最后以简谐运动为落点。'
  '其数学核心是"函数图像的变换实质是自变量的替换"，其应用核心是"用三角函数刻画周期现象"。'
  '本设计在教材基础上做两处增益：一是把单位圆与正弦曲线的"展开"做成联动图，'
  '让周期性可视化；二是在收尾引入傅里叶级数逼近方波的图像序列（不作证明要求），'
  '让学生在直观震撼中体会三角函数的普适地位，落实数学建模与直观想象素养。')

H('三、学情分析', 1)
P('知识基础：学生已掌握任意角三角函数定义、诱导公式、正余弦图像与基本性质，'
  '也在必修一学过 y=f(x+a)、y=Af(x) 的一般平移伸缩规律。')
P('能力基础：具备描点作图与整体代换的初步经验，但对"复合变换"的顺序敏感度不足。')
P('主要困难：① 混淆平移量 φ 与 φ/ω，这是历年最高频错误；'
  '② 变换顺序（先平移后伸缩 / 先伸缩后平移）对应的平移量不同，易张冠李戴；'
  '③ ω<0 时不会先用诱导公式化正；④ 实际问题中不会由物理量反推参数。')
P('对策：用"两条错开 π/6 的曲线"同图叠加打透易错点；用整体代换 u=ωx+φ 统一方法论。')

H('四、教学目标（三维）', 1)
P('1. 知识与技能', True)
P('（1）理解正弦曲线源于单位圆上动点纵坐标的变化，掌握周期性的本质；'
  '（2）熟练运用五点作图法画出 y=A sin(ωx+φ)+k 的图像；'
  '（3）掌握 A、ω、φ、k 对图像的作用，能由图像求解析式、求单调区间与对称轴/中心；'
  '（4）能用简谐运动模型解决振动、交流电、摩天轮、潮汐等实际问题。')
P('2. 过程与方法', True)
P('经历"观察—猜想—验证—归纳"的探究过程；通过对比图辨析变换顺序，'
  '发展直观想象与逻辑推理素养；通过实际建模，发展数学建模素养。')
P('3. 情感态度与价值观', True)
P('借助傅里叶级数逼近方波的图像序列，感受"简单元素叠加出复杂世界"的数学之美，'
  '体会三角函数作为描述一切周期现象通用语言的地位，激发进一步学习的兴趣。')

H('五、教学重点与难点', 1)
P('重点：y=A sin(ωx+φ)+k 中四个参数的几何意义与物理意义；五点作图法的规范流程；'
  '由图像求解析式。')
P('难点：图像变换的本质是自变量的替换——由 y=sin ωx 得到 y=sin(ωx+φ) 时，'
  '平移量是 φ/ω 而非 φ；不同变换顺序下平移量的相应调整。')
P('突破策略：先提取系数写成 sin[ω(x+φ/ω)]，从代数上说清；'
  '再用同图叠加的"正确曲线 vs 错误曲线"从图形上打透；最后用变式题当堂检测。')

H('六、教法与学法', 1)
P('教法：情境驱动 + 直观演示 + 变式训练。以单位圆动点情境导入，以对比图突破难点，'
  '以三例两变式滚动巩固，以傅里叶级数升华收束。')
P('学法：自主观察—合作辨析—独立演练。学生在参数探究环节以同桌为单位说出"我看到了什么变化"，'
  '在易错环节先独立判断再暴露错误，在建模环节完整走一遍"读题—定参—求解"流程。')

H('七、教学准备', 1)
P('教师：PPT 课件（31 页）、18 张 matplotlib 精确配图、几何画板/Desmos 备用、彩色粉笔（红笔专用于易错标注）。')
P('学生：直尺、圆规、方格纸（用于五点作图）、错题本。')

H('八、教学过程（分钟级时间轴）', 1)

rows = [
    ('0–2', '导入·2′', 'PPT1–2：展示课题与学习目标；提问"钟摆、心跳、潮汐、交流电有什么共同点？"',
     '思考并回答：都在重复、都有周期。', '以生活现象激活已有经验，点明本课主线：周期。'),
    ('2–8', '第1幕·6′', 'PPT3–5：演示单位圆动点投影→正弦曲线的展开联动图（图01）；'
     '引导得出 sin(x+2kπ)=sin x（图02）。',
     '观察动点纵坐标随角变化，口述"转一圈回到原处所以重复"。',
     '从定义源头理解周期性，避免死记；落实直观想象。'),
    ('8–14', '第1幕·6′', 'PPT6–7：讲解五点作图法五步流程（图03）；'
     '以 y=2sin(2x+π/6) 示范列表，强调令 u=ωx+φ 取五个特殊值。',
     '跟做列表，在方格纸上描点连线；发现五个 x 等差且公差为 T/4。',
     '规范作图流程，为后续所有图像题打基础。'),
    ('14–22', '第2幕·8′', 'PPT8–12：分别演示 A、ω、φ、k 单独变化的对比图（图04–07）；'
     '每讲一个参数追问"哪个量变了，哪个量没变"。',
     '同桌讨论并归纳：A 只改值域，ω 只改周期，φ 只改左右，k 只改中心线。',
     '控制变量法探究，培养归纳能力；建立四参数的完整认知框架。'),
    ('22–27', '第3幕·5′', 'PPT13–14：板书 sin(ωx+φ)=sin[ω(x+φ/ω)]；'
     '强调"先提系数，再看平移"。',
     '在错题本上抄下这一恒等变形，自己推导一遍。',
     '从代数上讲清难点本质，为图形辨析做铺垫。'),
    ('27–32', '第3幕·5′', 'PPT15–17：展示先平移后伸缩 vs 先伸缩后平移的四子图（图08）；'
     '再用同图叠加展示错开 π/6 的两条曲线（图09），红笔圈出差距；变换链演练（图16）。',
     '先独立判断"哪条对"，暴露错误后自我修正；口述两条路线的平移量差异。',
     '制造认知冲突，以图形冲击打透最大易错点，印象深刻。'),
    ('32–38', '第4幕·6′', 'PPT19：例1 由图像求解析式（图10）。板演三步：A 看幅度 → ω 看周期 → φ 代最高点。',
     '一名学生上台板演，其余在练习本同步完成；集体订正。',
     '固化"三步定参"的解题程序；强调代最高点避免增根。'),
    ('38–43', '第4幕·5′', 'PPT20：例2 单调区间·对称轴·对称中心（图11）。示范整体代换 u=2x+π/6。',
     '独立完成减区间与对称中心，两名学生报答案。',
     '掌握整体代换这把"万能钥匙"；提示 ω<0 需先化正。'),
    ('43–48', '第4幕·5′', 'PPT21–24：例3 弹簧振子（图12）；补充交流电背景（图13）；'
     '变式1 摩天轮建模（图17）；变式2 变换顺序选择题。',
     '小组完成变式1 的四步建模；变式2 举手表决后说明理由。',
     '把数学模型嵌回真实情境，发展数学建模素养；变式2 二次检测易错点。'),
    ('48–56', '第5幕·8′', 'PPT25–28：先展示三个正弦叠加成复杂波（图15）；'
     '再逐帧展示 1→2→3→5→10→40 项逼近方波（图14）；'
     '点明傅里叶级数与 MP3/JPEG/5G/CT 的关系。',
     '观察并惊叹：光滑的正弦竟能拼出尖角方波；提问"能拼出任何形状吗"。',
     '本课高光。以视觉震撼完成价值升华，回答"学三角函数有什么用"这一根本追问。'),
    ('56–60', '小结·4′', 'PPT29–31：用知识网络图（图18）收束；布置 A/B/C 三层作业；呈现板书提纲。',
     '用三句话复述本课主线；记录作业，按自身层次选做。',
     '结构化收束，分层作业保证下限、抬高上限。'),
]
tb = doc.add_table(rows=1, cols=5)
tb.style = 'Table Grid'
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ['时间', '环节', '教师活动', '学生活动', '设计意图']
for i, h in enumerate(hdr):
    c = tb.rows[0].cells[i]
    c.text = ''
    rr = c.paragraphs[0].add_run(h); rr.bold = True; rr.font.size = DPt(10)
    rr.font.name = 'Noto Serif CJK SC'
    rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for row in rows:
    cells = tb.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = ''
        rr = cells[i].paragraphs[0].add_run(v); rr.font.size = DPt(9.5)
        rr.font.name = 'Noto Serif CJK SC'
        rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for w, i in zip([Cm(1.7), Cm(2.1), Cm(5.6), Cm(3.6), Cm(3.6)], range(5)):
    for row in tb.rows:
        row.cells[i].width = w
P('合计：2 + 6 + 6 + 8 + 5 + 5 + 6 + 5 + 5 + 8 + 4 = 60 分钟。', True)

H('九、板书设计', 1)
P('【左栏·概念区】　单位圆 → sin x；T = 2π/|ω|；y = A sin(ωx+φ) + k；'
  '★（红笔）sin(ωx+φ) = sin[ω(x + φ/ω)]，平移量 = φ/ω')
P('【中栏·演算区】　五点法表格（u = 0, π/2, π, 3π/2, 2π ↔ x ↔ y）；'
  '例1 三步定参；例2 整体代换；例3 简谐四要素')
P('【右栏·图像区】　标准 sin x 图（全课保留）；三参数对比草图；'
  '（红笔）错开 π/6 的两条曲线；收尾手绘傅里叶逼近方波示意')
P('【角落·易错本】　① 平移量 ≠ φ　② ω<0 先化正　③ 五点是 u 的五点，不是 x 的五点')

H('十、分层作业', 1)
P('A 基础层（全体）：用五点法作 y=2sin(x/2+π/3) 一个周期的图像，写出 A、T、φ 与值域；教材 P243 练习 1–3。')
P('B 提高层（多数）：由图像求解析式 2 题（含 |φ|<π/2 的取舍讨论）；'
  '求 y=3sin(π/3−2x) 的单调递增区间与对称中心（先化正）；'
  '潮汐题：水深 h = 5 + 2sin(πt/6)，求吃水 6 m 的货轮可通行时段。')
P('C 挑战层（学有余力）：用 Desmos 画出 sin x + sin3x/3 + … 前 20 项和，'
  '观察并撰写 200 字说明吉布斯现象；查阅并解释"为什么钢琴与小提琴的中央 C 音色不同"。')

H('十一、教学反思（课后填写）', 1)
for t in ['1. 单位圆展开联动图是否真正帮助学生建立"圆—波"联系？学生能否自主复述？',
          '2. 易错对比图（平移量 φ/ω）当堂检测（变式2）的正确率：______%。若低于 80%，需在下节课以微专题回炉。',
          '3. 傅里叶环节的时间是否被前面环节挤占？高光环节必须留足 8 分钟，宁可压缩例3 讲授。',
          '4. 分层作业的 C 层实际完成人数：______。学生反馈：',
          '5. 其他改进设想：']:
    P(t)
    doc.add_paragraph('　')
    doc.add_paragraph('　')

path_doc = os.path.join(OUT, '教案_06_三角函数图像与变换.docx')
doc.save(path_doc)
print('docx:', path_doc)

# 校验
import glob
pngs = [p for p in glob.glob(os.path.join(FIG, '*.png')) if not os.path.basename(p).startswith('fx_')]
print('SLIDES=%d  MAINFIGS=%d  FORMULAS=%d' % (n_slides, len(pngs), _fc[0]))
assert 26 <= n_slides <= 32, n_slides
assert len(pngs) >= 13
print('OK')
