# -*- coding: utf-8 -*-
"""第04讲：数列与数学归纳法 —— 60分钟课件包构建脚本。"""
import sys, os
sys.path.insert(0, '/mnt/d/polaris/教师助手/mathkit')
import deckkit as k
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyArrowPatch, FancyBboxPatch, Circle

OUT = '/mnt/c/Users/mi/Desktop/数学名师课件包/04_数列与数学归纳法'
FIG = os.path.join(OUT, 'figures')
os.makedirs(FIG, exist_ok=True)
F = lambda n: os.path.join(FIG, n)

# ================= 配图 =================
# 画布宽度 = 该图在幻灯片上的最小展示宽度（不得缩小，否则图内字跌破 20pt）

def fig_types():
    """展示宽度 6.0in（右栏）"""
    fig, axes = plt.subplots(2, 2, figsize=(6.0, 5.0))
    n = np.arange(1, 13)
    data = [(2*n - 1, '等差 $2n\\!-\\!1$', k.M_ACC2),
            (1.35**n, '等比 $1.35^{\\,n}$', k.M_GRN),
            ((-1.0)**n * (1 + 3.0/n), '摆动 $(-1)^n$', k.M_ACC),
            (n + n*np.sin(n), '发散 无界', k.M_RED)]
    for ax, (y, t, c) in zip(axes.ravel(), data):
        ax.axhline(0, color=k.M_INK, lw=1.0)
        ax.vlines(n, 0, y, color=k.M_RULE, lw=1.4, zorder=1)
        ax.plot(n, y, 'o', color=c, ms=7, zorder=3)
        ax.set_title(t, color=k.M_INK, pad=8)
        ax.grid(True, color=k.M_RULE, lw=0.6, alpha=0.7, zorder=0)
        for s in ('top', 'right'):
            ax.spines[s].set_visible(False)
        ax.set_xticks([1, 6, 12])
        ax.yaxis.set_major_locator(plt.MaxNLocator(3))
        ax.tick_params(colors=k.M_SLATE)
    return k.save_fig(fig, F('01_数列四类型散点.png'))


def fig_fib():
    """展示宽度 10.6in（整幅）"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10.6, 4.4))
    f = [1, 1]
    for i in range(2, 14):
        f.append(f[-1] + f[-2])
    n = np.arange(1, 15)
    ax1.vlines(n, 0, f, color=k.M_RULE, lw=1.6)
    ax1.plot(n, f, 'o-', color=k.M_ACC2, ms=6, lw=1.8)
    ax1.text(1.0, 300, '$F_n=F_{n-1}+F_{n-2}$\n$F_1=F_2=1$', color=k.M_INK,
             bbox=dict(boxstyle='round,pad=0.4', fc='white', ec=k.M_ACC))
    ax1.set_ylim(0, 620)
    ax1.set_title('递推生成', color=k.M_INK, pad=8)
    ax1.grid(True, color=k.M_RULE, lw=0.6, alpha=0.7)
    ax1.set_xticks([1, 5, 10, 14])
    ax1.yaxis.set_major_locator(plt.MaxNLocator(4))
    for s in ('top', 'right'):
        ax1.spines[s].set_visible(False)
    ax1.tick_params(colors=k.M_SLATE)

    r = [f[i+1] / f[i] for i in range(12)]
    phi = (1 + 5 ** 0.5) / 2
    ax2.axhline(phi, color=k.M_RED, lw=1.8, ls='--')
    ax2.text(5.2, phi + 0.08, '$\\varphi\\approx1.618$', color=k.M_RED)
    ax2.plot(range(1, 13), r, 'o-', color=k.M_GRN, ms=6, lw=1.8)
    ax2.set_ylim(0.9, 2.25)
    ax2.set_title('相邻两项之比 $\\to\\varphi$', color=k.M_INK, pad=8)
    ax2.set_xticks([1, 4, 8, 12])
    ax2.yaxis.set_major_locator(plt.MaxNLocator(4))
    ax2.grid(True, color=k.M_RULE, lw=0.6, alpha=0.7)
    for s in ('top', 'right'):
        ax2.spines[s].set_visible(False)
    ax2.tick_params(colors=k.M_SLATE)
    return k.save_fig(fig, F('02_Fibonacci与黄金比.png'))


def fig_pairing():
    """最小展示宽度 6.6in（另在整幅页放大到 10.4in）"""
    fig, ax = k.new_fig(6.6, 2.95)
    n = 8
    xs = np.arange(1, n + 1)
    ax.bar(xs, xs, color=k.M_ACC2, width=0.74, zorder=2)
    ax.bar(xs, n + 1 - xs, bottom=xs, color=k.M_ACC, width=0.74, alpha=0.9, zorder=2)
    ax.axhline(n + 1, color=k.M_RED, lw=1.8, ls='--', zorder=3)
    ax.text(0.55, n + 1.35, '每列 $=n+1$', color=k.M_RED, va='bottom')
    ax.text(7.0, 3.5, '正序', color='white', ha='center', va='center', rotation=90)
    ax.text(2.0, 5.6, '倒序', color='white', ha='center', va='center', rotation=90)
    ax.set_ylim(0, n + 3.2)
    ax.set_xticks(xs)
    ax.set_yticks([0, 3, 6, 9])
    ax.grid(True, axis='y', color=k.M_RULE, lw=0.6, alpha=0.7, zorder=0)
    for s in ('top', 'right'):
        ax.spines[s].set_visible(False)
    ax.tick_params(colors=k.M_SLATE)
    return k.save_fig(fig, F('03_配对求和图证.png'))


def fig_triangle():
    """展示宽度 10.4in（整幅）"""
    fig, ax = k.new_fig(10.4, 4.5)
    n = 6
    for i in range(1, n + 1):
        for j in range(i):
            ax.add_patch(Circle((j, i - 1), 0.42, fc=k.M_ACC2, ec='white', lw=1.2))
    ax.text(2.5, -1.6, '$T_6=1+2+\\cdots+6=21$', ha='center', color=k.M_INK)
    ox = 9.4
    for i in range(1, n + 1):
        for j in range(n + 1):
            c = k.M_ACC2 if j < i else k.M_ACC
            ax.add_patch(Circle((ox + j, i - 1), 0.42, fc=c, ec='white', lw=1.2))
    ax.add_patch(Rectangle((ox - 0.62, -0.62), n + 1.24, n + 0.24,
                           fill=False, ec=k.M_RED, lw=2))
    ax.annotate('', xy=(ox - 0.85, 2.5), xytext=(6.2, 2.5),
                arrowprops=dict(arrowstyle='-|>', mutation_scale=20, color=k.M_RED, lw=2.4))
    ax.text(7.7, 2.85, '补一个\n倒三角', color=k.M_RED, ha='center', va='bottom')
    ax.text(ox + 3.0, -1.6, '$2T_6=6\\times 7=42$', ha='center', color=k.M_INK)
    ax.set_xlim(-1.2, ox + 7.2)
    ax.set_ylim(-2.4, 6.4)
    ax.set_aspect('equal')
    ax.axis('off')
    ax.set_title('两个三角形数拼成矩形：$2T_n=n(n+1)$', color=k.M_INK)
    return k.save_fig(fig, F('04_三角形数堆叠.png'))


def fig_domino():
    """展示宽度 10.6in（整幅）"""
    fig, ax = k.new_fig(10.6, 4.3)
    xs = np.linspace(1.9, 9.6, 9)
    for i, x in enumerate(xs):
        ang = -20 if i == 0 else 0
        r = Rectangle((x, 0.7), 0.34, 1.7, fc=k.M_ACC2 if i else k.M_ACC,
                      ec=k.M_INK, lw=1.4, angle=ang, rotation_point='xy')
        ax.add_patch(r)
        ax.text(x + 0.17, 0.12, f'{i+1}', ha='center', color=k.M_SLATE)
    for i in range(len(xs) - 1):
        ax.add_patch(FancyArrowPatch((xs[i] + 0.45, 2.7), (xs[i+1] + 0.05, 2.7),
                                     arrowstyle='-|>', mutation_scale=13,
                                     color=k.M_RED, lw=1.3))
    ax.add_patch(FancyArrowPatch((1.15, 3.05), (1.8, 2.25), arrowstyle='-|>',
                                 mutation_scale=18, color=k.M_ACC, lw=2.4))
    ax.text(0.1, 3.2, '奠基步：推倒第 1 块', color=k.M_ACC, weight='bold')
    ax.text(6.6, 3.2, '归纳步：第 $k$ 块倒 $\\Rightarrow$ 第 $k+1$ 块倒',
            color=k.M_RED, ha='center')
    ax.set_xlim(-0.1, 11.1)
    ax.set_ylim(-0.5, 4.0)
    ax.axis('off')
    ax.set_title('多米诺骨牌模型', color=k.M_INK)
    return k.save_fig(fig, F('05_多米诺骨牌.png'))


def fig_flow():
    """最小展示宽度 5.5in"""
    fig, ax = k.new_fig(5.5, 4.2)
    boxes = [(2.95, '① 奠基：验 $n=n_0$', k.M_ACC),
             (1.70, '② 归纳：$P(k)\\Rightarrow P(k+1)$', k.M_ACC2),
             (0.45, '③ 结论：$\\forall n\\geq n_0$ 成立', k.M_GRN)]
    for y, t, c in boxes:
        ax.add_patch(FancyBboxPatch((0.25, y), 5.0, 0.8, boxstyle='round,pad=0.05',
                                    fc='white', ec=c, lw=2.4))
        ax.text(0.45, y + 0.4, t, va='center', color=k.M_INK)
    for y in (2.90, 1.65):
        ax.add_patch(FancyArrowPatch((2.75, y), (2.75, y - 0.4), arrowstyle='-|>',
                                     mutation_scale=20, color=k.M_SLATE, lw=2.0))
    ax.set_xlim(0, 5.5)
    ax.set_ylim(0.1, 4.1)
    ax.axis('off')
    ax.set_title('三步骤：两步缺一不可', color=k.M_INK)
    return k.save_fig(fig, F('06_归纳三步流程图.png'))


def fig_wellorder():
    """展示宽度 7.4in"""
    fig, ax = k.new_fig(7.4, 4.0)
    xs = np.arange(1, 10)
    for x in xs:
        bad = x >= 6
        ax.add_patch(Circle((x, 1.7), 0.42, fc=(k.M_RED if bad else k.M_GRN),
                            ec='white', lw=1.4))
        ax.text(x, 1.7, str(x), ha='center', va='center', color='white')
    ax.text(3.0, 2.7, '$P(n)$ 真', color=k.M_GRN, ha='center')
    ax.text(7.7, 2.7, '反例集 $S$', color=k.M_RED, ha='center')
    ax.annotate('最小反例 $m=6$', xy=(6, 1.2), xytext=(4.4, 0.2),
                color=k.M_INK, ha='center',
                arrowprops=dict(arrowstyle='-|>', color=k.M_INK, lw=1.8))
    ax.annotate('', xy=(6.0, 2.32), xytext=(5.0, 2.32),
                arrowprops=dict(arrowstyle='-|>', color=k.M_ACC, lw=2.6))
    ax.text(0.1, -0.85, '取最小反例 $m$：$P(m-1)$ 真 $\\Rightarrow$ $P(m)$ 真，矛盾',
            color=k.M_INK)
    ax.set_xlim(0, 10.4)
    ax.set_ylim(-1.3, 3.3)
    ax.axis('off')
    ax.set_title('最小反例法', color=k.M_INK)
    return k.save_fig(fig, F('07_最小反例原理.png'))


def fig_nobase():
    """展示宽度 10.6in（整幅）"""
    fig, ax = k.new_fig(10.6, 4.3)
    xs = np.linspace(2.4, 9.8, 9)
    for i, x in enumerate(xs):
        ax.add_patch(Rectangle((x, 0.9), 0.34, 1.7, fc=k.M_RULE, ec=k.M_SLATE, lw=1.3))
        ax.text(x + 0.17, 0.35, f'{i+1}', ha='center', color=k.M_SLATE)
    for i in range(len(xs) - 1):
        ax.add_patch(FancyArrowPatch((xs[i] + 0.45, 2.85), (xs[i+1] + 0.05, 2.85),
                                     arrowstyle='-|>', mutation_scale=13,
                                     color=k.M_ACC2, lw=1.3))
    ax.text(6.2, 3.3, '归纳步完好：$P(k)\\Rightarrow P(k+1)$', color=k.M_ACC2, ha='center')
    ax.plot([1.3, 2.0], [1.4, 2.15], color=k.M_RED, lw=4)
    ax.plot([1.3, 2.0], [2.15, 1.4], color=k.M_RED, lw=4)
    ax.text(1.65, 0.5, '无人推倒\n第 1 块', color=k.M_RED, ha='center', va='top')
    ax.set_xlim(0, 11.2)
    ax.set_ylim(-1.2, 3.9)
    ax.axis('off')
    ax.set_title('错误 I：漏掉奠基步（$P(n)$：$n=n+1$）', color=k.M_INK)
    return k.save_fig(fig, F('08_漏基础步假证明.png'))


def fig_horses():
    """展示宽度 10.2in（整幅）"""
    fig, ax = k.new_fig(10.2, 5.6)

    def horse(x, y, c):
        ax.add_patch(FancyBboxPatch((x, y), 0.6, 0.42, boxstyle='round,pad=0.03',
                                    fc=c, ec=k.M_INK, lw=1.2))

    # 上：k>=2 两组有重叠
    ax.text(0.1, 4.6, '$k\\geq2$：两组有重叠', color=k.M_GRN)
    for i in range(5):
        horse(0.5 + i * 0.78, 3.55, k.M_ACC2)
    ax.add_patch(Rectangle((0.38, 3.42), 3.4, 0.68, fill=False, ec=k.M_GRN, lw=2, ls='--'))
    ax.add_patch(Rectangle((1.04, 3.32), 3.4, 0.68, fill=False, ec=k.M_ACC, lw=2, ls='--'))
    ax.text(4.65, 4.0, '前 $k$ 匹', color=k.M_GRN, va='center')
    ax.text(4.65, 3.2, '后 $k$ 匹', color=k.M_ACC, va='center')

    # 下：k=1 两组不相交
    ax.text(0.1, 2.5, '$k=1$：两组不相交', color=k.M_RED)
    horse(0.6, 1.5, k.M_ACC2)
    horse(2.3, 1.5, k.M_RED)
    ax.add_patch(Rectangle((0.48, 1.37), 0.84, 0.68, fill=False, ec=k.M_GRN, lw=2, ls='--'))
    ax.add_patch(Rectangle((2.18, 1.37), 0.84, 0.68, fill=False, ec=k.M_ACC, lw=2, ls='--'))
    ax.plot([1.55, 1.95], [1.5, 1.95], color=k.M_RED, lw=4)
    ax.plot([1.55, 1.95], [1.95, 1.5], color=k.M_RED, lw=4)
    ax.text(0.3, 0.75, '重叠为空！', color=k.M_RED)
    ax.text(3.35, 1.7, '$P(1)\\not\\Rightarrow P(2)$', color=k.M_RED, va='center')

    # 右：断裂的链条
    ax.text(5.9, 4.6, '链条断在 $n=1\\to n=2$', color=k.M_RED, weight='bold')
    cx = [6.1] + list(np.linspace(7.25, 9.9, 5))
    for i, x in enumerate(cx):
        c = k.M_GRN if i == 0 else k.M_RULE
        ax.add_patch(Circle((x, 3.55), 0.27, fc=c, ec='white', lw=1.2))
        ax.text(x, 3.55, str(i + 1), ha='center', va='center',
                color='white' if i == 0 else k.M_SLATE)
    ax.plot([6.47, 6.89], [3.34, 3.76], color=k.M_RED, lw=4)
    ax.plot([6.47, 6.89], [3.76, 3.34], color=k.M_RED, lw=4)
    ax.text(6.0, 2.5, '$n\\geq2$ 全部失去支撑', color=k.M_SLATE)
    ax.add_patch(FancyBboxPatch((5.9, 0.75), 4.15, 0.9, boxstyle='round,pad=0.05',
                                fc='white', ec=k.M_RED, lw=2))
    ax.text(6.1, 1.2, '归纳步须对每个 $k$ 成立', va='center', color=k.M_INK)

    ax.set_xlim(0, 10.2)
    ax.set_ylim(0.4, 5.15)
    ax.axis('off')
    ax.set_title('"所有马同色"悖论：断裂点在哪？', color=k.M_INK)
    return k.save_fig(fig, F('09_所有马同色悖论.png'))


def fig_bernoulli():
    """最小展示宽度 5.5in"""
    fig, ax = k.new_fig(5.5, 4.2)
    x = np.linspace(-0.8, 1.05, 300)
    ax.plot(x, (1 + x) ** 5, color=k.M_ACC2, lw=2.6)
    ax.plot(x, 1 + 5 * x, color=k.M_GRN, lw=2.0, ls='--')
    ax.fill_between(x, 1 + 5 * x, (1 + x) ** 5, color=k.M_ACC, alpha=0.16)
    ax.plot([0], [1], 'o', color=k.M_RED, ms=10, zorder=5)
    ax.text(0.54, 6.9, '$(1+x)^5$', color=k.M_ACC2)
    ax.text(0.55, 2.0, '$1+5x$', color=k.M_GRN)
    ax.annotate('$x=0$ 取等', xy=(0.02, 1.0), xytext=(-0.85, 4.4),
                color=k.M_RED,
                arrowprops=dict(arrowstyle='-|>', color=k.M_RED, lw=1.6))
    ax.set_xlim(-0.95, 1.2)
    ax.set_ylim(-1.8, 8.2)
    ax.set_xticks([-0.5, 0.5, 1.0])
    ax.set_yticks([2, 4, 6, 8])
    k.style_axes(ax, 'x', 'y')
    ax.set_title('曲线恒在切线上方', color=k.M_INK, pad=26)
    return k.save_fig(fig, F('10_Bernoulli不等式.png'))


def fig_divis():
    """展示宽度 10.2in（整幅）"""
    fig, ax = k.new_fig(10.2, 4.2)
    ax.text(5.1, 3.45, r'$3^{2(k+1)}-1 \;=\; 9\,(3^{2k}-1) \;+\; 8$',
            fontsize=30, ha='center', color=k.M_INK)
    ax.annotate('归纳假设：$8\\,|\\,(3^{2k}-1)$', xy=(5.5, 3.15), xytext=(3.0, 2.2),
                color=k.M_ACC2, ha='center',
                arrowprops=dict(arrowstyle='-|>', color=k.M_ACC2, lw=1.6))
    ax.annotate('显然 $8\\,|\\,8$', xy=(7.6, 3.15), xytext=(8.6, 2.2),
                color=k.M_GRN, ha='center',
                arrowprops=dict(arrowstyle='-|>', color=k.M_GRN, lw=1.6))
    ax.add_patch(FancyBboxPatch((0.6, 0.5), 9.0, 1.0, boxstyle='round,pad=0.06',
                                fc='white', ec=k.M_ACC2, lw=2.2))
    ax.text(5.1, 1.0, '通法："凑" —— 制造出归纳假设的原形',
            ha='center', va='center', color=k.M_INK)
    ax.set_xlim(0, 10.2)
    ax.set_ylim(0.1, 4.2)
    ax.axis('off')
    return k.save_fig(fig, F('11_整除性凑项.png'))


def fig_strong():
    """最小展示宽度 5.4in"""
    fig, ax = k.new_fig(5.4, 3.2)
    xs = np.linspace(0.55, 4.9, 7)
    y1, y2 = 2.42, 1.32
    ax.text(0.05, 2.82, '第一归纳法：只用 $P(k)$', color=k.M_ACC2)
    ax.text(0.05, 1.72, '第二归纳法：用 $P(1..k)$', color=k.M_GRN)
    for x in xs:
        ax.add_patch(Circle((x, y1), 0.13, fc=k.M_ACC2, ec='white', lw=1.0))
        ax.add_patch(Circle((x, y2), 0.13, fc=k.M_GRN, ec='white', lw=1.0))
    for i in range(len(xs) - 1):
        ax.add_patch(FancyArrowPatch((xs[i] + 0.15, y1), (xs[i+1] - 0.15, y1),
                                     arrowstyle='-|>', mutation_scale=12,
                                     color=k.M_SLATE, lw=1.4))
    for i in range(2, len(xs)):
        for j in range(i):
            ax.add_patch(FancyArrowPatch((xs[j], y2 - 0.13), (xs[i], y2 - 0.13),
                                         connectionstyle='arc3,rad=0.2',
                                         arrowstyle='-|>', mutation_scale=8,
                                         color=k.M_ACC, lw=0.8, alpha=0.55))
    ax.text(0.05, 0.15, '需要多远的过去，就假设多远', color=k.M_INK)
    ax.set_xlim(0, 5.4)
    ax.set_ylim(0.0, 3.2)
    ax.axis('off')
    return k.save_fig(fig, F('12_强归纳法依赖图.png'))


def fig_pitfall():
    """展示宽度 5.7in"""
    fig, ax = k.new_fig(5.7, 4.4)
    rows = [('只验证有限个 $n$', k.M_RED),
            ('归纳步没用上假设', k.M_RED),
            ('把归纳假设当结论', k.M_ACC),
            ('起点 $n_0$ 写错', k.M_ACC),
            ('$k+1$ 不与 $k$ 挂钩', k.M_ACC2)]
    ax.text(0.25, 3.95, '五大失分点', color=k.M_INK, weight='bold')
    for i, (t, c) in enumerate(rows):
        y = 3.15 - i * 0.72
        ax.add_patch(Rectangle((0.25, y - 0.28), 5.2, 0.62,
                               fc='white' if i % 2 else '#EFEDE6', ec='none'))
        ax.add_patch(Rectangle((0.25, y - 0.28), 0.1, 0.62, fc=c, ec='none'))
        ax.text(0.6, y + 0.03, t, va='center', color=c)
    ax.set_xlim(0, 5.7)
    ax.set_ylim(-0.1, 4.4)
    ax.axis('off')
    return k.save_fig(fig, F('13_失分点清单.png'))


def fig_map():
    """最小展示宽度 5.6in"""
    fig, ax = k.new_fig(5.6, 4.4)
    nodes = [(2.8, 3.95, '数列 $=\\mathbf{N^*}\\to\\mathbf{R}$', k.M_INK, 3.7),
             (1.35, 3.00, '通项 $a_n=f(n)$', k.M_ACC2, 2.5),
             (4.25, 3.00, '递推 $a_{n+1}$', k.M_ACC2, 2.3),
             (2.8, 2.05, '观察 $\\to$ 猜想', k.M_ACC, 2.7),
             (2.8, 1.15, '数学归纳法（证明）', k.M_RED, 4.0),
             (1.35, 0.30, '第一归纳法', k.M_GRN, 2.4),
             (4.25, 0.30, '第二归纳法', k.M_GRN, 2.4)]
    for x, y, t, c, w in nodes:
        ax.add_patch(FancyBboxPatch((x - w / 2, y - 0.25), w, 0.5,
                                    boxstyle='round,pad=0.05', fc='white', ec=c, lw=2))
        ax.text(x, y, t, ha='center', va='center', color=k.M_INK)
    edges = [((2.4, 3.68), (1.5, 3.30)), ((3.2, 3.68), (4.1, 3.30)),
             ((1.5, 2.72), (2.4, 2.35)), ((4.1, 2.72), (3.2, 2.35)),
             ((2.8, 1.77), (2.8, 1.45)),
             ((2.3, 0.87), (1.5, 0.60)), ((3.3, 0.87), (4.1, 0.60))]
    for a, b in edges:
        ax.add_patch(FancyArrowPatch(a, b, arrowstyle='-|>', mutation_scale=14,
                                     color=k.M_SLATE, lw=1.4))
    ax.set_xlim(0, 5.6)
    ax.set_ylim(-0.05, 4.35)
    ax.axis('off')
    return k.save_fig(fig, F('14_知识结构图.png'))


FIGS = {}
for fn in (fig_types, fig_fib, fig_pairing, fig_triangle, fig_domino, fig_flow,
           fig_wellorder, fig_nobase, fig_horses, fig_bernoulli, fig_divis,
           fig_strong, fig_pitfall, fig_map):
    FIGS[fn.__name__] = fn()
print('figures:', len(FIGS))

# ================= PPT =================
prs = k.new_deck()
_fc = [0]
def fx(tex, slide, **kw):
    _fc[0] += 1
    return k.formula(slide, tex, out=F('formula_%02d.png' % _fc[0]), **kw)

# 1 封面
k.title_slide(prs, '第04讲　数列与数学归纳法',
              '从"看出规律"到"证明规律"——数列的函数观点与归纳法的逻辑内核',
              '数学名师课件包', '60 分钟 · 高中数学 · 选择性必修')

# 2 学习目标
s = k.content_slide(prs, '学习目标与本讲路线', '导入')
k.bullets(s, [
    '理解数列是定义在 N* 上的函数，会用函数观点看通项与图像',
    ('通项公式 ↔ 递推关系：两种"生成"数列的方式', 1),
    '掌握数学归纳法的两步结构，理解"为什么两步就够"',
    ('多米诺骨牌直观 → 最小数原理的严格支撑', 1),
    '能识别并剖析归纳法的典型错误（漏奠基步、"所有马同色"）',
    '会用归纳法证明求和公式、整除性、不等式三类题',
    '了解第二数学归纳法及其必要性',
], y=1.55, w=6.3)
k.callout(s, '本讲高光：一个"看似完美"的证明，能推出"世上所有马同色"。\n错在哪？——这正是理解归纳法的钥匙。', y=5.75, w=6.3, h=1.25, kind='warn')
k.picture(s, FIGS['fig_map'], x=7.2, y=1.9, w=5.6)

# 3 第1幕
k.section_slide(prs, '第 1 幕', '数列：一个"离散的函数"', '0–12 min')

# 4 数列即函数
s = k.content_slide(prs, '数列的本质：定义域为 N* 的函数', '概念')
k.bullets(s, [
    '数列 {aₙ}：把每个正整数 n 对应到一个实数 aₙ',
    ('本质是函数 f: N* → R，只是定义域"离散"', 1),
    '图像：一串孤立的点，不能连成曲线',
    ('横坐标只取 1,2,3,…；"中间"没有值', 1),
    '数列的性质 = 函数性质的离散版：单调、有界、周期',
], y=1.55, w=5.9, size=17)
k.callout(s, '一句话：数列不是"一列数"，是"一个函数"。\n凡函数的眼光（单调/有界/极限），都可移植过来。', y=5.6, w=5.9, h=1.3)
k.picture(s, FIGS['fig_types'], x=6.85, y=1.5, w=6.0)

# 5 通项与递推
s = k.content_slide(prs, '两种生成方式：通项公式 vs 递推关系', '概念')
fx(r'$a_n=f(n)\qquad\Longleftrightarrow\qquad a_{n+1}=g(a_n),\quad a_1\ \mathrm{given}$', s, y=1.7, w=9.6, x=1.85)
k._tb(s, k.Inches(1.85), k.Inches(1.35), k.Inches(9.6), k.Inches(0.4), '通项公式　　　　　　　　　　　　　　　递推关系 + 初值', 15, k.SLATE)
k.bullets(s, [
    '通项公式：直接"点播"——要第 100 项，代 n=100 即得',
    '递推关系：只能"顺序播放"——必须从 a₁ 一步步算到 a₁₀₀',
    ('递推 + 初值 = 唯一确定一个数列（这就是归纳法的雏形！）', 1),
    '核心问题：由递推求通项（累加/累乘/构造/待定系数）',
    ('反向问题：由若干项猜通项 → 必须用归纳法验证', 1),
], y=3.0, w=11.4, size=17)
k.callout(s, '递推式的"存在唯一性"依赖的正是归纳原理：a₁ 定了（奠基），每步能推下一步（归纳），\n于是整个数列被唯一确定。', y=5.95, w=11.4, h=1.1, kind='note')

# 6 Fibonacci
s = k.content_slide(prs, '递推的力量：Fibonacci 与黄金比', '范例')
k.full_picture(s, FIGS['fig_fib'], y=1.45, w=10.6)
k.callout(s, 'Fᵢ 的通项（Binet 公式）存在但极不直观；而递推式 Fₙ=Fₙ₋₁+Fₙ₋₂ 一目了然。\n注意：本递推同时用到前两项 —— 这将成为"第二归纳法"的登场理由。', x=1.35, y=5.95, w=10.6, h=1.2, kind='note')

# 7 第2幕
k.section_slide(prs, '第 2 幕', '从"看出来"到"证出来"', '12–20 min')

# 8 配对求和
s = k.content_slide(prs, '引例：1+2+⋯+n = ?　高斯配对法', '图证')
k.full_picture(s, FIGS['fig_pairing'], y=1.4, w=10.4)
k.callout(s, '正序 + 倒序 逐项相加，每一列都是 n+1，共 n 列 ⇒ 2S = n(n+1)。', x=1.45, y=6.15, w=10.4, h=0.85)

# 9 三角形数
s = k.content_slide(prs, '同一结论的第二种"看见"：三角形数堆叠', '图证')
k.full_picture(s, FIGS['fig_triangle'], y=1.4, w=10.4)
k.callout(s, '图证很美，但它对"每一个 n"都成立吗？我们只画了 n=6。\n图形直觉给出猜想，逻辑演绎才给出证明 —— 数学归纳法登场。', x=1.45, y=5.95, w=10.4, h=1.1, kind='warn')

# 10 多米诺
s = k.content_slide(prs, '多米诺骨牌：归纳法的直观内核', '原理')
k.full_picture(s, FIGS['fig_domino'], y=1.5, w=10.6)
k.callout(s, '推倒第 1 块（奠基）＋ 任一块倒必带倒下一块（归纳）⇒ 全倒。\n两个条件缺一不可，且合起来"刚好够用"。', x=1.35, y=5.85, w=10.6, h=1.15)

# 11 三步流程
s = k.content_slide(prs, '数学归纳法：标准三步骤', '原理')
k.picture(s, FIGS['fig_flow'], x=0.6, y=1.5, w=7.0)
k.bullets(s, [
    '书写规范（评分点）：',
    ('① 当 n=n₀ 时，左=…，右=…，命题成立', 1),
    ('② 假设 n=k(k≥n₀) 时成立，即 P(k) 为真', 1),
    ('③ 则当 n=k+1 时：…（必须用到 P(k)）', 1),
    ('④ 由①②知，∀n≥n₀，命题成立', 1),
], x=7.85, y=1.6, w=5.0, size=16)
k.callout(s, '判分铁律：归纳步中若没有"用上"归纳假设，\n本题按 0 分处理。', x=7.85, y=5.75, w=5.0, h=1.25, kind='warn')

# 12 逻辑内核
s = k.content_slide(prs, '为什么"两步就够"？—— 逻辑内核', '内核')
fx(r'$[\,P(n_0)\ \wedge\ \forall k\geq n_0\,(P(k)\Rightarrow P(k+1))]\ \Longrightarrow\ \forall n\geq n_0,\ P(n)$', s, y=1.55, w=11.4, x=0.95, size=0.92)
k.picture(s, FIGS['fig_wellorder'], x=0.55, y=3.0, w=7.4)
k.bullets(s, [
    '皮亚诺公理视角：',
    ('归纳原理是 N 的定义性公理之一，不可再证', 1),
    '最小数原理视角：',
    ('N* 任一非空子集有最小元', 1),
    ('反证：设反例集非空，取最小反例 m', 1),
    ('m≠n₀（奠基），故 P(m−1) 真', 1),
    ('归纳步 ⇒ P(m) 真，矛盾', 1),
], x=8.2, y=2.95, w=4.7, size=15)

# 13 归纳假设的用法
s = k.content_slide(prs, '归纳假设：一根"必须踩上去"的踏板', '内核')
k.bullets(s, [
    '归纳假设 P(k) 是"临时赋予"的条件，不是已证事实',
    ('它的唯一使命：作为推出 P(k+1) 的跳板', 1),
    '正确用法：把 n=k+1 的表达式"拆"出 n=k 的表达式',
    ('求和：Sₖ₊₁ = Sₖ + aₖ₊₁', 1),
    ('整除：aₖ₊₁ = λ·aₖ + （可整除的余项）', 1),
    ('不等式：(1+x)^{k+1} = (1+x)^k · (1+x) ≥ (1+kx)(1+x)', 1),
    '典型错误：绕开假设，直接对 k+1 硬算 —— 那不是归纳法',
], y=1.55, w=6.6, size=16.5)
k.picture(s, FIGS['fig_pitfall'], x=7.15, y=1.7, w=5.7)
k.callout(s, '口诀："凑出前项，用上假设，配平余项。"', y=6.35, w=6.6, h=0.75)

# 14 第3幕
k.section_slide(prs, '第 3 幕', '当归纳法"证明"出荒谬结论', '20–30 min')

# 15 漏基础步
s = k.content_slide(prs, '错误 I：漏掉奠基步 —— "n = n+1" 的假证明', '陷阱')
k.full_picture(s, FIGS['fig_nobase'], y=1.5, w=10.6)
k.callout(s, '归纳步 100% 正确，结论 100% 荒谬。原因：P(1) 为假，链条根本没有起点。\n结论：奠基步不是"走过场"，它是整条链的地基。', x=1.35, y=5.9, w=10.6, h=1.15, kind='warn')

# 16 所有马同色 —— 提出
s = k.content_slide(prs, '错误 II：著名悖论"所有马都是同一颜色"', '高光')
k.bullets(s, [
    '命题 P(n)：任意 n 匹马，颜色都相同。',
    '奠基步：n=1。一匹马当然与自己同色。P(1) 真。 ✔',
    '归纳步：设 P(k) 真。取任意 k+1 匹马 h₁,…,h_{k+1}。',
    ('去掉最后一匹：{h₁,…,h_k} 共 k 匹，由假设同色', 1),
    ('去掉第一匹：{h₂,…,h_{k+1}} 共 k 匹，由假设同色', 1),
    ('两组共有 h₂,…,h_k，颜色被"串"起来 ⇒ k+1 匹全同色 ✔', 1),
    '结论：所有马同色。 —— 显然错了。错在哪一步？',
], y=1.55, w=12.0, size=17)
k.callout(s, '请学生独立思考 2 分钟：逐句检查，找出第一处不成立的地方。', y=6.3, w=12.0, h=0.8, kind='warn')

# 17 所有马同色 —— 剖析
s = k.content_slide(prs, '破案：链条断在 n = 1 → n = 2', '高光')
k.full_picture(s, FIGS['fig_horses'], y=1.4, w=10.2)

# 18 悖论小结
s = k.content_slide(prs, '悖论的启示：归纳步必须"对每一个 k"成立', '反思')
k.bullets(s, [
    '归纳步的真正含义：∀k≥n₀，P(k) ⇒ P(k+1)',
    ('是一个"全称命题"，不是"某个 k"', 1),
    '马悖论中，k≥2 时论证有效，唯独 k=1 时"重叠部分为空"',
    ('从 P(1) 推不出 P(2)，其后所有环节全部悬空', 1),
    '所以：写归纳步时，务必检查论证对最小的那个 k 是否成立',
    ('尤其当论证里出现"中间项""重叠部分""至少两个"时', 1),
], y=1.55, w=6.8, size=17)
k.picture(s, FIGS['fig_flow'], x=7.35, y=1.7, w=5.5)
k.callout(s, '奠基步 + 归纳步，二者都必须"全须全尾"。\n一个假的起点，或一个漏洞百出的 k，都足以毁掉整个证明。', y=6.15, w=6.8, h=1.05, kind='warn')

# 19 第4幕
k.section_slide(prs, '第 4 幕', '三类经典题型 · 完整板演', '30–50 min')

# 20 例1 题面
s = k.content_slide(prs, '例1（求和公式）用归纳法证明 1+2+⋯+n = n(n+1)/2', '例题')
fx(r'$\sum_{i=1}^{n} i \;=\; \frac{n(n+1)}{2}\qquad (n\in\mathbf{N^*})$', s, y=1.6, w=9.0, x=2.15)
k.bullets(s, [
    '【奠基】n=1：左 = 1，右 = 1·2/2 = 1，左=右，成立。',
    '【假设】设 n=k(k≥1) 时成立，即 1+2+⋯+k = k(k+1)/2。',
    '【递推】当 n=k+1 时：',
], y=3.2, w=11.6, size=17)
fx(r'$1+2+\cdots+k+(k+1)\;\overset{\rm IH}{=}\;\frac{k(k+1)}{2}+(k+1)=\frac{(k+1)(k+2)}{2}$', s, y=4.6, w=10.6, x=1.35, size=0.86)
k.callout(s, '恰为 n=k+1 时的右端形式 ⇒ P(k+1) 真。由①②，对一切 n∈N* 命题成立。∎', y=6.3, w=11.6, h=0.8)

# 21 例1 复盘 + 配对图
s = k.content_slide(prs, '例1 复盘：图证 vs 归纳证，各管什么', '例题')
k.picture(s, FIGS['fig_pairing'], x=0.5, y=1.6, w=6.6)
k.bullets(s, [
    '配对法（图证）：',
    ('优点：给出"为什么是 n(n+1)/2"的直觉来源', 1),
    ('缺点：论证依赖图形排布，n 为奇数时需另作说明', 1),
    '归纳法：',
    ('优点：逻辑封闭，对一切 n 一次性覆盖', 1),
    ('缺点：不告诉你公式从哪来（只验证，不发现）', 1),
    '教学观：先"猜"再"证"，二者互补，缺一不可。',
], x=7.4, y=1.6, w=5.5, size=15.5)
k.callout(s, '归纳法是"验证器"，不是"发现器"。', x=7.4, y=6.35, w=5.5, h=0.75, kind='note')

# 22 例2
s = k.content_slide(prs, '例2（整除性）求证 3²ⁿ − 1 能被 8 整除', '例题')
fx(r'$8\ \left|\ \left(3^{2n}-1\right)\right.,\qquad n\in\mathbf{N^*}$', s, y=1.6, w=8.6, x=2.35)
k.bullets(s, [
    '【奠基】n=1：3² − 1 = 8，被 8 整除。✔',
    '【假设】设 n=k 时 8 | (3^{2k} − 1)，即 3^{2k} − 1 = 8m（m∈Z）。',
    '【递推】n=k+1 时：3^{2(k+1)} − 1 = 9·3^{2k} − 1 = 9(3^{2k} − 1) + 8 = 9·8m + 8 = 8(9m+1)。',
    '9m+1 ∈ Z，故 8 | (3^{2(k+1)} − 1)，P(k+1) 真。',
    '由①②知，∀n∈N*，8 | (3²ⁿ − 1)。∎',
], y=3.05, w=12.0, size=16.5)
k.callout(s, '要害在"9(3^{2k}−1) + 8"这一步：把 9·3^{2k} 拆成 9(3^{2k}−1) + 9，再与 −1 合并。\n这就是"凑归纳假设"的标准动作。', y=6.05, w=12.0, h=1.05)

# 23 例2 通法图
s = k.content_slide(prs, '例2 通法："凑"出归纳假设的原形', '通法')
k.full_picture(s, FIGS['fig_divis'], y=1.55, w=10.2)
k.callout(s, '变式思路：证 aⁿ − bⁿ 被 (a−b) 整除，同样是拆成 a·(a^{k}−b^{k}) + b^{k}(a−b)。', x=1.55, y=5.85, w=10.2, h=0.85, kind='note')

# 24 例3 题面 + 图
s = k.content_slide(prs, '例3（不等式）Bernoulli 不等式 (1+x)ⁿ ≥ 1+nx', '例题')
k.picture(s, FIGS['fig_bernoulli'], x=0.5, y=1.55, w=6.5)
fx(r'$(1+x)^n\ \geq\ 1+nx\quad (x>-1,\ n\in\mathbf{N^*})$', s, x=7.15, y=1.6, w=5.7, size=0.62)
k.bullets(s, [
    '【奠基】n=1：左 = 1+x = 右，取等号，成立。',
    '【假设】设 n=k 时 (1+x)^k ≥ 1+kx。',
    '【递推】因 x > −1，故 1+x > 0，可保号相乘：',
], x=7.15, y=2.75, w=5.7, size=14.5)
fx(r'$(1+x)^{k+1}\geq(1+kx)(1+x)=1+(k+1)x+kx^2\ \geq\ 1+(k+1)x$', s, x=7.1, y=4.75, w=5.9, size=0.5)
k.callout(s, '关键：1+x>0 才能"不变号地"两边同乘 —— 条件 x>−1 正用于此；\n末尾丢掉 kx² ≥ 0 完成放缩。', x=7.15, y=5.85, w=5.7, h=1.25, kind='warn')

# 25 例3 复盘
s = k.content_slide(prs, '例3 复盘：不等式归纳的两个"动作"', '通法')
k.bullets(s, [
    '动作一：保号相乘 —— 把 (1+x)^{k+1} 拆成 (1+x)^k·(1+x)，',
    ('对前一因子用归纳假设，须确保后一因子为正', 1),
    '动作二：合理放缩 —— 得到 1+(k+1)x + kx² 后，',
    ('丢掉非负项 kx² ≥ 0，正好落到目标式 1+(k+1)x', 1),
    '放缩方向必须"朝目标走"：放大左端 / 缩小右端，不可反向',
    '取等条件：x=0 或 n=1 —— 答题时应交代',
], y=1.6, w=6.8, size=16.5)
k.picture(s, FIGS['fig_bernoulli'], x=7.35, y=1.9, w=5.5)
k.callout(s, '常见失分：忘记 1+x>0 就直接两边同乘（不等号可能反向）。', y=6.25, w=6.8, h=0.8, kind='warn')

# 26 变式1
s = k.content_slide(prs, '变式1　求证 1²+2²+⋯+n² = n(n+1)(2n+1)/6', '变式')
fx(r'$\sum_{i=1}^{n} i^2=\frac{n(n+1)(2n+1)}{6}$', s, y=1.55, w=7.4, x=2.95, size=0.85)
k.bullets(s, [
    '【奠基】n=1：左 = 1，右 = 1·2·3/6 = 1。✔',
    '【假设】n=k 时 Σᵏ i² = k(k+1)(2k+1)/6。',
    '【递推】Σ^{k+1} i² = k(k+1)(2k+1)/6 + (k+1)²',
    ('= (k+1)[k(2k+1) + 6(k+1)] / 6 = (k+1)(2k²+7k+6)/6', 1),
    ('= (k+1)(k+2)(2k+3)/6 = (k+1)[(k+1)+1][2(k+1)+1]/6 ✔', 1),
    '要点：提取公因式 (k+1) 后，把二次式因式分解，凑成目标形。',
], y=3.05, w=12.0, size=16.5)
k.callout(s, '目标意识：先写出 n=k+1 时"应该长什么样"，再朝它变形 —— 有靶子才好射。', y=6.35, w=12.0, h=0.8)

# 27 变式2
s = k.content_slide(prs, '变式2　求证 2ⁿ > n²（n ≥ 5）', '变式')
k.bullets(s, [
    '【留意起点】n=1: 2>1 ✔；n=2: 4=4 ✘；n=3: 8<9 ✘；n=4: 16=16 ✘；n=5: 32>25 ✔',
    ('⇒ 命题的正确起点是 n₀ = 5，奠基步必须验 n=5，不能验 n=1！', 1),
    '【假设】设 k≥5 时 2^k > k²。',
    '【递推】2^{k+1} = 2·2^k > 2k² = k² + k² ≥ k² + 5k = k² + 2k + 3k',
    ('因 k≥5，故 3k ≥ 15 > 1，得 2^{k+1} > k² + 2k + 1 = (k+1)² ✔', 1),
    '由①②，∀n≥5，2ⁿ > n²。∎',
], y=1.55, w=12.0, size=16.5)
k.callout(s, '本题正是"起点写错"这一失分点的靶场：奠基步必须落在 n₀ 上，\n且归纳步中的放缩用到了 k≥5 这一条件（k≥3 即可，但需与奠基一致）。', y=5.95, w=12.0, h=1.15, kind='warn')

# 28 第5幕
k.section_slide(prs, '第 5 幕', '强归纳法：当"一根拐杖"不够用', '50–56 min')

# 29 强归纳法
s = k.content_slide(prs, '第二数学归纳法（强归纳法）', '拓展')
fx(r'$[\,P(n_0)\ \wedge\ \forall k\ (P(n_0)\wedge\cdots\wedge P(k)\Rightarrow P(k+1))]\Rightarrow \forall n\geq n_0,\ P(n)$', s, y=1.5, w=11.6, x=0.85, size=0.78)
k.picture(s, FIGS['fig_strong'], x=0.5, y=2.75, w=7.2)
k.bullets(s, [
    '差别只在归纳假设：',
    ('第一归纳法：只假设 P(k)', 1),
    ('第二归纳法：假设 P(n₀),…,P(k) 全部为真', 1),
    '必要场景：',
    ('Fibonacci：F_{k+1} 需要 F_k 和 F_{k−1}', 1),
    ('素因数分解：n=ab，a,b 可远小于 n−1', 1),
    ('博弈/递归算法的正确性证明', 1),
    '两者等价（都可由最小数原理导出）。',
], x=8.0, y=2.7, w=4.9, size=14.5)

# 30 强归纳法范例
s = k.content_slide(prs, '强归纳法范例：每个 n≥2 都可分解为素数之积', '拓展')
k.bullets(s, [
    '【奠基】n=2：2 本身是素数，成立。',
    '【假设】设 2 ≤ j ≤ k 时，j 均可写成素数之积（用到全部前项！）。',
    '【递推】考察 k+1：',
    ('情形一：k+1 是素数 —— 它本身即为分解，成立。', 1),
    ('情形二：k+1 是合数 —— 则 k+1 = a·b，其中 2 ≤ a,b ≤ k。', 1),
    ('由强归纳假设，a、b 各自可分解为素数之积，相乘即得 k+1 的分解 ✔', 1),
    '第一归纳法为何不行？—— 只知 P(k)，但 a,b 未必等于 k！',
], y=1.55, w=6.9, size=16)
k.picture(s, FIGS['fig_strong'], x=7.45, y=2.0, w=5.4)
k.callout(s, '"需要多远的过去，就假设多远的过去。"', y=6.3, w=6.9, h=0.75, kind='note')

# 31 小结
s = k.content_slide(prs, '本讲小结', '小结')
k.picture(s, FIGS['fig_map'], x=0.45, y=1.55, w=7.0)
k.bullets(s, [
    '① 数列 = N* 上的函数；递推+初值唯一确定数列',
    '② 归纳法两步：奠基 + 归纳，缺一不可',
    '③ 逻辑内核：最小数原理 / 皮亚诺归纳公理',
    '④ 归纳假设必须"用上"，且归纳步须对每个 k 成立',
    '⑤ 马悖论：断裂点在 k=1（重叠为空）',
    '⑥ 三类题：求和 / 整除 / 不等式，各有"凑"法',
    '⑦ 强归纳法：需要多远的过去，就假设多远',
], x=7.75, y=1.6, w=5.2, size=14.5)
k.callout(s, '归纳法是"验证器"：\n先观察猜想，再严格证明。', x=7.75, y=6.15, w=5.2, h=1.0)

# 32 作业
s = k.content_slide(prs, '分层作业', '作业')
k.bullets(s, [
    'A 层（必做，全体）',
    ('1. 用归纳法证 1+3+5+⋯+(2n−1) = n²', 1),
    ('2. 证 5ⁿ − 1 能被 4 整除', 1),
    ('3. 写出归纳法三步骤的规范格式（默写）', 1),
    'B 层（提高，多数）',
    ('4. 证 1/(1·2)+1/(2·3)+⋯+1/[n(n+1)] = n/(n+1)', 1),
    ('5. 证 n! > 2ⁿ（n ≥ 4）', 1),
    ('6. 指出下列"证明"的错误：所有正整数都相等（提示：检查 k=1）', 1),
], y=1.55, w=6.4, size=15.5)
k.bullets(s, [
    'C 层（挑战，学有余力）',
    ('7. 用强归纳法证 Fₙ < (7/4)ⁿ', 1),
    ('8. 平面上 n 条直线两两相交、无三线共点，', 1),
    ('   求证：被分成 1 + n(n+1)/2 个区域', 1),
    ('9. 自编一个"看似正确实则错误"的归纳证明，', 1),
    ('   并指出其断裂点（下节课展示）', 1),
], x=7.05, y=1.55, w=5.8, size=15.5)
k.callout(s, '提交要求：B、C 层须写出完整的三步骤，凡未使用归纳假设者不予计分。', x=7.05, y=5.8, w=5.8, h=1.0, kind='warn')

# 33 板书提纲
s = k.content_slide(prs, '板书设计提纲', '板书')
k.bullets(s, [
    '【左栏｜概念区】',
    ('数列 = f: N* → R', 1),
    ('通项 aₙ=f(n) ⇄ 递推 aₙ₊₁=g(aₙ)', 1),
    ('归纳法：① 奠基 n=n₀  ② 假设 P(k)  ③ 推 P(k+1)', 1),
    ('逻辑根据：最小数原理（最小反例矛盾）', 1),
    '【中栏｜板演区】',
    ('例1 Σi = n(n+1)/2　（完整三步）', 1),
    ('例2 8 | 3²ⁿ−1　（凑：9(3^{2k}−1)+8）', 1),
    ('例3 (1+x)ⁿ ≥ 1+nx　（保号相乘 + 丢 kx²）', 1),
], y=1.5, w=6.4, size=15)
k.bullets(s, [
    '【右栏｜警示区】',
    ('✗ 漏奠基 → "n=n+1"', 1),
    ('✗ 归纳步漏 k → 马同色（k=1 重叠为空）', 1),
    ('✗ 不用假设 → 归纳法名存实亡', 1),
    ('✗ 起点写错 → 2ⁿ>n² 应从 n=5 起', 1),
    '【副板｜拓展】',
    ('第二归纳法：假设 P(n₀..k) 全真', 1),
    ('用例：Fibonacci、素因数分解', 1),
], x=7.05, y=1.5, w=5.8, size=15)
k.callout(s, '板书保留至下课：警示区四条是本讲的"错题防火墙"。', x=7.05, y=6.3, w=5.8, h=0.8)

path = k.save(prs, os.path.join(OUT, '04_数列与数学归纳法.pptx'))
print('pptx:', path, 'slides:', len(prs.slides.__iter__.__self__._sldIdLst))

# ================= 教案 docx =================
from docx import Document
from docx.shared import Pt as DPt, Cm, RGBColor as DRGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Noto Serif CJK SC'
st.font.size = DPt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = sec.bottom_margin = Cm(2.0)

def H(t, lv=1):
    p = doc.add_heading('', level=lv)
    r = p.add_run(t)
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    r.font.color.rgb = DRGB(0x1B, 0x2A, 0x4A)
    r.font.size = DPt(16 if lv == 1 else 13)
    r.font.bold = True
    return p

def P(t, bold=False, ind=0):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(ind)
    p.paragraph_format.space_after = DPt(4)
    r = p.add_run(t); r.bold = bold
    r.font.name = 'Noto Serif CJK SC'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
    return p

ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run('第04讲　数列与数学归纳法　教学设计')
r.font.size = DPt(20); r.bold = True
r.font.name = 'Noto Serif CJK SC'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('课型：新授课（概念 + 方法）　　课时：1 课时（60 分钟）　　授课对象：高二年级')
rs.font.size = DPt(10.5); rs.font.color.rgb = DRGB(0x4A, 0x55, 0x68)
rs.font.name = 'Noto Serif CJK SC'
rs._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')

H('一、教材分析')
P('本讲处于"数列"单元的枢纽位置。前接数列的概念与等差、等比数列，后启数列求和与极限思想。'
  '教材以"数列是定义在正整数集上的函数"统摄全章，而数学归纳法则是高中阶段唯一一种'
  '"处理无穷多个命题"的严格证明工具，是从有限走向无限的第一次思维跨越。', ind=0.74)
P('教材通常把数学归纳法安排在数列之后，原因有二：其一，数列的递推结构（由 a₁ 与 aₙ→aₙ₊₁ 唯一确定整个数列）'
  '本身就是归纳原理的具体化身；其二，数列命题（求和公式、整除性、不等式）为归纳法提供了最自然的练兵场。'
  '本讲将这一内在联系显性化，并进一步揭示归纳原理背后的逻辑根据（最小数原理 / 皮亚诺归纳公理），'
  '以及第二数学归纳法的必要性，使学生的理解从"会用"提升到"知其所以然"。', ind=0.74)

H('二、学情分析')
P('学生已掌握：函数的基本观念、等差等比数列的通项与求和、简单的不完全归纳（由几项猜规律）。', ind=0.74)
P('主要困难有三：（1）把"假设 n=k 成立"误认为"这就是要证的结论"，逻辑上循环论证的疑虑长期存在；'
  '（2）归纳步中不会"凑"出归纳假设的原形，写成对 k+1 的直接硬算，归纳法名存实亡；'
  '（3）轻视奠基步，认为"验一下 n=1 是走过场"。', ind=0.74)
P('对策：以多米诺骨牌建立直观，以最小反例法给出逻辑保证，再以两个"荒谬结论"的反例（n=n+1、所有马同色）'
  '制造强烈认知冲突，使两步的必要性从"老师说要写"变成"不写就会出事"。', ind=0.74)

H('三、教学目标（三维）')
P('1．知识与技能：理解数列的函数本质与递推—通项的关系；掌握数学归纳法的两步结构与规范书写；'
  '能用归纳法证明求和公式、整除性、不等式三类命题；了解第二数学归纳法。', ind=0.74)
P('2．过程与方法：经历"观察—猜想—证明"的完整数学发现过程；通过剖析错误证明，'
  '发展批判性审视与逻辑纠错能力；体会图形直观与逻辑演绎的分工与互补。', ind=0.74)
P('3．情感态度与价值观：在"所有马同色"悖论的思辨中体验数学的严谨之美与思维的乐趣；'
  '认识到"看起来对"与"确实对"之间的鸿沟，养成对论证负责的科学态度。', ind=0.74)

H('四、教学重点与难点')
P('重点：数学归纳法的两步结构、规范书写，以及三类典型题的"凑归纳假设"通法。', bold=False, ind=0.74)
P('难点：（1）理解"两步为何就够"——归纳原理的逻辑根据；'
  '（2）识别归纳步中的隐蔽漏洞（马悖论中 k=1 时重叠为空）；'
  '（3）第二归纳法的必要性判断。', ind=0.74)
P('突破策略：直观（多米诺）→ 逻辑（最小反例）→ 反例（两个荒谬证明）→ 通法（三类题）→ 拓展（强归纳）。', ind=0.74)

H('五、教法与学法')
P('教法：问题驱动 + 反例教学 + 变式训练。以"图证能算证明吗"引出归纳法，'
  '以"所有马同色"制造认知冲突，以板演示范规范书写。', ind=0.74)
P('学法：独立思考（找悖论断裂点）、同伴互评（互查归纳步是否用上假设）、'
  '结构化模仿（三步骤模板）、分层练习。', ind=0.74)

H('六、教学准备')
P('教师：PPT（33 页）、matplotlib 精绘配图 14 张、实物多米诺骨牌一组（可选）、板书分区规划。', ind=0.74)
P('学生：草稿纸、双色笔（红笔专用于标注归纳假设的使用处）。', ind=0.74)

H('七、教学过程（分钟级时间轴）')

rows = [
    ('0–2', '封面/学习目标\nPPT 1–2', '开门见山抛出"我能证明世上所有马都是同一颜色"，展示目标清单。',
     '产生疑惑与兴趣，浏览目标。', '以悖论作钩子，建立本课的核心悬念。'),
    ('2–6', '数列即函数\nPPT 3–4', '出示四类数列散点图（等差/等比/摆动/发散），追问：这四幅图的共同点是什么？'
     '引导得出"定义域是 N*，图像是孤立点"。',
     '观察比较，回答"都是离散的点"，归纳出数列是 N*→R 的函数。',
     '以函数观点统摄数列，为后续用函数性质（单调、有界）分析数列铺路。'),
    ('6–9', '通项与递推\nPPT 5', '对比"点播"与"顺序播放"：通项可直接求第 100 项，递推必须逐步推。'
     '追问：递推式 + 初值为什么能唯一确定整个数列？',
     '尝试回答，意识到"a₁ 定了，每步能推下一步"。',
     '让学生自己说出归纳原理的雏形，为归纳法作认知铺垫。'),
    ('9–12', 'Fibonacci\nPPT 6', '展示 Fibonacci 递推与相邻项之比收敛到黄金比的折线图。'
     '点出：本递推同时用到前两项。',
     '观察振荡收敛现象，记下"用到前两项"这一伏笔。',
     '埋下第二归纳法必要性的伏笔，同时呈现数学之美。'),
    ('12–16', '两幅图证\nPPT 7–9', '演示高斯配对法与三角形数堆叠，得出 S=n(n+1)/2。'
     '追问：我们只画了 n=8 和 n=6，凭什么说对一切 n 成立？',
     '欣赏图证，随即陷入困惑：图只画了有限个 n。',
     '制造第一次认知冲突：直观给出猜想，但不构成证明。归纳法的登场变得必要。'),
    ('16–19', '多米诺骨牌\nPPT 10', '出示骨牌链动画/实物：问"要让全部倒下，最少需要保证哪两件事？"'
     '板书两条件。',
     '讨论后答出：推倒第一块；任一块倒能带倒下一块。',
     '用生活直观锚定抽象结构，两步的"必要且充分"感由学生自己得出。'),
    ('19–22', '三步骤 + 逻辑内核\nPPT 11–12', '给出规范三步骤模板；再以最小反例法解释"两步为何就够"：'
     '若有反例，取最小反例 m，由奠基 m≠n₀，故 P(m−1) 真，归纳步推出 P(m) 真，矛盾。',
     '跟随推理，理解归纳原理并非"循环论证"，其根据是 N* 的最小数原理。',
     '突破难点一：从"老师规定"上升到"逻辑必然"，消除循环论证的疑虑。'),
    ('22–24', '归纳假设的用法\nPPT 13', '强调归纳假设是"临时条件"，其使命是当跳板；给出三类题的"凑"法口诀。',
     '记录口诀："凑出前项，用上假设，配平余项。"',
     '把方法论提前显性化，后面的例题即是口诀的落实。'),
    ('24–27', '错误 I：漏奠基\nPPT 14–15', '给出"n=n+1"的假证明：归纳步完全正确！请学生找茬。',
     '检查后发现 P(1) 为假，链条无起点。',
     '第二次认知冲突：让学生亲身体会"奠基步不是走过场"。'),
    ('27–33', '错误 II：所有马同色\nPPT 16–18', '逐句展示悖论证明，给足 2 分钟静思，再请学生指认断裂点；'
     '最后用双框重叠图揭示 k=1 时"重叠为空"。',
     '独立检查每一句 → 同伴讨论 → 指出 n=1→2 处失效。',
     '本讲高光：突破难点二。让学生理解归纳步是全称命题，必须对每一个 k（尤其最小的 k）成立。'),
    ('33–38', '例1 求和公式\nPPT 19–21', '完整板演三步骤，红笔圈出使用归纳假设之处；随后复盘"图证 vs 归纳证"的分工。',
     '同步在草稿上书写，用红笔标注归纳假设的使用处。',
     '示范规范书写，落实评分点；明确归纳法是验证器而非发现器。'),
    ('38–43', '例2 整除性\nPPT 22–23', '板演 8|(3²ⁿ−1)；重点讲解 9·3^{2k}−1 = 9(3^{2k}−1)+8 的"凑"法，'
     '并推广到 (a−b)|(aⁿ−bⁿ)。',
     '模仿"凑"的动作，尝试口述推广形式。',
     '把通法从个例中抽象出来，形成可迁移的解题模式。'),
    ('43–48', '例3 Bernoulli\nPPT 24–25', '板演不等式归纳：保号相乘（须 1+x>0）与丢项放缩（kx²≥0）；'
     '结合曲线—切线图说明几何含义。',
     '注意到条件 x>−1 的用处，理解放缩必须朝目标走。',
     '不等式是归纳法最易失分处，通过图形赋予"曲线在切线上方"的几何意义。'),
    ('48–52', '变式 1、2\nPPT 26–27', '变式1：平方和公式，示范"先写目标形，再朝它变形"；'
     '变式2：2ⁿ>n²，故意先验 n=1，让学生发现 n=2,3,4 均不成立 → 起点应为 n=5。',
     '两人一组互评：检查对方是否用上假设、起点是否正确。',
     '变式2 是"起点写错"这一失分点的靶场，在错误中确认奠基步的位置至关重要。'),
    ('52–56', '第二归纳法\nPPT 28–30', '回收 Fibonacci 伏笔：F_{k+1} 需要前两项 → 单一假设不够。'
     '再以素因数分解为例（n=ab，a、b 可远小于 k）演示强归纳。',
     '理解"需要多远的过去，就假设多远的过去"。',
     '突破难点三；同时闭合课首伏笔，形成结构完整的认知回路。'),
    ('56–59', '小结 + 作业\nPPT 31–32', '用知识结构图串联全课：数列→猜想→归纳法证明；布置 A/B/C 三层作业。',
     '口头复述七条要点；记录作业。',
     '结构化回顾，分层作业照顾差异，C 层"自编错误证明"指向高阶思维。'),
    ('59–60', '板书回看\nPPT 33', '指向板书右栏"警示区"四条，作为课堂收束。',
     '拍照/抄录警示区四条。',
     '以错题防火墙收尾，强化本讲最易失分处的记忆。'),
]

tb = doc.add_table(rows=1, cols=5)
tb.style = 'Table Grid'
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ['时间', '环节 / PPT', '教师活动', '学生活动', '设计意图']
for i, h in enumerate(hdr):
    c = tb.rows[0].cells[i]
    c.text = ''
    rr = c.paragraphs[0].add_run(h)
    rr.bold = True; rr.font.size = DPt(10)
    rr.font.name = 'Noto Serif CJK SC'
    rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
widths = [Cm(1.9), Cm(2.9), Cm(5.4), Cm(3.6), Cm(4.0)]
for r_ in rows:
    cells = tb.add_row().cells
    for i, v in enumerate(r_):
        cells[i].text = ''
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = DPt(2)
        rr = p.add_run(v)
        rr.font.size = DPt(9)
        rr.font.name = 'Noto Serif CJK SC'
        rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Serif CJK SC')
for row in tb.rows:
    for i, c in enumerate(row.cells):
        c.width = widths[i]
P('')
P('合计：2+4+3+3+4+3+3+2+3+6+5+5+5+4+4+3+1 = 60 分钟（精确覆盖一节 60 分钟课时）。', bold=True)

H('八、板书设计')
P('黑板三分区（保留至下课）：', bold=True)
P('【左栏·概念区】数列 = f: N*→R；通项 aₙ=f(n) ⇄ 递推 aₙ₊₁=g(aₙ)；'
  '归纳法三步：① 奠基 n=n₀　② 假设 P(k)　③ 推出 P(k+1)；逻辑根据：最小数原理（最小反例矛盾）。', ind=0.74)
P('【中栏·板演区】例1 Σi = n(n+1)/2（完整三步，红笔圈出归纳假设使用处）；'
  '例2 8|(3²ⁿ−1)（凑：9(3^{2k}−1)+8）；例3 (1+x)ⁿ≥1+nx（保号相乘 + 丢 kx²）。', ind=0.74)
P('【右栏·警示区】✗ 漏奠基 → "n=n+1"；✗ 归纳步漏 k → 马同色（k=1 重叠为空）；'
  '✗ 不用假设 → 归纳法名存实亡；✗ 起点写错 → 2ⁿ>n² 应从 n=5 起。', ind=0.74)
P('【副板·拓展】第二归纳法：假设 P(n₀..k) 全真；用例：Fibonacci、素因数分解。', ind=0.74)

H('九、分层作业')
P('A 层（必做，全体）：', bold=True)
P('1．用归纳法证 1+3+5+⋯+(2n−1)=n²。　2．证 5ⁿ−1 能被 4 整除。　3．默写归纳法三步骤规范格式。', ind=0.74)
P('B 层（提高，多数）：', bold=True)
P('4．证 1/(1·2)+1/(2·3)+⋯+1/[n(n+1)] = n/(n+1)。　5．证 n! > 2ⁿ（n≥4）。　'
  '6．指出"所有正整数都相等"这一伪证的错误（提示：检查 k=1）。', ind=0.74)
P('C 层（挑战，学有余力）：', bold=True)
P('7．用第二数学归纳法证 Fₙ < (7/4)ⁿ。　'
  '8．平面上 n 条直线两两相交且无三线共点，求证平面被分成 1+n(n+1)/2 个区域。　'
  '9．自编一个"看似正确实则错误"的归纳证明，并指出其断裂点（下节课展示）。', ind=0.74)
P('批改标准：B、C 层必须写出完整三步骤；凡归纳步中未实际使用归纳假设者，不予计分。', bold=True, ind=0.74)

H('十、教学反思（课后填写）')
for t in ['1．学生指认"所有马同色"断裂点所用时间与正确率：',
          '2．归纳步"未用上假设"的错误在课堂练习中出现的比例：',
          '3．变式2 起点判断（n₀=5）的学生自主发现情况：',
          '4．第二归纳法的必要性是否被真正理解（可用 Fibonacci 追问检测）：',
          '5．时间分配是否需要调整（哪一环节超时/不足）：',
          '6．改进设想：']:
    P(t)
    doc.add_paragraph('　').paragraph_format.space_after = DPt(0)
    pp = doc.add_paragraph('_' * 78)
    pp.paragraph_format.space_after = DPt(8)

dp = os.path.join(OUT, '教案_04_数列与数学归纳法.docx')
doc.save(dp)
print('docx:', dp)
